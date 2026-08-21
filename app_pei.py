import streamlit as st
from datetime import datetime, date, timedelta, timezone
from fpdf import FPDF
import datetime as dt_module
import io
import os   
import base64
import urllib.request
import json
import tempfile    
from PIL import Image
import pandas as pd
from supabase import create_client, Client
import time
import uuid
import threading
import random
import time
import zipfile
import io
from dados_curriculo import CURRICULO_DB
from dados_curriculo import LIBRAS_INFANTIL
from dados_curriculo import LIBRAS_FUNDAMENTAL
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from io import BytesIO
import calendar
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders




MIN_DATA = date(1900, 1, 1)
MAX_DATA = date(2100, 12, 31)

# --- CONFIGURAÇÃO INICIAL ---
# --- CONFIGURAÇÃO INICIAL ---
st.set_page_config(
    page_title="Integra | CEIEF Rafael Affonso Leite",
    layout="wide",
    page_icon="🧠",
    initial_sidebar_state="auto"
)

# --- OCULTAR TOOLBAR E MENU E RESPONSIVIDADE ---
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            .stAppDeployButton {display:none;}
            
            /* --- COMPORTAMENTO DESKTOP (Largura > 992px) --- */
            @media (min-width: 992px) {
                /* Esconde completamente o header */
                header {display: none !important;}
                [data-testid="stSidebarCollapseButton"] {display: none !important;}
                
                /* FORÇA A BARRA LATERAL A IR PARA O TOPO ABSOLUTO */
                section[data-testid="stSidebar"] {
                    top: 0px !important;
                    height: 100vh !important;
                }
            }
            
            /* --- COMPORTAMENTO MOBILE/TABLET (Largura <= 991px) --- */
            @media (max-width: 991px) {
                /* Header visível para acessar o menu hambúrguer */
                header {visibility: visible;}
                
                /* Ajustes para evitar que o conteúdo suba demais */
                .header-box {
                    margin-top: 0px !important;
                }
            }
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# --- FUNÇÕES AUXILIARES DE DESENHO (GLOBAIS) ---
def calc_lines(pdf, text, w):
    if not text: return 1
    lines = 0
    for p in str(text).split('\n'):
        words = p.split(' ')
        line_w = 0
        for word in words:
            word_w = pdf.get_string_width(word + ' ')
            if line_w + word_w > w - 2:
                lines += 1
                line_w = word_w
            else:
                line_w += word_w
        lines += 1
    return max(1, lines)

def draw_flex_row(pdf, col_data, line_h=6, font_size=9, fill_color=(240, 240, 240)):
    max_lines = 1
    x_start_measure = pdf.get_x()
    for w, text, weight, align, fill in col_data:
        pdf.set_font("Arial", weight, font_size)
        real_w = w if w > 0 else (210 - 15 - x_start_measure)
        lines = calc_lines(pdf, text, real_w)
        if lines > max_lines: max_lines = lines
        x_start_measure += real_w
        
    row_h = max_lines * line_h
    if pdf.get_y() + row_h > 275:
        pdf.add_page()
        
    x_start = pdf.get_x()
    y_start = pdf.get_y()
    for w, text, weight, align, fill in col_data:
        real_w = w if w > 0 else (210 - 15 - x_start)
        pdf.set_font("Arial", weight, font_size)
        if fill: pdf.set_fill_color(*fill_color)
        else: pdf.set_fill_color(255, 255, 255)
        
        pdf.set_xy(x_start, y_start)
        pdf.cell(real_w, row_h, "", border=1, fill=fill)
        
        y_text = y_start + 1
        if max_lines > 1 and calc_lines(pdf, text, real_w) == 1:
            y_text = y_start + (row_h - line_h) / 2
            
        pdf.set_xy(x_start + 1, y_text)
        pdf.multi_cell(real_w - 2, line_h, str(text), border=0, align=align)
        x_start += real_w
        
    pdf.set_xy(15, y_start + row_h)

# --- CONEXÃO COM SUPABASE ---
@st.cache_resource
def init_supabase():
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)

supabase = init_supabase()

def load_db(strict=False):
    try:
        res = supabase.table("Alunos").select("*").execute()
        df = pd.DataFrame(res.data)
        if df.empty: return pd.DataFrame(columns=["nome", "tipo_doc", "dados_json", "id", "ultima_atualizacao"])
        return df
    except Exception as e:
        if strict: raise Exception(f"Erro Supabase: {e}")
        return pd.DataFrame(columns=["nome", "tipo_doc", "dados_json", "id", "ultima_atualizacao"])

def safe_read(worksheet_name, columns):
    try:
        res = supabase.table(worksheet_name).select("*").execute()
        df = pd.DataFrame(res.data)
        if df.empty: return pd.DataFrame(columns=columns)
        return df
    except:
        return pd.DataFrame(columns=columns)

def safe_update(worksheet_name, data):
    """
    Sincroniza do Pandas para o Supabase de forma segura e adaptável para qualquer tabela.
    """
    try:
        # 1. Limpeza contra erros de JSON
        df_to_save = data.copy()
        df_to_save = df_to_save.fillna("") 
        
        # 2. SEPARAÇÃO INTELIGENTE POR TABELA
        if worksheet_name == "Alunos":
            df_to_save['id'] = df_to_save['nome'].astype(str).str.strip() + " (" + df_to_save['tipo_doc'].astype(str).str.strip() + ")"
            supabase.table(worksheet_name).delete().neq("nome", "FORCAR_LIMPEZA_TOTAL").execute()
            
        elif worksheet_name == "Atas_Conselho":
            supabase.table(worksheet_name).delete().neq("id_ata", "FORCAR_LIMPEZA_TOTAL").execute()
            
        elif worksheet_name in ["Recados", "Agenda", "Agendamentos"]:
            # Removemos a coluna 'id' para evitar o erro "1.0" (float -> int) gerado pelo Pandas
            if 'id' in df_to_save.columns:
                df_to_save = df_to_save.drop(columns=['id'])
                
            supabase.table(worksheet_name).delete().neq("Data", "FORCAR_LIMPEZA_TOTAL").execute()
            
        elif worksheet_name == "Config_Ata":
            if 'id' in df_to_save.columns:
                df_to_save = df_to_save.drop(columns=['id'])
            
            supabase.table(worksheet_name).delete().neq("chave", "FORCAR_LIMPEZA_TOTAL").execute()
            
        else:
            pass

        # 3. Gravação no Banco de Dados
        data_dict = df_to_save.to_dict(orient="records")
        if len(data_dict) > 0:
            supabase.table(worksheet_name).insert(data_dict).execute()
        
        return True
    except Exception as e:
        st.error(f"Erro crítico ao atualizar {worksheet_name}: {e}")
        return False
def create_backup(df_atual):
    pass # Backups agora são gerenciados nativamente pela infraestrutura do Supabase

def log_action(student_name, action, details):
    try:
        novo_log = {
            "Data_Hora": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
            "Aluno": student_name,
            "Usuario": st.session_state.get('usuario_nome', 'Desconhecido'),
            "Acao": action,
            "Detalhes": details
        }
        supabase.table("Historico").insert(novo_log).execute()
    except: pass

def save_student(doc_type, name, data, section="Geral"):
    is_monitor = st.session_state.get('user_role') == 'monitor'
    if is_monitor and doc_type != "DIARIO" and section != "Assinatura":
        st.error("Acesso negado: Monitores não podem editar este documento.")
        return

    try:
        id_registro = f"{name} ({doc_type})"
        if 'doc_uuid' not in data or not data['doc_uuid']: data['doc_uuid'] = str(uuid.uuid4()).upper()

        def serializar_datas(obj):
            if isinstance(obj, (date, datetime)): return obj.strftime("%Y-%m-%d")
            if isinstance(obj, dict): return {k: serializar_datas(v) for k, v in obj.items()}
            if isinstance(obj, list): return [serializar_datas(i) for i in obj]
            return obj
            
        data_limpa = serializar_datas(data)
        novo_json = json.dumps(data_limpa, ensure_ascii=False)
        fuso_br = timezone(timedelta(hours=-3))
        data_hora_agora = datetime.now(fuso_br).strftime("%d/%m/%Y %H:%M:%S")

        novo_registro = {
            "id": id_registro,
            "nome": name,
            "tipo_doc": doc_type,
            "dados_json": novo_json,
            "ultima_atualizacao": data_hora_agora
        }
        
        # O poderoso UPSERT substitui toda a sua lógica manual de lock e cópia de DFs
        supabase.table("Alunos").upsert(novo_registro).execute()
        st.toast(f"✅ Alterações em {name} salvas com segurança!", icon="💾")
    except Exception as e:
        st.error(f"❌ Falha ao salvar no banco Supabase. Erro: {e}")

def delete_student(student_name):
    is_monitor = st.session_state.get('user_role') == 'monitor'
    if is_monitor: return False
    try:
        supabase.table("Alunos").delete().eq("nome", student_name).execute()
        st.toast(f"🗑️ Registro de {student_name} excluído!", icon="🔥")
        return True
    except Exception as e:
        st.error(f"Erro ao excluir: {e}")
        return False


def load_carometro_db():
    """Carrega todos os registros da nova tabela Carometro"""
    try:
        response = supabase.table("Carometro").select("*").execute()
        return pd.DataFrame(response.data)
    except Exception as e:
        st.error(f"Erro ao carregar Carômetro: {e}")
        return pd.DataFrame()

def save_carometro_entry(nome, turma, foto_b64):
    """Salva ou atualiza um aluno na tabela Carometro"""
    try:
        # Tenta encontrar se o aluno já existe nessa turma para atualizar
        existente = supabase.table("Carometro").select("id").eq("nome", nome).eq("turma", turma).execute()
        
        dados = {"nome": nome, "turma": turma, "foto_base64": foto_b64}
        
        if existente.data:
            id_reg = existente.data[0]['id']
            supabase.table("Carometro").update(dados).eq("id", id_reg).execute()
        else:
            supabase.table("Carometro").insert(dados).execute()
        return True
    except Exception as e:
        st.error(f"Erro ao salvar no Carômetro: {e}")
        return False

def delete_carometro_entry(id_reg):
    """Remove um registro da tabela Carometro"""
    try:
        supabase.table("Carometro").delete().eq("id", id_reg).execute()
        return True
    except Exception:
        return False
# --- FIM DAS FUNÇÕES DE BANCO DE DADOS ---



import uuid
import base64
import json
import urllib.request # Importante para baixar as fotos depois

def migrar_base64_para_bucket():
    st.info("Iniciando a migração das fotos de toda a escola... Isso pode levar alguns minutos.")
    
    # 1. Puxa TODOS os alunos do banco
    res = supabase.table("Alunos").select("id, nome, dados_json").execute()
    
    # Tratamento de erro vital: dependendo da versão, os dados podem estar em res.data ou no próprio res
    registros = res.data if hasattr(res, 'data') else res
    
    sucessos = 0
    erros = 0
    ignorados = 0
    
    # Cria uma barra de progresso visual para você acompanhar
    barra_progresso = st.progress(0)
    total_alunos = len(registros)
    texto_status = st.empty()
    
    for i, row in enumerate(registros):
        try:
            texto_status.text(f"Analisando: {row.get('nome', 'Desconhecido')} ({i+1}/{total_alunos})")
            dados = json.loads(row["dados_json"])
            
            # Verifica se o aluno tem uma foto em Base64 salva
            if "foto_base64" in dados and dados["foto_base64"]:
                b64_string = dados.pop("foto_base64") 
                
                if "," in b64_string:
                    b64_string = b64_string.split(",")[1]
                    
                image_bytes = base64.b64decode(b64_string)
                nome_arquivo = f"{uuid.uuid4()}.jpg"
                
                # Faz o Upload para o Bucket
                supabase.storage.from_("fotos_alunos").upload(
                    file=image_bytes,
                    path=nome_arquivo,
                    file_options={"content-type": "image/jpeg"}
                )
                
                # Pega a URL
                url_publica = supabase.storage.from_("fotos_alunos").get_public_url(nome_arquivo)
                
                # Salva o link no JSON novo no lugar do Base64
                dados["foto_url"] = url_publica
                novo_json_limpo = json.dumps(dados, ensure_ascii=False)
                
                # Atualiza o banco de dados
                supabase.table("Alunos").update({"dados_json": novo_json_limpo}).eq("id", row["id"]).execute()
                
                sucessos += 1
            else:
                ignorados += 1 # Aluno não tem foto em base64 (ou não tem foto, ou já foi migrado)
                
        except Exception as e:
            st.error(f"Erro ao migrar a foto de {row.get('nome')}: {e}")
            erros += 1
            
        # Atualiza a barrinha
        barra_progresso.progress((i + 1) / total_alunos)
            
    texto_status.empty() # Limpa o texto "Analisando..."
    st.success(f"✅ Migração finalizada!\n- Convertidas: {sucessos}\n- Já eram links ou sem foto: {ignorados}\n- Erros: {erros}")







# --- HELPERS PARA PDF ---
def clean_pdf_text(text):
    if text is None or text is False: return ""
    if text is True: return "Sim"
    return str(text).encode('latin-1', 'replace').decode('latin-1')

def get_pdf_bytes(pdf_instance):
    try: return bytes(pdf_instance.output(dest='S').encode('latin-1'))
    except: return bytes(pdf_instance.output(dest='S'))

# --- CLASSE PDF CUSTOMIZADA COM ASSINATURA ---
class OfficialPDF(FPDF):
    def __init__(self, orientation='P', unit='mm', format='A4'):
        super().__init__(orientation, unit, format)
        self.signature_info = None # Texto da assinatura
        self.doc_uuid = None
        self.doc_type = None

    def header(self):
        # TIMBRADO DE FUNDO EXCLUSIVO PARA A ATA
        if self.doc_type == "Ata":
            try:
                # NOME DO FICHEIRO CORRIGIDO PARA "image.png"
                self.image("image.png", x=0, y=0, w=210, h=297)
            except Exception as e:
                pass # Se a imagem não for encontrada, gera sem quebrar o sistema

    def set_signature_footer(self, signatures_list, doc_uuid):
        """Prepara o texto de validação para o rodapé"""
        self.doc_uuid = doc_uuid
        if signatures_list and len(signatures_list) > 0:
            names = [s.get('name', '').upper() for s in signatures_list]
            names_str = ", ".join(names[:-1]) + " e " + names[-1] if len(names) > 1 else names[0]
            self.signature_info = f"Assinado por {len(names)} pessoas: {names_str}"
        else:
            self.signature_info = "Documento gerado sem assinaturas digitais."
            
    def footer(self):
        # O rodapé padrão só aparece se NÃO for a Ata
        if self.doc_type != "Ata":
            self.set_y(-20)
            self.set_font('Arial', '', 8)
            self.set_text_color(80, 80, 80)
            
            # Bloco de Assinatura Digital
            if self.doc_uuid:
                # Posicionamento dinâmico baseado na altura da página
                box_h = 9  # Altura reduzida para ~2 linhas
                margin_bottom = 20 # Distância da borda inferior
                
                y_box = self.h - margin_bottom 
                x_box = 10
                w_box = self.w - 20 # Largura total (menos margens laterais de 10mm)

                # Caixa cinza claro para validação
                self.set_fill_color(245, 245, 245)
                self.rect(x_box, y_box, w_box, box_h, 'F')
                
                # Texto
                self.set_xy(x_box + 2, y_box + 1.5)
                self.set_font('Arial', 'B', 7)
                if self.signature_info:
                    self.cell(0, 3, clean_pdf_text(self.signature_info), 0, 1, 'L')
                else:
                    self.ln(3) # Espaço caso não tenha texto de assinatura
                
                self.set_x(x_box + 2)
                self.set_font('Arial', '', 7)
                link_txt = f"Para verificar a validade das assinaturas, acesse https://integra.streamlit.app e informe o código {self.doc_uuid}"
                self.cell(0, 3, clean_pdf_text(link_txt), 0, 1, 'L')

            # Endereço Padrão (Abaixo da caixa)
            self.set_y(-10)
            self.set_font('Arial', '', 8)
            addr = "Secretaria Municipal de Educação | Centro de Formação do Professor - Limeira-SP"
            self.cell(0, 5, clean_pdf_text(addr), 0, 0, 'C')
            self.set_font('Arial', 'I', 8)
            self.cell(0, 5, clean_pdf_text(f'Página {self.page_no()}'), 0, 0, 'R')

    def section_title(self, title, width=0):
        self.set_font('Arial', 'B', 12); self.set_fill_color(240, 240, 240)
        self.cell(width, 8, clean_pdf_text(title), 1, 1, 'L', 1)

# --- FUNÇÃO DE LOGIN COMPLETA E ROBUSTA (SME LIMEIRA) ---
def login():
    # Inicializa o estado de autenticação se não existir
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    
    if "user_role" not in st.session_state:
        st.session_state.user_role = None

    if not st.session_state.authenticated:
        # --- CSS DA TELA DE LOGIN (NO-SCROLL LAYOUT) ---
        st.markdown("""
            <style>
                /* --- ESCONDER BARRA DO GITHUB, FORK, MENU E RODAPÉ --- */
                [data-testid="stHeader"] {
                    display: none !important;
                }
                .stApp > header {
                    display: none !important;
                    background-color: transparent !important;
                }
                .stDeployButton {
                    display: none !important;
                }
                #MainMenu {
                    visibility: hidden !important;
                }
                footer {
                    visibility: hidden !important;
                }

                /* Remove padding padrão do Streamlit para ocupar a tela toda */
                .block-container {
                    padding-top: 0rem !important;
                    padding-bottom: 0rem !important;
                    max-width: 100%;
                    min-height: 100vh;
                    display: flex;
                    flex-direction: column;
                    justify-content: center;
                }
                
                /* Fundo da Página */
                [data-testid="stAppViewContainer"] {
                    background: linear-gradient(135deg, #f0f4f8 0%, #d9e2ec 100%);
                }
                
                /* --- FORÇAR AS COLUNAS A ESTICAREM JUNTAS --- */
                div[data-testid="stHorizontalBlock"] {
                    align-items: stretch !important;
                }

                /* Container da Coluna da Arte (Coluna 2) */
                div[data-testid="column"]:nth-of-type(2),
                div[data-testid="stColumn"]:nth-of-type(2) {
                    display: flex;
                    flex-direction: column;
                }

                /* Painel Esquerdo (Arte) */
                .login-art-box {
                    background: linear-gradient(135deg, #2563eb 0%, #1e3a8a 100%);
                    height: 100% !important; 
                    border-radius: 16px 0 0 16px; 
                    display: flex;
                    flex-direction: column;
                    justify-content: center;
                    align-items: center;
                    color: white;
                    padding: 40px;
                    text-align: center;
                    box-shadow: -5px 10px 25px rgba(37, 99, 235, 0.2);
                }
                
                /* Painel Direito (Formulário) - Coluna 3 */
                div[data-testid="column"]:nth-of-type(3),
                div[data-testid="stColumn"]:nth-of-type(3) {
                    background-color: white;
                    padding: 2rem 3rem !important;
                    border-radius: 0 16px 16px 0; 
                    height: 100% !important; 
                    display: flex;
                    flex-direction: column;
                    justify-content: center; 
                    box-shadow: 5px 10px 25px rgba(0,0,0,0.05);
                }
                /* Tipografia */
                .welcome-title {
                    font-size: 1.8rem;
                    font-weight: 700;
                    color: #1e293b;
                    margin-bottom: 5px;
                }
                .welcome-sub {
                    font-size: 0.95rem;
                    color: #64748b;
                    margin-bottom: 20px;
                }
                
                /* Inputs Customizados */
                .stTextInput label {
                    font-size: 0.85rem;
                    color: #475569;
                    font-weight: 600;
                }
                
                /* Aviso LGPD */
                .lgpd-box {
                    background-color: #fff7ed;
                    border-left: 4px solid #f97316;
                    padding: 10px;
                    margin-top: 15px;
                    margin-bottom: 15px;
                    border-radius: 6px;
                }
                .lgpd-title {
                    color: #9a3412;
                    font-weight: 700;
                    font-size: 0.75rem;
                    display: flex; 
                    align-items: center; 
                    gap: 6px;
                }
                .lgpd-text {
                    color: #9a3412;
                    font-size: 0.7rem;
                    margin-top: 2px;
                    line-height: 1.2;
                    text-align: justify; /* Texto justificado */
                }

/* --- RESPONSIVIDADE PARA MOBILE (FORÇADO E INFALÍVEL) --- */
                @media (max-width: 768px) {
                    /* Destrói o comportamento de flexbox nativo para forçar empilhamento */
                    div[data-testid="stHorizontalBlock"] {
                        display: block !important; 
                        width: 100% !important;
                    }

                    /* Força CADA coluna a ocupar 100% da tela como um bloco isolado */
                    div[data-testid="column"], 
                    div[data-testid="stColumn"] {
                        width: 100% !important;
                        min-width: 100% !important;
                        max-width: 100% !important;
                        display: block !important;
                        margin: 0 !important;
                        padding: 0 !important;
                    }

                    /* Esconde os espaços laterais (1 e 4) */
                    div[data-testid="column"]:nth-of-type(1),
                    div[data-testid="stColumn"]:nth-of-type(1),
                    div[data-testid="column"]:nth-of-type(4),
                    div[data-testid="stColumn"]:nth-of-type(4) {
                        display: none !important;
                    }

                    /* Banner Azul (Vai ficar no Topo) */
                    .login-art-box {
                        border-radius: 16px 16px 0 0 !important; 
                        padding: 30px 20px !important;
                        height: auto !important; 
                        min-height: auto !important;
                        width: 100% !important;
                        box-sizing: border-box !important;
                    }
                    
                    /* Ajuste de fontes para não quebrar a tela do celular */
                    .login-art-box h1 { font-size: 2.2rem !important; margin-top: 5px !important; }
                    .login-art-box div[style*="font-size: 6rem"] { font-size: 3.5rem !important; margin-bottom: 0 !important; }
                    .login-art-box hr, .login-art-box p:last-child { display: none !important; } 

                    /* Formulário Branco (Vai ficar Embaixo) */
                    div[data-testid="column"]:nth-of-type(3),
                    div[data-testid="stColumn"]:nth-of-type(3) {
                        border-radius: 0 0 16px 16px !important; 
                        padding: 1.5rem !important;
                        height: auto !important;
                        width: 100% !important;
                        box-sizing: border-box !important;
                        box-shadow: none !important;
                    }
                }
                /* ------------------------------ */

            </style>
        """, unsafe_allow_html=True)
        
        # Espaçamento para centralizar verticalmente na tela
        st.write("")
        st.write("")

        # Layout em Colunas: Spacer, Arte, Form, Spacer
        # Ajuste de proporção para ficar elegante
        c_pad1, c_art, c_form, c_pad2 = st.columns([1, 4, 4, 1])

        
# --- LADO ESQUERDO (ARTE AZUL) ---
        with c_art:
            st.markdown("""
<div class="login-art-box" style="display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center; height: 100%;">
<div style="font-size: 6rem; margin-bottom: 1rem; filter: drop-shadow(0px 4px 6px rgba(0,0,0,0.2));">🧠</div>
<p style="color: white; font-weight: 800; font-size: 3.5rem; margin: 0; line-height: 1;">INTEGRA</p>
<p style="font-size: 1.2rem; opacity: 0.9; font-weight: 300; margin-top: 10px; margin-bottom: 0;">Ensino Regular<br>Educação Especial Inclusiva<br>Sala de Leitura - Empréstimos<br>Gestão de Suplementos</p>
<div style="margin-top: 40px; width: 100%; display: flex; flex-direction: column; align-items: center;">
<hr style="border-color: rgba(255,255,255,0.3); margin-bottom: 20px; width: 100%;">
<p style="font-style: italic; font-size: 1rem; opacity: 0.9; margin: 0;">
"A inclusão acontece quando se aprende com as diferenças e não com as igualdades."
</p>
</div>
</div>
""", unsafe_allow_html=True)
            
        # --- LADO DIREITO (FORMULÁRIO BRANCO) ---
        with c_form:
            # Abas de Login e Validação
            tab_login, tab_validar = st.tabs(["🔐 Acesso ao Sistema", "✅ Validar Documento"])
            
            with tab_login:
                with st.form("login_form"):
                    # Layout Header: Texto à esquerda, Logo à direita (menor)
                    c_head_txt, c_head_logo = st.columns([3, 1.2])
                    
                    with c_head_txt:
                        st.markdown('<div class="welcome-title" style="margin-top: 0px;">Bem-vindo(a)</div>', unsafe_allow_html=True)
                        st.markdown('<div class="welcome-sub">Insira suas credenciais para acessar o sistema.</div>', unsafe_allow_html=True)
                    
                    with c_head_logo:
                        if os.path.exists("logo_escola.png"):
                            st.image("logo_escola.png", use_container_width=True)
                    
                    st.write("") # Espaço
                    
                    user_id = st.text_input("Matrícula", placeholder="Ex: 12345")
                    password = st.text_input("Senha", type="password", placeholder="••••••")
                    
                    st.markdown("""
                        <div class="lgpd-box">
                            <div class="lgpd-title">🔒 CONFIDENCIALIDADE E SIGILO</div>
                            <div class="lgpd-text">
                                Acesso Monitorado. Protegido pela LGPD. Uso estritamente profissional.
                            </div>
                        </div>
                    """, unsafe_allow_html=True)
                    
                    submit = st.form_submit_button("ACESSAR SISTEMA", type="primary")
                    
                    if submit:
                        try:
                            SENHA_MESTRA = st.secrets.get("credentials", {}).get("password", "admin")
                            user_id_limpo = str(user_id).strip()
                            authenticated = False
                            
                            # 1. TENTATIVA: PROFESSORES (Equipe Gestora / Docentes)
                            df_professores = safe_read("Professores", ["matricula", "nome"])
                            if not df_professores.empty:
                                df_professores['matricula'] = df_professores['matricula'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                                if password == SENHA_MESTRA and user_id_limpo in df_professores['matricula'].values:
                                    registro = df_professores[df_professores['matricula'] == user_id_limpo]
                                    nome_prof = registro['nome'].values[0]
                                    
                                    st.session_state.authenticated = True
                                    st.session_state.usuario_nome = nome_prof
                                    st.session_state.user_role = 'professor'
                                    st.session_state.usuario_matricula = user_id_limpo
                                    authenticated = True
                                    
                                    st.toast(f"Acesso Docente autorizado. Bem-vindo(a), {nome_prof}!", icon="🔓")
                                    time.sleep(1)
                                    st.rerun()

                            # 2. TENTATIVA: MONITORES / FUNCIONÁRIOS
                            if not authenticated:
                                df_monitores = safe_read("Monitores", ["matricula", "nome"])
                                if not df_monitores.empty:
                                    df_monitores['matricula'] = df_monitores['matricula'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                                    if password == "123" and user_id_limpo in df_monitores['matricula'].values:
                                        registro = df_monitores[df_monitores['matricula'] == user_id_limpo]
                                        nome_mon = registro['nome'].values[0]
                                        
                                        st.session_state.authenticated = True
                                        st.session_state.usuario_nome = nome_mon
                                        st.session_state.user_role = 'monitor'
                                        st.session_state.usuario_matricula = user_id_limpo
                                        authenticated = True
                                        
                                        st.toast(f"Acesso Monitor autorizado. Bem-vindo(a), {nome_mon}!", icon="🛡️")
                                        time.sleep(1)
                                        st.rerun()
                            
                            # 3. TENTATIVA: ESTUDANTES (Para o Álbum de Figurinhas)
                            if not authenticated:
                                try:
                                    # Consulta o RA diretamente na tabela estudantes do PostgreSQL (Supabase)
                                    res_estudantes = supabase.table("estudantes").select("ra, nome").eq("ra", user_id_limpo).execute()
                                    
                                    if res_estudantes.data:
                                        # Defini a senha padrão das crianças como "123" ou o próprio R.A. para facilitar
                                        if password == "123" or password == user_id_limpo: 
                                            aluno_nome = res_estudantes.data[0]['nome']
                                            
                                            st.session_state.authenticated = True
                                            st.session_state.usuario_nome = aluno_nome
                                            st.session_state.user_role = 'estudante'
                                            st.session_state.usuario_ra = user_id_limpo # Fundamental para o Álbum!
                                            st.session_state.modulo_atuacao = "Álbum do Estudante"
                                            authenticated = True
                                            st.toast(f"Acesso de Estudante autorizado! Bora colecionar, {aluno_nome}!", icon="🎒")
                                            time.sleep(1)
                                            st.rerun()
                                except Exception as e_banco:
                                    pass # Se a tabela não existir ou der erro, segue para falha de login
                                    
                            # SE PASSOU PELAS 3 ETAPAS E NÃO ACHOU NINGUÉM
                            if not authenticated:
                                st.error("Credenciais inválidas. Verifique sua Matrícula/R.A. e a Senha.")

                        except Exception as e:
                            st.error(f"Erro técnico no servidor: {e}")

            with tab_validar:
                st.markdown("### Validação Pública")
                st.caption("Insira o código UUID presente no rodapé do documento para verificar sua autenticidade e assinaturas.")
                uuid_input = st.text_input("Código do Documento (UUID)", placeholder="Ex: 7D2B-5135...")
                if st.button("Verificar Autenticidade", type="primary"):
                    if uuid_input:
                        try:
                            df_alunos = load_db()
                            encontrado = False
                            for _, row in df_alunos.iterrows():
                                try:
                                    d = json.loads(row['dados_json'])
                                    if d.get('doc_uuid') == uuid_input.strip():
                                        encontrado = True
                                        st.success("✅ DOCUMENTO VÁLIDO E AUTÊNTICO")
                                        st.markdown(f"**Aluno:** {d.get('nome', 'N/A')}")
                                        st.markdown(f"**Tipo:** {row['tipo_doc']}")
                                        
                                        assinaturas = d.get('signatures', [])
                                        if assinaturas:
                                            st.markdown("---")
                                            st.markdown("### Assinaturas Digitais:")
                                            for sig in assinaturas:
                                                st.info(f"✍️ **{sig['name']}** ({sig.get('role', 'Profissional')})\n\n📅 Assinado em: {sig['date']}")
                                        else:
                                            st.warning("Este documento ainda não possui assinaturas digitais registradas.")
                                        break
                                except: pass
                            if not encontrado:
                                st.error("❌ Documento não encontrado ou código inválido.")
                        except Exception as e:
                            st.error(f"Erro na busca: {e}")
        
        # Interrompe o carregamento do restante do app até que o login seja feito
        st.stop()


# --- INICIALIZAÇÃO DO CONTROLE DE MÓDULO (PORTAL) ---
if "modulo_atuacao" not in st.session_state:
    st.session_state.modulo_atuacao = None


# --- ATIVAÇÃO DO LOGIN ---
login()


# ==============================================================================
# TELA DE PORTAL (ESCOLHA DO MÓDULO APÓS LOGIN)
# ==============================================================================
if st.session_state.authenticated:

    # 1. AQUI O ALUNO É INTERCEPTADO E PULA O PORTAL:
    if st.session_state.get('user_role') == 'estudante':
        pass # Pula todos os menus de professor e segue direto para o final do arquivo
        
    # 2. AQUI TEM DE SER 'elif' PARA QUE O ALUNO NÃO ENTRE NESTE BLOCO!
    elif st.session_state.modulo_atuacao is None:
        # Removido os <br> para o conteúdo subir e ficar mais centralizado
        st.write("") 
        
        # Título alterado para <div> para não aparecer o ícone de link
        st.markdown("<div style='text-align: center; color: #1e293b; font-size: 32px; font-weight: bold; margin-bottom: 40px;'>Seja bem-vindo(a)! Escolha seu ambiente de trabalho:</div>", unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4) # Alterado para 4 colunas
        
        
        # --- BOTÃO 1: ENSINO REGULAR (ESQUERDA) ---
        with col1:
            st.markdown("""
            <div style='background-color: #eff6ff; padding: 35px 20px; border-radius: 12px; text-align: center; border: 2px solid #3b82f6; box-shadow: 0 4px 6px rgba(0,0,0,0.05);'>
                <div style='font-size: 55px; margin-bottom: 15px;'>🏫</div>
                <div style='color: #1d4ed8; font-size: 24px; font-weight: 700; margin: 0;'>Ensino Regular</div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("Acessar Ensino Regular", type="primary", use_container_width=True, key="btn_er"):
                st.session_state.modulo_atuacao = "🏫 Ensino Regular"
                st.rerun()

        # --- BOTÃO 2: EDUCAÇÃO ESPECIAL (DIREITA) ---
        with col2:
            st.markdown("""
            <div style='background-color: #f0fdf4; padding: 35px 20px; border-radius: 12px; text-align: center; border: 2px solid #22c55e; box-shadow: 0 4px 6px rgba(0,0,0,0.05);'>
                <div style='font-size: 55px; margin-bottom: 15px;'>🧠</div>
                <div style='color: #15803d; font-size: 24px; font-weight: 700; margin: 0;'>Educação Especial</div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            if st.button("Acessar Educação Especial", type="primary", use_container_width=True, key="btn_ee"):
                st.session_state.modulo_atuacao = "🧠 Educação Especial Inclusiva"
                st.rerun()
                
# --- BOTÃO 3: BIBLIOTECA (DIREITA) ---
        with col3:
            st.markdown("""
            <div style='background-color: #fdfbf0; padding: 35px 20px; border-radius: 12px; text-align: center; border: 2px solid #c5b722; box-shadow: 0 4px 6px rgba(0,0,0,0.05);'>
                <div style='font-size: 55px; margin-bottom: 15px;'>📚 </div>
                <div style='color: #c5b722; font-size: 24px; font-weight: 700; margin: 0;'>Sala de Leitura</div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            # ALERTA DE CORREÇÃO: key alterada para "btn_sl"
            if st.button("Acessar Sala de Leitura", type="primary", use_container_width=True, key="btn_sl"):
                st.session_state.modulo_atuacao = "📚  Sala de Leitura"
                st.rerun()

        # --- BOTÃO 4: ADM (DIREITA) ---
        with col4:
            st.markdown("""
            <div style='background-color: #fdf1f0; padding: 35px 20px; border-radius: 12px; text-align: center; border: 2px solid #c53522; box-shadow: 0 4px 6px rgba(0,0,0,0.05);'>
                <div style='font-size: 55px; margin-bottom: 15px;'>📂</div>
                <div style='color: #c53522; font-size: 24px; font-weight: 700; margin: 0;'>Administrativo</div>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            # ALERTA DE CORREÇÃO: key alterada para "btn_adm"
            if st.button("Acessar Administrativo", type="primary", use_container_width=True, key="btn_adm"):
                st.session_state.modulo_atuacao = "📂 Administrativo"
                st.rerun()

        
        # st.stop() bloqueia o carregamento da sidebar até a pessoa clicar num botão
        st.stop()


# --- DEFINIÇÃO DE PERMISSÕES ---
user_role = st.session_state.get('user_role', 'professor')
is_monitor = (user_role == 'monitor') # Flag para bloquear edições

# --- ESTILO VISUAL DA INTERFACE (CSS MELHORADO E RESPONSIVO) ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
    
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    .stApp { background-color: #f8fafc; }
    
    /* --- CORREÇÃO: PUXA O CONTEÚDO PARA CIMA --- */
    .block-container {
        padding-top: 1.5rem !important; /* Reduz o espaço em branco no topo */
    }
    
    /* Melhoria da Sidebar */
    [data-testid="stSidebar"] [data-testid="stImage"] {
        display: flex;
        justify-content: center;
        margin-left: auto;
        margin-right: auto;
    }
    
    /* Centralizar o container de texto da sidebar */
    .sidebar-header {
        display: flex;
        flex-direction: column;
        align-items: center;
        text-align: center;
        width: 100%;
        padding-bottom: 20px;
    }
    
    .sidebar-title {
        color: #1e3a8a; /* Azul Institucional */
        font-weight: 800;
        font-size: 1.4rem;
        margin-top: 10px;
        line-height: 1.2;
    }
    .sidebar-subtitle {
        color: #64748b;
        font-size: 0.85rem;
        font-weight: 400;
    }

    /* Estilo dos Cards Principais */
    .header-box {
        background: white; padding: 2rem; border-radius: 12px;
        border-left: 6px solid #2563eb; /* Borda lateral azul */
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        margin-bottom: 2rem;
        margin-top: 0px; 
    }
    
    .header-title { color: #1e293b; font-weight: 700; font-size: 1.8rem; margin: 0; }
    .header-subtitle { color: #64748b; font-size: 1rem; margin-top: 5px; }
    
/* Dashboard Cards */
    .metric-card {
        background-color: white;
        padding: 1rem 0.2rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        border: 1px solid #e2e8f0;
        text-align: center;
        white-space: nowrap; 
        overflow: hidden; 
        text-overflow: ellipsis; 
    }
    .metric-value {
        font-size: 1.8rem; 
        font-weight: 700;
        line-height: 1.2;
    }
    .metric-label {
        color: #64748b;
        font-size: 0.72rem; 
        font-weight: 600;
        text-transform: uppercase;
        margin-top: 5px;
    }
    
    /* Botões */
    .stButton button { width: 100%; border-radius: 8px; }
    
    /* --- MEDIA QUERIES PARA MOBILE --- */
    @media (max-width: 991px) {
        .header-box {
            margin-top: 10px !important; 
            padding: 1.5rem !important;
        }
        .header-title {
            font-size: 1.5rem !important;
        }
        
        .stBlock {
            padding-top: 1rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# --- INICIALIZAÇÃO DE ESTADO ---
if 'data_declaracao' not in st.session_state:
    st.session_state.data_declaracao = {}
# --- NOVAS VARIÁVEIS PARA O PORTÃO DE ENTRADA ---
if 'ee_aluno_confirmado' not in st.session_state:
    st.session_state.ee_aluno_confirmado = None
if 'ee_doc_confirmado' not in st.session_state:
    st.session_state.ee_doc_confirmado = None
if 'data_pei' not in st.session_state: 
    st.session_state.data_pei = {
        'terapias': {}, 'avaliacao': {}, 'flex': {}, 'plano_ensino': {},
        'comunicacao_tipo': [], 'permanece': []
    }
if 'data_conduta' not in st.session_state:
    st.session_state.data_conduta = {}
if 'data_avaliacao' not in st.session_state:
    st.session_state.data_avaliacao = {}
if 'data_avaliacao2' not in st.session_state: # NOVA LINHA
    st.session_state.data_avaliacao2 = {}     # NOVA LINHA
if 'data_diario' not in st.session_state:
    st.session_state.data_diario = {}
if 'data_pdi' not in st.session_state:
    st.session_state.data_pdi = {
        'metas': [{'objetivo': '', 'prazo': '', 'estrategia': '', 'status': 'Pendente'} for _ in range(5)],
        'pdi_fortalezas': '',
        'pdi_desafios': '',
        'pdi_recursos': '',
        'pdi_periodo': 'Trimestral',
        'pdi_obs': ''
    }
if 'data_declaracao' not in st.session_state:
    st.session_state.data_declaracao = {}

def carregar_dados_aluno():
    selecao = st.session_state.get('aluno_selecionado')
    
    # Init empty
    st.session_state.data_pei = {'terapias': {}, 'avaliacao': {}, 'flex': {}, 'plano_ensino': {}, 'comunicacao_tipo': [], 'permanece': []}
    st.session_state.data_case = {'irmaos': [{'nome': '', 'idade': '', 'esc': ''} for _ in range(4)], 'checklist': {}, 'clinicas': []}
    st.session_state.data_conduta = {}
    st.session_state.data_avaliacao = {}
    st.session_state.data_diario = {}
    st.session_state.data_avaliacao2 = {}
    st.session_state.data_pdi = {
        'metas': [{'objetivo': '', 'prazo': '', 'estrategia': '', 'status': 'Pendente'} for _ in range(5)],
        'pdi_fortalezas': '', 'pdi_desafios': '', 'pdi_recursos': '', 'pdi_periodo': 'Trimestral', 'pdi_obs': ''
    }
    st.session_state.nome_original_salvamento = None


    if not selecao or selecao == "-- Novo Registro --":
        return

    try:
        df_db = load_db()
        # Filter by name
        if "nome" in df_db.columns:
            rows = df_db[df_db["nome"] == selecao]
            if rows.empty: return
            
            st.session_state.nome_original_salvamento = selecao
            st.session_state.data_pei['nome'] = selecao
            st.session_state.data_case['nome'] = selecao
            st.session_state.data_conduta['nome'] = selecao
            st.session_state.data_avaliacao['nome'] = selecao
            st.session_state.data_diario['nome'] = selecao
            st.session_state.data_pdi['nome'] = selecao

            for _, row in rows.iterrows():
                try:
                    dados = json.loads(row["dados_json"])
                    # Date conversion
                    for k, v in dados.items():
                        if isinstance(v, str) and len(v) == 10 and v.count('-') == 2:
                            try: dados[k] = datetime.strptime(v, '%Y-%m-%d').date()
                            except: pass
                    
                    dtype = row["tipo_doc"]
                    if dtype == "PEI":
                        st.session_state.data_pei.update(dados)
                    elif dtype == "CASO":
                        st.session_state.data_case.update(dados)
                    elif dtype == "CONDUTA":
                        st.session_state.data_conduta.update(dados)
                    elif dtype == "AVALIACAO":
                        st.session_state.data_avaliacao.update(dados)
                    elif dtype == "AVALIACAO2": # NOVO BLOCO
                        st.session_state.data_avaliacao2.update(dados)
                    elif dtype == "DIARIO":
                        st.session_state.data_diario.update(dados)
                    elif dtype == "PDI":
                        st.session_state.data_pdi.update(dados)
                except: pass
            
            st.toast(f"✅ {selecao} carregado.")
            
    except Exception as e:
        st.info("Pronto para novo preenchimento.")

# --- BARRA LATERAL ULTRA-COMPACTA ---
with st.sidebar:
    # CSS PARA "ESPREMER" O LAYOUT
    st.markdown("""
    <style>
        section[data-testid="stSidebar"] > div {
            padding-top: 1rem !important;
            padding-bottom: 2rem !important;
        }
        [data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
            gap: 1.2rem !important;
        }
        .sidebar-title {
            font-size: 1.1rem;
            font-weight: 800;
            color: #1e3a8a;
            margin: 0;
            margin-top: 0px !important;
            text-align: center;
            line-height: 1.2;
        }
        .sidebar-sub {
            font-size: 0.7rem;
            color: #64748b;
            text-align: center;
            margin-bottom: 8px;
        }
        .section-label {
            font-size: 0.8rem;
            font-weight: 700;
            color: #475569;
            margin-top: 8px;
            margin-bottom: 0px;
        }
        .user-slim {
            background-color: #f8fafc;
            border: 1px solid #e2e8f0;
            border-radius: 4px;
            padding: 4px;
            font-size: 0.8rem;
            color: #334155;
            text-align: center;
        }
        .role-tag {
            background-color: #e0f2fe;
            color: #0369a1;
            padding: 2px 6px;
            border-radius: 4px;
            font-size: 0.7rem;
            font-weight: 600;
            margin-top: 2px;
            display: inline-block;
        }
        .stRadio { margin-top: -5px; }
        div[data-baseweb="select"] { min-height: 32px; }
        hr { margin: 0.5em 0 !important; }
    </style>
    """, unsafe_allow_html=True)

    # 1. TÍTULO
    st.markdown('<div class="sidebar-title">SISTEMA INTEGRA</div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-sub">Gestão Escolar</div>', unsafe_allow_html=True)

    # 2. USUÁRIO
    nome_prof = st.session_state.get('usuario_nome', 'Usuário')
    role_label = "Monitor(a)" if is_monitor else "Docente/Admin"
    nomes = nome_prof.split()
    nome_curto = f"{nomes[0]} {nomes[-1]}" if len(nomes) > 1 else nomes[0]
    
    st.markdown(f"""
        <div style="text-align: center;">
            <div class="user-slim">👤 <b>{nome_curto}</b></div>
            <span class="role-tag">{role_label}</span>
        </div>
    """, unsafe_allow_html=True)
    
    st.divider()
    
    # 3. MÓDULO DE ATUAÇÃO
    modulo_atuacao = st.session_state.modulo_atuacao
    
    # Exibe em qual módulo estamos e cria um botão para voltar ao Portal
    st.markdown(f'<p class="section-label" style="color:#2563eb; font-weight:bold; text-align:center;">{modulo_atuacao}</p>', unsafe_allow_html=True)
    if st.session_state.get('user_role') != 'estudante':
        if st.button("🔄 Trocar Ambiente", use_container_width=True):
            st.session_state.modulo_atuacao = None
            st.rerun()

    st.divider()

    # Variáveis padrão de controle do sistema
    selected_student = "-- Novo Registro --"
    pei_level = "Fundamental" 
    doc_mode = "Dashboard"

# --- NAVEGAÇÃO CONDICIONAL BASEADA NO MÓDULO ---
    app_mode = None
    app_mode_regular = None
    app_mode_adm = None  # <-- NOVA VARIÁVEL ADICIONADA AQUI

    if modulo_atuacao == "🧠 Educação Especial Inclusiva":
        st.markdown('<p class="section-label">📌 Navegação</p>', unsafe_allow_html=True)
        # ADICIONADA A CHAVE: key="nav_especial"
        app_mode = st.radio(
            "Navegação", 
            ["📊 Painel de Gestão", "👥 Gestão de Alunos", "🖼️ Carômetro"], 
            label_visibility="collapsed",
            key="nav_especial"
        )
        
    elif modulo_atuacao == "🏫 Ensino Regular":
        
        st.markdown('<p class="section-label">📌 Navegação</p>', unsafe_allow_html=True)
        
        opcoes_regular = [
            "🖼️ Carômetro Escolar", 
            "💻 Agendamento Informática", 
            "📝 Nova Ata de Conselho", 
            "📂 Histórico de Atas", 
            "📖 Planejamento Curricular"
        ]
        
        # A opção do Álbum aparece para si
        if st.session_state.get('usuario_nome') == "José Victor Souza Gallo":
            opcoes_regular.append("⚙️ Configurações")
            opcoes_regular.append("💾 Cofre de Segurança")
            opcoes_regular.append("🏆 Álbum de Figurinhas")
            
        app_mode_regular = st.radio(
            "Navegação", 
            opcoes_regular, 
            label_visibility="collapsed",
            key="nav_regular"
        )

        
    # ====== ADICIONE ESTE BLOCO ABAIXO ======
    elif modulo_atuacao == "📂 Administrativo":
        st.markdown('<p class="section-label">📌 Navegação</p>', unsafe_allow_html=True)
        app_mode_adm = st.radio(
            "Navegação", 
            ["🏷️ Patrimônio e Inventário", "📦 Almoxarifado Escolar", "🖨️ Emissão de Boletins"], 
            label_visibility="collapsed",
            key="nav_adm"
        )


# --- SEÇÃO GESTÃO DE ALUNOS ---
    # --- SEÇÃO GESTÃO DE ALUNOS (SIDEBAR LIMPA) ---
    if app_mode == "👥 Gestão de Alunos":
        st.divider()
        st.info("👉 Selecione o aluno e o documento na tela principal ao lado.")
        
        # Puxamos as variáveis confirmadas para a sidebar saber o que exibir de foto
        selected_student = st.session_state.get('ee_aluno_confirmado', "-- Novo Registro --")
        doc_mode = st.session_state.get('ee_doc_confirmado', "Dashboard")

        # Foto na Sidebar (só renderiza se o aluno estiver confirmado)
        # Foto na Sidebar (só renderiza se o aluno estiver confirmado)
        if selected_student and selected_student != "-- Novo Registro --":
            dados_atuais = st.session_state.get('data_pei', {}) if st.session_state.get('data_pei', {}).get('nome') == selected_student else st.session_state.get('data_case', {})
            
            if dados_atuais.get('foto_url'):
                st.image(dados_atuais['foto_url'], use_container_width=True)
            elif dados_atuais.get('foto_base64'):
                try:
                    img_bytes_sb = base64.b64decode(dados_atuais['foto_base64'])
                    st.image(img_bytes_sb, use_container_width=True)
                except: pass
        
        st.markdown('<div style="flex-grow: 1;"></div>', unsafe_allow_html=True)


    # 4. RODAPÉ FIXO (já existe no seu código)

    # 4. RODAPÉ FIXO
    if st.sidebar.button("🚪 Sair", use_container_width=True):
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.rerun()


    # Confirmação de exclusão
    if st.session_state.get("confirm_delete"):
        if is_monitor:
             st.session_state.confirm_delete = False
             st.error("Monitores não podem excluir alunos.")
        else:
            st.warning(f"Excluir {selected_student}?")
            col_d1, col_d2 = st.columns(2)
            if col_d1.button("✅ Sim"):
                delete_student(selected_student)
                st.session_state.confirm_delete = False
                st.rerun()
            if col_d2.button("❌ Não"):
                st.session_state.confirm_delete = False
                st.rerun()

# ==============================================================================
# VIEW: DASHBOARD
# ==============================================================================
if app_mode == "📊 Painel de Gestão":
    # Data e Hora (Fuso BR)
    fuso_br = timezone(timedelta(hours=-3))
    agora = datetime.now(fuso_br)
    
    dias_semana = ["Segunda-feira", "Terça-feira", "Quarta-feira", "Quinta-feira", "Sexta-feira", "Sábado", "Domingo"]
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    
    dia_str = dias_semana[agora.weekday()]
    mes_str = meses[agora.month - 1]
    data_formatada = f"{dia_str}, {agora.day} de {mes_str} de {agora.year}"
    
    st.markdown(f"""
    <div class="header-box" style="margin-top:-50px;">
        <div class="header-title">Painel de Gestão</div>
        <div class="header-subtitle">{data_formatada} | {agora.strftime('%H:%M')}</div>
    </div>
    """, unsafe_allow_html=True)
    
    df_dash = load_db()
    
    # --- CHECK DE ASSINATURAS PENDENTES ---
    pending_docs = []
    user_name_lower = st.session_state.get('usuario_nome', '').strip().lower()
    
    if not df_dash.empty and user_name_lower:
        for idx, row in df_dash.iterrows():
            try:
                d = json.loads(row['dados_json'])
                doc_uuid = d.get('doc_uuid')
                signatures = d.get('signatures', [])
                signed_names = [s.get('name', '').strip().lower() for s in signatures]
                
                # Check fields for possible citation
                fields_to_check = [
                    'prof_poli', 'prof_aee', 'prof_arte', 'prof_ef', 'prof_tec', 'gestor', 'coord', # PEI
                    'resp_sala', 'resp_ee', 'resp_dir', # Avaliação
                    'acompanhante' # Diario
                ]
                
                found_role = None
                for f in fields_to_check:
                    val = d.get(f)
                    if val and isinstance(val, str) and user_name_lower in val.strip().lower():
                        found_role = f
                        break
                
                if found_role and user_name_lower not in signed_names:
                    pending_docs.append(f"{row['nome']} - {row['tipo_doc']}")
            except: pass

    if pending_docs:
        st.warning(f"⚠️ **Atenção:** Você foi citado em {len(pending_docs)} documento(s) e necessita assinar digitalmente.")
        with st.expander("Ver documentos pendentes"):
            for p in pending_docs:
                st.write(f"- {p}")
        st.divider()
    
    # --- CÁLCULO DE MÉTRICAS ---
    # Contagem de alunos únicos
    if not df_dash.empty and "nome" in df_dash.columns:
        total_alunos = df_dash["nome"].nunique()
    else:
        total_alunos = 0
        
    total_pei = len(df_dash[df_dash["tipo_doc"] == "PEI"])
    total_caso = len(df_dash[df_dash["tipo_doc"] == "CASO"])
    total_pdi = len(df_dash[df_dash["tipo_doc"] == "PDI"])
    
    # Função Auxiliar de Progresso
    def calc_progress(row_json, keys_check):
        try:
            data = json.loads(row_json)
            filled = 0
            for k in keys_check:
                val = data.get(k)
                if val:
                    if isinstance(val, list) and len(val) > 0: filled += 1
                    elif isinstance(val, dict) and len(val) > 0: filled += 1
                    elif isinstance(val, str) and val.strip() != "": filled += 1
                    elif isinstance(val, (int, float)): filled += 1
                    elif val is True: filled += 1
            return int((filled / len(keys_check)) * 100)
        except: return 0

    # --- DEFINIÇÃO DAS CHAVES ESSENCIAIS PARA CADA DOCUMENTO ---
    
    keys_pei = [
        'prof_poli', 'prof_aee',       # 1. Identificação
        'defic_txt', 'saude_extra',    # 2. Saúde
        'beh_interesses', 'beh_desafios', # 3. Conduta
        'dev_afetivo',                 # 4. Escolar
        'aval_port', 'aval_ling_verbal', # 5. Acadêmico (um dos dois)
        'meta_social_obj', 'meta_acad_obj', # 6. Metas
        'plano_obs_geral'              # Final
    ]

    keys_caso = [
        'endereco', 'quem_mora',                   # Identificação e Família
        'hist_idade_entrou', 'gest_parentesco',    # Histórico e Gestação
        'saude_prob', 'med_uso',                   # Saúde
        'entrevista_prof', 'entrevista_resp'       # Comportamento / Entrevista
    ]

    keys_aval = [
        'aspectos_gerais', 'defic_chk',            # Identificação
        'alim_nivel', 'hig_nivel', 'loc_nivel',    # Parte I
        'comportamento', 'part_grupo', 'interacao',# Parte II
        'rotina', 'ativ_pedag',                    # Parte III
        'atencao_sust', 'linguagem',               # Parte IV
        'conclusao_nivel', 'resp_ee'               # Conclusão
    ]

    keys_pdi = [
        'potencialidades', 'areas_interesse',      # Avaliação Inicial
        'acao_escola', 'acao_sala', 'acao_familia',# Ações Necessárias
        'aee_tempo', 'aee_tipo',                   # Organização AEE
        'goals_specific'                           # Objetivos Detalhados
    ]
    
    concluidos = 0
    deficiencies_count = {}
    
    # --- INICIALIZAÇÃO DAS LISTAS DE PROGRESSO ---
    pei_progress_list = []
    caso_progress_list = []
    apoio_progress_list = []
    pdi_progress_list = []

    # --- LOOP DE CÁLCULO GERAL ---
    for idx, row in df_dash.iterrows():
        try:
            d = json.loads(row['dados_json'])
            
            # Gráfico de Deficiências
            for dtype in d.get('diag_tipo', []):
                deficiencies_count[dtype] = deficiencies_count.get(dtype, 0) + 1
            if "Deficiência" in d.get('diag_tipo', []) and d.get('defic_txt'):
                d_txt = d.get('defic_txt').upper().strip()
                deficiencies_count[d_txt] = deficiencies_count.get(d_txt, 0) + 1
            
            # Separação por Tipo de Documento e Cálculo
            tipo_documento = row['tipo_doc']
            nome_aluno = row['nome']
            
            if tipo_documento == "PEI":
                prog = calc_progress(row['dados_json'], keys_pei)
                pei_progress_list.append({"Aluno": nome_aluno, "Progresso": prog})
                if prog >= 90: concluidos += 1
                
            elif tipo_documento == "CASO":
                prog = calc_progress(row['dados_json'], keys_caso)
                caso_progress_list.append({"Aluno": nome_aluno, "Progresso": prog})
                
            elif tipo_documento == "AVALIACAO":
                prog = calc_progress(row['dados_json'], keys_aval)
                apoio_progress_list.append({"Aluno": nome_aluno, "Progresso": prog})
                
            elif tipo_documento == "PDI":
                prog = calc_progress(row['dados_json'], keys_pdi)
                pdi_progress_list.append({"Aluno": nome_aluno, "Progresso": prog})
                
        except: pass

# --- CÁLCULO DE NOVAS MÉTRICAS DE GESTÃO ---
    
    # 1. Total de Alunos Únicos
    total_alunos = df_dash["nome"].nunique() if not df_dash.empty and "nome" in df_dash.columns else 0
    
    # 2. Alunos com Laudo Médico / Diagnóstico Conclusivo (NOVA MÉTRICA)
    alunos_com_laudo = set()
    if not df_dash.empty:
        for _, row in df_dash.iterrows():
            try:
                d_laudo = json.loads(row['dados_json'])
                # Checa no PEI se marcou "Sim" para diagnóstico conclusivo
                if row['tipo_doc'] == "PEI" and d_laudo.get('diag_status') == "Sim":
                    alunos_com_laudo.add(row['nome'])
                # Ou checa no Estudo de Caso se o campo "Possui diagnóstico" foi preenchido
                elif row['tipo_doc'] == "CASO" and d_laudo.get('diag_possui') and str(d_laudo.get('diag_possui')).strip():
                    alunos_com_laudo.add(row['nome'])
            except: pass
    total_laudos = len(alunos_com_laudo)
    
    # 3. Documentos em Elaboração (PEIs e PDIs abaixo de 100%)
    docs_em_elaboracao = sum(1 for p in pei_progress_list + pdi_progress_list if p['Progresso'] < 100)
    
    # 4. Alunos com necessidade de Profissional de Apoio (Extraído da Avaliação)
    total_apoio = 0
    if not df_dash.empty:
        df_aval = df_dash[df_dash["tipo_doc"] == "AVALIACAO"]
        for _, row in df_aval.iterrows():
            try:
                d_aval = json.loads(row['dados_json'])
                nivel = d_aval.get('conclusao_nivel', '')
                if "Nível 2" in nivel or "Nível 3" in nivel or d_aval.get('apoio_existente'):
                    total_apoio += 1
            except: pass

    # 5. Estudos de Caso Realizados 
    total_caso = len(df_dash[df_dash["tipo_doc"] == "CASO"]) if not df_dash.empty else 0


    # --- CARDS DE MÉTRICAS ---
    # CSS inline para dar destaque aos números que exigem atenção
    cor_elaboracao = "#ea580c" if docs_em_elaboracao > 0 else "#64748b"

    col_m1, col_m2, col_m3, col_m4, col_m5 = st.columns(5)
    
    col_m1.markdown(f'<div class="metric-card"><div class="metric-value">{total_alunos}</div><div class="metric-label">👥 Total AEE</div></div>', unsafe_allow_html=True)
    
    col_m2.markdown(f'<div class="metric-card"><div class="metric-value">{total_apoio}</div><div class="metric-label">🤝 Apoio Escolar</div></div>', unsafe_allow_html=True)
    
    col_m3.markdown(f'<div class="metric-card"><div class="metric-value" style="color: {cor_elaboracao};">{docs_em_elaboracao}</div><div class="metric-label">⏳ Em elaboração</div></div>', unsafe_allow_html=True)
    
    col_m4.markdown(f'<div class="metric-card"><div class="metric-value" style="color: #0284c7;">{total_laudos}</div><div class="metric-label">📄 Com Laudo</div></div>', unsafe_allow_html=True)
    
    col_m5.markdown(f'<div class="metric-card"><div class="metric-value" style="color: #1e3a8a;">{total_caso}</div><div class="metric-label">📋 Estudos Caso</div></div>', unsafe_allow_html=True)
    
    st.divider()

# --- ABAS DO DASHBOARD ---
    tab_graf, tab_concluidos, tab_com = st.tabs(["📊 Estatísticas & Progresso", "✅ Documentos Concluídos", "📢 Comunicação & Agenda"])
    
    with tab_graf:
        c_chart, c_prog = st.columns([1, 1])
        with c_chart:
            st.subheader("Tipos de Deficiência")
            if deficiencies_count:
                df_def = pd.DataFrame(list(deficiencies_count.items()), columns=["Tipo", "Qtd"])
                st.bar_chart(df_def.set_index("Tipo"), color="#1e3a8a")
            else:
                st.info("Sem dados suficientes.")
        
        with c_prog:
            st.subheader("Progresso de Preenchimento")
            
            # 1. Cria o seletor de documentos
            tipo_doc = st.selectbox(
                "Selecione o documento:",
                ["PEI", "Estudo de Caso", "Avaliação de Apoio", "PDI"],
                label_visibility="collapsed" # Esconde o rótulo para ficar mais limpo
            )
            
            # 2. Define qual lista usar baseado na seleção
            lista_progresso_atual = []
            
            if tipo_doc == "PEI":
                # Sua lista original que já funciona
                lista_progresso_atual = pei_progress_list 
            elif tipo_doc == "Estudo de Caso":
                # Você precisará ter essa lista calculada no seu backend
                lista_progresso_atual = caso_progress_list 
            elif tipo_doc == "Avaliação de Apoio":
                # Você precisará ter essa lista calculada no seu backend
                lista_progresso_atual = apoio_progress_list 
            elif tipo_doc == "PDI":
                # Você precisará ter essa lista calculada no seu backend
                lista_progresso_atual = pdi_progress_list 

            # 3. Renderiza os gráficos da lista escolhida
            if lista_progresso_atual:
                # Opcional: ascending=False deixa os mais completos no topo
                df_prog = pd.DataFrame(lista_progresso_atual).sort_values("Progresso", ascending=False) 
                with st.container(height=300):
                    for _, row in df_prog.iterrows():
                        st.caption(f"{row['Aluno']} ({row['Progresso']}%)")
                        st.progress(row['Progresso'] / 100)
            else:
                st.info(f"Nenhum {tipo_doc} calculado ainda.")

    with tab_com:
        c_aviso, c_agenda = st.columns([1, 1])

        # --- ABA DE CONCLUÍDOS (Agora dentro da variável correta) ---
    with tab_concluidos:
        st.subheader("Documentos Prontos para Emissão")
        
        if not df_dash.empty:
            # Criamos a tabela visual dos concluídos
            lista_concluidos = []
            for idx, row in df_dash.iterrows():
                try:
                    d = json.loads(row['dados_json'])
                    if d.get('status_elaboracao') == "Concluído":
                        lista_concluidos.append(row)
                except: continue
    
            if lista_concluidos:
                for row in lista_concluidos:
                    c_nome, c_tipo, c_btn = st.columns([3, 2, 1])
                    
                    c_nome.write(f"👤 **{row['nome']}**")
                    c_tipo.caption(f"{row['tipo_doc']}")
                    
                    # O TRUQUE: O botão abaixo apenas seleciona o aluno e muda o modo
                    if c_btn.button("📄 Abrir para Baixar", key=f"go_{row['id']}"):
                        st.session_state.ee_aluno_confirmado = row['nome']
                        # Define o tipo de documento correto para o sistema carregar
                        if row['tipo_doc'] == "PEI":
                            st.session_state.ee_doc_confirmado = "PEI - Ensino Fundamental" # Ou Infantil conforme o banco
                        else:
                            st.session_state.ee_doc_confirmado = row['tipo_doc']
                        
                        st.session_state.aluno_selecionado = row['nome']
                        # Força o sistema a ir para a aba de Gestão de Alunos onde o PDF já está pronto
                        st.session_state.app_mode = "👥 Gestão de Alunos" 
                        st.rerun()
            else:
                st.info("Nenhum documento concluído encontrado.")
            
        
        # --- MURAL DE AVISOS ---
        with c_aviso:
            st.markdown("### 📌 Mural de Avisos")
            if not is_monitor:
                with st.form("form_recado"):
                    txt_recado = st.text_area("Novo Recado", height=80)
                    if st.form_submit_button("Publicar"):
                        df_recados = safe_read("Recados", ["Data", "Autor", "Mensagem"])
                        novo_recado = {
                            "Data": datetime.now().strftime("%d/%m %H:%M"),
                            "Autor": st.session_state.get('usuario_nome', 'Admin'),
                            "Mensagem": txt_recado
                        }
                        df_recados = pd.concat([pd.DataFrame([novo_recado]), df_recados], ignore_index=True)
                        safe_update("Recados", df_recados)
                        st.cache_data.clear() # Limpa cache para atualizar
                        time.sleep(1) # Aguarda propagação
                        st.rerun()
            else:
                st.info("Apenas Docentes podem publicar avisos.")
            
            # Listar Recados
            df_recados = safe_read("Recados", ["Data", "Autor", "Mensagem"])
            if not df_recados.empty:
                with st.container(height=300):
                    for index, row in df_recados.iterrows():
                        c_msg, c_del = st.columns([0.85, 0.15])
                        with c_msg:
                            st.info(f"**{row['Autor']}** ({row['Data']}):\n\n{row['Mensagem']}")
                        with c_del:
                            if not is_monitor:
                                if st.button("🗑️", key=f"del_rec_{index}", help="Excluir recado"):
                                    df_recados = df_recados.drop(index)
                                    safe_update("Recados", df_recados)
                                    st.cache_data.clear()
                                    time.sleep(0.5)
                                    st.rerun()
            else:
                st.write("Nenhum recado.")

        # --- AGENDA DA EQUIPE ---
        with c_agenda:
            st.markdown("### 📅 Agenda da Equipe")
            if not is_monitor:
                with st.form("form_agenda"):
                    c_d, c_e = st.columns([1, 2])
                    data_evento = c_d.date_input("Data", format="DD/MM/YYYY")
                    desc_evento = c_e.text_input("Evento")
                    if st.form_submit_button("Agendar"):
                        df_agenda = safe_read("Agenda", ["Data", "Evento", "Autor"])
                        novo_evento = {
                            "Data": data_evento.strftime("%Y-%m-%d"),
                            "Evento": desc_evento,
                            "Autor": st.session_state.get('usuario_nome', 'Admin')
                        }
                        df_agenda = pd.concat([df_agenda, pd.DataFrame([novo_evento])], ignore_index=True)
                        # Ordenar por data
                        df_agenda = df_agenda.sort_values(by="Data", ascending=False)
                        safe_update("Agenda", df_agenda)
                        st.cache_data.clear() # Limpa cache para atualizar
                        time.sleep(1) # Aguarda propagação
                        st.rerun()
            else:
                st.info("Apenas Docentes podem adicionar eventos.")
            
            # Listar Agenda
            df_agenda = safe_read("Agenda", ["Data", "Evento", "Autor"])
            if not df_agenda.empty:
                with st.container(height=300):
                    for index, row in df_agenda.iterrows():
                        try:
                            d_fmt = datetime.strptime(str(row['Data']), "%Y-%m-%d").strftime("%d/%m")
                        except:
                            d_fmt = str(row['Data'])
                        
                        c_evt, c_del_evt = st.columns([0.85, 0.15])
                        with c_evt:
                            st.write(f"🗓️ **{d_fmt}** - {row['Evento']} _({row['Autor']})_")
                        with c_del_evt:
                            if not is_monitor:
                                if st.button("🗑️", key=f"del_agd_{index}", help="Excluir evento"):
                                    df_agenda = df_agenda.drop(index)
                                    safe_update("Agenda", df_agenda)
                                    st.cache_data.clear()
                                    time.sleep(0.5)
                                    st.rerun()
            else:
                st.write("Agenda vazia.")

# ==============================================================================
# VIEW: GESTÃO DE ALUNOS (PEI / CASO)
# ==============================================================================
elif app_mode == "👥 Gestão de Alunos":
    
    # --- PORTÃO DE ENTRADA ---
    if not st.session_state.ee_aluno_confirmado or not st.session_state.ee_doc_confirmado:
        st.markdown("""
        <div style='background-color: white; padding: 2rem; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); border: 1px solid #e2e8f0;'>
            <h3 style='color: #1e293b; margin-top:0;'>🚪 Selecione o Estudante e o Documento</h3>
            <p style='color: #64748b;'>Para iniciar o preenchimento ou edição, selecione abaixo o estudante e o tipo de documento desejado.</p>
        </div>
        <br>
        """, unsafe_allow_html=True)

        df_db = load_db()
        
        # A MÁGICA ACONTECE AQUI: Adicionamos o 'sorted()' para organizar de A a Z
        lista_nomes = sorted(df_db["nome"].dropna().unique().tolist()) if not df_db.empty else []
        
        opcoes_nomes = ["-- Novo Registro --"] + lista_nomes

        c_aluno, c_doc = st.columns(2)
        aluno_sel = c_aluno.selectbox("1. Selecione o Estudante:", opcoes_nomes)

        # Se for novo registro, abre campo para digitar o nome
        nome_novo = ""
        if aluno_sel == "-- Novo Registro --":
            nome_novo = c_aluno.text_input("Digite o nome completo do novo estudante:")

        # LISTA DE DOCUMENTOS ATUALIZADA (PEI Dividido)
        docs_disponiveis = [
            "Estudo de Caso", 
            "PEI - Ensino Fundamental", 
            "PEI - Educação Infantil", 
            "PDI - Pré Escola e Ens. Fundamental", 
            "Protocolo de Conduta", 
            "Avaliação de Apoio", 
            "Avaliação de Apoio 2.0", # <--- ADICIONE ESTA LINHA AQUI
            "Relatório de Acompanhamento", 
            "Declaração de Matrícula"
        ]
        doc_sel = c_doc.selectbox("2. Selecione o Documento:", docs_disponiveis)

        st.write("")
        if st.button("✅ Confirmar e Acessar Documento", type="primary", use_container_width=True):
            nome_final = nome_novo if aluno_sel == "-- Novo Registro --" else aluno_sel
            
            if nome_final and nome_final.strip() != "":
                
                # ==========================================================
                # TRAVA DE SEGURANÇA (PADRÃO OURO): BLINDAGEM DO BANCO DE DADOS
                # ==========================================================
                if doc_sel in ["PEI - Ensino Fundamental", "PEI - Educação Infantil"]:
                    df_val = df_db[df_db['nome'] == nome_final]
                    if not df_val.empty:
                        row_pei = df_val[df_val['tipo_doc'] == 'PEI']
                        if not row_pei.empty:
                            try:
                                dados_salvos = json.loads(row_pei.iloc[0]['dados_json'])
                                # Tenta descobrir o modelo pelo carimbo oculto ou pelos campos já preenchidos
                                modelo_salvo = dados_salvos.get('modelo_pei_salvo')
                                
                                if not modelo_salvo:
                                    if 'aval_ling_verbal' in dados_salvos and str(dados_salvos['aval_ling_verbal']).strip() != "":
                                        modelo_salvo = "Infantil"
                                    elif 'aval_port' in dados_salvos and str(dados_salvos['aval_port']).strip() != "":
                                        modelo_salvo = "Fundamental"

                                # Se descobrir que já existe um PEI incompatível, trava a entrada!
                                if modelo_salvo:
                                    if doc_sel == "PEI - Ensino Fundamental" and modelo_salvo == "Infantil":
                                        st.error("⚠️ **RISCO DE PERDA DE DADOS!** Este estudante já possui um PEI da **Educação Infantil** salvo no sistema. Por favor, troque a seleção acima para 'PEI - Educação Infantil' para não apagar o histórico.")
                                        st.stop()
                                    if doc_sel == "PEI - Educação Infantil" and modelo_salvo == "Fundamental":
                                        st.error("⚠️ **RISCO DE PERDA DE DADOS!** Este estudante já possui um PEI do **Ensino Fundamental** salvo no sistema. Por favor, troque a seleção acima para 'PEI - Ensino Fundamental' para não apagar o histórico.")
                                        st.stop()
                            except:
                                pass
                # ==========================================================

                st.session_state.ee_aluno_confirmado = nome_final
                st.session_state.ee_doc_confirmado = doc_sel
                
                # Seta o nome para a função carregar_dados_aluno puxar do banco
                st.session_state.aluno_selecionado = nome_final 
                carregar_dados_aluno()
                st.rerun()
            else:
                st.warning("⚠️ Por favor, informe o nome do estudante.")

        # Interrompe a renderização para não mostrar os docs em branco embaixo
        st.stop() 

    # --- BARRA DE INFORMAÇÃO E NAVEGAÇÃO ---
    c_info, c_btn, c_del = st.columns([6, 2, 1])
    c_info.success(f"📌 **Documento em edição:** {st.session_state.ee_doc_confirmado} - do estudante {st.session_state.ee_aluno_confirmado}")
    
    if c_btn.button("⬅️ Trocar Aluno/Documento", use_container_width=True):
        st.session_state.ee_aluno_confirmado = None
        st.session_state.ee_doc_confirmado = None
        st.session_state.aluno_selecionado = None
        st.rerun()
        
    if not is_monitor:
        if c_del.button("🗑️ Excluir", use_container_width=True, help="Atenção: Isso excluirá o estudante e TODOS os seus documentos do banco."):
            st.session_state.confirm_delete = True

    # --- LÓGICA DE INTERPRETAÇÃO DO DOCUMENTO (PEI Fundamental vs Infantil) ---
    escolha_doc = st.session_state.ee_doc_confirmado
    selected_student = st.session_state.ee_aluno_confirmado
    
    if escolha_doc == "PEI - Ensino Fundamental":
        doc_mode = "PEI"
        pei_level = "Fundamental"
    elif escolha_doc == "PEI - Educação Infantil":
        doc_mode = "PEI"
        pei_level = "Infantil"
    else:
        doc_mode = escolha_doc
        pei_level = "Fundamental" # Padrão para os outros documentos

    st.divider()

# --- SEÇÃO PEI ---
    if doc_mode == "PEI":
        st.markdown(f"""<div class="header-box"><div class="header-title">Plano Educacional Individualizado - PEI ({pei_level})</div></div>""", unsafe_allow_html=True)
        st.markdown("""<style>div[data-testid="stFormSubmitButton"] > button {width: 100%; background-color: #dcfce7; color: #166534; border: 1px solid #166534;}</style>""", unsafe_allow_html=True)

        # Inicializa o dicionário de dados da sessão
        data = st.session_state.data_pei

        data['modelo_pei_salvo'] = pei_level

        # --- NOVO BLOCO DE STATUS (ENTRE O TÍTULO E O FORMULÁRIO) ---
        # Verificação de segurança: só exibe se houver um aluno carregado
        if 'data_pei' in st.session_state and st.session_state.data_pei:
            
            # Garante que as chaves de status existam no dicionário
            if 'status_elaboracao' not in st.session_state.data_pei:
                st.session_state.data_pei['status_elaboracao'] = "Em elaboração"
            
            st.markdown("---")
            col_icon, col_status, col_trim = st.columns([0.1, 0.4, 0.5])

            with col_icon:
                # O ícone muda visualmente de acordo com a escolha
                if st.session_state.data_pei.get('status_elaboracao') == "Concluído":
                    st.markdown("### ✅")
                else:
                    st.markdown("### 📝")

            with col_status:
                status_selecionado = st.radio(
                    "**Situação do PEI:**",
                    ["Em elaboração", "Concluído"],
                    index=0 if st.session_state.data_pei['status_elaboracao'] == "Em elaboração" else 1,
                    horizontal=True,
                    key="radio_status_pei_final"
                )
                st.session_state.data_pei['status_elaboracao'] = status_selecionado

            # Só mostra a seleção de trimestre se estiver "Concluído"
            if status_selecionado == "Concluído":
                with col_trim:
                    trim_opcoes = ["1º Trimestre", "2º Trimestre", "3º Trimestre"]
                    val_p = st.session_state.data_pei.get('trimestre_finalizado', "1º Trimestre")
                    
                    # Garante que o index seja válido
                    idx_p = trim_opcoes.index(val_p) if val_p in trim_opcoes else 0
                    
                    st.session_state.data_pei['trimestre_finalizado'] = st.selectbox(
                        "Indique o trimestre concluído:",
                        trim_opcoes,
                        index=idx_p,
                        key="select_trim_concluido"
                    )
            
            st.markdown("---")
        # --- FIM DO BLOCO DE STATUS ---

        # Definição das abas (Mantenha o alinhamento destas linhas com o 'if doc_mode')
        tabs = st.tabs(["1. Identificação", "2. Saúde", "3. Conduta", "4. Escolar", "5. Acadêmico", "6. Metas/Flex", "7. Assinaturas", "8. Emissão", "9. Histórico"])

        with tabs[0]:
            with st.form("form_pei_identificacao") if not is_monitor else st.container():
                st.subheader("1. Identificação")
                
                # --- LAYOUT COM FOTO ---
                col_img, col_data = st.columns([1, 4])
                
                with col_img:
                    st.markdown("📷 **Foto**")
                    # Se ja tiver foto, mostra
                    # Se já tiver foto (URL ou Base64), mostra
                    if data.get('foto_url'):
                        st.image(data['foto_url'], use_container_width=True)
                        if not is_monitor:
                            if st.checkbox("Remover", key="rem_foto_pei"):
                                data['foto_url'] = None
                                
                    elif data.get('foto_base64'):
                        try:
                            b = base64.b64decode(data['foto_base64'])
                            st.image(b, use_container_width=True)
                            if not is_monitor:
                                if st.checkbox("Remover", key="rem_foto_pei"):
                                    data['foto_base64'] = None
                        except:
                            st.error("Erro ao carregar foto Base64")
                    
                    # Upload
                    uploaded_photo = st.file_uploader("Carregar", type=["jpg", "jpeg", "png"], label_visibility="collapsed", key="up_foto_pei", disabled=is_monitor)
                    if uploaded_photo and not is_monitor:
                        try:
                            img = Image.open(uploaded_photo)
                            if img.mode != 'RGB': img = img.convert('RGB')
                            # Resize para não pesar no banco
                            img.thumbnail((300, 400))
                            buf = io.BytesIO()
                            img.save(buf, format="JPEG", quality=85)
                            data['foto_base64'] = base64.b64encode(buf.getvalue()).decode()
                            st.success("OK!")
                        except Exception as e:
                            st.error(f"Erro: {e}")
                
                with col_data:
                    c1, c2 = st.columns([3, 1])
                    data['nome'] = c1.text_input("Nome", value=data.get('nome', ''), disabled=True)
                    
                    # --- TRATAMENTO ROBUSTO ---
                    d_val = data.get('nasc')
                    if isinstance(d_val, str): 
                        try: d_val = datetime.strptime(d_val, '%Y-%m-%d').date()
                        except: d_val = date.today()
                    
                    if not isinstance(d_val, date):
                        d_val = date.today()

                    # Garante que o valor inicial não quebre o widget
                    d_val = max(MIN_DATA, min(d_val, MAX_DATA))

                    # CHAVE ÚNICA PARA FORÇAR O RESET
                    # Mudando o nome da chave de 'data_nasc_unique' para 'nasc_fix_v3'
                    # o Streamlit ignora o bloqueio anterior de 2016.
                    input_key = f"nasc_fix_{data.get('nome', 'novo').replace(' ', '_')}"

                    data['nasc'] = c2.date_input(
                        "Nascimento", 
                        value=d_val,
                        min_value=MIN_DATA, 
                        max_value=MAX_DATA,
                        format="DD/MM/YYYY", 
                        disabled=is_monitor,
                        key=input_key
                    )
                    
                    c3, c4 = st.columns(2)
                    data['idade'] = c3.text_input("Idade", value=data.get('idade', ''), disabled=is_monitor)
                    data['ano_esc'] = c4.text_input("Ano Escolar", value=data.get('ano_esc', ''), disabled=is_monitor)
                    
                    data['mae'] = st.text_input("Nome da Mãe", value=data.get('mae', ''), disabled=is_monitor)
                    data['pai'] = st.text_input("Nome do Pai", value=data.get('pai', ''), disabled=is_monitor)
                    data['tel'] = st.text_input("Telefone", value=data.get('tel', ''), disabled=is_monitor)
                
                st.markdown("**Docentes Responsáveis**")
                d1, d2, d3 = st.columns(3)
                data['prof_poli'] = d1.text_input("Polivalente/Regente", value=data.get('prof_poli', ''), disabled=is_monitor)
                data['prof_aee'] = d2.text_input("Prof. AEE", value=data.get('prof_aee', ''), disabled=is_monitor)
                data['prof_arte'] = d3.text_input("Arte", value=data.get('prof_arte', ''), disabled=is_monitor)
                
                d4, d5, d6 = st.columns(3)
                data['prof_ef'] = d4.text_input("Ed. Física", value=data.get('prof_ef', ''), disabled=is_monitor)
                data['prof_tec'] = d5.text_input("Tecnologia", value=data.get('prof_tec', ''), disabled=is_monitor)
                data['gestor'] = d6.text_input("Gestor Escolar", value=data.get('gestor', ''), disabled=is_monitor)
                
                dg1, dg2 = st.columns(2)
                data['coord'] = dg1.text_input("Coordenação", value=data.get('coord', ''), disabled=is_monitor)
                data['revisoes'] = st.text_input("Revisões", value=data.get('revisoes', ''), disabled=is_monitor)
                
                elab_opts = ["1º Trimestre", "2º Trimestre", "3º Trimestre", "Anual"]
                idx_elab = elab_opts.index(data['elab_per']) if data.get('elab_per') in elab_opts else 0
                data['elab_per'] = st.selectbox("Período", elab_opts, index=idx_elab, disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Identificação"):
                        save_student("PEI", data.get('nome'), data, "Identificação")

        # --- ABA 2: SAÚDE ---
        with tabs[1]:
            with st.form("form_pei_saude") if not is_monitor else st.container():
                st.subheader("Informações de Saúde")
                diag_idx = 0 if data.get('diag_status') == "Sim" else 1
                data['diag_status'] = st.radio("Diagnóstico conclusivo?", ["Sim", "Não"], horizontal=True, index=diag_idx, disabled=is_monitor)
                
                c_l1, c_l2 = st.columns(2)
                ld_val = data.get('laudo_data')
                
                # Tenta converter a string em data; se falhar ou estiver vazio, define como None
                if isinstance(ld_val, str) and ld_val.strip(): 
                    try: 
                        ld_val = datetime.strptime(ld_val, '%Y-%m-%d').date()
                    except ValueError: 
                        ld_val = None
                elif not ld_val:
                    ld_val = None

                data['laudo_data'] = c_l1.date_input(
                    "Data do Laudo Médico", 
                    value=ld_val,  # Passando None, o campo permanecerá em branco
                    format="DD/MM/YYYY", 
                    disabled=is_monitor
                )
                
                st.markdown("Categorias de Diagnóstico:")
                cats = ["Deficiência", "Transtorno do Neurodesenvolvimento", "Transtornos Aprendizagem", "AH/SD", "Outros"]
                if 'diag_tipo' not in data: data['diag_tipo'] = []
                
                c_c1, c_c2 = st.columns(2)
                for i, cat in enumerate(cats):
                    col = c_c1 if i % 2 == 0 else c_c2
                    is_checked = cat in data['diag_tipo']
                    if col.checkbox(cat, value=is_checked, key=f"pei_chk_{i}", disabled=is_monitor):
                        if cat not in data['diag_tipo']: data['diag_tipo'].append(cat)
                    else:
                        if cat in data['diag_tipo']: data['diag_tipo'].remove(cat)
                
                data['defic_txt'] = st.text_input("Descrição da Deficiência", value=data.get('defic_txt', ''), disabled=is_monitor)
                data['neuro_txt'] = st.text_input("Descrição do Transtorno Neuro", value=data.get('neuro_txt', ''), disabled=is_monitor)
                data['aprend_txt'] = st.text_input("Descrição do Transtorno de Aprendizagem", value=data.get('aprend_txt', ''), disabled=is_monitor)

                st.divider()
                st.markdown("**Terapias que realiza**")
                especs = ["Psicologia", "Fonoaudiologia", "Terapia Ocupacional", "Psicopedagogia", "Fisioterapia", "Outros"]
                if 'terapias' not in data: data['terapias'] = {}
                
                for esp in especs:
                    st.markdown(f"**{esp}**")
                    if esp not in data['terapias']: data['terapias'][esp] = {'realiza': False, 'dias': [], 'horario': ''}
                    
                    c_t1, c_t2, c_t3 = st.columns([1, 2, 2])
                    data['terapias'][esp]['realiza'] = c_t1.checkbox("Realiza?", value=data['terapias'][esp].get('realiza', False), key=f"pei_terapias_{esp}", disabled=is_monitor)
                    
                    data['terapias'][esp]['dias'] = c_t2.multiselect("Dias", ["2ª", "3ª", "4ª", "5ª", "6ª", "Sábado", "Domingo"], default=data['terapias'][esp].get('dias', []), key=f"pei_dias_{esp}", disabled=is_monitor)
                    data['terapias'][esp]['horario'] = c_t3.text_input("Horário", value=data['terapias'][esp].get('horario', ''), key=f"pei_hor_{esp}", disabled=is_monitor)
                    
                    if esp == "Outros":
                        data['terapias'][esp]['nome_custom'] = st.text_input("Especifique (Outros):", value=data['terapias'][esp].get('nome_custom', ''), key="pei_custom_name", disabled=is_monitor)
                    st.divider()

                data['med_nome'] = st.text_area("Nome da(s) Medicação(ões)", value=data.get('med_nome', ''), disabled=is_monitor)
                m1, m2 = st.columns(2)
                data['med_hor'] = m1.text_input("Horário(s)", value=data.get('med_hor', ''), disabled=is_monitor)
                data['med_doc'] = m2.text_input("Médico Responsável (Medicação)", value=data.get('med_doc', ''), disabled=is_monitor)
                data['med_obj'] = st.text_area("Objetivo da medicação", value=data.get('med_obj', ''), disabled=is_monitor)
                data['saude_extra'] = st.text_area("Outras informações de saúde:", value=data.get('saude_extra', ''), disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Saúde"):
                        save_student("PEI", data.get('nome'), data, "Saúde")

        # --- ABA 3: CONDUTA ---
        with tabs[2]:
            with st.form("form_pei_conduta") if not is_monitor else st.container():
                st.subheader("3. Protocolo de Conduta")
                st.markdown("### 🗣️ COMUNICAÇÃO")
                com_opts = ["Oralmente", "Não comunica", "Não se aplica", "Comunicação alternativa"]
                idx_com = com_opts.index(data['com_tipo']) if data.get('com_tipo') in com_opts else 0
                data['com_tipo'] = st.selectbox("Como o estudante se comunica?", com_opts, index=idx_com, disabled=is_monitor)
                data['com_alt_espec'] = st.text_input("Especifique (Comunicação alternativa):", value=data.get('com_alt_espec', ''), disabled=is_monitor)
                
                nec_idx = 0 if data.get('com_necessidades') == 'Sim' else 1
                data['com_necessidades'] = st.radio("Expressa necessidades/desejos?", ["Sim", "Não"], horizontal=True, index=nec_idx, disabled=is_monitor)
                data['com_necessidades_espec'] = st.text_input("Especifique necessidades:", value=data.get('com_necessidades_espec', ''), disabled=is_monitor)
                
                cha_idx = 0 if data.get('com_chamado') == 'Sim' else 1
                data['com_chamado'] = st.radio("Atende quando é chamado?", ["Sim", "Não"], horizontal=True, index=cha_idx, disabled=is_monitor)
                data['com_chamado_espec'] = st.text_input("Especifique chamado:", value=data.get('com_chamado_espec', ''), disabled=is_monitor)
                
                cmd_idx = 0 if data.get('com_comandos') == 'Sim' else 1
                data['com_comandos'] = st.radio("Responde a comandos simples?", ["Sim", "Não"], horizontal=True, index=cmd_idx, disabled=is_monitor)
                data['com_comandos_espec'] = st.text_input("Especifique comandos:", value=data.get('com_comandos_espec', ''), disabled=is_monitor)

                st.divider()
                st.markdown("### 🚶 LOCOMOÇÃO")
                loc_r_idx = 1 if data.get('loc_reduzida') == 'Sim' else 0
                data['loc_reduzida'] = st.radio("Possui mobilidade reduzida?", ["Não", "Sim"], horizontal=True, index=loc_r_idx, disabled=is_monitor)
                data['loc_reduzida_espec'] = st.text_input("Especifique mobilidade:", value=data.get('loc_reduzida_espec', ''), disabled=is_monitor)
                
                c_l1, c_l2 = st.columns(2)
                amb_idx = 0 if data.get('loc_ambiente') == 'Sim' else 1
                data['loc_ambiente'] = c_l1.radio("Locomove-se pela casa?", ["Sim", "Não"], horizontal=True, index=amb_idx, disabled=is_monitor)
                helper_idx = 0 if data.get('loc_ambiente_ajuda') == 'Com autonomia' else 1
                data['loc_ambiente_ajuda'] = c_l2.selectbox("Grau:", ["Com autonomia", "Com ajuda"], index=helper_idx, disabled=is_monitor)
                data['loc_ambiente_espec'] = st.text_input("Especifique locomoção:", value=data.get('loc_ambiente_espec', ''), disabled=is_monitor)

                st.divider()
                st.markdown("### 🧼 AUTOCUIDADO E HIGIENE")
                c_h1, c_h2 = st.columns(2)
                wc_idx = 0 if data.get('hig_banheiro') == 'Sim' else 1
                data['hig_banheiro'] = c_h1.radio("Utiliza o banheiro?", ["Sim", "Não"], horizontal=True, index=wc_idx, disabled=is_monitor)
                wc_help_idx = 0 if data.get('hig_banheiro_ajuda') == 'Com autonomia' else 1
                data['hig_banheiro_ajuda'] = c_h2.selectbox("Ajuda banheiro:", ["Com autonomia", "Com ajuda"], index=wc_help_idx, disabled=is_monitor)
                data['hig_banheiro_espec'] = st.text_input("Especifique banheiro:", value=data.get('hig_banheiro_espec', ''), disabled=is_monitor)
                
                c_h3, c_h4 = st.columns(2)
                tooth_idx = 0 if data.get('hig_dentes') == 'Sim' else 1
                data['hig_dentes'] = c_h3.radio("Escova os dentes?", ["Sim", "Não"], horizontal=True, index=tooth_idx, disabled=is_monitor)
                tooth_help_idx = 0 if data.get('hig_dentes_ajuda') == 'Com autonomia' else 1
                data['hig_dentes_ajuda'] = c_h4.selectbox("Ajuda dentes:", ["Com autonomia", "Com ajuda"], index=tooth_help_idx, disabled=is_monitor)
                data['hig_dentes_espec'] = st.text_input("Especifique dentes:", value=data.get('hig_dentes_espec', ''), disabled=is_monitor)

                st.divider()
                st.markdown("### 🧩 COMPORTAMENTO")
                data['beh_interesses'] = st.text_area("Interesses do estudante:", value=data.get('beh_interesses', ''), disabled=is_monitor)
                data['beh_objetos_gosta'] = st.text_area("Objetos que gosta / Apego:", value=data.get('beh_objetos_gosta', ''), disabled=is_monitor)
                data['beh_objetos_odeia'] = st.text_area("Objetos que não gosta / Aversão:", value=data.get('beh_objetos_odeia', ''), disabled=is_monitor)
                data['beh_toque'] = st.text_area("Gosta de toque/abraço?", value=data.get('beh_toque', ''), disabled=is_monitor)
                data['beh_calmo'] = st.text_area("O que o acalma?", value=data.get('beh_calmo', ''), disabled=is_monitor)
                data['beh_atividades'] = st.text_area("Atividades prazerosas:", value=data.get('beh_atividades', ''), disabled=is_monitor)
                data['beh_gatilhos'] = st.text_area("Gatilhos de crise:", value=data.get('beh_gatilhos', ''), disabled=is_monitor)
                data['beh_crise_regula'] = st.text_area("Como se regula na crise?", value=data.get('beh_crise_regula', ''), disabled=is_monitor)
                data['beh_desafios'] = st.text_area("Comportamentos desafiadores / Manejo:", value=data.get('beh_desafios', ''), disabled=is_monitor)
                
                c_b1, c_b2 = st.columns([1, 2])
                food_idx = 1 if data.get('beh_restricoes') == 'Sim' else 0
                data['beh_restricoes'] = c_b1.radio("Restrições alimentares?", ["Não", "Sim"], horizontal=True, index=food_idx, disabled=is_monitor)
                data['beh_restricoes_espec'] = c_b2.text_input("Especifique alimentação:", value=data.get('beh_restricoes_espec', ''), disabled=is_monitor)
                
                c_b3, c_b4 = st.columns([1, 2])
                water_idx = 0 if data.get('beh_autonomia_agua') == 'Sim' else 1
                data['beh_autonomia_agua'] = c_b3.radio("Autonomia (água/comida)?", ["Sim", "Não"], horizontal=True, index=water_idx, disabled=is_monitor)
                data['beh_autonomia_agua_espec'] = c_b4.text_input("Especifique autonomia:", value=data.get('beh_autonomia_agua_espec', ''), disabled=is_monitor)
                
                data['beh_pertinentes'] = st.text_area("Outras informações:", value=data.get('beh_pertinentes', ''), disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Conduta"):
                        save_student("PEI", data.get('nome'), data, "Conduta")

        # --- ABA 4: ESCOLAR ---
        with tabs[3]:
            with st.form("form_pei_escolar") if not is_monitor else st.container():
                st.subheader("4. Desenvolvimento Escolar")
                
                c_p1, c_p2 = st.columns([1, 2])
                perm_opts = ["Sim - Por longo período", "Sim - Por curto período", "Não"]
                idx_perm = perm_opts.index(data.get('dev_permanece')) if data.get('dev_permanece') in perm_opts else 0
                data['dev_permanece'] = c_p1.selectbox("Permanece em sala?", perm_opts, index=idx_perm, disabled=is_monitor)
                data['dev_permanece_espec'] = c_p2.text_input("Obs Permanência:", value=data.get('dev_permanece_espec', ''), disabled=is_monitor)

                c_i1, c_i2 = st.columns([1, 2])
                int_idx = 0 if data.get('dev_integrado') == 'Sim' else 1
                data['dev_integrado'] = c_i1.radio("Integrado ao ambiente?", ["Sim", "Não"], horizontal=True, index=int_idx, disabled=is_monitor)
                data['dev_integrado_espec'] = c_i2.text_input("Obs Integração:", value=data.get('dev_integrado_espec', ''), disabled=is_monitor)

                c_l1, c_l2 = st.columns([1, 2])
                loc_opts = ["Sim - Com autonomia", "Sim - Com ajuda", "Não"]
                idx_loc = loc_opts.index(data.get('dev_loc_escola')) if data.get('dev_loc_escola') in loc_opts else 0
                data['dev_loc_escola'] = c_l1.selectbox("Locomove-se pela escola?", loc_opts, index=idx_loc, disabled=is_monitor)
                data['dev_loc_escola_espec'] = c_l2.text_input("Obs Locomoção:", value=data.get('dev_loc_escola_espec', ''), disabled=is_monitor)

                c_t1, c_t2 = st.columns([1, 2])
                tar_opts = ["Sim - Com autonomia", "Sim - Com ajuda", "Não"]
                idx_tar = tar_opts.index(data.get('dev_tarefas')) if data.get('dev_tarefas') in tar_opts else 0
                data['dev_tarefas'] = c_t1.selectbox("Realiza tarefas?", tar_opts, index=idx_tar, disabled=is_monitor)
                data['dev_tarefas_espec'] = c_t2.text_input("Obs Tarefas:", value=data.get('dev_tarefas_espec', ''), disabled=is_monitor)

                c_a1, c_a2 = st.columns([1, 2])
                amg_idx = 0 if data.get('dev_amigos') == 'Sim' else 1
                data['dev_amigos'] = c_a1.radio("Tem amigos?", ["Sim", "Não"], horizontal=True, index=amg_idx, disabled=is_monitor)
                data['dev_amigos_espec'] = c_a2.text_input("Obs Amigos:", value=data.get('dev_amigos_espec', ''), disabled=is_monitor)

                data['dev_colega_pref'] = st.radio("Tem colega predileto?", ["Sim", "Não"], horizontal=True, index=0 if data.get('dev_colega_pref') == 'Sim' else 1, disabled=is_monitor)

                c_ia1, c_ia2 = st.columns([1, 2])
                ia_idx = 0 if data.get('dev_participa') == 'Sim' else 1
                data['dev_participa'] = c_ia1.radio("Participa/Interage?", ["Sim", "Não"], horizontal=True, index=ia_idx, disabled=is_monitor)
                data['dev_participa_espec'] = c_ia2.text_input("Obs Interação:", value=data.get('dev_participa_espec', ''), disabled=is_monitor)

                data['dev_afetivo'] = st.text_area("Envolvimento afetivo/social da turma:", value=data.get('dev_afetivo', ''), disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Escolar"):
                        save_student("PEI", data.get('nome'), data, "Escolar")

        # --- ABA 5: ACADÊMICO ---
        with tabs[4]:
            with st.form("form_pei_academico") if not is_monitor else st.container():
                st.subheader("5. Avaliação Acadêmica")
                
                if pei_level == "Fundamental":
                    c_f1, c_f2 = st.columns(2)
                    data['aval_port'] = c_f1.text_area("Língua Portuguesa", value=data.get('aval_port', ''), disabled=is_monitor)
                    data['aval_mat'] = c_f2.text_area("Matemática", value=data.get('aval_mat', ''), disabled=is_monitor)
                    data['aval_con_gerais'] = st.text_area("Conhecimentos Gerais", value=data.get('aval_con_gerais', ''), disabled=is_monitor)

                    st.markdown("**ARTE**")
                    data['aval_arte_visuais'] = st.text_area("Artes Visuais", value=data.get('aval_arte_visuais', ''), disabled=is_monitor)
                    data['aval_arte_musica'] = st.text_area("Música", value=data.get('aval_arte_musica', ''), disabled=is_monitor)
                    c_a1, c_a2 = st.columns(2)
                    data['aval_arte_teatro'] = c_a1.text_area("Teatro", value=data.get('aval_arte_teatro', ''), disabled=is_monitor)
                    data['aval_arte_danca'] = c_a2.text_area("Dança", value=data.get('aval_arte_danca', ''), disabled=is_monitor)

                    st.markdown("**EDUCAÇÃO FÍSICA**")
                    c_ef1, c_ef2 = st.columns(2)
                    data['aval_ef_motoras'] = c_ef1.text_area("Habilidades Motoras", value=data.get('aval_ef_motoras', ''), disabled=is_monitor)
                    data['aval_ef_corp_conhec'] = c_ef2.text_area("Conhecimento Corporal", value=data.get('aval_ef_corp_conhec', ''), disabled=is_monitor)
                    data['aval_ef_exp'] = st.text_area("Exp. Corporais e Expressividade", value=data.get('aval_ef_exp', ''), disabled=is_monitor)
                    
                    st.markdown("**LINGUAGENS E TECNOLOGIAS**")
                    data['aval_ling_tec'] = st.text_area("Avaliação da disciplina:", value=data.get('aval_ling_tec', ''), disabled=is_monitor)
                else:
                    # Infantil
                    data['aval_ling_verbal'] = st.text_area("Linguagem Verbal", value=data.get('aval_ling_verbal', ''), disabled=is_monitor)
                    data['aval_ling_mat'] = st.text_area("Linguagem Matemática", value=data.get('aval_ling_mat', ''), disabled=is_monitor)
                    data['aval_ind_soc'] = st.text_area("Indivíduo e Sociedade", value=data.get('aval_ind_soc', ''), disabled=is_monitor)
                    
                    st.markdown("**ARTE**")
                    data['aval_arte_visuais'] = st.text_area("Artes Visuais", value=data.get('aval_arte_visuais', ''), disabled=is_monitor)
                    data['aval_arte_musica'] = st.text_area("Música", value=data.get('aval_arte_musica', ''), disabled=is_monitor)
                    data['aval_arte_teatro'] = st.text_area("Teatro", value=data.get('aval_arte_teatro', ''), disabled=is_monitor)

                    st.markdown("**EDUCAÇÃO FÍSICA**")
                    c_ef1, c_ef2, c_ef3 = st.columns(3)
                    data['aval_ef_jogos'] = c_ef1.text_area("Jogos/Brincadeiras", value=data.get('aval_ef_jogos', ''), disabled=is_monitor)
                    data['aval_ef_ritmo'] = c_ef2.text_area("Ritmo", value=data.get('aval_ef_ritmo', ''), disabled=is_monitor)
                    data['aval_ef_corp'] = c_ef3.text_area("Conhecimento Corporal", value=data.get('aval_ef_corp', ''), disabled=is_monitor)
                    
                    st.markdown("**LINGUAGEM E TECNOLOGIAS**")
                    data['aval_ling_tec'] = st.text_area("Avaliação da disciplina:", value=data.get('aval_ling_tec', ''), disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Acadêmico"):
                        save_student("PEI", data.get('nome'), data, "Acadêmico")

        # --- ABA 6: METAS E FLEXIBILIZAÇÃO (VERSÃO CORRIGIDA E SEGURA) ---
        with tabs[5]:
            # Identificador único para as keys (evita que dados de um aluno fiquem presos na tela de outro)
            aluno_id = data.get('nome', 'default')

            with st.form("form_pei_metas") if not is_monitor else st.container():
                st.header("6. Metas Específicas")
                
                st.subheader("Habilidades Sociais")
                data['meta_social_obj'] = st.text_area("Metas (Sociais):", value=data.get('meta_social_obj', ''), key=f"ms_obj_{aluno_id}", disabled=is_monitor)
                data['meta_social_est'] = st.text_area("Estratégias (Sociais):", value=data.get('meta_social_est', ''), key=f"ms_est_{aluno_id}", disabled=is_monitor)

                st.divider(); st.subheader("Autocuidado e Vida Prática")
                data['meta_auto_obj'] = st.text_area("Metas (Autocuidado):", value=data.get('meta_auto_obj', ''), key=f"ma_obj_{aluno_id}", disabled=is_monitor)
                data['meta_auto_est'] = st.text_area("Estratégias (Autocuidado):", value=data.get('meta_auto_est', ''), key=f"ma_est_{aluno_id}", disabled=is_monitor)

                st.divider(); st.subheader("Habilidades Acadêmicas")
                data['meta_acad_obj'] = st.text_area("Metas (Acadêmicas):", value=data.get('meta_acad_obj', ''), key=f"mac_obj_{aluno_id}", disabled=is_monitor)
                data['meta_acad_est'] = st.text_area("Estratégias (Acadêmicas):", value=data.get('meta_acad_est', ''), key=f"mac_est_{aluno_id}", disabled=is_monitor)

                st.header("7. Flexibilização Curricular")
                if pei_level == "Fundamental":
                    disciplinas_flex = ["Língua Portuguesa", "Matemática", "História", "Geografia", "Ciências", "Arte", "Educação Física", "Linguagens e Tecnologia"]
                else:
                    disciplinas_flex = ["Linguagem Verbal", "Linguagem Matemática", "Indivíduo e Sociedade", "Arte", "Educação Física", "Linguagens e Tecnologia"]

                if 'flex_matrix' not in data: data['flex_matrix'] = {}
                
                st.markdown("**7.1 Disciplinas que necessitam de adaptação**")
                c_h1, c_h2, c_h3 = st.columns([2, 1, 1])
                c_h1.write("**Disciplina**")
                c_h2.write("**Conteúdo?**")
                c_h3.write("**Metodologia?**")
                
                for disc in disciplinas_flex:
                    if disc not in data['flex_matrix']: data['flex_matrix'][disc] = {'conteudo': False, 'metodologia': False}
                    
                    c1, c2, c3 = st.columns([2, 1, 1])
                    c1.write(disc)
                    data['flex_matrix'][disc]['conteudo'] = c2.checkbox("Sim", key=f"flex_c_{aluno_id}_{disc}", value=data['flex_matrix'][disc]['conteudo'], disabled=is_monitor)
                    data['flex_matrix'][disc]['metodologia'] = c3.checkbox("Sim", key=f"flex_m_{aluno_id}_{disc}", value=data['flex_matrix'][disc]['metodologia'], disabled=is_monitor)

                st.divider()
                st.subheader("7.2 Plano de Ensino Anual")
                trimestres = ["1º Trimestre", "2º Trimestre", "3º Trimestre"]
                if 'plano_ensino_tri' not in data: data['plano_ensino_tri'] = {}

                for tri in trimestres:
                    st.markdown(f"### 🗓️ {tri}")
                    if tri not in data['plano_ensino_tri']: data['plano_ensino_tri'][tri] = {}
                    
                    for disc in disciplinas_flex:
                        with st.expander(f"{tri} - {disc}", expanded=False):
                            if disc not in data['plano_ensino_tri'][tri]:
                                data['plano_ensino_tri'][tri][disc] = {'obj': '', 'cont': '', 'met': ''}
                            
                            p_ref = data['plano_ensino_tri'][tri][disc]
                            p_ref['obj'] = st.text_area(f"Objetivos ({disc})", value=p_ref.get('obj', ''), key=f"obj_{aluno_id}_{tri}_{disc}", disabled=is_monitor)
                            p_ref['cont'] = st.text_area(f"Conteúdos ({disc})", value=p_ref.get('cont', ''), key=f"cont_{aluno_id}_{tri}_{disc}", disabled=is_monitor)
                            p_ref['met'] = st.text_area(f"Metodologia ({disc})", value=p_ref.get('met', ''), key=f"met_{aluno_id}_{tri}_{disc}", disabled=is_monitor)

                    # --- CORREÇÃO DA OBS/RECOMENDAÇÕES ---
                    obs_valor_banco = data['plano_ensino_tri'][tri].get('obs', '')
                    
                    obs_input = st.text_area(
                        f"Obs/Recomendações {tri}:", 
                        value=obs_valor_banco, 
                        key=f"obs_tri_{aluno_id}_{tri}", 
                        disabled=is_monitor
                    )
                    data['plano_ensino_tri'][tri]['obs'] = obs_input
                    st.markdown("---")

                st.markdown("Considerações finais:")
                data['plano_obs_geral'] = st.text_area("", value=data.get('plano_obs_geral', ''), key=f"obs_geral_{aluno_id}", disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Metas e Plano"):
                        save_student("PEI", data.get('nome'), data, "Metas e Plano")
        
        # --- ABA 7: ASSINATURAS (NOVO) ---
        with tabs[6]:
            st.subheader("Assinaturas Digitais")
            st.caption(f"Código Único do Documento: {data.get('doc_uuid', 'Não gerado ainda')}")
            
            # Identify required signers based on content
            required_roles = []
            if data.get('prof_poli'): required_roles.append({'role': 'Prof. Polivalente', 'name': data['prof_poli']})
            if data.get('prof_aee'): required_roles.append({'role': 'Prof. AEE', 'name': data['prof_aee']})
            if data.get('prof_arte'): required_roles.append({'role': 'Prof. Arte', 'name': data['prof_arte']})
            if data.get('prof_ef'): required_roles.append({'role': 'Prof. Ed. Física', 'name': data['prof_ef']})
            if data.get('prof_tec'): required_roles.append({'role': 'Prof. Tecnologia', 'name': data['prof_tec']})
            if data.get('gestor'): required_roles.append({'role': 'Gestor Escolar', 'name': data['gestor']})
            if data.get('coord'): required_roles.append({'role': 'Coordenação', 'name': data['coord']})
            
            # Show list of signatories
            if required_roles:
                st.markdown("##### Profissionais Citados no Documento")
                for r in required_roles:
                    st.write(f"- **{r['role']}:** {r['name']}")
            else:
                st.info("Nenhum profissional identificado automaticamente nos campos.")

            st.divider()
            
            # Current Signatures
            current_signatures = data.get('signatures', [])
            if current_signatures:
                st.success("✅ Documento assinado por:")
                for sig in current_signatures:
                    st.write(f"✍️ **{sig['name']}** ({sig.get('role', 'Profissional')}) em {sig['date']}")
            else:
                st.warning("Nenhuma assinatura registrada.")

            st.divider()
            
            # Signing Action
            user_name = st.session_state.get('usuario_nome', '')
            user_role_sys = "Monitor" if is_monitor else "Docente/Gestor"
            
            # Check if user matches any role
            match_role = "Profissional"
            is_cited = False
            for r in required_roles:
                if user_name.strip().lower() in r['name'].strip().lower():
                    is_cited = True
                    match_role = r['role']
                    break
            
            st.markdown(f"**Assinar como:** {user_name} ({match_role})")
            
            # Check if already signed
            already_signed = any(s['name'] == user_name for s in current_signatures)
            
            if already_signed:
                st.info("Você já assinou este documento.")
            else:
                if st.button("🖊️ Assinar Digitalmente"):
                    new_sig = {
                        "name": user_name,
                        "role": match_role,
                        "date": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                        "hash": str(uuid.uuid4())
                    }
                    if 'signatures' not in data: data['signatures'] = []
                    data['signatures'].append(new_sig)
                    
                    # Salva apenas a assinatura
                    save_student("PEI", data.get('nome'), data, "Assinatura")
                    st.rerun()

        # --- ABA 8: EMISSÃO ---
        with tabs[7]:
            if not is_monitor:
                st.info("Antes de gerar o PDF, certifique-se de ter clicado em 'Salvar' nas abas anteriores.")
                if st.button("💾 SALVAR PEI COMPLETO", type="primary"): save_student("PEI", data['nome'], data, "Completo")
            else:
                st.info("Modo Visualização.")

            if st.button("👁️ GERAR PDF COMPLETO"):
                # Registrar ação de gerar PDF
                log_action(data.get('nome'), "Gerou PDF", "PEI Completo")
                
                pdf = OfficialPDF('L', 'mm', 'A4'); pdf.add_page(); pdf.set_margins(10, 10, 10)
                
                # SET SIGNATURE FOOTER
                pdf.set_signature_footer(data.get('signatures', []), data.get('doc_uuid', ''))
                
                # --- PÁGINA 1 ---
                if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 10, 8, 25)
                if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 252, 4, 37) 
                pdf.set_xy(0, 12); pdf.set_font("Arial", "", 14)
                pdf.cell(305, 6, clean_pdf_text("      PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
                pdf.ln(6); pdf.set_font("Arial", "B", 12)
                pdf.cell(297, 6, clean_pdf_text("CEIEF RAFAEL AFFONSO LEITE"), 0, 1, 'C')
                pdf.ln(8); pdf.set_font("Arial", "B", 14)
                pdf.cell(297, 8, clean_pdf_text("PLANO EDUCACIONAL ESPECIALIZADO - PEI"), 0, 1, 'C')
                
                # --- FOTO ---
                # Retângulo da foto: x=256, y=53, w=30, h=40
                if data.get('foto_url'):
                    try:
                        # Baixa a foto do link do bucket
                        req = urllib.request.Request(data['foto_url'], headers={'User-Agent': 'Mozilla/5.0'})
                        with urllib.request.urlopen(req) as response:
                            img_data = response.read()
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                            tmp_file.write(img_data)
                            tmp_path = tmp_file.name
                            
                        pdf.image(tmp_path, 256, 53, 30, 40)
                        os.unlink(tmp_path)
                        pdf.rect(256, 53, 30, 40) # Borda
                    except:
                        pdf.rect(256, 53, 30, 40)
                        pdf.set_xy(255.5, 70); pdf.set_font("Arial", "", 8); pdf.cell(30, 5, "Erro URL", 0, 0, 'C')
                        
                elif data.get('foto_base64'):
                    try:
                        img_data = base64.b64decode(data.get('foto_base64'))
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                            tmp_file.write(img_data)
                            tmp_path = tmp_file.name
                        pdf.image(tmp_path, 256, 53, 30, 40)
                        os.unlink(tmp_path)
                        pdf.rect(256, 53, 30, 40) # Borda
                    except:
                        pdf.rect(256, 53, 30, 40)
                        pdf.set_xy(255.5, 70); pdf.set_font("Arial", "", 8); pdf.cell(30, 5, "Erro", 0, 0, 'C')
                else:
                    pdf.rect(256, 53, 30, 40)
                    pdf.set_xy(255.5, 70); pdf.set_font("Arial", "", 9); pdf.cell(30, 5, "FOTO", 0, 0, 'C')
                
                pdf.set_xy(10, 48); table_w = 240; h = 9 
                pdf.section_title("1. IDENTIFICAÇÃO DO ESTUDANTE", width=table_w) 
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Estudante:", 1); pdf.set_font("Arial", "", 12); pdf.cell(table_w-40, h, clean_pdf_text(data.get('nome', '')), 1, 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Nascimento:", 1); pdf.set_font("Arial", "", 12); pdf.cell(40, h, clean_pdf_text(str(data.get('nasc', ''))), 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(20, h, "Idade:", 1); pdf.set_font("Arial", "", 12); pdf.cell(20, h, clean_pdf_text(data.get('idade', '')), 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(30, h, "Ano:", 1); pdf.set_font("Arial", "", 12); pdf.cell(table_w - 150, h, clean_pdf_text(data.get('ano_esc', '')), 1, 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Mãe:", 1); pdf.set_font("Arial", "", 12); pdf.cell(table_w - 40, h, clean_pdf_text(data.get('mae', '')), 1, 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Pai:", 1); pdf.set_font("Arial", "", 12); pdf.cell(table_w - 40, h, clean_pdf_text(data.get('pai', '')), 1, 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Telefone:", 1); pdf.set_font("Arial", "", 12); pdf.cell(table_w - 40, h, clean_pdf_text(data.get('tel', '')), 1, 1)
                
                pdf.ln(5); full_w = 277 
                pdf.set_font("Arial", "B", 12); pdf.cell(full_w, h, "Docentes Responsáveis", 1, 1, 'L', 1)
                docs = [("Polivalente:", data.get('prof_poli')), ("Arte:", data.get('prof_arte')), ("Ed. Física:", data.get('prof_ef')), ("Tecnologia:", data.get('prof_tec')), ("AEE:", data.get('prof_aee')), ("Gestor:", data.get('gestor')), ("Coordenação:", data.get('coord')), ("Revisões:", data.get('revisoes'))]
                for l, v in docs:
                    pdf.set_font("Arial", "B", 12); pdf.cell(60, h, clean_pdf_text(l), 1); pdf.set_font("Arial", "", 12); pdf.cell(full_w-60, h, clean_pdf_text(v), 1, 1)

                # --- PÁGINA 2 ---
                pdf.add_page(); pdf.section_title("2. INFORMAÇÕES DE SAÚDE", width=0); h = 10
                pdf.set_font("Arial", "B", 12); pdf.cell(100, h, clean_pdf_text("O estudante tem diagnóstico conclusivo:"), 1, 0, 'L')
                status_sim = "[ X ]" if data.get('diag_status') == "Sim" else "[   ]"
                status_nao = "[ X ]" if data.get('diag_status') == "Não" else "[   ]"
                pdf.set_font("Arial", "", 12); pdf.cell(0, h, f"  {status_sim} Sim      {status_nao} Não", 1, 1, 'L')
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Data do Laudo:", 1, 0, 'L')
                pdf.set_font("Arial", "", 12); pdf.cell(60, h, clean_pdf_text(str(data.get('laudo_data', ''))), 1, 0, 'L')
                pdf.set_font("Arial", "B", 12); pdf.cell(40, h, "Médico Respons.:", 1, 0, 'L')
                pdf.set_font("Arial", "", 12); pdf.cell(0, h, clean_pdf_text(data.get('laudo_medico', '---')), 1, 1, 'L')

                pdf.ln(2); diag_list = data.get('diag_tipo', []); diag_ativos = []
                if "Deficiência" in diag_list and data.get('defic_txt'): diag_ativos.append(("Deficiência:", data.get('defic_txt')))
                if "Transtorno do Neurodesenvolvimento" in diag_list and data.get('neuro_txt'): diag_ativos.append(("Transtorno Neuro:", data.get('neuro_txt')))
                if "Transtornos Aprendizagem" in diag_list and data.get('aprend_txt'): diag_ativos.append(("Transt. Aprendizagem:", data.get('aprend_txt')))
                if "AH/SD" in diag_list: diag_ativos.append(("Destaque:", "Altas Habilidades / Superdotação"))
                if "Outros" in diag_list: diag_ativos.append(("Outros Diagnósticos:", "Conforme prontuário"))

                if diag_ativos:
                    for l_diag, t_diag in diag_ativos:
                        pdf.set_font("Arial", "B", 11); pdf.cell(60, h, clean_pdf_text(l_diag), "LTB", 0, 'L')
                        pdf.set_font("Arial", "", 11); pdf.cell(0, h, clean_pdf_text(t_diag), "RTB", 1, 'L')
                else: pdf.set_font("Arial", "I", 11); pdf.cell(0, h, "Nenhum diagnóstico selecionado.", 1, 1, 'C')

                pdf.ln(6); pdf.set_font("Arial", "B", 12); pdf.set_fill_color(245, 245, 245); pdf.cell(277, 10, "Terapias que realiza", 1, 1, 'C', 1)
                pdf.set_font("Arial", "B", 11); pdf.cell(80, 10, "Especialidades", 1, 0, 'L', 1); pdf.cell(0, 10, clean_pdf_text("Frequência e Horário de Atendimento"), 1, 1, 'L', 1)
                for esp in ["Psicologia", "Fonoaudiologia", "Terapia Ocupacional", "Psicopedagogia", "Fisioterapia", "Outros"]:
                    info = data.get('terapias', {}).get(esp, {'realiza': False, 'dias': [], 'horario': ''})
                    chk = "[ X ]" if info['realiza'] else "[   ]"
                    label_esp = f"  {chk} {esp}"
                    if esp == "Outros" and info.get('nome_custom'): label_esp = f"  {chk} Outros ({info['nome_custom']})"
                    pdf.set_font("Arial", "B", 11); pdf.cell(80, 12, clean_pdf_text(label_esp), 1, 0, 'L')
                    x_start = pdf.get_x(); y_start = pdf.get_y(); pdf.set_font("Arial", "", 10)
                    if info['realiza']:
                        pdf.set_xy(x_start + 5, y_start + 2); pdf.cell(0, 4, clean_pdf_text("Dias: " + ", ".join(info['dias'])), 0, 1)
                        pdf.set_x(x_start + 5); pdf.set_font("Arial", "B", 10); pdf.cell(16, 4, "Horário:", 0); pdf.set_font("Arial", "", 10); pdf.cell(0, 4, clean_pdf_text(info['horario']), 0, 1)
                    else:
                        pdf.set_xy(x_start + 5, y_start + 4); pdf.set_font("Arial", "I", 10); pdf.set_text_color(150, 0, 0)
                        pdf.cell(0, 4, "NÃO REALIZA ATENDIMENTO NESTA ESPECIALIDADE.", 0, 1); pdf.set_text_color(0, 0, 0)
                    pdf.set_xy(x_start, y_start); pdf.cell(0, 12, "", 1, 1)

                pdf.ln(5); pdf.set_font("Arial", "B", 12); pdf.cell(0, 10, "Medicação e Horários:", "LTR", 1, 'L', 1)
                pdf.set_font("Arial", "", 12); pdf.multi_cell(0, 8, clean_pdf_text(f"{data.get('med_nome', 'Não utiliza')}\nHorários: {data.get('med_hor', 'N/A')}"), "LRB")
                pdf.ln(5); pdf.set_font("Arial", "B", 12); pdf.cell(50, 8, clean_pdf_text("Médico Responsável:"), 1, 0); pdf.set_font("Arial", "", 12); pdf.cell(0, 8, clean_pdf_text(data.get('med_doc', 'N/A')), 1, 1)
                pdf.set_font("Arial", "B", 12); pdf.cell(0, 8, "Objetivo da medicação:", "LTR", 1, 'L', 1); pdf.set_font("Arial", "", 12); pdf.multi_cell(0, 8, clean_pdf_text(data.get('med_obj', 'Não informado.')), "LRB")
                pdf.ln(3); pdf.set_font("Arial", "B", 12); pdf.cell(0, 8, clean_pdf_text("Outras informações de saúde consideradas relevantes:"), "LTR", 1, 'L', 1)
                pdf.set_font("Arial", "", 12); pdf.multi_cell(0, 8, clean_pdf_text(data.get('saude_extra', 'Nenhuma informação adicional.')), "LRB")

                # --- 3. PROTOCOLO DE CONDUTA ---
                pdf.ln(5); pdf.section_title("3. PROTOCOLO DE CONDUTA", width=0); h = 8
                pdf.set_font("Arial", "B", 11); pdf.set_fill_color(245, 245, 245); pdf.cell(0, 8, "COMUNICAÇÃO, LOCOMOÇÃO E HIGIENE", 1, 1, 'C', 1)
                rows_cond = [
                    ("Como o estudante se comunica?", f"{data.get('com_tipo')} {data.get('com_alt_espec')}"),
                    ("Capaz de expressar necessidades, desejos e interesses?", f"{data.get('com_necessidades')} - {data.get('com_necessidades_espec')}"),
                    ("Atende quando é chamado?", f"{data.get('com_chamado')} - {data.get('com_chamado_espec')}"),
                    ("Responde a comandos simples?", f"{data.get('com_comandos')} - {data.get('com_comandos_espec')}"),
                    ("Possui mobilidade reduzida?", f"{data.get('loc_reduzida')} - {data.get('loc_reduzida_espec')}"),
                    ("Locomove-se pela casa e ambientes?", f"{data.get('loc_ambiente')} ({data.get('loc_ambiente_ajuda')}) - {data.get('loc_ambiente_espec')}"),
                    ("Utiliza o banheiro?", f"{data.get('hig_banheiro')} ({data.get('hig_banheiro_ajuda')}) - {data.get('hig_banheiro_espec')}"),
                    ("Escova os dentes?", f"{data.get('hig_dentes')} ({data.get('hig_dentes_ajuda')}) - {data.get('hig_dentes_espec')}")
                ]
                for l, v in rows_cond:
                    pdf.set_font("Arial", "B", 10); pdf.cell(95, h, clean_pdf_text(l), 1, 0, 'L'); pdf.set_font("Arial", "", 10); pdf.cell(0, h, clean_pdf_text(v), 1, 1, 'L')
                
                pdf.ln(4); pdf.set_font("Arial", "B", 11); pdf.set_fill_color(245, 245, 245); pdf.cell(0, 8, "COMPORTAMENTO E INTERESSES", 1, 1, 'C', 1)

                verbatims = [
                    ("Quais são os interesses do estudante?", data.get('beh_interesses')),
                    ("Quais objetos que gosta? Tem um objeto de apego?", data.get('beh_objetos_gosta')),
                    ("Quais objetos o estudante não gosta e/ou causam aversão?", data.get('beh_objetos_odeia')),
                    ("Gosta de toque, abraço, beijo?", data.get('beh_toque')),
                    ("O que o deixa calmo e relaxado?", data.get('beh_calmo')),
                    ("Quais atividades são mais prazerosas?", data.get('beh_atividades')),
                    ("Quais são os gatilhos já identificados para episódios de crise?", data.get('beh_gatilhos')),
                    ("Quando o estudante está em crise como normalmente se regula?", data.get('beh_crise_regula')),
                    ("O estudante costuma apresentar comportamentos desafiadores? Manejo?", data.get('beh_desafios')),
                    ("Tem restrições alimentares / Seletividade?", f"{data.get('beh_restricoes')} - {data.get('beh_restricoes_espec')}"),
                    ("Tem autonomia para tomar água e se alimentar?", f"{data.get('beh_autonomia_agua')} - {data.get('beh_autonomia_agua_espec')}"),
                    ("Outras informações julgadas pertinentes:", data.get('beh_pertinentes'))
                ]
                
                for l, v in verbatims:
                    if pdf.get_y() > 250: 
                        pdf.add_page()
                        
                    pdf.set_x(10)
                    pdf.set_font("Arial", "B", 10)
                    pdf.multi_cell(0, 7, clean_pdf_text(l), border="LTR", align='L', fill=True) 
                    
                    pdf.set_x(10)
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 6, clean_pdf_text(v if v else "---"), border="LBR", align='L', fill=False)

                # --- 4. DESENVOLVIMENTO ESCOLAR ---
                pdf.ln(5); pdf.section_title("4. DESENVOLVIMENTO ESCOLAR", width=0); h = 8
                dev_rows = [
                    ("Permanece em sala e aula?", f"{data.get('dev_permanece')} - {data.get('dev_permanece_espec')}"),
                    ("Está integrado ao ambiente escolar?", f"{data.get('dev_integrado')} - {data.get('dev_integrado_espec')}"),
                    ("Locomove-se pela escola?", f"{data.get('dev_loc_escola')} - {data.get('dev_loc_escola_espec')}"),
                    ("Realiza tarefas escolares?", f"{data.get('dev_tarefas')} - {data.get('dev_tarefas_espec')}"),
                    ("Tem amigos?", f"{data.get('dev_amigos')} - {data.get('dev_amigos_espec')}"),
                    ("Tem um colega predileto?", f"{data.get('dev_colega_pref')}"),
                    ("Participa das atividades e interage em diferentes espaços?", f"{data.get('dev_participa')} - {data.get('dev_participa_espec')}")
                ]
                for l, v in dev_rows:
                    pdf.set_font("Arial", "B", 10); pdf.cell(100, h, clean_pdf_text(l), 1, 0, 'L'); pdf.set_font("Arial", "", 10); pdf.cell(0, h, clean_pdf_text(v), 1, 1, 'L')
                
                pdf.ln(2); pdf.set_font("Arial", "B", 10); pdf.cell(0, 7, clean_pdf_text("Envolvimento afetivo e social da turma com o estudante:"), "LTR", 1, 'L', 1)
                pdf.set_font("Arial", "", 10); pdf.multi_cell(0, 6, clean_pdf_text(data.get('dev_afetivo', '---')), "LRB")

                # --- 5. AVALIAÇÃO ACADÊMICA ---
                pdf.ln(5)
                if pdf.get_y() > 220: pdf.add_page()

                pdf.section_title("5. AVALIAÇÃO ACADÊMICA DO ESTUDANTE", width=0)
                pdf.ln(2)
                
                areas_aval = []
                
                if pei_level == "Fundamental":
                    areas_aval = [
                        ("LÍNGUA PORTUGUESA", data.get('aval_port')),
                        ("MATEMÁTICA", data.get('aval_mat')),
                        ("CONHECIMENTOS GERAIS", data.get('aval_con_gerais')),
                        ("ARTE - Artes Visuais", data.get('aval_arte_visuais')),
                        ("ARTE - Música", data.get('aval_arte_musica')),
                        ("ARTE - Teatro", data.get('aval_arte_teatro')),
                        ("ARTE - Dança", data.get('aval_arte_danca')),
                        ("EDUCAÇÃO FÍSICA - Habilidades Motoras", data.get('aval_ef_motoras')),
                        ("EDUCAÇÃO FÍSICA - Conhecimento Corporal", data.get('aval_ef_corp_conhec')),
                        ("EDUCAÇÃO FÍSICA - Exp. Corporais e Expressividade", data.get('aval_ef_exp')),
                        ("LINGUAGENS E TECNOLOGIA", data.get('aval_ling_tec'))
                    ]
                else: # Infantil
                    areas_aval = [
                        ("LINGUAGEM VERBAL", data.get('aval_ling_verbal')),
                        ("LINGUAGEM MATEMÁTICA", data.get('aval_ling_mat')),
                        ("INDÍVIDUO E SOCIEDADE", data.get('aval_ind_soc')),
                        ("ARTE - Artes Visuais", data.get('aval_arte_visuais')),
                        ("ARTE - Música", data.get('aval_arte_musica')),
                        ("ARTE - Teatro", data.get('aval_arte_teatro')),
                        ("EDUCAÇÃO FÍSICA - Jogos e Brincadeiras", data.get('aval_ef_jogos')),
                        ("EDUCAÇÃO FÍSICA - Ritmo e Expressividade", data.get('aval_ef_ritmo')),
                        ("EDUCAÇÃO FÍSICA - Conhecimento Corporal", data.get('aval_ef_corp')),
                        ("LINGUAGEM E TECNOLOGIAS", data.get('aval_ling_tec'))
                    ]
                
                for titulo, texto in areas_aval:
                    if pdf.get_y() > 230: pdf.add_page()
                    
                    pdf.set_font("Arial", "B", 10); pdf.set_fill_color(240, 240, 240)
                    pdf.cell(0, 7, clean_pdf_text(titulo), "LTR", 1, 'L', 1)
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 6, clean_pdf_text(texto if texto else "---"), "LRB")
                    pdf.ln(2)

                # --- 6. METAS ---
                pdf.ln(5)
                if pdf.get_y() > 220: pdf.add_page()
                
                pdf.section_title("6. METAS ESPECÍFICAS PARA O ANO EM CURSO", width=0)
                pdf.ln(2)
                
                def print_meta_row(titulo, meta, estrategia):
                    if pdf.get_y() > 220: pdf.add_page()
                    pdf.set_font("Arial", "B", 11); pdf.set_fill_color(230, 230, 230)
                    pdf.cell(0, 8, clean_pdf_text(titulo), 1, 1, 'L', 1)
                    pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "Metas / Habilidades:", "LTR", 1)
                    pdf.set_font("Arial", "", 10); pdf.multi_cell(0, 5, clean_pdf_text(meta if meta else "---"), "LRB")
                    pdf.set_x(10); pdf.set_font("Arial", "B", 10); pdf.cell(0, 5, clean_pdf_text("Estratégias:"), "LTR", 1)
                    pdf.set_x(10); pdf.set_font("Arial", "", 10); pdf.multi_cell(0, 5, clean_pdf_text(estrategia if estrategia else "---"), "LRB")
                    pdf.ln(2)

                print_meta_row("Habilidades Sociais", data.get('meta_social_obj'), data.get('meta_social_est'))
                print_meta_row("Habilidades de Autocuidado e Vida Prática", data.get('meta_auto_obj'), data.get('meta_auto_est'))
                print_meta_row("Habilidades Acadêmicas", data.get('meta_acad_obj'), data.get('meta_acad_est'))

                # --- 7. FLEXIBILIZAÇÃO ---
                pdf.ln(5)
                if pdf.get_y() > 230: pdf.add_page()
                
                pdf.section_title("7. FLEXIBILIZAÇÃO CURRICULAR", width=0)
                pdf.ln(4)
                
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 6, clean_pdf_text("7.1 DISCIPLINAS QUE NECESSITAM DE ADAPTAÇÃO"), 0, 1)
                pdf.ln(2)

                pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", "B", 9)
                pdf.cell(80, 8, "DISCIPLINA", 1, 0, 'C', 1)
                pdf.cell(80, 8, clean_pdf_text("CONTEÚDO"), 1, 0, 'C', 1)
                pdf.cell(0, 8, "METODOLOGIA", 1, 1, 'C', 1)

                if pei_level == "Fundamental":
                    disciplinas_flex = ["Língua Portuguesa", "Matemática", "História", "Geografia", "Ciências", "Arte", "Educação Física", "Linguagens e Tecnologia"]
                else:
                    disciplinas_flex = ["Linguagem Verbal", "Linguagem Matemática", "Indivíduo e Sociedade", "Arte", "Educação Física", "Linguagens e Tecnologia"]
                
                pdf.set_font("Arial", "", 10)
                for disc in disciplinas_flex:
                    vals = data.get('flex_matrix', {}).get(disc, {'conteudo': False, 'metodologia': False})
                    chk_c_sim = "[X] Sim  [  ] Não" if vals['conteudo'] else "[  ] Sim  [X] Não"
                    chk_m_sim = "[X] Sim  [  ] Não" if vals['metodologia'] else "[  ] Sim  [X] Não"
                    pdf.cell(80, 8, clean_pdf_text(f" {disc}"), 1, 0, 'L')
                    pdf.cell(80, 8, chk_c_sim, 1, 0, 'C')
                    pdf.cell(0, 8, chk_m_sim, 1, 1, 'C')
                    pdf.set_x(10) # Força o cursor a voltar para a margem esquerda (10mm)
                    pdf.ln(0)     # Dá um pequeno espaçamento vertical

                # --- 7.2 PLANO DE ENSINO (TRIMESTRES) ---
                trimestres = ["1º Trimestre", "2º Trimestre", "3º Trimestre"]
                
                for tri in trimestres:
                    dados_tri = data.get('plano_ensino_tri', {}).get(tri, {})
                    has_content = False
                    if dados_tri.get('obs', '').strip(): has_content = True
                    for disc in disciplinas_flex:
                        d_dados = dados_tri.get(disc, {'obj': '', 'cont': '', 'met': ''})
                        if d_dados['obj'].strip() or d_dados['cont'].strip() or d_dados['met'].strip():
                            has_content = True; break
                    
                    if has_content:
                        pdf.ln(8)
                        if pdf.get_y() > 180: pdf.add_page()
                        pdf.set_font("Arial", "B", 12)
                        pdf.cell(0, 8, clean_pdf_text(f"7.2 PLANO DE ENSINO - {tri.upper()}"), 0, 1, 'L')
                        pdf.ln(2)

                        for disc in disciplinas_flex:
                            plan = dados_tri.get(disc, {'obj': '', 'cont': '', 'met': ''})
                            
                            if plan['obj'].strip() or plan['cont'].strip() or plan['met'].strip():
                                if pdf.get_y() > 180: pdf.add_page()
                            
                            pdf.set_x(10)
                            largura_util = 277
                            pdf.set_font("Arial", "B", 10); pdf.set_fill_color(230, 230, 230)  # Trocado 0 por largura_util para o título não vazar
                            pdf.cell(largura_util, 7, clean_pdf_text(disc), 1, 1, 'L', 1)
        
                            pdf.set_font("Arial", "B", 10); pdf.set_fill_color(230, 230, 230)
                            pdf.set_x(10)
                            pdf.cell(largura_util, 6, "Objetivos:", "LTR", 1, 'L', 1)
                            
                            pdf.set_font("Arial", "", 9)
                            pdf.set_x(10)
                            pdf.multi_cell(largura_util, 5, clean_pdf_text(plan['obj'] if plan['obj'] else "---"), "LRB")
                            
                            pdf.set_font("Arial", "B", 9)
                            pdf.set_x(10)
                            pdf.cell(largura_util, 6, clean_pdf_text("Conteúdos Específicos:"), "LTR", 1, 'L', 1)
                            
                            pdf.set_font("Arial", "", 9)
                            pdf.set_x(10)
                            pdf.multi_cell(largura_util, 5, clean_pdf_text(plan['cont'] if plan['cont'] else "---"), "LRB")
                            
                            pdf.set_font("Arial", "B", 9)
                            pdf.set_x(10)
                            pdf.cell(largura_util, 6, "Metodologia:", "LTR", 1, 'L', 1)
                            
                            pdf.set_font("Arial", "", 9)
                            pdf.set_x(10)
                            pdf.multi_cell(largura_util, 5, clean_pdf_text(plan['met'] if plan['met'] else "---"), "LRB")
                            
                            pdf.ln(2)

                        if dados_tri.get('obs'):
                            if pdf.get_y() > 240: pdf.add_page()
                            pdf.ln(2)
                            pdf.set_font("Arial", "B", 10)
                            pdf.cell(0, 6, clean_pdf_text(f"Observações do {tri}:"), "LTR", 1, 'L')
                            pdf.set_font("Arial", "", 10)
                            pdf.multi_cell(0, 6, clean_pdf_text(dados_tri.get('obs')), "LRB")

                # --- OBSERVAÇÕES FINAIS ---
                if data.get('plano_obs_geral'):
                    pdf.ln(5)
                    if pdf.get_y() > 230: pdf.add_page()
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(0, 6, clean_pdf_text("Considerações e/ou recomendações finais:"), "LTR", 1, 'L')
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 6, clean_pdf_text(data.get('plano_obs_geral')), "LRB")

                # --- ASSINATURAS ---
                pdf.ln(15)
                if pdf.get_y() > 230: pdf.add_page(); pdf.ln(15)
                pdf.set_font("Arial", "", 8)
                
                # Exibe assinaturas tradicionais (linhas)
                def draw_signature(x_pos, y_pos, nome, cargo):
                    pdf.line(x_pos, y_pos, x_pos + 70, y_pos)
                    pdf.set_xy(x_pos, y_pos + 2)
                    pdf.cell(70, 4, clean_pdf_text(nome if nome else "____________________"), 0, 2, 'C')
                    pdf.set_font("Arial", "B", 7)
                    pdf.cell(70, 3, clean_pdf_text(cargo), 0, 0, 'C')
                    pdf.set_font("Arial", "", 8)

                y = pdf.get_y()
                draw_signature(15, y, data.get('prof_poli', ''), "Prof. Polivalente / Regente")
                draw_signature(113, y, data.get('prof_arte', ''), "Prof. Arte")
                draw_signature(211, y, data.get('prof_ef', ''), "Prof. Ed. Física")
                
                pdf.ln(18)
                y = pdf.get_y()
                draw_signature(65, y, data.get('prof_aee', ''), "Prof. Ed. Especial (AEE)")
                draw_signature(162, y, data.get('prof_tec', ''), "Prof. Linguagens e Tec.")
                
                pdf.ln(18)
                y = pdf.get_y()
                draw_signature(65, y, data.get('coord', ''), "Coordenador Pedagógico")
                draw_signature(162, y, data.get('gestor', ''), "Gestor Escolar")

                st.session_state.pdf_bytes = get_pdf_bytes(pdf)
                st.rerun()

            if 'pdf_bytes' in st.session_state:
                st.download_button("📥 BAIXAR PEI COMPLETO", st.session_state.pdf_bytes, f"PEI_{data.get('nome','aluno')}.pdf", "application/pdf", type="primary")

        # --- ABA 9: HISTÓRICO ---
        with tabs[8]:
            st.subheader("Histórico de Atividades")
            st.caption("Registro de alterações, salvamentos e geração de documentos.")
            
            df_hist = safe_read("Historico", ["Data_Hora", "Aluno", "Usuario", "Acao", "Detalhes"])
            
            if not df_hist.empty and data.get('nome'):
                # Filtrar pelo aluno atual
                student_hist = df_hist[df_hist["Aluno"] == data.get('nome')]
                
                if not student_hist.empty:
                    # Ordenar por data (mais recente primeiro)
                    student_hist = student_hist.iloc[::-1]
                    st.dataframe(student_hist, use_container_width=True, hide_index=True)
                else:
                    st.info("Nenhum histórico encontrado para este aluno.")
            else:
                st.info("O histórico está vazio ou aluno não selecionado.")

   



# --- PDI - PLANO DE DESENVOLVIMENTO INDIVIDUAL (ATUALIZADO) ---
    if doc_mode == "PDI - Pré Escola e Ens. Fundamental":
        st.markdown(f"""<div class="header-box"><div class="header-title">PDI - Plano de Desenvolvimento Individual</div></div>""", unsafe_allow_html=True)
        st.markdown("""<style>div[data-testid="stFormSubmitButton"] > button {width: 100%; background-color: #dcfce7; color: #166534; border: 1px solid #166534;}</style>""", unsafe_allow_html=True)

        data_pdi = st.session_state.data_pdi
        data_case = st.session_state.get('data_case', {})

        # Tabs de Navegação (Reestruturadas para o novo PAEE)
        tabs = st.tabs([
            "Item 3: Avaliação Pedagógica",
            "Item 4: PAEE",
            "Assinaturas",
            "PDF Final",
            "Histórico"
        ])
        
        st.info("ℹ️ Os dados de **Identificação**, **Família**, **Histórico** e **Avaliação Geral** são importados automaticamente do módulo **Estudo de Caso** (Item 1).")

        # --- ABA 1: AVALIAÇÃO PEDAGÓGICA ---
        with tabs[0]:
            st.header("3. Avaliação Pedagógica")
            
            with st.form("pdi_avaliacao_form"):
                
                st.subheader("3.1 Potencialidades e 3.2 Áreas de Interesse")
                data_pdi['potencialidades'] = st.text_area("3.1 Potencialidades do Estudante", value=data_pdi.get('potencialidades', ''), disabled=is_monitor)
                data_pdi['areas_interesse'] = st.text_area("3.2 Áreas de Interesse", value=data_pdi.get('areas_interesse', ''), disabled=is_monitor)
                st.divider()

                def render_percurso_final_row(label, key_base, options=["Sim", "Não", "Parcialmente"]):
                    st.markdown(f"**{label}**")
                    c1, c2 = st.columns(2)
                    
                    # Recupera o valor. Se for None ou string vazia, não seleciona nada.
                    v_proc = data_pdi.get(f"{key_base}_proc")
                    v_final = data_pdi.get(f"{key_base}_final")
                    
                    # O segredo é passar None para index se o valor não existir
                    idx_proc = options.index(v_proc) if v_proc in options else None
                    idx_final = options.index(v_final) if v_final in options else None
                    
                    data_pdi[f"{key_base}_proc"] = c1.radio(
                        "Avaliação de Percurso", options, horizontal=True, 
                        key=f"p_{key_base}", index=idx_proc, disabled=is_monitor
                    )
                    data_pdi[f"{key_base}_final"] = c2.radio(
                        "Avaliação Final", options, horizontal=True, 
                        key=f"f_{key_base}", index=idx_final, disabled=is_monitor
                    )
                    st.divider()

                st.subheader("3.3 DESENVOLVIMENTO COGNITIVO")
                
                with st.expander("Atenção", expanded=True):
                    render_percurso_final_row("Concentrada", "atencao_conc")
                    render_percurso_final_row("Sustentada", "atencao_sust")
                    render_percurso_final_row("Seletiva", "atencao_sel")
                    render_percurso_final_row("Alternada", "atencao_alt")

                with st.expander("Percepção", expanded=False):
                    render_percurso_final_row("Memória Visual/Percepção de diferenças e semelhanças", "perc_vis")
                    render_percurso_final_row("Percepção e Discriminação Auditiva", "perc_aud")
                    render_percurso_final_row("Percepção Tátil", "perc_tat")
                    render_percurso_final_row("Orientação temporal", "perc_temp")
                    render_percurso_final_row("Orientação espacial", "perc_esp")

                with st.expander("Memória e Funções Executivas", expanded=False):
                    st.markdown("#### Memória")
                    render_percurso_final_row("Curto Prazo", "mem_curto")
                    render_percurso_final_row("Médio Prazo", "mem_medio")
                    render_percurso_final_row("Longo Prazo", "mem_longo")
                    
                    st.markdown("#### Funções Executivas")
                    render_percurso_final_row("Controle Inibitório (inibição de impulsos)", "fe_inib")
                    render_percurso_final_row("Memória de Trabalho (capacidade de reter e manipular)", "fe_trab")
                    render_percurso_final_row("Flexibilidade Cognitiva (adaptação a mudanças)", "fe_flex")
                    render_percurso_final_row("Planejamento/Organização", "fe_plan")

                st.subheader("LINGUAGEM E COMUNICAÇÃO")
                
                with st.expander("Linguagem e Comunicação Oral", expanded=False):
                    render_percurso_final_row("Utiliza palavras para se comunicar?", "ling_palavras")
                    
                    st.markdown("**Outras formas de comunicação utilizadas:**")
                    c1, c2 = st.columns(2)
                    with c1:
                        st.markdown("*Avaliação de Percurso*")
                        data_pdi['ling_gestos_p'] = st.checkbox("Comunica-se por gestos (Percurso)", value=data_pdi.get('ling_gestos_p', False))
                        data_pdi['ling_aponta_p'] = st.checkbox("Através de apontamentos (Percurso)", value=data_pdi.get('ling_aponta_p', False))
                        data_pdi['ling_pisca_p'] = st.checkbox("Piscar dos olhos (Percurso)", value=data_pdi.get('ling_pisca_p', False))
                        data_pdi['ling_ca_p'] = st.text_input("Comunicação Alternativa. Qual? (Percurso)", value=data_pdi.get('ling_ca_p', ''))
                        
                        opts_lib = ["Não", "Básico", "Fluente"]
                        v_lib_p = data_pdi.get('ling_libras_p', None)
                        idx_lib_p = opts_lib.index(v_lib_p) if v_lib_p in opts_lib else None
                        data_pdi['ling_libras_p'] = st.radio("LIBRAS (Percurso)", opts_lib, horizontal=True, index=idx_lib_p)
                        
                        data_pdi['ling_outros_p'] = st.text_input("Outros (Percurso)", value=data_pdi.get('ling_outros_p', ''))
                    with c2:
                        st.markdown("*Avaliação Final*")
                        data_pdi['ling_gestos_f'] = st.checkbox("Comunica-se por gestos (Final)", value=data_pdi.get('ling_gestos_f', False))
                        data_pdi['ling_aponta_f'] = st.checkbox("Através de apontamentos (Final)", value=data_pdi.get('ling_aponta_f', False))
                        data_pdi['ling_pisca_f'] = st.checkbox("Piscar dos olhos (Final)", value=data_pdi.get('ling_pisca_f', False))
                        data_pdi['ling_ca_f'] = st.text_input("Comunicação Alternativa. Qual? (Final)", value=data_pdi.get('ling_ca_f', ''))
                        
                        v_lib_f = data_pdi.get('ling_libras_f', None)
                        idx_lib_f = opts_lib.index(v_lib_f) if v_lib_f in opts_lib else None
                        data_pdi['ling_libras_f'] = st.radio("LIBRAS (Final)", opts_lib, horizontal=True, index=idx_lib_f)
                        
                        data_pdi['ling_outros_f'] = st.text_input("Outros (Final)", value=data_pdi.get('ling_outros_f', ''))
                    st.divider()

                    render_percurso_final_row("Apresenta trocas fonéticas orais?", "ling_trocas")
                    render_percurso_final_row("Estabelece diálogo com troca de turno?", "ling_dialogo")
                    render_percurso_final_row("Inventa frases ou histórias?", "ling_inventa")
                    render_percurso_final_row("Descreve cenas com sentido?", "ling_descreve")
                    render_percurso_final_row("Consegue expressar e explicar pensamentos, ideias e desejos?", "ling_expressa")
                    render_percurso_final_row("Reconta história com sentido e/ou faz relatos numa sequência lógica?", "ling_reconta")

                with st.expander("Linguagem Compreensiva", expanded=False):
                    render_percurso_final_row("Compreende e processa informações orais simples?", "ling_comp_simp")
                    render_percurso_final_row("Compreende e processa informações orais complexas?", "ling_comp_comp")

                with st.expander("Linguagem Escrita", expanded=False):
                    opts_escrita = [
                        "Não distingue desenho, letras e números.", "Identifica e nomeia as letras.", 
                        "Escreve seu nome.", "Escreve letras de forma aleatória.", 
                        "Relaciona som/grafia.", "Escreve apenas palavras canônicas.", "Escreve palavras não-canônicas."
                    ]
                    render_percurso_final_row("Escreve convencionalmente?", "escrita_conv", options=opts_escrita)
                    
                    opts_org = ["Não escreve textos convencionais.", "Escreve frases simples.", "Escreve textos."]
                    render_percurso_final_row("Apresenta organização textual?", "escrita_org", options=opts_org)

                with st.expander("Leitura", expanded=False):
                    opts_leitura = [
                        "Não realiza leitura.", "Realiza leitura apenas de palavras canônicas.", 
                        "Realiza leitura de palavras canônicas e não-canônicas.", "Realiza leitura de frases e textos com dificuldade.", 
                        "Realiza leitura de frases e textos com fluência.", "Compreende o que lê com apoio.", "Compreende o que lê com autonomia."
                    ]
                    render_percurso_final_row("Nível de Leitura", "leitura_nivel", options=opts_leitura)
                    
                    st.markdown("**Leitura em BRAILLE e Outros:**")
                    c1, c2 = st.columns(2)
                    opts_braille = ["Não utiliza", "Com autonomia.", "Com apoio.", "Com dificuldade."]
                    with c1:
                        v_br_p = data_pdi.get('braille_p', None)
                        idx_br_p = opts_braille.index(v_br_p) if v_br_p in opts_braille else None
                        data_pdi['braille_p'] = st.radio("BRAILLE (Percurso)", opts_braille, index=idx_br_p)
                        data_pdi['leitura_outros_p'] = st.text_input("Outros (Leitura - Percurso)", value=data_pdi.get('leitura_outros_p', ''))
                    with c2:
                        v_br_f = data_pdi.get('braille_f', None)
                        idx_br_f = opts_braille.index(v_br_f) if v_br_f in opts_braille else None
                        data_pdi['braille_f'] = st.radio("BRAILLE (Final)", opts_braille, index=idx_br_f)
                        data_pdi['leitura_outros_f'] = st.text_input("Outros (Leitura - Final)", value=data_pdi.get('leitura_outros_f', ''))
                
                data_pdi['ling_obs'] = st.text_area("Observações Gerais (Linguagem e Comunicação)", value=data_pdi.get('ling_obs', ''), disabled=is_monitor)

                st.subheader("RACIOCÍNIO E OUTROS DOMÍNIOS")
                with st.expander("Raciocínio e Resolução de Problemas", expanded=False):
                    render_percurso_final_row("a) Planeja, antecipa, argumenta?", "rac_plan")
                    render_percurso_final_row("b) Compara, classifica, categoriza, sequencia, inferi?", "rac_comp")
                    render_percurso_final_row("c) Conhece conceitos básicos do vocabulário matemático?", "rac_mat")
                    render_percurso_final_row("d) Tem capacidade de conclusões lógicas?", "rac_log")

                with st.expander("3.4 Desenvolvimento Motor", expanded=False):
                    render_percurso_final_row("Realiza os diversos tipos de locomoção?", "mot_loc")
                    render_percurso_final_row("Manipula bola e outros objetos?", "mot_bola")
                    render_percurso_final_row("Desenvolvimento esperado da lateralidade?", "mot_lat")
                    render_percurso_final_row("Capacidades físicas (equilíbrio, força, flexibilidade)?", "mot_fis")
                    render_percurso_final_row("Preensão trípode e rotação de punho?", "mot_preensao")

                with st.expander("3.5 Habilidades Pessoais e de Socialização", expanded=False):
                    opts_hab = ["Com autonomia", "Com ajuda", "Não realiza"]
                    st.markdown("#### Pessoais")
                    render_percurso_final_row("Alimentação", "hab_alim", options=opts_hab)
                    render_percurso_final_row("Higiene", "hab_hig", options=opts_hab)
                    render_percurso_final_row("Uso funcional dos objetos", "hab_obj", options=opts_hab)
                    render_percurso_final_row("Locomoção", "hab_loc", options=opts_hab)
                    
                    st.markdown("#### Socialização")
                    render_percurso_final_row("Interage com os adultos?", "soc_adul")
                    render_percurso_final_row("Interage com os colegas?", "soc_col")
                    render_percurso_final_row("Tem tolerância a frustração?", "soc_frust")

                with st.expander("3.6 Função do Brincar", expanded=False):
                    opts_brincar = ["Sim", "Não", "Com modelo"]
                    render_percurso_final_row("Faz uso dos brinquedos de maneira funcional?", "bri_func", options=opts_brincar)
                    render_percurso_final_row("Explora os brinquedos espontaneamente?", "bri_explora", options=opts_brincar)
                    render_percurso_final_row("Utiliza objetos atribuindo diferentes funções?", "bri_dif", options=opts_brincar)
                    render_percurso_final_row("Estrutura uma brincadeira de forma criativa?", "bri_cria", options=opts_brincar)

                if st.form_submit_button("💾 Salvar Avaliação Pedagógica"):
                    save_student("PDI", data_pdi.get('nome'), data_pdi, "Avaliação Pedagógica")

        # --- ABA 2: PAEE (NOVA) ---
        with tabs[1]:
            st.header("4. Plano de Atendimento Educacional Especializado (PAEE)")
            
            # 1. Definindo as listas de opções
            opcoes_freq = ["Não realiza", "1x", "2x", "3x", "4x ou mais"]
            opcoes_atend = ["Individual", "Grupo"]
            
            with st.form("pdi_paee_form"):
                st.subheader("4.1 Organização do AEE")
                
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("**Sala de Recursos**")
                    
                    # Lógica para Frequência Semanal
                    val_freq_sr = data_pdi.get('paee_freq_sr')
                    idx_freq_sr = opcoes_freq.index(val_freq_sr) if val_freq_sr in opcoes_freq else 0
                    data_pdi['paee_freq_sr'] = st.radio("Frequência Semanal", opcoes_freq, index=idx_freq_sr, horizontal=True)
                    
                    # Lógica para Atendimento
                    val_atend_sr = data_pdi.get('paee_atend_sr')
                    idx_atend_sr = opcoes_atend.index(val_atend_sr) if val_atend_sr in opcoes_atend else 0
                    data_pdi['paee_atend_sr'] = st.radio("Atendimento", opcoes_atend, index=idx_atend_sr, horizontal=True)
                    
                with c2:
                    st.markdown("**Colaborativo**")
                    
                    # Lógica para Frequência Semanal (Colab.)
                    val_freq_colab = data_pdi.get('paee_freq_colab')
                    idx_freq_colab = opcoes_freq.index(val_freq_colab) if val_freq_colab in opcoes_freq else 0
                    data_pdi['paee_freq_colab'] = st.radio("Frequência Semanal (Colab.)", opcoes_freq, index=idx_freq_colab, horizontal=True)
        
                
                st.divider()
                # Nota: Você colocou o subheader 4.2 duas vezes no seu código, mantive a sua estrutura abaixo.
                st.subheader("4.2 Objetivos da ação educativa")
                
                def render_semestres(label, key_base):
                    st.markdown(f"**{label}**")
                    cs1, cs2 = st.columns(2)
                    
                    # Adicionado o parâmetro 'key' para garantir que cada text_area seja único
                    data_pdi[f"obj_{key_base}_1sem"] = cs1.text_area(
                        "1º Semestre", 
                        value=data_pdi.get(f"obj_{key_base}_1sem", ""), 
                        height=100,
                        key=f"txt_1sem_{key_base}"
                    )
                    
                    data_pdi[f"obj_{key_base}_2sem"] = cs2.text_area(
                        "2º Semestre", 
                        value=data_pdi.get(f"obj_{key_base}_2sem", ""), 
                        height=100,
                        key=f"txt_2sem_{key_base}"
                    )
                    st.divider()

                # ---- COLOQUE ESTAS 4 LINHAS AQUI ----
                render_semestres("DESENVOLVIMENTO COGNITIVO", "cog")
                render_semestres("DESENVOLVIMENTO MOTOR", "mot")
                render_semestres("HABILIDADES PESSOAIS E DE SOCIALIZAÇÃO", "soc")
                render_semestres("FUNÇÃO DO BRINCAR", "bri")
                # -------------------------------------
                
                if st.form_submit_button("💾 Salvar PAEE"):
                    save_student("PDI", data_pdi.get('nome'), data_pdi, "Plano AEE (PAEE)")

# --- ABA 3: ASSINATURAS (NOVA) ---
        with tabs[2]:
            st.subheader("Assinaturas Digitais")
            st.caption(f"Código Único do Documento: {data_pdi.get('doc_uuid', 'Não gerado ainda')}")
            
            # Identify required signers based on content (Puxando do data_case)
            required_roles = []
            if data_case.get('prof_poli'): required_roles.append({'role': 'Prof. Polivalente', 'name': data_case['prof_poli']})
            if data_case.get('prof_aee'): required_roles.append({'role': 'Prof. AEE', 'name': data_case['prof_aee']})
            if data_case.get('prof_arte'): required_roles.append({'role': 'Prof. Arte', 'name': data_case['prof_arte']})
            if data_case.get('prof_ef'): required_roles.append({'role': 'Prof. Ed. Física', 'name': data_case['prof_ef']})
            if data_case.get('prof_tec'): required_roles.append({'role': 'Prof. Tecnologia', 'name': data_case['prof_tec']})
            if data_case.get('gestor'): required_roles.append({'role': 'Gestor Escolar', 'name': data_case['gestor']})
            if data_case.get('coord'): required_roles.append({'role': 'Coordenação', 'name': data_case['coord']})
            
            # Show list of signatories
            if required_roles:
                st.markdown("##### Profissionais Citados no Documento")
                for r in required_roles:
                    st.write(f"- **{r['role']}:** {r['name']}")
            else:
                st.info("Nenhum profissional identificado automaticamente nos campos.")

            st.divider()
            
            # Current Signatures
            current_signatures = data_pdi.get('signatures', [])
            if current_signatures:
                st.success("✅ Documento assinado por:")
                for sig in current_signatures:
                    st.write(f"✍️ **{sig['name']}** ({sig.get('role', 'Profissional')}) em {sig['date']}")
            else:
                st.warning("Nenhuma assinatura registrada.")

            st.divider()

            # Ação de assinar
            opcoes_cargos = ["Professor AEE", "Prof. Polivalente", "Gestão Escolar", "Coordenação Pedagógica", "Outro"]
            cargo_selecionado = st.selectbox("Assinar como:", opcoes_cargos)
            
            if st.button("🖊️ Assinar Digitalmente"):
                new_sig = {"name": st.session_state.get('usuario_nome',''), "date": datetime.now().strftime("%d/%m/%Y"), "role": cargo_selecionado}
                if 'signatures' not in data_pdi: data_pdi['signatures'] = []
                data_pdi['signatures'].append(new_sig)
                save_student("PDI", data_pdi.get('nome'), data_pdi, "Assinatura")
                st.rerun()

        # --- ABA 4: PDF ---
        with tabs[3]:
            st.subheader("Geração do Documento")
            st.info("O rodapé de cada página no PDF trará a validação das assinaturas digitais, garantindo a autenticidade.")
            
            if st.button("👁️ GERAR PDI COMPLETO (PDF)"):
                log_action(data_pdi.get('nome'), "Gerou PDF", "PDI Completo")
                
                pdf = OfficialPDF('P', 'mm', 'A4')
                pdf.set_auto_page_break(auto=True, margin=15)
                
                # A mágica da assinatura de rodapé
                pdf.set_signature_footer(data_pdi.get('signatures', []), data_pdi.get('doc_uuid', ''))
                
                # --- CAPA PRINCIPAL ---
                pdf.add_page()
                if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 10, 10, 25)
                if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 175, 10, 25)

                pdf.set_y(15); pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, clean_pdf_text("PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
                pdf.cell(0, 10, clean_pdf_text("SECRETARIA MUNICIPAL DE EDUCAÇÃO"), 0, 1, 'C')
                
                pdf.ln(40)
                pdf.set_font("Arial", "B", 30)
                pdf.cell(0, 20, "PDI", 0, 1, 'C')
                pdf.set_font("Arial", "B", 20)
                pdf.cell(0, 15, "PLANO DE DESENVOLVIMENTO", 0, 1, 'C')
                pdf.cell(0, 15, "INDIVIDUAL", 0, 1, 'C')
                
                pdf.ln(20)
                pdf.set_font("Arial", "", 16)
                pdf.cell(0, 10, "Estudo de Caso e Plano de AEE", 0, 1, 'C')
                pdf.ln(40)
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, f"ANO: {datetime.now().year}", 0, 1, 'C')

                # --- CAPA SECUNDÁRIA: ESTUDO DE CASO ---
                pdf.add_page()
                if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 10, 10, 25)
                if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 175, 10, 25)
                
                pdf.set_y(120) 
                pdf.set_font("Arial", "B", 24)
                pdf.cell(0, 10, "ESTUDO DE CASO", 0, 1, 'C')

                # ==========================================================
                # INÍCIO DO CONTEÚDO DO ESTUDO DE CASO
                # ==========================================================
                data = data_case  
                pdf.set_margins(15, 15, 15)
                # --- 1.1 DADOS GERAIS ---
                pdf.add_page()
                pdf.section_title("1.1 DADOS GERAIS DO ESTUDANTE", width=0)
                pdf.ln(4)
                
                # 1.1.1 IDENTIFICAÇÃO
                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, "1.1.1 - IDENTIFICAÇÃO", 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [
                    (30, "Nome:", "B", "L", True),
                    (110, clean_pdf_text(data.get('nome', '')), "", "L", False),
                    (15, "D.N.:", "B", "C", True),
                    (25, clean_pdf_text(str(data.get('d_nasc', ''))), "", "C", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (30, "Escolaridade:", "B", "L", True),
                    (25, clean_pdf_text(data.get('ano_esc', '')), "", "L", False),
                    (20, "Período:", "B", "C", True),
                    (20, clean_pdf_text(data.get('periodo', '')), "", "C", False),
                    (20, "Unidade:", "B", "C", True),
                    (65, clean_pdf_text(data.get('unidade', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (30, "Endereço:", "B", "L", True),
                    (150, clean_pdf_text(data.get('endereco', '')), "", "L", False)
                ], line_h=7, font_size=10)

                draw_flex_row(pdf, [
                    (20, "Bairro:", "B", "L", True),
                    (70, clean_pdf_text(data.get('bairro', '')), "", "L", False),
                    (20, "Cidade:", "B", "C", True),
                    (70, clean_pdf_text(data.get('cidade', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (20, "Telefone:", "B", "L", True),
                    (160, clean_pdf_text(data.get('telefones', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                # 1.1.2 DADOS FAMILIARES
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, "1.1.2 - DADOS FAMILIARES", 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [
                    (20, "Pai:", "B", "L", True),
                    (80, clean_pdf_text(data.get('pai_nome', '')), "", "L", False),
                    (25, "Profissão:", "B", "C", True),
                    (55, clean_pdf_text(data.get('pai_prof', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (20, "Mãe:", "B", "L", True),
                    (80, clean_pdf_text(data.get('mae_nome', '')), "", "L", False),
                    (25, "Profissão:", "B", "C", True),
                    (55, clean_pdf_text(data.get('mae_prof', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                # Irmãos
                pdf.ln(2)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, clean_pdf_text("Irmãos (Nome | Idade | Escolaridade)"), 1, 1, 'L', 1)
                for i, irmao in enumerate(data.get('irmaos', [])):
                    if irmao['nome']:
                        txt = f"{irmao['nome']}  |  {irmao['idade']}  |  {irmao['esc']}"
                        draw_flex_row(pdf, [(180, clean_pdf_text(txt), "", "L", False)], line_h=6, font_size=9)
                
                pdf.ln(2)
                draw_flex_row(pdf, [
                    (40, "Com quem mora:", "B", "L", True),
                    (140, clean_pdf_text(data.get('quem_mora', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (40, "Convênio Médico:", "B", "L", True),
                    (50, clean_pdf_text(data.get('convenio')), "", "L", False),
                    (20, "Qual:", "B", "C", True),
                    (70, clean_pdf_text(data.get('convenio_qual')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (40, "Benefício Social:", "B", "L", True),
                    (50, clean_pdf_text(data.get('social')), "", "L", False),
                    (20, "Qual:", "B", "C", True),
                    (70, clean_pdf_text(data.get('social_qual')), "", "L", False)
                ], line_h=7, font_size=10)

                # 1.1.3 HISTÓRIA ESCOLAR
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, clean_pdf_text("1.1.3 - HISTÓRIA ESCOLAR"), 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [(50, "Idade entrou na escola:", "B", "L", True), (130, clean_pdf_text(data.get('hist_idade_entrou')), "", "L", False)], line_h=7, font_size=10)
                draw_flex_row(pdf, [(50, "Outras escolas:", "B", "L", True), (130, clean_pdf_text(data.get('hist_outra_escola')), "", "L", False)], line_h=7, font_size=10)
                draw_flex_row(pdf, [(50, "Motivo transferência:", "B", "L", True), (130, clean_pdf_text(data.get('hist_motivo_transf')), "", "L", False)], line_h=7, font_size=10)
                
                if data.get('hist_obs'):
                    pdf.ln(2)
                    pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "Observações Escolares:", 0, 1)
                    draw_flex_row(pdf, [(180, clean_pdf_text(data.get('hist_obs')), "", "L", False)], line_h=6, font_size=9)

                # --- 1.2 GESTAÇÃO, PARTO E DESENVOLVIMENTO ---
                pdf.add_page()
                pdf.section_title("1.2 GESTAÇÃO, PARTO E DESENVOLVIMENTO", width=0)
                pdf.ln(4)
                
                def print_data_row(label, value):
                    draw_flex_row(pdf, [
                        (80, clean_pdf_text(label), "B", "L", True),
                        (100, clean_pdf_text(str(value) if value else ""), "", "L", False)
                    ], line_h=6, font_size=9)

                rows_gest = [
                    ("Parentesco entre pais:", data.get('gest_parentesco')),
                    ("Doença/Trauma na gestação:", data.get('gest_doenca')),
                    ("Uso de substâncias (mãe):", data.get('gest_substancias')),
                    ("Uso de medicamentos (mãe):", data.get('gest_medicamentos')),
                    ("Ocorrência no parto:", data.get('parto_ocorrencia')),
                    ("Necessitou de incubadora:", data.get('parto_incubadora')),
                    ("Prematuro?", f"{data.get('parto_prematuro')}  |  UTI: {data.get('parto_uti')}"),
                    ("Tempo de gestação / Peso:", f"{data.get('dev_tempo_gest')}  /  {data.get('dev_peso')}"),
                    ("Desenvolvimento normal no 1º ano:", data.get('dev_normal_1ano')),
                    ("Apresentou atraso importante?", data.get('dev_atraso')),
                    ("Idade que andou / falou:", f"{data.get('dev_idade_andar')}  /  {data.get('dev_idade_falar')}"),
                    ("Possui diagnóstico?", data.get('diag_possui')),
                    ("Reação da família ao diagnóstico:", data.get('diag_reacao')),
                    ("Data / Origem do diagnóstico:", f"{data.get('diag_data')}  |  {data.get('diag_origem')}"),
                    ("Pessoa com deficiência na família:", data.get('fam_deficiencia')),
                    ("Pessoa com AH/SD na família:", data.get('fam_altas_hab'))
                ]
                
                for label, value in rows_gest:
                    print_data_row(label, value)

                # --- 1.3 INFORMAÇÕES SOBRE SAÚDE ---
                pdf.add_page()
                pdf.section_title("1.3 INFORMAÇÕES SOBRE SAÚDE", width=0)
                pdf.ln(4)
                
                saude_rows = [
                    ("Problemas de saúde:", data.get('saude_prob')),
                    ("Já necessitou de internação:", data.get('saude_internacao')),
                    ("Restrição/Seletividade alimentar:", data.get('saude_restricao')),
                    ("Uso de medicamentos controlados:", f"{data.get('med_uso')} - Quais: {data.get('med_quais')}"),
                    ("Horário / Dosagem / Início:", f"{data.get('med_hor')}  |  {data.get('med_dos')}  |  {data.get('med_ini')}"),
                    ("Qualidade do sono:", data.get('sono')),
                    ("Última visita ao médico:", data.get('medico_ultimo'))
                ]
                for label, value in saude_rows:
                    print_data_row(label, value)
                
                esf = []
                if data.get('esf_urina'): esf.append("Urina")
                if data.get('esf_fezes'): esf.append("Fezes")
                print_data_row("Controle de Esfíncter:", f"{', '.join(esf) if esf else 'Não'}  (Idade: {data.get('esf_idade')})")
                
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10); pdf.set_fill_color(240, 240, 240)
                pdf.cell(0, 8, "Atendimentos Clínicos Extraescolares", 1, 1, 'L', 1)
                
                clins = data.get('clinicas', [])
                print_data_row("Realiza atendimento em:", ", ".join(clins) if clins else "Não realiza")
                print_data_row("Especialidade médica:", data.get('clinicas_med_esp'))
                print_data_row("Nome da Clínica/Profissional:", data.get('clinicas_nome'))
                
                if data.get('saude_obs_geral'):
                    pdf.ln(2)
                    pdf.set_font("Arial", "B", 9); pdf.cell(0, 6, "Outras observações de saúde:", 0, 1)
                    draw_flex_row(pdf, [(180, clean_pdf_text(data.get('saude_obs_geral')), "", "L", False)], line_h=5, font_size=9)

                # --- 1.4 COMPREENSÃO DA FAMÍLIA (CHECKLIST) ---
                pdf.add_page()
                pdf.section_title("1.4 COMPREENSÃO DA FAMÍLIA (CHECKLIST)", width=0)
                pdf.ln(4)
                
                draw_flex_row(pdf, [
                    (110, "PERGUNTA / ASPECTO OBSERVADO", "B", "C", True),
                    (25, "SIM/NÃO", "B", "C", True),
                    (45, "OBSERVAÇÕES DA FAMÍLIA", "B", "C", True)
                ], line_h=8, font_size=9, fill_color=(220, 220, 220))
                
                checklist_items = [
                    "Relata fatos do dia a dia? Apresentando boa memória?",
                    "É organizado com seus pertences?",
                    "Aceita regras de forma tranquila?",
                    "Busca e aceita ajuda quando não sabe ou não consegue algo?",
                    "Aceita alterações no ambiente?",
                    "Tem algum medo?",
                    "Tem alguma mania?",
                    "Tem alguma área/assunto, brinquedo ou hiperfoco?",
                    "Prefere brincar sozinho ou com outras crianças? Tem amigos?",
                    "Qual a expectativa da família em relação à escolaridade da criança?"
                ]
                
                pdf.set_font("Arial", "", 9)
                
                for i, item in enumerate(checklist_items):
                    key_base = f"itemcomport_{i}"
                    opt = data_case.get('checklist', {}).get(f"{key_base}_opt", "Não")
                    obs = data_case.get('checklist', {}).get(f"{key_base}_obs", "")
                    
                    line_height = 6
                    num_lines = pdf.get_string_width(obs) / 50 
                    cell_height = max(line_height, (int(num_lines) + 1) * line_height)
                    
                    x_start = pdf.get_x(); y_start = pdf.get_y()
                    
                    pdf.multi_cell(110, line_height, clean_pdf_text(item), 1, 'L')
                    
                    pdf.set_xy(x_start + 110, y_start)
                    pdf.cell(25, cell_height, clean_pdf_text(opt), 1, 0, 'C')
                    
                    pdf.set_xy(x_start + 135, y_start)
                    pdf.multi_cell(0, line_height, clean_pdf_text(obs), 1, 'L')
                    
                    pdf.set_xy(x_start, y_start + cell_height)
                
                # ==========================================================
                # FIM DO CONTEÚDO DO ESTUDO DE CASO
                # RETOMADA DO PDI
                pdf.set_margins(10, 10, 10)
                # ==========================================================

                # --- CAPA SECUNDÁRIA: PLANO DE AEE ---
                pdf.add_page()
                if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 10, 10, 25)
                if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 175, 10, 25)
                
                pdf.set_y(100)
                pdf.set_font("Arial", "B", 20)
                pdf.cell(0, 10, clean_pdf_text("PLANO DE AEE"), 0, 1, 'C')
                pdf.ln(5)
                pdf.cell(0, 10, clean_pdf_text("ATENDIMENTO EDUCACIONAL"), 0, 1, 'C')
                pdf.cell(0, 10, clean_pdf_text("ESPECIALIZADO"), 0, 1, 'C')

                # --- 3. AVALIAÇÃO PEDAGÓGICA DO ESTUDANTE ---
                pdf.add_page()
                pdf.section_title("3. AVALIAÇÃO PEDAGÓGICA DO ESTUDANTE", width=0)
                pdf.ln(5)
                
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "3.1 POTENCIALIDADES:", 0, 1)
                pdf.set_font("Arial", "", 10)
                pdf.multi_cell(0, 5, clean_pdf_text(data_pdi.get('potencialidades', '')), 1)

                pdf.ln(5)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "3.2 ÁREAS DE INTERESSE:", 0, 1)
                pdf.set_font("Arial", "", 10)
                pdf.multi_cell(0, 5, clean_pdf_text(data_pdi.get('areas_interesse', '')), 1)

                pdf.ln(5)

                # ==========================================================
                # MUDANÇA PARA MODO PAISAGEM (LANDSCAPE) A PARTIR DO 3.3
                # ==========================================================
                pdf.add_page(orientation='L')
                
                # Configurações de colunas para Landscape (Largura útil ~277mm)
                col_dominio = 77
                col_percurso = 100
                col_final = 100

                def print_tabela_cabecalho(pdf_obj):
                    pdf_obj.set_fill_color(220, 220, 220)
                    pdf_obj.set_font("Arial", "B", 10)
                    pdf_obj.cell(col_dominio, 8, "Domínios", 1, 0, 'C', True)
                    pdf_obj.cell(col_percurso, 8, "Resultados da Avaliação de Percurso", 1, 0, 'C', True)
                    pdf_obj.cell(col_final, 8, "Resultados da Avaliação Final", 1, 1, 'C', True)

                def print_linha_paisagem(pdf_obj, label, key_base, options=["Sim", "Não", "Parcialmente"]):
                    pdf_obj.set_font("Arial", "B", 9)
                    
                    v_proc = data_pdi.get(f"{key_base}_proc", "")
                    v_final = data_pdi.get(f"{key_base}_final", "")

                    str_proc = "  ".join([f"(X) {opt}" if opt == v_proc else f"( ) {opt}" for opt in options])
                    str_final = "  ".join([f"(X) {opt}" if opt == v_final else f"( ) {opt}" for opt in options])

                    if pdf_obj.get_y() > 175:
                        pdf_obj.add_page(orientation='L')
                        print_tabela_cabecalho(pdf_obj)
                    
                    x_start = pdf_obj.get_x()
                    y_start = pdf_obj.get_y()
                    
                    pdf_obj.multi_cell(col_dominio, 8, clean_pdf_text(label), 1, 'L')
                    y_end = pdf_obj.get_y()
                    h_row = max(8, y_end - y_start)
                    
                    pdf_obj.set_xy(x_start + col_dominio, y_start)
                    pdf_obj.set_font("Arial", "", 9)
                    pdf_obj.cell(col_percurso, h_row, clean_pdf_text(str_proc), 1, 0, 'L')
                    pdf_obj.cell(col_final, h_row, clean_pdf_text(str_final), 1, 1, 'L')

                def print_linha_paisagem_multiline(pdf_obj, label, str_proc, str_final):
                    """Imprime blocos complexos onde as opções precisam ser quebradas em várias linhas."""
                    if pdf_obj.get_y() > 165:
                        pdf_obj.add_page(orientation='L')
                        print_tabela_cabecalho(pdf_obj)
                        
                    x_start = pdf_obj.get_x()
                    y_start = pdf_obj.get_y()
                    
                    pdf_obj.set_font("Arial", "", 8)
                    h_proc = (str_proc.count('\n') + 1) * 5 + 4
                    h_final = (str_final.count('\n') + 1) * 5 + 4
                    h_row = max(h_proc, h_final, 10)
                    
                    # Domínio
                    pdf_obj.set_font("Arial", "B", 9)
                    pdf_obj.multi_cell(col_dominio, 5, clean_pdf_text(label), 0, 'L')
                    
                    # Percurso
                    pdf_obj.set_xy(x_start + col_dominio, y_start + 2)
                    pdf_obj.set_font("Arial", "", 8)
                    pdf_obj.multi_cell(col_percurso, 5, clean_pdf_text(str_proc), 0, 'L')
                    
                    # Final
                    pdf_obj.set_xy(x_start + col_dominio + col_percurso, y_start + 2)
                    pdf_obj.multi_cell(col_final, 5, clean_pdf_text(str_final), 0, 'L')
                    
                    # Bordas
                    pdf_obj.rect(x_start, y_start, col_dominio, h_row)
                    pdf_obj.rect(x_start + col_dominio, y_start, col_percurso, h_row)
                    pdf_obj.rect(x_start + col_dominio + col_percurso, y_start, col_final, h_row)
                    pdf_obj.set_xy(x_start, y_start + h_row)

                def build_options_string(selected_val, options_list):
                    """Auxiliar para gerar o texto multilinhas."""
                    return "\n".join([f"(X) {opt}" if opt == selected_val else f"( ) {opt}" for opt in options_list])

                print_tabela_cabecalho(pdf)

                # 3.3 COGNITIVO
                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "3.3 DESENVOLVIMENTO COGNITIVO - ATENÇÃO", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "Concentrada", "atencao_conc")
                print_linha_paisagem(pdf, "Sustentada", "atencao_sust")
                print_linha_paisagem(pdf, "Seletiva", "atencao_sel")
                print_linha_paisagem(pdf, "Alternada", "atencao_alt")

                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "PERCEPÇÃO", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "Memória Visual/Percepção de dif. e semelhanças", "perc_vis")
                print_linha_paisagem(pdf, "Percepção e Discriminação Auditiva", "perc_aud")
                print_linha_paisagem(pdf, "Percepção Tátil", "perc_tat")
                print_linha_paisagem(pdf, "Orientação temporal", "perc_temp")
                print_linha_paisagem(pdf, "Orientação espacial", "perc_esp")

                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "MEMÓRIA", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "Curto Prazo", "mem_curto")
                print_linha_paisagem(pdf, "Médio Prazo", "mem_medio")
                print_linha_paisagem(pdf, "Longo Prazo", "mem_longo")

                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "FUNÇÕES EXECUTIVAS", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "Controle Inibitório (inibição de impulsos)", "fe_inib")
                print_linha_paisagem(pdf, "Memória de Trabalho (reter e manipular info)", "fe_trab")
                print_linha_paisagem(pdf, "Flexibilidade Cognitiva (adaptação a mudanças)", "fe_flex")
                print_linha_paisagem(pdf, "Planejamento/Organização", "fe_plan")

                # LINGUAGEM
                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "LINGUAGEM E COMUNICAÇÃO", 1, 1, 'C', True)
                
                print_linha_paisagem(pdf, "Utiliza palavras para se comunicar?", "ling_palavras")
                
                str_ca_p = f"(X) Gestos\n" if data_pdi.get('ling_gestos_p') else "( ) Gestos\n"
                str_ca_p += f"(X) Apontamentos\n" if data_pdi.get('ling_aponta_p') else "( ) Apontamentos\n"
                str_ca_p += f"(X) Piscar dos olhos\n" if data_pdi.get('ling_pisca_p') else "( ) Piscar dos olhos\n"
                str_ca_p += f"Com. Alternativa: {data_pdi.get('ling_ca_p', '')}\n" if data_pdi.get('ling_ca_p') else "( ) Com. Alternativa\n"
                str_ca_p += f"LIBRAS: {data_pdi.get('ling_libras_p', 'Não')}\n"
                str_ca_p += f"Outros: {data_pdi.get('ling_outros_p', '')}"

                str_ca_f = f"(X) Gestos\n" if data_pdi.get('ling_gestos_f') else "( ) Gestos\n"
                str_ca_f += f"(X) Apontamentos\n" if data_pdi.get('ling_aponta_f') else "( ) Apontamentos\n"
                str_ca_f += f"(X) Piscar dos olhos\n" if data_pdi.get('ling_pisca_f') else "( ) Piscar dos olhos\n"
                str_ca_f += f"Com. Alternativa: {data_pdi.get('ling_ca_f', '')}\n" if data_pdi.get('ling_ca_f') else "( ) Com. Alternativa\n"
                str_ca_f += f"LIBRAS: {data_pdi.get('ling_libras_f', 'Não')}\n"
                str_ca_f += f"Outros: {data_pdi.get('ling_outros_f', '')}"

                print_linha_paisagem_multiline(pdf, "Outras formas de comunicação / LIBRAS", str_ca_p, str_ca_f)

                print_linha_paisagem(pdf, "Apresenta trocas fonéticas orais?", "ling_trocas")
                print_linha_paisagem(pdf, "Estabelece diálogo com troca de turno?", "ling_dialogo")
                print_linha_paisagem(pdf, "Inventa frases ou histórias?", "ling_inventa")
                print_linha_paisagem(pdf, "Descreve cenas com sentido?", "ling_descreve")
                print_linha_paisagem(pdf, "Consegue expressar pensamentos e desejos?", "ling_expressa")
                print_linha_paisagem(pdf, "Reconta história / relatos do cotidiano?", "ling_reconta")

                pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "LINGUAGEM COMPREENSIVA", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "Compreende informações orais simples?", "ling_comp_simp")
                print_linha_paisagem(pdf, "Compreende informações orais complexas?", "ling_comp_comp")

                pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "LINGUAGEM ESCRITA", 1, 1, 'C', True)
                
                opts_escrita = [
                    "Não distingue desenho, letras e números.", "Identifica e nomeia as letras.", 
                    "Escreve seu nome.", "Escreve letras de forma aleatória.", 
                    "Relaciona som/grafia.", "Escreve apenas palavras canônicas.", "Escreve palavras não-canônicas."
                ]
                str_escrita_p = build_options_string(data_pdi.get('escrita_conv_proc', ''), opts_escrita)
                str_escrita_f = build_options_string(data_pdi.get('escrita_conv_final', ''), opts_escrita)
                print_linha_paisagem_multiline(pdf, "Escreve convencionalmente?", str_escrita_p, str_escrita_f)

                opts_org = ["Não escreve textos convencionais.", "Escreve frases simples.", "Escreve textos."]
                str_org_p = build_options_string(data_pdi.get('escrita_org_proc', ''), opts_org)
                str_org_f = build_options_string(data_pdi.get('escrita_org_final', ''), opts_org)
                print_linha_paisagem_multiline(pdf, "Apresenta organização textual?", str_org_p, str_org_f)

                pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "LEITURA", 1, 1, 'C', True)
                opts_leitura = [
                    "Não realiza leitura.", "Realiza leitura apenas de palavras canônicas.", 
                    "Realiza leitura de palavras canônicas e não-canônicas.", "Realiza leitura de frases e textos com dificuldade.", 
                    "Realiza leitura de frases e textos com fluência.", "Compreende o que lê com apoio.", "Compreende o que lê com autonomia."
                ]
                str_leit_p = build_options_string(data_pdi.get('leitura_nivel_proc', ''), opts_leitura)
                str_leit_f = build_options_string(data_pdi.get('leitura_nivel_final', ''), opts_leitura)
                print_linha_paisagem_multiline(pdf, "Nível de Leitura", str_leit_p, str_leit_f)

                opts_braille = ["Não utiliza", "Com autonomia.", "Com apoio.", "Com dificuldade."]
                str_br_p = "BRAILLE:\n" + build_options_string(data_pdi.get('braille_p', ''), opts_braille) + f"\nOutros: {data_pdi.get('leitura_outros_p', '')}"
                str_br_f = "BRAILLE:\n" + build_options_string(data_pdi.get('braille_f', ''), opts_braille) + f"\nOutros: {data_pdi.get('leitura_outros_f', '')}"
                print_linha_paisagem_multiline(pdf, "Leitura em BRAILLE / Outros", str_br_p, str_br_f)
                
                if data_pdi.get('ling_obs'):
                    if pdf.get_y() > 175: pdf.add_page(orientation='L')
                    pdf.set_font("Arial", "B", 9)
                    pdf.multi_cell(0, 8, clean_pdf_text(f"Observações: {data_pdi.get('ling_obs')}"), 1, 'L')

                # RACIOCÍNIO E DESENVOLVIMENTO MOTOR E OUTROS
                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "RACIOCÍNIO E RESOLUÇÃO DE PROBLEMAS", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "a) Planeja, antecipa, argumenta?", "rac_plan")
                print_linha_paisagem(pdf, "b) Compara, classifica, categoriza, sequencia, inferi?", "rac_comp")
                print_linha_paisagem(pdf, "c) Conhece conceitos básicos do vocabulário matemático?", "rac_mat")
                print_linha_paisagem(pdf, "d) Tem capacidade de conclusões lógicas?", "rac_log")

                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "3.4 DESENVOLVIMENTO MOTOR", 1, 1, 'C', True)
                print_linha_paisagem(pdf, "Realiza os diversos tipos de locomoção?", "mot_loc")
                print_linha_paisagem(pdf, "Manipula bola e outros objetos?", "mot_bola")
                print_linha_paisagem(pdf, "Desenvolvimento esperado da lateralidade?", "mot_lat")
                print_linha_paisagem(pdf, "Capacidades físicas (equilíbrio, força, flex)?", "mot_fis")
                print_linha_paisagem(pdf, "Preensão trípode e rotação de punho?", "mot_preensao")

                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "3.5 HABILIDADES PESSOAIS E DE SOCIALIZAÇÃO", 1, 1, 'C', True)
                opts_hab = ["Com autonomia", "Com ajuda", "Não realiza"]
                print_linha_paisagem(pdf, "Alimentação", "hab_alim", options=opts_hab)
                print_linha_paisagem(pdf, "Higiene", "hab_hig", options=opts_hab)
                print_linha_paisagem(pdf, "Uso funcional dos objetos", "hab_obj", options=opts_hab)
                print_linha_paisagem(pdf, "Locomoção", "hab_loc", options=opts_hab)
                print_linha_paisagem(pdf, "Interage com os adultos?", "soc_adul")
                print_linha_paisagem(pdf, "Interage com os colegas?", "soc_col")
                print_linha_paisagem(pdf, "Tem tolerância a frustração?", "soc_frust")

                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "3.6 FUNÇÃO DO BRINCAR", 1, 1, 'C', True)
                opts_brincar = ["Sim", "Não", "Com modelo"]
                print_linha_paisagem(pdf, "Uso dos brinquedos de maneira funcional?", "bri_func", options=opts_brincar)
                print_linha_paisagem(pdf, "Explora os brinquedos espontaneamente?", "bri_explora", options=opts_brincar)
                print_linha_paisagem(pdf, "Utiliza objetos atribuindo diferentes funções?", "bri_dif", options=opts_brincar)
                print_linha_paisagem(pdf, "Estrutura uma brincadeira de forma criativa?", "bri_cria", options=opts_brincar)

                # ==========================================================
                # NOVO ITEM 4 - PAEE E OBJETIVOS DE AÇÃO EDUCATIVA
                # ==========================================================
                if pdf.get_y() > 160: pdf.add_page(orientation='L')
                pdf.ln(10)
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, "4. PLANO DE ATENDIMENTO EDUCACIONAL ESPECIALIZADO (PAEE)", 0, 1, 'C')
                
                # 4.1 ORGANIZAÇÃO DO AEE
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "4.1 ORGANIZAÇÃO DO AEE", 0, 1, 'L')
                
                pdf.set_font("Arial", "B", 9)
                pdf.set_fill_color(240, 240, 240)
                pdf.cell(138, 8, "Sala de Recursos", 1, 0, 'C', True)
                pdf.cell(139, 8, "Colaborativo", 1, 1, 'C', True)
                
                pdf.set_font("Arial", "", 9)
                freq_rec = data_pdi.get('paee_freq_sr', '')
                atend_rec = data_pdi.get('paee_atend_sr', '')
                freq_colab = data_pdi.get('paee_freq_colab', '')
                
                opts_freq = ["Não realiza", "1x", "2x", "3x", "4x ou mais"]
                opts_atend = ["Individual", "Grupo"]
                
                str_freq_rec = "Frequência: " + "  ".join([f"(X) {o}" if o == freq_rec else f"( ) {o}" for o in opts_freq])
                str_atend_rec = "Atendimento: " + "  ".join([f"(X) {o}" if o == atend_rec else f"( ) {o}" for o in opts_atend])
                str_freq_colab = "Frequência: " + "  ".join([f"(X) {o}" if o == freq_colab else f"( ) {o}" for o in opts_freq])
                
                pdf.cell(138, 8, clean_pdf_text(str_freq_rec), 'LR', 0, 'L')
                pdf.cell(139, 8, clean_pdf_text(str_freq_colab), 'LR', 1, 'L')
                pdf.cell(138, 8, clean_pdf_text(str_atend_rec), 'LRB', 0, 'L')
                pdf.cell(139, 8, "", 'LRB', 1, 'L')
                
                # 4.2 OBJETIVOS DA AÇÃO EDUCATIVA
                pdf.ln(5)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "4.2 Objetivos da ação educativa", 0, 1, 'L')
                
                col_dom = 57
                col_sem = 110
                pdf.set_fill_color(220, 220, 220)
                pdf.cell(col_dom, 8, "Domínios", 1, 0, 'C', True)
                pdf.cell(col_sem, 8, "1º SEMESTRE", 1, 0, 'C', True)
                pdf.cell(col_sem, 8, "2º SEMESTRE", 1, 1, 'C', True)
                
                domains = [
                    ("DESENVOLVIMENTO COGNITIVO", "cog"),
                    ("DESENVOLVIMENTO MOTOR", "mot"),
                    ("HABILIDADES PESSOAIS E DE SOCIALIZAÇÃO", "soc"),
                    ("FUNÇÃO DO BRINCAR", "bri")
                ]
                
                for dom_label, key in domains:
                    val1 = data_pdi.get(f"obj_{key}_1sem", "")
                    val2 = data_pdi.get(f"obj_{key}_2sem", "")
                    
                    pdf.set_font("Arial", "", 9)
                    
                    w1 = pdf.get_string_width(clean_pdf_text(val1))
                    w2 = pdf.get_string_width(clean_pdf_text(val2))
                    
                    lines1 = max(1, int(w1 / (col_sem - 4)) + 1) + val1.count('\n')
                    lines2 = max(1, int(w2 / (col_sem - 4)) + 1) + val2.count('\n')
                    
                    h_row = max(12, lines1 * 5 + 4, lines2 * 5 + 4)
                    
                    if pdf.get_y() + h_row > 180:
                        pdf.add_page(orientation='L')
                        pdf.set_fill_color(220, 220, 220)
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(col_dom, 8, "Domínios", 1, 0, 'C', True)
                        pdf.cell(col_sem, 8, "1º SEMESTRE", 1, 0, 'C', True)
                        pdf.cell(col_sem, 8, "2º SEMESTRE", 1, 1, 'C', True)
                        
                    x_start = pdf.get_x()
                    y_start = pdf.get_y()
                    
                    pdf.set_font("Arial", "B", 9)
                    pdf.multi_cell(col_dom, 6, clean_pdf_text(dom_label), 0, 'C')
                    
                    pdf.set_xy(x_start + col_dom, y_start + 2)
                    pdf.set_font("Arial", "", 9)
                    pdf.multi_cell(col_sem, 5, clean_pdf_text(val1), 0, 'L')
                    
                    pdf.set_xy(x_start + col_dom + col_sem, y_start + 2)
                    pdf.multi_cell(col_sem, 5, clean_pdf_text(val2), 0, 'L')
                    
                    pdf.rect(x_start, y_start, col_dom, h_row)
                    pdf.rect(x_start + col_dom, y_start, col_sem, h_row)
                    pdf.rect(x_start + col_dom + col_sem, y_start, col_sem, h_row)
                    
                    pdf.set_xy(x_start, y_start + h_row)

                # ==========================================================
                # BLOCO VISUAL DE ASSINATURAS NO FINAL DO DOCUMENTO
                # ==========================================================
                # Verifica o espaço na página (se estiver muito no fim, joga pra próxima)
                if pdf.get_y() > 150: 
                    pdf.add_page(orientation='L')
                    
                pdf.ln(15)
                pdf.set_font("Arial", "B", 11)
                pdf.cell(0, 8, "REGISTRO DE ASSINATURAS", 0, 1, 'L')
                
                signatures = data_pdi.get('signatures', [])
                if signatures:
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(0, 6, clean_pdf_text("Este documento foi assinado digitalmente no sistema Integra pelos seguintes profissionais:"), 0, 1, 'L')
                    pdf.ln(5)
                    
                    for sig in signatures:
                        x_start = pdf.get_x()
                        
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(15, 6, "Nome:", 0, 0, 'L')
                        pdf.set_font("Arial", "", 10)
                        pdf.cell(100, 6, clean_pdf_text(sig.get('name', '')), 0, 0, 'L')
                        
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(15, 6, "Cargo:", 0, 0, 'L')
                        pdf.set_font("Arial", "", 10)
                        pdf.cell(80, 6, clean_pdf_text(sig.get('role', '')), 0, 0, 'L')
                        
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(12, 6, "Data:", 0, 0, 'L')
                        pdf.set_font("Arial", "", 10)
                        pdf.cell(0, 6, clean_pdf_text(sig.get('date', '')), 0, 1, 'L')
                        
                        pdf.ln(2) # Espaço entre as assinaturas
                else:
                    # Se ninguém assinou no sistema, deixa a linha física
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(0, 6, clean_pdf_text("Nenhuma assinatura digital foi registrada neste documento até o momento da geração."), 0, 1, 'L')
                    pdf.ln(20)
                    
                    # Linhas para assinatura manual de Prof. AEE e Gestor/Coordenação
                    pdf.cell(120, 6, "___________________________________________________", 0, 0, 'L')
                    pdf.cell(0, 6, "___________________________________________________", 0, 1, 'L')
                    pdf.cell(120, 6, "Professor(a) AEE", 0, 0, 'L')
                    pdf.cell(0, 6, "Coordenação / Gestão Escolar", 0, 1, 'L')
                
                # ==========================================================
                # FINALIZAÇÃO PDF
                # ==========================================================
                st.session_state.pdf_bytes_pdi = get_pdf_bytes(pdf)
                st.rerun()

            if 'pdf_bytes_pdi' in st.session_state:
                st.download_button("📥 BAIXAR PDI COMPLETO", st.session_state.pdf_bytes_pdi, f"PDI_{data_pdi.get('nome','aluno')}.pdf", "application/pdf", type="primary")

        # --- ABA 4: HISTÓRICO ---
        with tabs[3]:
            st.subheader("Histórico de Atividades")
            df_hist = safe_read("Historico", ["Data_Hora", "Aluno", "Usuario", "Acao", "Detalhes"])
            if not df_hist.empty and data_pdi.get('nome'):
                student_hist = df_hist[df_hist["Aluno"] == data_pdi.get('nome')]
                if not student_hist.empty:
                    st.dataframe(student_hist.iloc[::-1], use_container_width=True, hide_index=True)
                else: st.info("Sem histórico.")
            else: st.info("Histórico vazio.")
    



    


    # ESTUDO DE CASO COM FORMULÁRIOS
    elif doc_mode == "Estudo de Caso":
        st.markdown("""<div class="header-box"><div class="header-title">Estudo de Caso</div></div>""", unsafe_allow_html=True)
        
        if 'data_case' not in st.session_state: 
            st.session_state.data_case = {
                'irmaos': [{'nome': '', 'idade': '', 'esc': ''} for _ in range(4)], 
                'checklist': {},
                'clinicas': []
            }
        
        data = st.session_state.data_case
        
        st.markdown("""<style>div[data-testid="stFormSubmitButton"] > button {width: 100%; background-color: #dcfce7; color: #166534; border: 1px solid #166534;}</style>""", unsafe_allow_html=True)

        tabs = st.tabs(["1. Identificação", "2. Família", "3. Histórico", "4. Saúde", "5. Comportamento", "6. Assinaturas", "7. Gerar PDF", "8. Histórico"])

        # --- ABA 1: IDENTIFICAÇÃO ---
        with tabs[0]:
            with st.form("form_caso_identificacao") if not is_monitor else st.container():
                st.subheader("1.1 Dados Gerais do Estudante")
                data['nome'] = st.text_input("Nome Completo", value=data.get('nome', ''), disabled=True)
                
                c1, c2, c3 = st.columns([1, 1, 2])
                data['ano_esc'] = c1.text_input("Ano Escolaridade", value=data.get('ano_esc', ''), disabled=is_monitor)
                
                p_val = data.get('periodo') if data.get('periodo') in ["Manhã", "Tarde", "Integral"] else "Manhã"
                idx_per = ["Manhã", "Tarde", "Integral"].index(p_val)
                data['periodo'] = c2.selectbox("Período", ["Manhã", "Tarde", "Integral"], index=idx_per, disabled=is_monitor)
                data['unidade'] = c3.text_input("Unidade Escolar", value=data.get('unidade', ''), disabled=is_monitor)

                c4, c5 = st.columns([1, 1])
                data['sexo'] = c4.radio("Sexo", ["Feminino", "Masculino"], horizontal=True, index=0 if data.get('sexo') == 'Feminino' else 1, disabled=is_monitor)
                
                d_nasc = data.get('d_nasc')
                if isinstance(d_nasc, str):
                    try: d_nasc = datetime.strptime(d_nasc, '%Y-%m-%d').date()
                    except: d_nasc = date.today()
                data['d_nasc'] = c5.date_input("Data de Nascimento", value=d_nasc if d_nasc else date.today(), format="DD/MM/YYYY", disabled=is_monitor)

                data['endereco'] = st.text_input("Endereço", value=data.get('endereco', ''), disabled=is_monitor)
                c6, c7, c8 = st.columns([2, 2, 2])
                data['bairro'] = c6.text_input("Bairro", value=data.get('bairro', ''), disabled=is_monitor)
                data['cidade'] = c7.text_input("Cidade", value=data.get('cidade', 'Limeira'), disabled=is_monitor)
                data['telefones'] = c8.text_input("Telefones", value=data.get('telefones', ''), disabled=is_monitor)
                
                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Dados de Identificação"):
                        save_student("CASO", data.get('nome'), data, "Identificação")

        # --- ABA 2: DADOS FAMILIARES ---
        with tabs[1]:
            with st.form("form_caso_familia") if not is_monitor else st.container():
                st.subheader("1.1.2 Dados Familiares")
                
                st.markdown("**Pai**")
                c_p1, c_p2, c_p3, c_p4 = st.columns([3, 2, 2, 2])
                data['pai_nome'] = c_p1.text_input("Nome do Pai", value=data.get('pai_nome', ''), disabled=is_monitor)
                data['pai_prof'] = c_p2.text_input("Profissão Pai", value=data.get('pai_prof', ''), disabled=is_monitor)
                data['pai_esc'] = c_p3.text_input("Escolaridade Pai", value=data.get('pai_esc', ''), disabled=is_monitor)
                data['pai_dn'] = c_p4.text_input("D.N. Pai", value=data.get('pai_dn', ''), disabled=is_monitor) 

                st.markdown("**Mãe**")
                c_m1, c_m2, c_m3, c_m4 = st.columns([3, 2, 2, 2])
                data['mae_nome'] = c_m1.text_input("Nome da Mãe", value=data.get('mae_nome', ''), disabled=is_monitor)
                data['mae_prof'] = c_m2.text_input("Profissão Mãe", value=data.get('mae_prof', ''), disabled=is_monitor)
                data['mae_esc'] = c_m3.text_input("Escolaridade Mãe", value=data.get('mae_esc', ''), disabled=is_monitor)
                data['mae_dn'] = c_m4.text_input("D.N. Mãe", value=data.get('mae_dn', ''), disabled=is_monitor)

                st.divider()
                st.markdown("**Irmãos**")
                if 'irmaos' not in data: data['irmaos'] = [{'nome': '', 'idade': '', 'esc': ''} for _ in range(4)]
                
                for i in range(4):
                    c_i1, c_i2, c_i3 = st.columns([3, 1, 2])
                    data['irmaos'][i]['nome'] = c_i1.text_input(f"Nome Irmão {i+1}", value=data['irmaos'][i]['nome'], disabled=is_monitor)
                    data['irmaos'][i]['idade'] = c_i2.text_input(f"Idade {i+1}", value=data['irmaos'][i]['idade'], disabled=is_monitor)
                    data['irmaos'][i]['esc'] = c_i3.text_input(f"Escolaridade {i+1}", value=data['irmaos'][i]['esc'], disabled=is_monitor)

                data['outros_familia'] = st.text_area("Outros (Moradores da casa):", value=data.get('outros_familia', ''), disabled=is_monitor)
                data['quem_mora'] = st.text_input("Com quem mora?", value=data.get('quem_mora', ''), disabled=is_monitor)
                
                c_conv1, c_conv2 = st.columns([1, 3])
                data['convenio'] = c_conv1.radio("Possui convênio?", ["Sim", "Não"], horizontal=True, index=1 if data.get('convenio') == "Não" else 0, disabled=is_monitor)
                data['convenio_qual'] = c_conv2.text_input("Qual convênio?", value=data.get('convenio_qual', ''), disabled=is_monitor)
                
                c_soc1, c_soc2 = st.columns([1, 3])
                data['social'] = c_soc1.radio("Recebe benefício social?", ["Sim", "Não"], horizontal=True, index=1 if data.get('social') == "Não" else 0, disabled=is_monitor)
                data['social_qual'] = c_soc2.text_input("Qual benefício?", value=data.get('social_qual', ''), disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Dados Familiares"):
                        save_student("CASO", data.get('nome'), data, "Família")

        # --- ABA 3: HISTÓRICO ---
        with tabs[2]:
            with st.form("form_caso_historico") if not is_monitor else st.container():
                st.subheader("1.1.3 História Escolar")
                data['hist_idade_entrou'] = st.text_input("Idade que entrou na escola:", value=data.get('hist_idade_entrou', ''), disabled=is_monitor)
                data['hist_outra_escola'] = st.text_input("Já estudou em outra escola? Quais?", value=data.get('hist_outra_escola', ''), disabled=is_monitor)
                data['hist_motivo_transf'] = st.text_input("Motivo da transferência:", value=data.get('hist_motivo_transf', ''), disabled=is_monitor)
                data['hist_obs'] = st.text_area("Outras observações escolares:", value=data.get('hist_obs', ''), disabled=is_monitor)

                st.divider()
                st.subheader("1.2 Informações sobre Gestação")
                
                c_g1, c_g2 = st.columns(2)
                data['gest_parentesco'] = c_g1.radio("Parentesco entre pais?", ["Sim", "Não"], horizontal=True, index=1 if data.get('gest_parentesco') == "Não" else 0, disabled=is_monitor)
                data['gest_doenca'] = c_g2.text_input("Doença/trauma na gestação? Quais?", value=data.get('gest_doenca', ''), disabled=is_monitor)
                
                c_g3, c_g4 = st.columns(2)
                data['gest_substancias'] = c_g3.radio("Uso de álcool/fumo/drogas?", ["Sim", "Não"], horizontal=True, index=1 if data.get('gest_substancias') == "Não" else 0, disabled=is_monitor)
                data['gest_medicamentos'] = c_g4.text_input("Uso de medicamentos? Quais?", value=data.get('gest_medicamentos', ''), disabled=is_monitor)

                data['parto_ocorrencia'] = st.text_input("Ocorrência no parto? Quais?", value=data.get('parto_ocorrencia', ''), disabled=is_monitor)
                data['parto_incubadora'] = st.text_input("Incubadora? Motivo?", value=data.get('parto_incubadora', ''), disabled=is_monitor)
                
                c_p1, c_p2 = st.columns(2)
                data['parto_prematuro'] = c_p1.radio("Prematuro?", ["Sim", "Não"], horizontal=True, index=1 if data.get('parto_prematuro') == "Não" else 0, disabled=is_monitor)
                data['parto_uti'] = c_p2.radio("Ficou em UTI?", ["Sim", "Não"], horizontal=True, index=1 if data.get('parto_uti') == "Não" else 0, disabled=is_monitor)

                c_d1, c_d2, c_d3 = st.columns(3)
                data['dev_tempo_gest'] = c_d1.text_input("Tempo Gestação", value=data.get('dev_tempo_gest', ''), disabled=is_monitor)
                data['dev_peso'] = c_d2.text_input("Peso", value=data.get('dev_peso', ''), disabled=is_monitor)
                data['dev_normal_1ano'] = c_d3.radio("Desenv. normal 1º ano?", ["Sim", "Não"], horizontal=True, index=0 if data.get('dev_normal_1ano') == "Sim" else 1, disabled=is_monitor)
                
                data['dev_atraso'] = st.text_input("Atraso importante? Quais?", value=data.get('dev_atraso', ''), disabled=is_monitor)
                c_m1, c_m2 = st.columns(2)
                data['dev_idade_andar'] = c_m1.text_input("Idade começou a andar?", value=data.get('dev_idade_andar', ''), disabled=is_monitor)
                data['dev_idade_falar'] = c_m2.text_input("Idade começou a falar?", value=data.get('dev_idade_falar', ''), disabled=is_monitor)

                st.markdown("---")
                data['diag_possui'] = st.text_input("Possui diagnóstico? Qual?", value=data.get('diag_possui', ''), disabled=is_monitor)
                data['diag_reacao'] = st.text_input("Reação da família:", value=data.get('diag_reacao', ''), disabled=is_monitor)
                c_dx1, c_dx2 = st.columns(2)
                data['diag_data'] = c_dx1.text_input("Data do diagnóstico:", value=data.get('diag_data', ''), disabled=is_monitor)
                data['diag_origem'] = c_dx2.text_input("Origem da informação:", value=data.get('diag_origem', ''), disabled=is_monitor)
                
                c_fam1, c_fam2 = st.columns(2)
                data['fam_deficiencia'] = c_fam1.text_input("Pessoa com deficiência na família?", value=data.get('fam_deficiencia', ''), disabled=is_monitor)
                data['fam_altas_hab'] = c_fam2.radio("Pessoa com AH/SD na família?", ["Sim", "Não"], horizontal=True, index=1 if data.get('fam_altas_hab') == "Não" else 0, disabled=is_monitor)
                
                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Dados de Histórico"):
                        save_student("CASO", data.get('nome'), data, "Histórico")

        # --- ABA 4: SAÚDE ---
        with tabs[3]:
            with st.form("form_caso_saude") if not is_monitor else st.container():
                st.subheader("1.3 Informações sobre Saúde")
                data['saude_prob'] = st.text_input("Problema de saúde? Quais?", value=data.get('saude_prob', ''), disabled=is_monitor)
                data['saude_internacao'] = st.text_input("Internação? Motivos?", value=data.get('saude_internacao', ''), disabled=is_monitor)
                data['saude_restricao'] = st.text_input("Restrição/Seletividade alimentar?", value=data.get('saude_restricao', ''), disabled=is_monitor)
                
                st.markdown("**Medicamentos Controlados**")
                data['med_uso'] = st.radio("Faz uso?", ["Sim", "Não"], horizontal=True, index=1 if data.get('med_uso') == "Não" else 0, disabled=is_monitor)
                data['med_quais'] = st.text_input("Quais medicamentos?", value=data.get('med_quais', ''), disabled=is_monitor)
                c_med1, c_med2, c_med3 = st.columns(3)
                data['med_hor'] = c_med1.text_input("Horário", value=data.get('med_hor', ''), disabled=is_monitor)
                data['med_dos'] = c_med2.text_input("Dosagem", value=data.get('med_dos', ''), disabled=is_monitor)
                data['med_ini'] = c_med3.text_input("Início", value=data.get('med_ini', ''), disabled=is_monitor)

                st.divider()
                c_esf1, c_esf2 = st.columns(2)
                data['esf_urina'] = c_esf1.checkbox("Controla Urina", value=data.get('esf_urina', False), disabled=is_monitor)
                data['esf_fezes'] = c_esf2.checkbox("Controla Fezes", value=data.get('esf_fezes', False), disabled=is_monitor)
                data['esf_idade'] = st.text_input("Com qual idade controlou?", value=data.get('esf_idade', ''), disabled=is_monitor)
                data['sono'] = st.text_input("Dorme bem? Obs:", value=data.get('sono', ''), disabled=is_monitor)
                data['medico_ultimo'] = st.text_input("Última visita ao médico:", value=data.get('medico_ultimo', ''), disabled=is_monitor)

                st.markdown("**Atendimento Clínico Extraescolar**")
                clinicas_opts = ["APAE", "ARIL", "CEMA", "Família Azul", "CAPS", "Amb. Saúde Mental", "João Fischer D.A.", "João Fischer D.V."]
                prof_opts = ["Fonoaudiólogo", "Terapeuta Ocupacional", "Psicólogo", "Psicopedagogo", "Fisioterapeuta"]
                
                data['clinicas'] = st.multiselect("Selecione os atendimentos:", clinicas_opts + prof_opts, default=data.get('clinicas', []), disabled=is_monitor)
                data['clinicas_med_esp'] = st.text_input("Área médica (Especialidade):", value=data.get('clinicas_med_esp', ''), disabled=is_monitor)
                data['clinicas_nome'] = st.text_input("Nome da Clínica/Profissional:", value=data.get('clinicas_nome', ''), disabled=is_monitor)
                
                data['saude_obs_geral'] = st.text_area("Outras observações de saúde:", value=data.get('saude_obs_geral', ''), disabled=is_monitor)

                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Dados de Saúde"):
                        save_student("CASO", data.get('nome'), data, "Saúde")

# --- ABA 5: COMPORTAMENTO ---
        with tabs[4]:
            with st.form("form_caso_comportamento") if not is_monitor else st.container():
                st.subheader("1.4 Compreensão da Família (Checklist)")
                
                checklist_items = [
                    "Relata fatos do dia a dia? Apresentando boa memória?",
                    "É organizado com seus pertences?",
                    "Aceita regras de forma tranquila?",
                    "Busca e aceita ajuda quando não sabe ou não consegue algo?",
                    "Aceita alterações no ambiente?",
                    "Tem algum medo?",
                    "Tem alguma mania?",
                    "Tem alguma área/assunto, brinquedo ou hiperfoco?",
                    "Prefere brincar sozinho ou com outras crianças? Tem amigos?",
                    "Qual a expectativa da família em relação à escolaridade da criança?"
                ]
                if 'checklist' not in data: data['checklist'] = {}
                
                # Pegamos um ID único do aluno para evitar o cache do Streamlit
                aluno_id = data.get('doc_uuid', data.get('nome', 'novo_aluno'))
                
                for i, item in enumerate(checklist_items):
                    st.markdown(f"**{item}**")
                    col_a, col_b = st.columns([1, 3])
                    
                    key_base = f"itemcomport_{i}" 
                    
                    # Lemos a opção e a observação que estão no JSON (banco de dados)
                    opt_salva = data['checklist'].get(f"{key_base}_opt", "Não")
                    obs_salva = data['checklist'].get(f"{key_base}_obs", "")
                    
                    # Atualizamos o dicionário com o widget, forçando o Streamlit a ler o index/value correto
                    data['checklist'][f"{key_base}_opt"] = col_a.radio(
                        "Opção", 
                        ["Sim", "Não"], 
                        key=f"rad_{aluno_id}_{i}", # ID único impede bug visual
                        horizontal=True, 
                        label_visibility="collapsed", 
                        index=0 if opt_salva == "Sim" else 1, 
                        disabled=is_monitor
                    )
                    
                    data['checklist'][f"{key_base}_obs"] = col_b.text_input(
                        "Obs:", 
                        value=obs_salva, 
                        key=f"obs_{aluno_id}_{i}", # ID único impede bug visual
                        disabled=is_monitor
                    )
                    st.divider()

                st.subheader("Dados da Entrevista")
                c_e1, c_e2, c_e3 = st.columns(3)
                data['entrevista_prof'] = c_e1.text_input("Prof. Responsável", value=data.get('entrevista_prof', ''), disabled=is_monitor)
                data['entrevista_resp'] = c_e2.text_input("Responsável info", value=data.get('entrevista_resp', ''), disabled=is_monitor)
                
                d_ent = data.get('entrevista_data')
                if isinstance(d_ent, str): 
                     try: d_ent = datetime.strptime(d_ent, '%Y-%m-%d').date()
                     except: d_ent = date.today()
                
                input_data = c_e3.date_input("Data", value=d_ent if d_ent else date.today(), format="DD/MM/YYYY", disabled=is_monitor)
                
                # Convertendo a data para string (YYYY-MM-DD) para salvar sem erros
                data['entrevista_data'] = input_data.strftime('%Y-%m-%d') 
                
                data['entrevista_extra'] = st.text_area("Outras informações relevantes:", value=data.get('entrevista_extra', ''), disabled=is_monitor)
                
                st.markdown("---")
                if not is_monitor:
                    if st.form_submit_button("💾 Salvar Comportamento"):
                        save_student("CASO", data.get('nome'), data, "Comportamento")
                        st.success("Dados de comportamento salvos com sucesso!")
                        
        # --- ABA 6: ASSINATURAS (NOVO) ---
        with tabs[5]:
            st.subheader("Assinaturas Digitais")
            st.caption(f"Código Único do Documento: {data.get('doc_uuid', 'Não gerado ainda')}")
            
            # Roles for Caso
            required_roles = []
            if data.get('entrevista_prof'): required_roles.append({'role': 'Prof. Entrevistador', 'name': data.get('entrevista_prof')})
            if data.get('entrevista_resp'): required_roles.append({'role': 'Responsável (Família)', 'name': data.get('entrevista_resp')})
            
            # Show list of signatories
            if required_roles:
                st.markdown("##### Profissionais/Responsáveis Citados")
                for r in required_roles:
                    st.write(f"- **{r['role']}:** {r['name']}")
            else:
                st.info("Nenhum profissional identificado automaticamente.")

            st.divider()
            
            # Current Signatures
            current_signatures = data.get('signatures', [])
            if current_signatures:
                st.success("✅ Documento assinado por:")
                for sig in current_signatures:
                    st.write(f"✍️ **{sig['name']}** ({sig.get('role', 'Profissional')}) em {sig['date']}")
            else:
                st.warning("Nenhuma assinatura registrada.")

            st.divider()
            
            user_name = st.session_state.get('usuario_nome', '')
            match_role = "Profissional"
            already_signed = any(s['name'] == user_name for s in current_signatures)
            
            if already_signed:
                st.info("Você já assinou este documento.")
            else:
                if st.button("🖊️ Assinar Digitalmente", key="btn_sign_caso"):
                    # Tenta descobrir o papel
                    for r in required_roles:
                        if user_name.strip().lower() in r['name'].strip().lower():
                            match_role = r['role']
                            break
                    
                    new_sig = {
                        "name": user_name,
                        "role": match_role,
                        "date": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                        "hash": str(uuid.uuid4())
                    }
                    if 'signatures' not in data: data['signatures'] = []
                    data['signatures'].append(new_sig)
                    save_student("CASO", data.get('nome'), data, "Assinatura")
                    st.rerun()
# --- ABA 7: GERAR PDF (ESTUDO DE CASO) ---
        with tabs[6]:
            if not is_monitor:
                if st.button("💾 SALVAR ESTUDO DE CASO", type="primary"): 
                    save_student("CASO", data.get('nome', 'aluno'), data, "Completo")
            else:
                st.info("Modo Visualização.")

            if st.button("👁️ GERAR PDF"):
                # Registrar ação de gerar PDF
                log_action(data.get('nome'), "Gerou PDF", "Estudo de Caso")
                
                # --- FUNÇÕES AUXILIARES PARA TABELAS DINÂMICAS ---
                def calc_lines(pdf, text, w):
                    """Calcula quantas linhas o texto vai ocupar na largura especificada."""
                    if not text: return 1
                    lines = 0
                    for p in str(text).split('\n'):
                        words = p.split(' ')
                        line_w = 0
                        for word in words:
                            word_w = pdf.get_string_width(word + ' ')
                            if line_w + word_w > w - 2: # 2 de margem interna
                                lines += 1
                                line_w = word_w
                            else:
                                line_w += word_w
                        lines += 1
                    return max(1, lines)

                def draw_flex_row(pdf, col_data, line_h=6, font_size=9, fill_color=(240, 240, 240)):
                    """
                    Desenha uma linha garantindo que todas as colunas tenham a mesma altura.
                    col_data = [(largura, texto, estilo_fonte, alinhamento, preenchimento_bool), ...]
                    """
                    max_lines = 1
                    x_start_measure = pdf.get_x()
                    
                    # 1. Medir qual coluna precisará de mais linhas
                    for w, text, weight, align, fill in col_data:
                        pdf.set_font("Arial", weight, font_size)
                        real_w = w if w > 0 else (210 - 15 - x_start_measure)
                        lines = calc_lines(pdf, text, real_w)
                        if lines > max_lines: max_lines = lines
                        x_start_measure += real_w
                        
                    row_h = max_lines * line_h
                    
                    # 2. Quebra de página automática se a linha não couber
                    if pdf.get_y() + row_h > 275:
                        pdf.add_page()
                        
                    x_start = pdf.get_x()
                    y_start = pdf.get_y()
                    
                    # 3. Desenhar de fato a linha
                    for w, text, weight, align, fill in col_data:
                        real_w = w if w > 0 else (210 - 15 - x_start)
                        pdf.set_font("Arial", weight, font_size)
                        
                        if fill: pdf.set_fill_color(*fill_color)
                        else: pdf.set_fill_color(255, 255, 255)
                        
                        # Desenha a caixa (borda e fundo) com a altura MÁXIMA da linha
                        pdf.set_xy(x_start, y_start)
                        pdf.cell(real_w, row_h, "", border=1, fill=fill)
                        
                        # Centraliza verticalmente o texto se for apenas 1 linha em uma caixa grande
                        y_text = y_start + 1
                        if max_lines > 1 and calc_lines(pdf, text, real_w) == 1:
                            y_text = y_start + (row_h - line_h) / 2
                            
                        # Insere o texto com multi_cell (sem bordas, já que o cell atrás já fez a borda)
                        pdf.set_xy(x_start + 1, y_text)
                        pdf.multi_cell(real_w - 2, line_h, str(text), border=0, align=align)
                        
                        x_start += real_w
                        
                    # Move o cursor para o início da próxima linha
                    pdf.set_xy(15, y_start + row_h)

                # Cria PDF em Retrato ('P')
                pdf = OfficialPDF('P', 'mm', 'A4')
                pdf.add_page()
                pdf.set_margins(15, 15, 15)
                
                # SET SIGNATURE FOOTER
                pdf.set_signature_footer(data.get('signatures', []), data.get('doc_uuid', ''))
                
                # --- CABEÇALHO ---
                if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 15, 10, 25)
                if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 170, 6, 25)

                # Títulos Centralizados
                pdf.set_xy(0, 15); pdf.set_font("Arial", "B", 12)
                pdf.cell(210, 6, clean_pdf_text("PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
                pdf.cell(180, 6, clean_pdf_text("CEIEF RAFAEL AFFONSO LEITE"), 0, 1, 'C')
                pdf.ln(8)
                pdf.set_font("Arial", "B", 16); pdf.cell(0, 10, "ESTUDO DE CASO", 0, 1, 'C')
                pdf.ln(5)
                
                # --- 1.1 DADOS GERAIS ---
                pdf.section_title("1.1 DADOS GERAIS DO ESTUDANTE", width=0)
                pdf.ln(4)
                
                # 1.1.1 IDENTIFICAÇÃO
                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, "1.1.1 - IDENTIFICAÇÃO", 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [
                    (30, "Nome:", "B", "L", True),
                    (110, clean_pdf_text(data.get('nome', '')), "", "L", False),
                    (15, "D.N.:", "B", "C", True),
                    (25, clean_pdf_text(str(data.get('d_nasc', ''))), "", "C", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (30, "Escolaridade:", "B", "L", True),
                    (25, clean_pdf_text(data.get('ano_esc', '')), "", "L", False),
                    (20, "Período:", "B", "C", True),
                    (20, clean_pdf_text(data.get('periodo', '')), "", "C", False),
                    (20, "Unidade:", "B", "C", True),
                    (65, clean_pdf_text(data.get('unidade', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (30, "Endereço:", "B", "L", True),
                    (150, clean_pdf_text(data.get('endereco', '')), "", "L", False)
                ], line_h=7, font_size=10)

                draw_flex_row(pdf, [
                    (20, "Bairro:", "B", "L", True),
                    (70, clean_pdf_text(data.get('bairro', '')), "", "L", False),
                    (20, "Cidade:", "B", "C", True),
                    (70, clean_pdf_text(data.get('cidade', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (20, "Telefone:", "B", "L", True),
                    (160, clean_pdf_text(data.get('telefones', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                # 1.1.2 DADOS FAMILIARES
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, "1.1.2 - DADOS FAMILIARES", 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [
                    (20, "Pai:", "B", "L", True),
                    (80, clean_pdf_text(data.get('pai_nome', '')), "", "L", False),
                    (25, "Profissão:", "B", "C", True),
                    (55, clean_pdf_text(data.get('pai_prof', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (20, "Mãe:", "B", "L", True),
                    (80, clean_pdf_text(data.get('mae_nome', '')), "", "L", False),
                    (25, "Profissão:", "B", "C", True),
                    (55, clean_pdf_text(data.get('mae_prof', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                # Irmãos
                pdf.ln(2)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, clean_pdf_text("Irmãos (Nome | Idade | Escolaridade)"), 1, 1, 'L', 1)
                for i, irmao in enumerate(data.get('irmaos', [])):
                    if irmao['nome']:
                        txt = f"{irmao['nome']}  |  {irmao['idade']}  |  {irmao['esc']}"
                        draw_flex_row(pdf, [(180, clean_pdf_text(txt), "", "L", False)], line_h=6, font_size=9)
                
                pdf.ln(2)
                draw_flex_row(pdf, [
                    (40, "Com quem mora:", "B", "L", True),
                    (140, clean_pdf_text(data.get('quem_mora', '')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (40, "Convênio Médico:", "B", "L", True),
                    (50, clean_pdf_text(data.get('convenio')), "", "L", False),
                    (20, "Qual:", "B", "C", True),
                    (70, clean_pdf_text(data.get('convenio_qual')), "", "L", False)
                ], line_h=7, font_size=10)
                
                draw_flex_row(pdf, [
                    (40, "Benefício Social:", "B", "L", True),
                    (50, clean_pdf_text(data.get('social')), "", "L", False),
                    (20, "Qual:", "B", "C", True),
                    (70, clean_pdf_text(data.get('social_qual')), "", "L", False)
                ], line_h=7, font_size=10)

                # 1.1.3 HISTÓRIA ESCOLAR
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10); pdf.cell(0, 8, clean_pdf_text("1.1.3 - HISTÓRIA ESCOLAR"), 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [(50, "Idade entrou na escola:", "B", "L", True), (130, clean_pdf_text(data.get('hist_idade_entrou')), "", "L", False)], line_h=7, font_size=10)
                draw_flex_row(pdf, [(50, "Outras escolas:", "B", "L", True), (130, clean_pdf_text(data.get('hist_outra_escola')), "", "L", False)], line_h=7, font_size=10)
                draw_flex_row(pdf, [(50, "Motivo transferência:", "B", "L", True), (130, clean_pdf_text(data.get('hist_motivo_transf')), "", "L", False)], line_h=7, font_size=10)
                
                if data.get('hist_obs'):
                    pdf.ln(2)
                    pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "Observações Escolares:", 0, 1)
                    draw_flex_row(pdf, [(180, clean_pdf_text(data.get('hist_obs')), "", "L", False)], line_h=6, font_size=9)

                # --- 1.2 GESTAÇÃO, PARTO E DESENVOLVIMENTO ---
                pdf.add_page()
                pdf.section_title("1.2 GESTAÇÃO, PARTO E DESENVOLVIMENTO", width=0)
                pdf.ln(4)
                
                # A função print_data_row agora usa o Helper para não encavalar NADA!
                def print_data_row(label, value):
                    draw_flex_row(pdf, [
                        (80, clean_pdf_text(label), "B", "L", True),
                        (100, clean_pdf_text(str(value) if value else ""), "", "L", False)
                    ], line_h=6, font_size=9)

                rows_gest = [
                    ("Parentesco entre pais:", data.get('gest_parentesco')),
                    ("Doença/Trauma na gestação:", data.get('gest_doenca')),
                    ("Uso de substâncias (mãe):", data.get('gest_substancias')),
                    ("Uso de medicamentos (mãe):", data.get('gest_medicamentos')),
                    ("Ocorrência no parto:", data.get('parto_ocorrencia')),
                    ("Necessitou de incubadora:", data.get('parto_incubadora')),
                    ("Prematuro?", f"{data.get('parto_prematuro')}  |  UTI: {data.get('parto_uti')}"),
                    ("Tempo de gestação / Peso:", f"{data.get('dev_tempo_gest')}  /  {data.get('dev_peso')}"),
                    ("Desenvolvimento normal no 1º ano:", data.get('dev_normal_1ano')),
                    ("Apresentou atraso importante?", data.get('dev_atraso')),
                    ("Idade que andou / falou:", f"{data.get('dev_idade_andar')}  /  {data.get('dev_idade_falar')}"),
                    ("Possui diagnóstico?", data.get('diag_possui')),
                    ("Reação da família ao diagnóstico:", data.get('diag_reacao')),
                    ("Data / Origem do diagnóstico:", f"{data.get('diag_data')}  |  {data.get('diag_origem')}"),
                    ("Pessoa com deficiência na família:", data.get('fam_deficiencia')),
                    ("Pessoa com AH/SD na família:", data.get('fam_altas_hab'))
                ]
                
                for label, value in rows_gest:
                    print_data_row(label, value)

                # --- 1.3 INFORMAÇÕES SOBRE SAÚDE ---
                pdf.add_page()
                pdf.section_title("1.3 INFORMAÇÕES SOBRE SAÚDE", width=0)
                pdf.ln(4)
                
                saude_rows = [
                    ("Problemas de saúde:", data.get('saude_prob')),
                    ("Já necessitou de internação:", data.get('saude_internacao')),
                    ("Restrição/Seletividade alimentar:", data.get('saude_restricao')),
                    ("Uso de medicamentos controlados:", f"{data.get('med_uso')} - Quais: {data.get('med_quais')}"),
                    ("Horário / Dosagem / Início:", f"{data.get('med_hor')}  |  {data.get('med_dos')}  |  {data.get('med_ini')}"),
                    ("Qualidade do sono:", data.get('sono')),
                    ("Última visita ao médico:", data.get('medico_ultimo'))
                ]
                for label, value in saude_rows:
                    print_data_row(label, value)
                
                esf = []
                if data.get('esf_urina'): esf.append("Urina")
                if data.get('esf_fezes'): esf.append("Fezes")
                print_data_row("Controle de Esfíncter:", f"{', '.join(esf) if esf else 'Não'}  (Idade: {data.get('esf_idade')})")
                
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10); pdf.set_fill_color(240, 240, 240)
                pdf.cell(0, 8, "Atendimentos Clínicos Extraescolares", 1, 1, 'L', 1)
                
                clins = data.get('clinicas', [])
                print_data_row("Realiza atendimento em:", ", ".join(clins) if clins else "Não realiza")
                print_data_row("Especialidade médica:", data.get('clinicas_med_esp'))
                print_data_row("Nome da Clínica/Profissional:", data.get('clinicas_nome'))
                
                if data.get('saude_obs_geral'):
                    pdf.ln(2)
                    pdf.set_font("Arial", "B", 9); pdf.cell(0, 6, "Outras observações de saúde:", 0, 1)
                    draw_flex_row(pdf, [(180, clean_pdf_text(data.get('saude_obs_geral')), "", "L", False)], line_h=5, font_size=9)

                # --- 1.4 COMPREENSÃO DA FAMÍLIA (CHECKLIST) ---
                pdf.add_page()
                pdf.section_title("1.4 COMPREENSÃO DA FAMÍLIA (CHECKLIST)", width=0)
                pdf.ln(4)
                
                # Cabeçalho da tabela de checklist adaptado
                draw_flex_row(pdf, [
                    (110, "PERGUNTA / ASPECTO OBSERVADO", "B", "C", True),
                    (25, "SIM/NÃO", "B", "C", True),
                    (45, "OBSERVAÇÕES DA FAMÍLIA", "B", "C", True)
                ], line_h=8, font_size=9, fill_color=(220, 220, 220))
                
                checklist_items = [
                    "Relata fatos do dia a dia? Apresentando boa memória?",
                    "É organizado com seus pertences?",
                    "Aceita regras de forma tranquila?",
                    "Busca e aceita ajuda quando não sabe ou não consegue algo?",
                    "Aceita alterações no ambiente?",
                    "Tem algum medo?",
                    "Tem alguma mania?",
                    "Tem alguma área/assunto, brinquedo ou hiperfoco?",
                    "Prefere brincar sozinho ou com outras crianças? Tem amigos?",
                    "Qual a expectativa da família em relação à escolaridade da criança?"
                ]
                
                for i, item in enumerate(checklist_items):
                    key_base = f"itemcomport_{i}"
                    opt = data.get('checklist', {}).get(f"{key_base}_opt", "Não")
                    obs = data.get('checklist', {}).get(f"{key_base}_obs", "")
                    
                    # Usa o Helper e automaticamente ajusta a altura das três colunas de uma vez
                    draw_flex_row(pdf, [
                        (110, clean_pdf_text(item), "", "L", False),
                        (25, clean_pdf_text(opt), "", "C", False),
                        (45, clean_pdf_text(obs), "", "L", False)
                    ], line_h=6, font_size=9)

                # --- FINALIZAÇÃO ---
                pdf.ln(5)
                pdf.set_font("Arial", "B", 10); pdf.set_fill_color(240, 240, 240)
                pdf.cell(0, 8, clean_pdf_text("OUTRAS INFORMAÇÕES RELEVANTES"), 1, 1, 'L', 1)
                
                draw_flex_row(pdf, [(180, clean_pdf_text(data.get('entrevista_extra', '---')), "", "L", False)], line_h=6, font_size=9)
                
                pdf.ln(10)
                if pdf.get_y() > 240: pdf.add_page()
                
                pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "DADOS DA ENTREVISTA", 1, 1, 'L', 1)
                
                print_data_row("Responsável pelas informações:", data.get('entrevista_resp'))
                print_data_row("Profissional Entrevistador:", data.get('entrevista_prof'))
                print_data_row("Data da Entrevista:", str(data.get('entrevista_data', '')))
                
                pdf.ln(25) 
                
                y = pdf.get_y()
                pdf.line(20, y, 90, y); pdf.line(110, y, 190, y)
                pdf.set_font("Arial", "", 9)
                pdf.set_xy(20, y+2); pdf.cell(70, 5, "Assinatura do Responsável Legal", 0, 0, 'C')
                pdf.set_xy(110, y+2); pdf.cell(80, 5, "Assinatura do Docente/Gestor", 0, 1, 'C')

                st.session_state.pdf_bytes_caso = get_pdf_bytes(pdf)
                st.rerun()

            if 'pdf_bytes_caso' in st.session_state:
                st.download_button("📥 BAIXAR PDF ESTUDO DE CASO", st.session_state.pdf_bytes_caso, f"Caso_{data.get('nome','estudante')}.pdf", "application/pdf", type="primary")
        # --- ABA 8: HISTÓRICO ---
        with tabs[7]:
            st.subheader("Histórico de Atividades")
            st.caption("Registro de alterações, salvamentos e geração de documentos.")
            
            df_hist = safe_read("Historico", ["Data_Hora", "Aluno", "Usuario", "Acao", "Detalhes"])
            
            if not df_hist.empty and data.get('nome'):
                # Filtrar pelo aluno atual
                student_hist = df_hist[df_hist["Aluno"] == data.get('nome')]
                
                if not student_hist.empty:
                    # Ordenar por data (mais recente primeiro)
                    student_hist = student_hist.iloc[::-1]
                    st.dataframe(student_hist, use_container_width=True, hide_index=True)
                else:
                    st.info("Nenhum histórico encontrado para este aluno.")
            else:
                st.info("O histórico está vazio ou aluno não selecionado.")

    # --- PROTOCOLO DE CONDUTA ---
    elif doc_mode == "Protocolo de Conduta":
        st.markdown("""<div class="header-box"><div class="header-title">Protocolo de Conduta</div></div>""", unsafe_allow_html=True)
        st.markdown("""<style>div[data-testid="stFormSubmitButton"] > button {width: 100%; background-color: #dcfce7; color: #166534; border: 1px solid #166534;}</style>""", unsafe_allow_html=True)
        
        tabs = st.tabs(["📝 Preenchimento e Emissão", "🕒 Histórico"])
        
        data_conduta = st.session_state.data_conduta
        data_pei = st.session_state.data_pei
        
        with tabs[0]:
            with st.form("form_conduta") if not is_monitor else st.container():
                st.subheader("Configuração do Protocolo")
                st.caption("Preencha manualmente ou utilize o botão abaixo para importar informações do PEI do aluno, convertendo-as automaticamente para a 1ª pessoa.")
                
                if not is_monitor:
                    if st.form_submit_button("🔄 Preencher Automaticamente com dados do PEI"):
                        # Mapeamento e conversão simples para 1ª pessoa
                        if data_pei:
                            # Sobre Mim
                            defic = data_pei.get('defic_txt', '') or data_pei.get('neuro_txt', '')
                            data_conduta['conduta_sobre_mim'] = f"Olá, meu nome é {data_pei.get('nome', '')}. Tenho {data_pei.get('idade', '')} anos. Estou matriculado no {data_pei.get('ano_esc', '')} ano. {defic}"
                            
                            # Coisas que eu gosto
                            gostos = []
                            if data_pei.get('beh_interesses'): gostos.append(data_pei.get('beh_interesses'))
                            if data_pei.get('beh_objetos_gosta'): gostos.append(data_pei.get('beh_objetos_gosta'))
                            if data_pei.get('beh_atividades'): gostos.append(data_pei.get('beh_atividades'))
                            data_conduta['conduta_gosto'] = "\n".join(gostos)
                            
                            # Coisas que não gosto
                            nao_gosto = []
                            if data_pei.get('beh_objetos_odeia'): nao_gosto.append(data_pei.get('beh_objetos_odeia'))
                            if data_pei.get('beh_gatilhos'): nao_gosto.append(f"Fico chateado/nervoso quando: {data_pei.get('beh_gatilhos')}")
                            data_conduta['conduta_nao_gosto'] = "\n".join(nao_gosto)
                            
                            # Como me comunico
                            data_conduta['conduta_comunico'] = f"Eu me comunico: {data_pei.get('com_tipo', '')}. {data_pei.get('com_alt_espec', '')}"
                            
                            # Como me ajudar
                            ajuda = []
                            if data_pei.get('beh_crise_regula'): ajuda.append(f"Para me regular: {data_pei.get('beh_crise_regula')}")
                            if data_pei.get('beh_calmo'): ajuda.append(f"O que me acalma: {data_pei.get('beh_calmo')}")
                            data_conduta['conduta_ajuda'] = "\n".join(ajuda)
                            
                            # Habilidades
                            habs = []
                            if data_pei.get('hig_banheiro'): habs.append(f"Uso do banheiro: {data_pei.get('hig_banheiro')}")
                            if data_pei.get('hig_dentes'): habs.append(f"Escovação: {data_pei.get('hig_dentes')}")
                            if data_pei.get('dev_tarefas'): habs.append(f"Tarefas: {data_pei.get('dev_tarefas')}")
                            data_conduta['conduta_habilidades'] = "\n".join(habs)
                            
                            st.success("Dados importados do PEI com sucesso! Revise abaixo.")
                        else:
                            st.warning("Dados do PEI não encontrados para este aluno.")

                # Campos do Formulário
                c1, c2 = st.columns([3, 1])
                data_conduta['nome'] = c1.text_input("Nome", value=data_pei.get('nome', data_conduta.get('nome','')), disabled=True)
                
                d_val = data_conduta.get('nasc') or data_pei.get('nasc')
                if isinstance(d_val, str): 
                    try: d_val = datetime.strptime(d_val, '%Y-%m-%d').date()
                    except: d_val = date.today()
                data_conduta['nasc'] = c2.date_input("Nascimento", value=d_val if d_val else date.today(), format="DD/MM/YYYY", disabled=is_monitor)
                
                data_conduta['ano_esc'] = st.text_input("Ano de Escolaridade", value=data_pei.get('ano_esc', data_conduta.get('ano_esc','')), disabled=is_monitor)
                
                st.divider()
                
                c_g, c_s = st.columns(2)
                data_conduta['conduta_gosto'] = c_g.text_area("Coisas que eu gosto (Laranja)", value=data_conduta.get('conduta_gosto', ''), height=150, disabled=is_monitor)
                data_conduta['conduta_sobre_mim'] = c_s.text_area("Sobre mim (Verde)", value=data_conduta.get('conduta_sobre_mim', ''), height=150, disabled=is_monitor)
                
                c_ng, c_com = st.columns(2)
                data_conduta['conduta_nao_gosto'] = c_ng.text_area("Coisas que eu não gosto (Vermelho)", value=data_conduta.get('conduta_nao_gosto', ''), height=150, disabled=is_monitor)
                data_conduta['conduta_comunico'] = c_com.text_area("Como me comunico (Roxo)", value=data_conduta.get('conduta_comunico', ''), height=150, disabled=is_monitor)
                
                c_aj, c_hab = st.columns(2)
                data_conduta['conduta_ajuda'] = c_aj.text_area("Como me ajudar (Azul)", value=data_conduta.get('conduta_ajuda', ''), height=150, disabled=is_monitor)
                data_conduta['conduta_habilidades'] = c_hab.text_area("Habilidades / Eu posso (Amarelo)", value=data_conduta.get('conduta_habilidades', ''), height=150, disabled=is_monitor)

                st.markdown("---")
                c_save, c_pdf = st.columns(2)
                
                if not is_monitor:
                    if c_save.form_submit_button("💾 Salvar Protocolo"):
                        save_student("CONDUTA", data_conduta.get('nome', 'aluno'), data_conduta, "Protocolo")
                
                # Check button type depending on context (Form vs Container)
                gen_pdf = False
                if is_monitor:
                    if c_pdf.button("👁️ Gerar PDF"): gen_pdf = True
                else:
                    if c_pdf.form_submit_button("👁️ Gerar PDF"): gen_pdf = True

                if gen_pdf:
                    log_action(data_conduta.get('nome'), "Gerou PDF", "Protocolo de Conduta")
                    
                    pdf = OfficialPDF('P', 'mm', 'A4')
                    pdf.add_page(); pdf.set_margins(10, 10, 10)
                    
                    # SET SIGNATURE FOOTER
                   # pdf.set_signature_footer(data.get('signatures', []), data.get('doc_uuid', ''))
                    
                    # --- CABEÇALHO ---
                    if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 10, 8, 20)
                    pdf.set_xy(35, 10); pdf.set_font("Arial", "", 12)
                    pdf.cell(0, 6, clean_pdf_text("Secretaria Municipal de"), 0, 1)
                    pdf.set_x(35); pdf.set_font("Arial", "B", 16)
                    pdf.cell(0, 8, clean_pdf_text("EDUCAÇÃO"), 0, 1)
                    
                    # Box Titulo
                    pdf.set_xy(130, 8)
                    pdf.set_font("Arial", "", 12)
                    pdf.cell(70, 10, "Protocolo de conduta", 1, 1, 'C')
                    
                    # --- IDENTIFICAÇÃO (FOTO E DADOS) ---
                    start_y = 35
                    
                    # FOTO (Placeholder circular visual - quadrado com label por simplicidade do FPDF)
                    pdf.set_xy(10, start_y)
                    # Tenta carregar foto do PEI se não tiver no conduta (usa mesma ref)
                    foto_b64 = data_pei.get('foto_base64')
                    if foto_b64:
                        try:
                            img_data = base64.b64decode(foto_b64)
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                                tmp_file.write(img_data)
                                tmp_path = tmp_file.name
                            pdf.image(tmp_path, 15, start_y, 40, 50) # Imagem retangular
                            os.unlink(tmp_path)
                        except:
                            pdf.rect(15, start_y, 40, 50)
                            pdf.set_xy(15, start_y+20); pdf.set_font("Arial", "", 8); pdf.cell(40, 5, "ERRO FOTO", 0, 0, 'C')
                    else:
                        pdf.rect(15, start_y, 40, 50) # Moldura
                        pdf.set_xy(15, start_y+20); pdf.set_font("Arial", "", 8); pdf.cell(40, 5, "FOTO DO ESTUDANTE", 0, 0, 'C')
                    
                    # Campos ao lado da foto
                    pdf.set_font("Arial", "", 10)
                    
                    # Nome (Borda Vermelha)
                    pdf.set_draw_color(255, 69, 0) # Red
                    pdf.set_line_width(0.8)
                    pdf.set_xy(70, start_y)
                    pdf.cell(130, 8, clean_pdf_text(f"Meu nome: {data_conduta.get('nome','')}"), 1, 1, 'L')
                    
                    # Data Nasc (Borda Azul)
                    pdf.set_draw_color(0, 191, 255) # Cyan/Blue
                    pdf.set_xy(70, start_y + 12)
                    pdf.multi_cell(40, 6, clean_pdf_text(f"Data de\nNascimento:\n{str(data_conduta.get('nasc',''))}"), 1, 'C')
                    
                    # Ano escolar (Borda Rosa)
                    pdf.set_draw_color(255, 105, 180) # Pink
                    pdf.set_xy(115, start_y + 12)
                    pdf.multi_cell(50, 9, clean_pdf_text(f"Ano de escolaridade:\n{data_conduta.get('ano_esc','')}") , 1, 'C')
                    

        # --- CONFIGURAÇÃO PARA CARTAZ DE PÁGINA ÚNICA ---
                    # Desliga a quebra de página automática para termos controle total do espaço
                    pdf.set_auto_page_break(False)
                    
                    # --- CAIXAS DE CONTEÚDO ---
                    
                    def draw_colored_box(x, y, w, target_h, r, g, b, title, content):
                        texto_limpo = clean_pdf_text(str(content) if content else "")
                        pdf.set_font("Arial", "", 9)
                        
                        w_text = w - 4 
                        line_height = 5
                        linhas = 0
                        
                        for paragrafo in texto_limpo.split('\n'):
                            largura = pdf.get_string_width(paragrafo)
                            if largura == 0:
                                linhas += 1
                            else:
                                linhas += int(largura / w_text) + 1
                                
                        h_texto = 8 + (max(1, linhas) * line_height) + 4
                        h_final = max(target_h, h_texto)
                        
                        # Removemos o add_page() manual daqui para forçar a ficar na mesma página
                            
                        # Desenha o Retângulo Externo
                        pdf.set_draw_color(r, g, b)
                        pdf.set_line_width(0.8)
                        pdf.rect(x, y, w, h_final)
                        
                        # Imprime o Título
                        pdf.set_xy(x, y+2)
                        pdf.set_text_color(0, 0, 0)
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(w, 5, clean_pdf_text(title), 0, 1, 'C')
                        
                        # Imprime o Conteúdo
                        pdf.set_xy(x+2, y+8)
                        pdf.set_font("Arial", "", 9)
                        pdf.multi_cell(w_text, line_height, texto_limpo, 0, 'L')
                        
                        # Retorna a posição da próxima caixa (com 3mm de respiro para alinhamento perfeito)
                        return y + h_final + 3

                    # --- LÓGICA DE ORGANIZAÇÃO PARA PREENCHER O A4 ---
                    
                    y_esquerdo = 90 
                    y_direito = 75 
                    
                    # LADO DIREITO (Altura ideal reduzida para 65mm para não bater na margem limite)
                    y_direito = draw_colored_box(100, y_direito, 100, 65, 154, 205, 50, "Sobre mim", data_conduta.get('conduta_sobre_mim', ''))
                    y_direito = draw_colored_box(130, y_direito, 70, 65, 255, 69, 0, "Coisas que eu não gosto", data_conduta.get('conduta_nao_gosto', ''))
                    y_direito = draw_colored_box(130, y_direito, 70, 65, 255, 215, 0, "Habilidades (eu posso...)", data_conduta.get('conduta_habilidades', ''))
                    
                    # LADO ESQUERDO (Altura ideal reduzida para 60mm)
                    y_esquerdo = draw_colored_box(10, y_esquerdo, 85, 60, 255, 165, 0, "Coisas que eu gosto", data_conduta.get('conduta_gosto', ''))
                    y_esquerdo = draw_colored_box(10, y_esquerdo, 110, 60, 147, 112, 219, "Como me comunico", data_conduta.get('conduta_comunico', ''))
                    y_esquerdo = draw_colored_box(10, y_esquerdo, 110, 60, 0, 191, 255, "Como me ajudar", data_conduta.get('conduta_ajuda', ''))

                    # Religando a quebra de página por precaução para não afetar o resto do app
                    pdf.set_auto_page_break(True, margin=15)

                    st.session_state.pdf_bytes_conduta = get_pdf_bytes(pdf)
                    st.rerun()

            if 'pdf_bytes_conduta' in st.session_state:
                st.download_button("📥 BAIXAR PROTOCOLO PDF", st.session_state.pdf_bytes_conduta, f"Conduta_{data_conduta.get('nome','aluno')}.pdf", "application/pdf", type="primary")
        # --- ABA 7: HISTÓRICO ---
        with tabs[1]:
            st.subheader("Histórico de Atividades")
            st.caption("Registro de alterações, salvamentos e geração de documentos.")
            
            df_hist = safe_read("Historico", ["Data_Hora", "Aluno", "Usuario", "Acao", "Detalhes"])
            
            # CORREÇÃO DE BUG: Usar data_conduta ao invés de data
            if not df_hist.empty and data_conduta.get('nome'):
                # Filtrar pelo aluno atual
                student_hist = df_hist[df_hist["Aluno"] == data_conduta.get('nome')]
                
                if not student_hist.empty:
                    # Ordenar por data (mais recente primeiro)
                    student_hist = student_hist.iloc[::-1]
                    st.dataframe(student_hist, use_container_width=True, hide_index=True)
                else:
                    st.info("Nenhum histórico encontrado para este aluno.")
            else:
                st.info("O histórico está vazio ou aluno não selecionado.")


    # --- AVALIAÇÃO PEDAGÓGICA ---
    elif doc_mode == "Avaliação de Apoio":
        st.markdown("""<div class="header-box"><div class="header-title">Avaliação Pedagógica: Apoio Escolar</div></div>""", unsafe_allow_html=True)
        st.markdown("""<style>div[data-testid="stFormSubmitButton"] > button {width: 100%; background-color: #dcfce7; color: #166534; border: 1px solid #166534;}</style>""", unsafe_allow_html=True)
        
        tabs = st.tabs(["📝 Preenchimento e Emissão", "🕒 Histórico"])
        
        # Inicialização de variáveis de estado se não existirem
        if 'data_avaliacao' not in st.session_state: st.session_state.data_avaliacao = {}
        if 'data_pei' not in st.session_state: st.session_state.data_pei = {}
        if 'data_case' not in st.session_state: st.session_state.data_case = {}
        
        data_aval = st.session_state.data_avaliacao
        data_pei = st.session_state.data_pei
        data_caso = st.session_state.data_case
        
        # --- DEFINIÇÃO DAS LISTAS DE OPÇÕES (GLOBAL PARA O CONTEXTO) ---
        defs_opts = ["Deficiência auditiva/surdez", "Deficiência física", "Deficiência intelectual", "Deficiência múltipla", "Deficiência visual", "Transtorno do Espectro Autista", "Síndrome de Down"]
        
        opts_alim = ["É independente.", "Necessita de apoio parcial.", "Necessita de apoio total."]
        opts_hig = ["É independente.", "Usa fralda.", "Necessita de apoio parcial.", "Necessita de apoio total."]
        opts_loc = ["é independente.", "cai ou tropeça com frequência.", "faz uso de cadeira de rodas de forma independente", "faz uso de cadeira de rodas, necessitando ser conduzido.", "possui prótese/órtese.", "faz uso de andador.", "faz uso de bengala."]
        
        opts_comp = [
            "Demonstra comportamento adequado em relação às situações escolares cotidianas (sala de aula, refeitório, quadra etc).",
            "Apresenta alguns comportamentos inadequados (choro, recusa verbal, se jogar no chão) em momentos específicos , mas a recuperação é rápida.",
            "diariamente apresenta comportamentos inadequados que envolvem choro, recusa verbal, birras, saídas sem autorização, correr incontido não atendimento às solicitações dos docentes e funcionários.",
            "Frequentemente a criança emite comportamento inadequado severo que é perigoso a si própria ou outras pessoas (ex: agressões, autolesivos)."
        ]
        
        opts_part = [
            "participa de atividades em grupo da rotina escolar, interagindo com os estudantes",
            "é capaz de participar de atividades em grupo somente em momentos de curta duração",
            "não é capaz de participar de atividades em grupo de forma autônoma, dependendo de apoio para essa interação",
            "Mesmo com apoio, não é capaz de participar de atividades em grupo."
        ]
        
        opts_int = ["Adequada com as crianças e adultos.", "Satisfatória.", "Inadequada.", "Outros"]
        
        opts_rot = [
            "Compreende e atende as orientações oferecidas pelo docente de forma autônoma",
            "Precisa de intervenções pontuais do docente para compreender e atender as orientações.",
            "Mesmo com apoio apresenta severas dificuldades quanto à compreensão para atendimento de solicitações."
        ]
        
        opts_ativ = [
            "não há necessidade de flexibilização curricular",
            "precisa de flexibilização curricular em relação à metodologia de ensino, mantendo-se os conteúdos previstos para o ano de escolaridade",
            "precisa de flexibilização curricular em relação à metodologia de ensino e ao conteúdo curricular, adequando às potencialidades do estudantes",
            "há a necessidade de um currículo funcional, envolvendo as atividades de vida prática e diária."
        ]
        
        opts_at_sust = [
            "Mantém atenção por longo período de tempo.",
            "Mantém atenção por longo período de tempo com apoio.",
            "Não mantém atenção por longo período de tempo."
        ]
        
        opts_at_div = [
            "Mantém atenção em dois estímulos diferentes.",
            "Mantém atenção em dois estímulos diferentes em algumas situações.",
            "Não mantém atenção em dois estímulos differentes."
        ]
        
        opts_at_sel = [
            "Mantém atenção na tarefa ignorando estímulos externos.",
            "Mantém atenção na tarefa ignorando estímulos externos com apoio.",
            "Não mantém atenção na tarefa com a presença de outros"
        ]
        
        opts_ling = [
            "Faz uso de palavras para se comunicar, expressando seus pensamentos e desejos.",
            "Faz uso de palavras para se comunicar, apresentando trocas fonéticas orais.",
            "Utiliza palavras e frases desconexas, não conseguindo se expressar.",
            "Não faz uso de palavras para se comunicar, expressando seus desejos por meio de gestos e comportamentos",
            "Não faz uso de palavras e de gestos para se comunicar."
        ]

        with tabs[0]:
            with st.form("form_avaliacao") if not is_monitor else st.container():
                st.subheader("Configuração da Avaliação")
                st.caption("Utilize o botão abaixo para importar informações já preenchidas no PEI e Estudo de Caso.")
                
                if not is_monitor:
                    if st.form_submit_button("🔄 Preencher Automaticamente"):
                        if data_pei or data_caso:
                            data_aval['nome'] = data_pei.get('nome') or data_caso.get('nome', '')
                            data_aval['nasc'] = data_pei.get('nasc') or data_caso.get('d_nasc', '')
                            data_aval['ano_esc'] = data_pei.get('ano_esc') or data_caso.get('ano_esc', '')
                            
                            # --- CORREÇÃO 1: FILTRAGEM INTELIGENTE DE DIAGNÓSTICO ---
                            diag_tipo_pei = data_pei.get('diag_tipo', [])
                            # Mantém apenas os que batem exatamente com as opções da avaliação
                            data_aval['defic_chk'] = [d for d in diag_tipo_pei if d in defs_opts]
                            
                            # Puxa os textos descritivos do PEI para o campo "Outra"
                            descricoes_outras = []
                            if "Deficiência" in diag_tipo_pei and data_pei.get('defic_txt'):
                                descricoes_outras.append(data_pei.get('defic_txt'))
                            if "Transtorno do Neurodesenvolvimento" in diag_tipo_pei and data_pei.get('neuro_txt'):
                                descricoes_outras.append(data_pei.get('neuro_txt'))
                            if "Transtornos Aprendizagem" in diag_tipo_pei and data_pei.get('aprend_txt'):
                                descricoes_outras.append(data_pei.get('aprend_txt'))
                            
                            if descricoes_outras:
                                data_aval['defic_outra'] = " / ".join(descricoes_outras)
                            # --------------------------------------------------------
                            
                            aspectos = []
                            if data_pei.get('prof_poli'): aspectos.append(f"Polivalente: {data_pei.get('prof_poli')}")
                            if data_pei.get('prof_aee'): aspectos.append(f"AEE: {data_pei.get('prof_aee')}")
                            if data_pei.get('flex_matrix'): aspectos.append("Possui flexibilização curricular registrada no PEI.")
                            data_aval['aspectos_gerais'] = "\n".join(aspectos)
                            
                            if data_pei.get('beh_autonomia_agua') == 'Sim': data_aval['alim_nivel'] = opts_alim[0]
                            if data_pei.get('hig_banheiro') == 'Sim': data_aval['hig_nivel'] = opts_hig[0]
                            if data_pei.get('loc_reduzida') == 'Não': data_aval['loc_nivel'] = [opts_loc[0]]
                            
                            st.success("Dados importados com sucesso!")
                        else:
                            st.warning("Sem dados prévios para importar.")

                # --- CAMPOS DO FORMULÁRIO ---
                st.markdown("### Identificação")
                c_nom, c_ano = st.columns([3, 1])
                data_aval['nome'] = c_nom.text_input("Estudante", value=data_aval.get('nome', ''), disabled=True)
                data_aval['ano_esc'] = c_ano.text_input("Ano Escolaridade", value=data_aval.get('ano_esc', ''), disabled=is_monitor)
                
                st.markdown("**Deficiências (Marque as opções):**")
                
                # --- CORREÇÃO 2: BLINDAGEM DO WIDGET MULTISELECT ---
                valores_salvos = data_aval.get('defic_chk', [])
                if not isinstance(valores_salvos, list): 
                    valores_salvos = []
                # Garante que os defaults passados ao Streamlit existam em defs_opts
                valores_validos = [v for v in valores_salvos if v in defs_opts]
                
                data_aval['defic_chk'] = st.multiselect("Selecione:", defs_opts, default=valores_validos, disabled=is_monitor)
                # ---------------------------------------------------
                
                data_aval['defic_outra'] = st.text_input("Outra:", value=data_aval.get('defic_outra', ''), disabled=is_monitor)
                
                st.markdown("---")
                st.markdown("### Aspectos Gerais da Vida Escolar")
                
                # --- CAMPO INSERIDO PARA DIGITAÇÃO MANUAL ---
                data_aval['aspectos_gerais'] = st.text_area(
                    "Descrição geral da vida escolar do estudante:", 
                    value=data_aval.get('aspectos_gerais', ''), 
                    height=100,
                    disabled=is_monitor
                )
                
                with st.expander("Parte I - Habilidades de Vida Diária", expanded=True):
                    c_a, c_h = st.columns(2)
                    with c_a:
                        st.markdown("**1. Alimentação**")
                        idx_alim = opts_alim.index(data_aval.get('alim_nivel')) if data_aval.get('alim_nivel') in opts_alim else 0
                        data_aval['alim_nivel'] = st.radio("Nível Alimentação", opts_alim, index=idx_alim, key="rad_alim", disabled=is_monitor)
                        data_aval['alim_obs'] = st.text_input("Obs Alimentação:", value=data_aval.get('alim_obs', ''), disabled=is_monitor)
                    
                    with c_h:
                        st.markdown("**2. Higiene**")
                        idx_hig = opts_hig.index(data_aval.get('hig_nivel')) if data_aval.get('hig_nivel') in opts_hig else 0
                        data_aval['hig_nivel'] = st.radio("Nível Higiene", opts_hig, index=idx_hig, key="rad_hig", disabled=is_monitor)
                        data_aval['hig_obs'] = st.text_input("Obs Higiene:", value=data_aval.get('hig_obs', ''), disabled=is_monitor)
                    
                    st.markdown("**3. Locomoção (Selecione todos que se aplicam)**")
                    data_aval['loc_nivel'] = st.multiselect("Itens:", opts_loc, default=data_aval.get('loc_nivel', []), disabled=is_monitor)
                    data_aval['loc_obs'] = st.text_input("Obs Locomoção:", value=data_aval.get('loc_obs', ''), disabled=is_monitor)

                with st.expander("Parte II - Habilidades Sociais e de Interação"):
                    st.markdown("**4. Comportamento**")
                    idx_comp = opts_comp.index(data_aval.get('comportamento')) if data_aval.get('comportamento') in opts_comp else 0
                    data_aval['comportamento'] = st.radio("Nível Comportamento", opts_comp, index=idx_comp, disabled=is_monitor)
                    data_aval['comp_obs'] = st.text_input("Obs Comportamento:", value=data_aval.get('comp_obs', ''), disabled=is_monitor)
                    
                    st.divider()
                    st.markdown("**5. Participação em Grupo**")
                    idx_part = opts_part.index(data_aval.get('part_grupo')) if data_aval.get('part_grupo') in opts_part else 0
                    data_aval['part_grupo'] = st.radio("Nível Participação", opts_part, index=idx_part, disabled=is_monitor)
                    data_aval['part_obs'] = st.text_input("Obs Participação:", value=data_aval.get('part_obs', ''), disabled=is_monitor)
                    
                    st.divider()
                    st.markdown("**6. Interação**")
                    idx_int = opts_int.index(data_aval.get('interacao')) if data_aval.get('interacao') in opts_int else 0
                    data_aval['interacao'] = st.radio("Nível Interação", opts_int, index=idx_int, disabled=is_monitor)
                    if data_aval['interacao'] == "Outros":
                        data_aval['interacao_outros'] = st.text_input("Especifique (Interação):", value=data_aval.get('interacao_outros', ''), disabled=is_monitor)

                with st.expander("Parte III - Habilidades Pedagógicas"):
                    st.markdown("**7. Rotina Sala de Aula**")
                    idx_rot = opts_rot.index(data_aval.get('rotina')) if data_aval.get('rotina') in opts_rot else 0
                    data_aval['rotina'] = st.radio("Nível Rotina", opts_rot, index=idx_rot, disabled=is_monitor)
                    data_aval['rotina_obs'] = st.text_input("Obs Rotina:", value=data_aval.get('rotina_obs', ''), disabled=is_monitor)
                    
                    st.divider()
                    st.markdown("**8. Atividades Pedagógicas**")
                    idx_ativ = opts_ativ.index(data_aval.get('ativ_pedag')) if data_aval.get('ativ_pedag') in opts_ativ else 0
                    data_aval['ativ_pedag'] = st.radio("Nível Atividades", opts_ativ, index=idx_ativ, disabled=is_monitor)

                with st.expander("Parte IV - Habilidades de Comunicação e Atenção"):
                    c_com1, c_com2 = st.columns(2)
                    with c_com1:
                        st.markdown("**9. Atenção Sustentada**")
                        idx_as = opts_at_sust.index(data_aval.get('atencao_sust')) if data_aval.get('atencao_sust') in opts_at_sust else 0
                        data_aval['atencao_sust'] = st.radio("Sustentada", opts_at_sust, index=idx_as, key="at_sust", disabled=is_monitor)
                        
                        st.markdown("**11. Atenção Seletiva**")
                        idx_asel = opts_at_sel.index(data_aval.get('atencao_sel')) if data_aval.get('atencao_sel') in opts_at_sel else 0
                        data_aval['atencao_sel'] = st.radio("Seletiva", opts_at_sel, index=idx_asel, key="at_sel", disabled=is_monitor)
                    
                    with c_com2:
                        st.markdown("**10. Atenção Dividida**")
                        idx_ad = opts_at_div.index(data_aval.get('atencao_div')) if data_aval.get('atencao_div') in opts_at_div else 0
                        data_aval['atencao_div'] = st.radio("Dividida", opts_at_div, index=idx_ad, key="at_div", disabled=is_monitor)
                    
                    st.divider()
                    st.markdown("**12. Linguagem (Marque todas que se aplicam)**")
                    data_aval['linguagem'] = st.multiselect("Linguagem:", opts_ling, default=data_aval.get('linguagem', []), disabled=is_monitor)
                    data_aval['ling_obs'] = st.text_input("Obs Linguagem:", value=data_aval.get('ling_obs', ''), disabled=is_monitor)

                st.markdown("### Conclusão e Responsáveis")
                
                # 1. Defina a lista de opções para garantir que o texto seja idêntico ao do banco
                opcoes_apoio = ["Não necessita de apoio", "Nível 1", "Nível 2", "Nível 3"]
                
                # 2. Busque o valor salvo. Se não houver nada, o padrão será "Não necessita de apoio"
                valor_salvo = data_aval.get('conclusao_nivel', "Não necessita de apoio")
                
                # 3. Descubra o índice (posição) desse valor na lista
                # Se por algum motivo o valor do banco não estiver na lista, ele volta para o 0
                try:
                    indice_salvo = opcoes_apoio.index(valor_salvo)
                except ValueError:
                    indice_salvo = 0
                
                # 4. Agora use o indice_salvo no componente
                data_aval['conclusao_nivel'] = st.selectbox(
                    "Nível de Apoio Concluído", 
                    opcoes_apoio, 
                    index=indice_salvo, 
                    disabled=is_monitor
                )
                
                data_aval['apoio_existente'] = st.text_input(
                    "Se este apoio já é oferecido, explicitar aqui:", 
                    value=data_aval.get('apoio_existente', ''), 
                    disabled=is_monitor
                )
                
                c_resp1, c_resp2 = st.columns(2)
                data_aval['resp_sala'] = c_resp1.text_input("Prof. Sala Regular", value=data_aval.get('resp_sala', ''), disabled=is_monitor)
                data_aval['resp_arte'] = c_resp2.text_input("Prof. Arte", value=data_aval.get('resp_arte', ''), disabled=is_monitor)
                data_aval['resp_ef'] = c_resp1.text_input("Prof. Ed. Física", value=data_aval.get('resp_ef', ''), disabled=is_monitor)
                data_aval['resp_ee'] = c_resp2.text_input("Prof. Ed. Especial", value=data_aval.get('resp_ee', ''), disabled=is_monitor)
                data_aval['resp_dir'] = c_resp1.text_input("Direção Escolar", value=data_aval.get('resp_dir', ''), disabled=is_monitor)
                data_aval['resp_coord'] = c_resp2.text_input("Coordenação", value=data_aval.get('resp_coord', ''), disabled=is_monitor)
                
                data_aval['data_emissao'] = st.date_input("Data Emissão", value=date.today(), format="DD/MM/YYYY", disabled=is_monitor)

                st.markdown("---")
                c_sv, c_pd = st.columns(2)
                if not is_monitor:
                    if c_sv.form_submit_button("💾 Salvar Avaliação"):
                        save_student("AVALIACAO", data_aval.get('nome', 'aluno'), data_aval, "Avaliação")
                
                gen_pdf_aval = False
                if is_monitor:
                    if c_pd.button("👁️ Gerar PDF Avaliação"): gen_pdf_aval = True
                else:
                    if c_pd.form_submit_button("👁️ Gerar PDF Avaliação"): gen_pdf_aval = True

                if gen_pdf_aval:
                    def check_page_break(pdf, required_height=30):
                        # A folha tem 297mm. Deixamos uma margem inferior segura (ex: até o Y=275)
                        if pdf.get_y() + required_height > 275:
                            pdf.add_page()
                    # --- PDF GENERATION EXPERT MODE ---
                    pdf = OfficialPDF('P', 'mm', 'A4')
                    pdf.add_page(); pdf.set_margins(15, 15, 15)
                    
                    # SET SIGNATURE FOOTER
                    #pdf.set_signature_footer(data.get('signatures', []), data.get('doc_uuid', ''))
                    
                    # 1. HEADER (FIXED CEIEF RAFAEL AFFONSO LEITE)
                    if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 15, 10, 25)
                    if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 170, 6, 25)

                    pdf.set_xy(0, 15); pdf.set_font("Arial", "B", 12)
                    pdf.cell(210, 6, clean_pdf_text("PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
                    pdf.cell(180, 6, clean_pdf_text("CEIEF RAFAEL AFFONSO LEITE"), 0, 1, 'C')
                    pdf.ln(8)
                    pdf.set_font("Arial", "B", 12); pdf.cell(0, 10, clean_pdf_text("AVALIAÇÃO PEDAGÓGICA: APOIO ESCOLAR PARA ESTUDANTE COM DEFICIÊNCIA"), 0, 1, 'C')
                    pdf.ln(5)
                    
                    # 2. IDENTIFICATION
                    pdf.set_font("Arial", "B", 10); pdf.cell(20, 6, "Estudante:", 0, 0)
                    pdf.set_font("Arial", "", 10); pdf.cell(100, 6, clean_pdf_text(data_aval.get('nome', '')), "B", 0)
                    pdf.set_font("Arial", "B", 10); pdf.cell(35, 6, "Ano escolaridade:", 0, 0)
                    pdf.set_font("Arial", "", 10); pdf.cell(0, 6, clean_pdf_text(data_aval.get('ano_esc', '')), "B", 1)
                    pdf.ln(4)
                    
                    # 3. DEFICIENCIES
                    pdf.set_font("Arial", "", 9)
                    selected_defs = data_aval.get('defic_chk', [])
                    
                    def draw_check_option_simple(pdf, text, checked):
                        pdf.set_x(15) 
                        x, y = pdf.get_x(), pdf.get_y()
                        pdf.set_draw_color(0,0,0)
                        pdf.rect(x, y + 1, 3, 3)
                        if checked:
                            pdf.line(x, y + 1, x + 3, y + 4)
                            pdf.line(x, y + 4, x + 3, y + 1)
                        pdf.set_xy(x + 5, y)
                        # Width 175 ensures it ends at 15+5+175 = 195 (Right margin boundary)
                        pdf.multi_cell(175, 5, clean_pdf_text(text), 0, 'L')

                    if selected_defs:
                        for d in selected_defs:
                            draw_check_option_simple(pdf, d, True)
                        if data_aval.get('defic_outra'):
                            draw_check_option_simple(pdf, f"Outra: {data_aval.get('defic_outra')}", True)
                    else:
                        pdf.cell(0, 5, clean_pdf_text("Nenhuma deficiência selecionada."), 0, 1)
                    pdf.ln(3)
                    
                    # 4. LEGAL TEXT (INTEGRAL) - Justified
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(0, 6, clean_pdf_text("PRESSUPOSTOS LEGAIS:"), 0, 1, 'L')
                    pdf.set_font("Arial", "", 10)
                    
                    # Full width (0) uses 180mm. 
                    pdf.multi_cell(0, 4, clean_pdf_text("1- Lei nº 12.764/2012, em seu artigo 3º que trata dos direitos da pessoa com transtorno do espectro autista indica:"), 0, 'J')
                    
                    pdf.set_x(25)
                    # Indent 25 (10 more than margin). Max width to right margin (195): 195 - 25 = 170.
                    pdf.multi_cell(170, 4, clean_pdf_text("Parágrafo único. Em casos de comprovada necessidade, a pessoa com transtorno do espectro autista incluída nas classes comuns de ensino regular, nos termos do inciso IV do art. 2º, terá direito a acompanhante especializado."), 0, 'J')
                    pdf.ln(2)

                    pdf.multi_cell(0, 4, clean_pdf_text("2- Lei Brasileira de Inclusão da Pessoa com Deficiência (LBI) no art. 3º, inciso XIII, descreve as ações referentes ao apoio:"), 0, 'J')
                    
                    pdf.set_x(25)
                    pdf.multi_cell(170, 4, clean_pdf_text("XIII - profissional de apoio escolar: pessoa que exerce atividades de alimentação, higiene e locomoção do estudante com deficiência e atua em todas as atividades escolares nas quais se fizer necessária, em todos os níveis e modalidades de ensino, em instituições públicas e privadas, excluídas as técnicas ou os procedimentos identificados com profissões legalmente estabelecidas;"), 0, 'J')
                    pdf.ln(2)

                    pdf.multi_cell(0, 4, clean_pdf_text("3- CNE/CEB nº 02/01, do Conselho Nacional de Educação, que Instituiu as Diretrizes Nacionais para a Educação Especial na Educação Básica, cujo artigo 6º assim dispõe:"), 0, 'J')
                    
                    pdf.set_x(25)
                    pdf.multi_cell(170, 4, clean_pdf_text("Art. 6º Para a identificação das necessidades educacionais especiais dos alunos e a tomada de decisões quanto ao atendimento necessário, a escola deve realizar, com assessoramento técnico, avaliação do aluno no processo de ensino e aprendizagem, contando, para tal, com:"), 0, 'J')
                    
                    pdf.set_x(35)
                    # Indent 35. Max width to right margin (195): 195 - 35 = 160.
                    pdf.multi_cell(160, 4, clean_pdf_text("I - a experiência de seu corpo docente, seus diretores, coordenadores, orientadores e supervisores educacionais;\nII - o setor responsável pela educação especial do respectivo sistema;\nIII - a colaboração da família e a cooperação dos serviços de Saúde, Assistência Social, Trabalho, Justiça e Esporte, bem como do Ministério Público, quando necessário."), 0, 'J')
                    pdf.ln(4)

                    # 5. GENERAL ASPECTS
                    pdf.set_fill_color(240, 240, 240)
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(0, 7, clean_pdf_text("ASPECTOS GERAIS DA VIDA ESCOLAR DO ESTUDANTE"), 1, 1, 'L', True)
                    pdf.set_font("Arial", "", 10); pdf.set_fill_color(255, 255, 255)
                    text_general = data_aval.get('aspectos_gerais') if data_aval.get('aspectos_gerais') else " "
                    # Use 0 for auto width (margin to margin), Justified 'J'
                    pdf.multi_cell(0, 5, clean_pdf_text(text_general), 1, 'J')
                    pdf.ln(5)

                    def print_section_header_fix(pdf, title):
                        pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", "B", 10)
                        pdf.cell(0, 8, clean_pdf_text(title), 1, 1, 'L', True)
                        pdf.ln(1)

                    def print_question_options_fix(pdf, question_title, options, selected_value, obs=None):
                        # 1. Calcula o espaço que a pergunta inteira vai ocupar
                        # ~8mm pro título + 5mm por opção + 15mm se tiver observação
                        espaco_necessario = 8 + (len(options) * 5) + (15 if obs else 0)
                        
                        # 2. Verifica se a pergunta cabe inteira na página atual
                        check_page_break(pdf, espaco_necessario)
                        
                        pdf.set_x(15)
                        pdf.set_font("Arial", "B", 10)
                        pdf.cell(0, 6, clean_pdf_text(question_title), 0, 1)
                        pdf.set_font("Arial", "", 10)
                        
                        for opt in options:
                            is_checked = (selected_value == opt) or (isinstance(selected_value, list) and opt in selected_value)
                            
                            # Trava de segurança extra para cada linha
                            check_page_break(pdf, 6) 
                            
                            pdf.set_x(15)
                            x, y = pdf.get_x(), pdf.get_y()
                            
                            # Desenha o checkbox
                            pdf.rect(x, y+1, 3, 3)
                            if is_checked:
                                pdf.line(x, y+1, x+3, y+4)
                                pdf.line(x, y+4, x+3, y+1)
                            
                            # Escreve o texto da opção
                            pdf.set_xy(x + 5, y)
                            pdf.multi_cell(175, 5, clean_pdf_text(opt), 0, 'L')
                            
                        if obs:
                            pdf.set_x(15)
                            pdf.multi_cell(0, 5, clean_pdf_text(f"Obs: {obs}"), 0, 'J')
                        pdf.ln(2)

                    # PART I
                    print_section_header_fix(pdf, "PARTE I - HABILIDADES DE VIDA DIÁRIA")
                    print_question_options_fix(pdf, "1. ALIMENTAÇÃO:", opts_alim, data_aval.get('alim_nivel'), data_aval.get('alim_obs'))
                    print_question_options_fix(pdf, "2. HIGIENE:", opts_hig, data_aval.get('hig_nivel'), data_aval.get('hig_obs'))
                    print_question_options_fix(pdf, "3. LOCOMOÇÃO:", opts_loc, data_aval.get('loc_nivel'), data_aval.get('loc_obs'))
                    
                    # PART II
                    check_page_break(pdf, 40) # Verifica se tem 40mm livres, se não, pula a página
                    print_section_header_fix(pdf, "PARTE II - HABILIDADE SOCIAIS E DE INTERAÇÃO")
                    print_question_options_fix(pdf, "4. COMPORTAMENTO:", opts_comp, data_aval.get('comportamento'), data_aval.get('comp_obs'))
                    
                    check_page_break(pdf, 40)
                    print_question_options_fix(pdf, "5. PARTICIPAÇÃO EM GRUPO:", opts_part, data_aval.get('part_grupo'), data_aval.get('part_obs'))
                    
                    check_page_break(pdf, 35)
                    pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "6. INTERAÇÃO:", 0, 1)
                    pdf.set_font("Arial", "", 10)
                    for opt in opts_int[:-1]:
                        draw_check_option_simple(pdf, opt, data_aval.get('interacao') == opt)
                    is_outros = (data_aval.get('interacao') == "Outros")
                    txt_outros = f"Outros: {data_aval.get('interacao_outros') if data_aval.get('interacao_outros') else '____________________'}"
                    draw_check_option_simple(pdf, txt_outros, is_outros)
                    pdf.ln(4)

                    # PART III
                    check_page_break(pdf, 40)
                    print_section_header_fix(pdf, "PARTE III - HABILIDADES PEDAGÓGICAS")
                    print_question_options_fix(pdf, "7. ROTINA EM SALA:", opts_rot, data_aval.get('rotina'), data_aval.get('rotina_obs'))
                    print_question_options_fix(pdf, "8. ATIVIDADES PEDAGÓGICAS:", opts_ativ, data_aval.get('ativ_pedag'))

                    # PART IV
                    check_page_break(pdf, 40)
                    print_section_header_fix(pdf, "PARTE IV - HABILIDADES DE COMUNICAÇÃO E ATENÇÃO")
                    print_question_options_fix(pdf, "9. ATENÇÃO SUSTENTADA:", opts_at_sust, data_aval.get('atencao_sust'))
                    print_question_options_fix(pdf, "10. ATENÇÃO DIVIDIDA:", opts_at_div, data_aval.get('atencao_div'))
                    
                    check_page_break(pdf, 40)
                    print_question_options_fix(pdf, "11. ATENÇÃO SELETIVA:", opts_at_sel, data_aval.get('atencao_sel'))
                    print_question_options_fix(pdf, "12. LINGUAGEM:", opts_ling, data_aval.get('linguagem'), data_aval.get('ling_obs'))

                    # 6. ZEBRA STRIPED TABLE - IMPROVED
                    check_page_break(pdf, 60) # A tabela precisa de mais espaço
                    pdf.ln(2); pdf.set_font("Arial", "B", 10)
                    pdf.set_fill_color(200, 200, 200)
                    # Use width 180 total (60+120)
                    pdf.cell(60, 8, clean_pdf_text("NÍVEIS DE APOIO"), 1, 0, 'C', True)
                    pdf.cell(120, 8, clean_pdf_text("CARACTERÍSTICAS"), 1, 1, 'C', True)
                    
                    def print_zebra_row_fix(pdf, col1, col2, fill):
                        # Mapeamento cirúrgico de linhas baseado no visual real do PDF
                        if "Não há necessidade" in col1: 
                            lines_left = 1; lines_right = 2
                        elif "Nível 1" in col1: 
                            lines_left = 1; lines_right = 1
                        elif "Nível 2" in col1: 
                            lines_left = 1; lines_right = 1
                        elif "Nível 3" in col1: 
                            lines_left = 1; lines_right = 2
                        else:
                            lines_left = 1; lines_right = 1

                        # Calcula a altura exata: 5mm por linha + 4mm de margem total
                        max_lines = max(lines_left, lines_right)
                        row_height = (max_lines * 5) + 4 
                        
                        # Verifica se cabe na página usando a nossa função inteligente
                        check_page_break(pdf, row_height)
                        
                        x, y = 15, pdf.get_y()
                        
                        # Define a cor de fundo (zebrado)
                        if fill:
                            pdf.set_fill_color(240, 240, 240)
                        else:
                            pdf.set_fill_color(255, 255, 255)
                        
                        # Desenha os retângulos de fundo e as bordas
                        pdf.rect(x, y, 60, row_height, 'DF') 
                        pdf.rect(x+60, y, 120, row_height, 'DF')
                        
                        # Imprime a Coluna Esquerda (Níveis) - Centralizado
                        pdf.set_font("Arial", "B", 9)
                        y_off1 = (row_height - (lines_left * 5)) / 2
                        pdf.set_xy(x, y + y_off1)
                        pdf.multi_cell(60, 5, clean_pdf_text(col1), 0, 'C')
                        
                        # Imprime a Coluna Direita (Características) - Justificado
                        pdf.set_font("Arial", "", 9)
                        y_off2 = (row_height - (lines_right * 5)) / 2
                        pdf.set_xy(x+60, y + y_off2)
                        pdf.multi_cell(120, 5, clean_pdf_text(col2), 0, 'J')
                        
                        # Move o cursor para baixo da linha que acabou de ser desenhada
                        pdf.set_xy(x, y + row_height)

                    print_zebra_row_fix(pdf, "Não há necessidade de apoio", "O estudante apresenta autonomia. As ações disponibilizadas aos demais estudantes são suficientes, acrescidas de ações do AEE.", False)
                    print_zebra_row_fix(pdf, "Nível 1 - Apoio pouco substancial", "Não há necessidade de apoio constante, apenas em ações pontuais.", True)
                    print_zebra_row_fix(pdf, "Nível 2 - Apoio substancial", "Há necessidade de apoio constante ao estudante.", False)
                    print_zebra_row_fix(pdf, "Nível 3 - Apoio muito substancial", "Casos severos com necessidade de monitor e ações específicas: flexibilização de horário e espaços.", True)

                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 11); pdf.cell(0, 8, clean_pdf_text("CONCLUSÃO DA EQUIPE PEDAGÓGICA"), 0, 1)
                    pdf.set_font("Arial", "", 10)
                    pdf.multi_cell(0, 5, clean_pdf_text("Diante dos aspectos avaliados, a equipe pedagógica verificou que o estudante corresponde ao Nível:"), 0, 'L')
                    
                    level_result = data_aval.get('conclusao_nivel', 'NÃO NECESSITA DE APOIO').upper()
                    pdf.set_font("Arial", "B", 12); pdf.ln(2); pdf.cell(0, 8, clean_pdf_text(level_result), 1, 1, 'C')
                    
                    pdf.ln(3); pdf.set_font("Arial", "", 10)
                    apoio_txt = data_aval.get('apoio_existente') if data_aval.get('apoio_existente') else "______________________________________________________"
                    pdf.multi_cell(0, 5, clean_pdf_text(f"Profissional de Apoio Escolar (se houver): {apoio_txt}"), 0, 'J')

                    pdf.ln(10)
                    if pdf.get_y() > 240: pdf.add_page()
                    pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, clean_pdf_text("Responsáveis pela avaliação:"), 0, 1); pdf.ln(5)
                    
                    # Signatures formatted with Name on one line, Role below
                    def draw_signature_block(pdf, x, y, width, name, role):
                        pdf.line(x, y, x + width, y)
                        pdf.set_xy(x, y + 2)
                        pdf.set_font("Arial", "", 9)
                        pdf.multi_cell(width, 4, clean_pdf_text(name), 0, 'C')
                        pdf.set_xy(x, pdf.get_y())
                        pdf.set_font("Arial", "I", 8)
                        pdf.multi_cell(width, 4, clean_pdf_text(role), 0, 'C')

                    y_sig_1 = pdf.get_y()
                    draw_signature_block(pdf, 10, y_sig_1, 55, data_aval.get('resp_sala',''), "Prof. Sala Regular")
                    draw_signature_block(pdf, 75, y_sig_1, 55, data_aval.get('resp_ef',''), "Prof. Ed. Física")
                    draw_signature_block(pdf, 140, y_sig_1, 55, data_aval.get('resp_arte',''), "Prof. Arte")
                    
                    # Add space for next row
                    pdf.set_xy(10, y_sig_1 + 25)
                    y_sig_2 = pdf.get_y()
                    
                    draw_signature_block(pdf, 10, y_sig_2, 55, data_aval.get('resp_dir',''), "Equipe Gestora")
                    draw_signature_block(pdf, 75, y_sig_2, 55, data_aval.get('resp_ee',''), "Prof. Ed. Especial")
                    draw_signature_block(pdf, 140, y_sig_2, 55, data_aval.get('resp_coord',''), "Coordenação")
                    
                    pdf.ln(25); pdf.set_font("Arial", "", 10)
                    # Left aligned date ('L')
                    pdf.cell(0, 6, clean_pdf_text(f"Limeira, {data_aval.get('data_emissao', date.today()).strftime('%d/%m/%Y')}."), 0, 1, 'L')

                    st.session_state.pdf_bytes_aval = get_pdf_bytes(pdf)
                    st.rerun()

            if 'pdf_bytes_aval' in st.session_state:
                st.download_button("📥 BAIXAR PDF AVALIAÇÃO", st.session_state.pdf_bytes_aval, f"Avaliacao_{data_aval.get('nome','aluno')}.pdf", "application/pdf", type="primary")

        # --- ABA HISTÓRICO ---
        with tabs[1]:
            st.subheader("Histórico de Atividades")
            df_hist = safe_read("Historico", ["Data_Hora", "Aluno", "Usuario", "Acao", "Detalhes"])
            if not df_hist.empty and data_aval.get('nome'):
                student_hist = df_hist[df_hist["Aluno"] == data_aval.get('nome')]
                if not student_hist.empty:
                    st.dataframe(student_hist.iloc[::-1], use_container_width=True, hide_index=True)
                else: st.info("Sem histórico.")
            else: st.info("Histórico vazio.")



















# ==============================================================================
    # --- AVALIAÇÃO PEDAGÓGICA PARA APOIO ESCOLAR 2.0 ---
    # ==============================================================================
    elif doc_mode == "Avaliação de Apoio 2.0":
        st.markdown("""<div class="header-box"><div class="header-title">Avaliação Pedagógica para Apoio Escolar 2.0</div></div>""", unsafe_allow_html=True)
        
        tabs = st.tabs(["📝 Preenchimento e Emissão", "🕒 Histórico"])

        if 'data_avaliacao2' not in st.session_state: st.session_state.data_avaliacao2 = {}
        data_aval2 = st.session_state.data_avaliacao2
        data_pei = st.session_state.data_pei
        data_caso = st.session_state.data_case
        
        with tabs[0]:
            st.subheader("Configuração e Preenchimento")
            
            # Botão Importar (Restrito para Gestão/Professores)
            if not is_monitor:
                if st.button("🔄 Importar Dados Básicos (PEI/Estudo de Caso)"):
                    if data_pei or data_caso:
                        data_aval2['nome'] = data_pei.get('nome') or data_caso.get('nome', '')
                        data_aval2['nasc'] = data_pei.get('nasc') or data_caso.get('d_nasc', '')
                        data_aval2['ano_esc'] = data_pei.get('ano_esc') or data_caso.get('ano_esc', '')
                        
                        diag_tipo = data_pei.get('diag_tipo', [])
                        deficiencias_formatadas = []
                        if "Deficiência" in diag_tipo and data_pei.get('defic_txt'):
                            deficiencias_formatadas.append(data_pei.get('defic_txt'))
                        if "Transtorno do Neurodesenvolvimento" in diag_tipo and data_pei.get('neuro_txt'):
                            deficiencias_formatadas.append(data_pei.get('neuro_txt'))
                        if "Transtornos Aprendizagem" in diag_tipo and data_pei.get('aprend_txt'):
                            deficiencias_formatadas.append(data_pei.get('aprend_txt'))
                        
                        data_aval2['diagnostico_outra'] = " / ".join(deficiencias_formatadas)
                        
                        # Importar professores automaticamente do PEI
                        if data_pei.get('gestor'): data_aval2['resp_diretor'] = data_pei.get('gestor')
                        if data_pei.get('prof_poli'): data_aval2['resp_poli'] = data_pei.get('prof_poli')
                        if data_pei.get('prof_aee'): data_aval2['resp_aee'] = data_pei.get('prof_aee')
                        
                        st.success("Dados importados com sucesso!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.warning("Sem dados prévios para importar.")

            # --- 1. Identificação ---
            st.markdown("### 1. Identificação")
            c_nom, c_ano = st.columns([3, 1])
            data_aval2['nome'] = c_nom.text_input("Estudante", value=data_aval2.get('nome', st.session_state.get('aluno_selecionado', '')), disabled=True)
            data_aval2['ano_esc'] = c_ano.text_input("Ano/Etapa", value=data_aval2.get('ano_esc', ''), disabled=is_monitor)
            
            st.markdown("**Assinale o diagnóstico (se houver):**")
            defs_opts = ["Deficiência auditiva/surdez", "Deficiência física", "Deficiência intelectual", "Deficiência múltipla", "Deficiência visual", "Transtorno do Espectro Autista", "Síndrome de Down", "Altas habilidades/Superdotação"]
            
            valores_salvos = data_aval2.get('defic_chk', [])
            if not isinstance(valores_salvos, list): valores_salvos = []
            valores_validos = [v for v in valores_salvos if v in defs_opts]
            
            data_aval2['defic_chk'] = st.multiselect("Selecione:", defs_opts, default=valores_validos, disabled=is_monitor)
            data_aval2['diagnostico_outra'] = st.text_input("Outra (Diagnóstico):", value=data_aval2.get('diagnostico_outra', ''), disabled=is_monitor)

            st.divider()
            versao = st.radio("Selecione a Etapa:", ["Creche", "Pré-escola e Ensino Fundamental"], index=0 if data_aval2.get('versao') == 'Creche' else 1, horizontal=True, disabled=is_monitor)
            data_aval2['versao'] = versao

            st.markdown("### 2. Dimensões Avaliadas (Cálculo Automático)")
            
            # --- Definição das Opções Conforme Marcos Legais ---
            if versao == "Creche":
                st.info("Na educação infantil (creche), o cuidado integra o currículo. A avaliação considera descompassos significativos em relação ao esperado para a idade.")
                opcoes = {
                    "2.1 Desenvolvimento Motor": ["Compatível com a idade", "Pequeno atraso", "Atraso moderado com necessidade de suporte eventual", "Atraso importante com necessidade de suporte constante"],
                    "2.2 Comunicação e Interação": ["Interação típica para a idade", "Dificuldades leves", "Interação limitada", "Ausência significativa de interação funcional"],
                    "2.3 Regulação Sensorial/Comportamental": ["Regulação compatível com a idade", "Sensibilidade leve", "Crises frequentes e desproporcionais", "Crises intensas com risco"],
                    "2.4 Alimentação e Cuidados com higiene": ["Cuidado típico da faixa etária", "Pequenas adaptações", "Procedimentos diferenciados frequentes", "Necessidade individualizada constante fora do padrão da turma"],
                    "2.5 Segurança Física": ["Sem risco adicional", "Necessita de supervisão ampliada", "Risco frequente", "Risco grave e permanente"],
                    "2.6 Participação Pedagógica": ["Participa das atividades compatíveis com idade", "Precisa de estimulação adicional", "Precisa de mediação frequente", "Não participa sem apoio integral"]
                }
            else:
                st.info("Critérios para Indicação: Descompasso significativo e persistente; necessidade que extrapole o padrão da turma; risco permanente ou impossibilidade de participação.")
                opcoes = {
                    "Comunicação": ["Independente", "Mediação eventual", "Mediação frequente", "Dependência significativa"],
                    "Mobilidade": ["Independente", "Mediação eventual", "Mediação frequente", "Dependência significativa"],
                    "Autocuidado": ["Independente", "Mediação eventual", "Ajuda parcial", "Dependência permanente"],
                    "Interação Social": ["Adequada", "Dificuldades leves", "Mediação constante", "Isolamento grave ou conflitos recorrentes"],
                    "Autorregulação Comportamental": ["Adequada", "Oscilações leves", "Crises frequentes", "Risco à integridade"],
                    "Participação Pedagógica": ["Participa com autonomia", "Participa com adaptações", "Necessita de mediação frequente", "Não participa sem apoio integral"]
                }

            # --- Renderização Dinâmica e Cálculo ---
            pontos = 0
            for label, list_opts in opcoes.items():
                key = f"val_{label}"
                val = st.selectbox(label, list_opts, index=list_opts.index(data_aval2.get(key, list_opts[0])) if data_aval2.get(key) in list_opts else 0, disabled=is_monitor)
                data_aval2[key] = val
                pontos += list_opts.index(val)
            
            # --- Lógica de Indicação ---
            if versao == "Creche":
                if pontos <= 5: ind, desc = "Não há indicação de apoio contínuo", "O estudante apresenta autonomia e desenvolvimento compatível à idade cronológica. As ações disponibilizadas aos demais estudantes da classe comum são suficientes."
                elif pontos <= 9: ind, desc = "Apoio intermitente", "O estudante apresenta dificuldade em algumas atividades e precisa do apoio pontual disponibilizado pelos educadores da turma."
                elif pontos <= 13: ind, desc = "Apoio parcial", "O estudante apresenta muitas dificuldades no desenvolvimento global. Há necessidade de ampliação no número de educadores para além daquele preconizado para a turma possibilitando o apoio escolar nas atividades escolares e nas atividades de vida prática e diária, acrescidas de ações do atendimento educacional especializado (AEE)."
                else: ind, desc = "Apoio contínuo", "Existe comprometimento acentuado na compreensão e/ou execução das atividades pedagógicas, necessitando da atuação de um monitor específico para o apoio físico, visual e verbal ao estudante, além de adaptações curriculares e metodológicas com ações do Atendimento Educacional Especializado."
            else:
                if pontos <= 5: ind, desc = "Não há necessidade de apoio escolar", "Não há necessidade de apoio além daquele dispensado ao coletivo da turma. Não há necessidade de atuação do monitor."
                elif pontos <= 9: ind, desc = "Apoio intermitente (Pouco Substancial)", "O estudante apresenta dificuldades em relação ao desenvolvimento da aprendizagem. Flexibilizações metodológicas, ações de recuperação contínua e ajustes na rotina dispensados pelos docentes da classe comum e do AEE são suficientes. O monitor permanece fora da sala de aula para apoio quando necessário."
                elif pontos <= 13: ind, desc = "Apoio parcial (Substancial)", "O estudante apresenta dificuldades em relação ao desenvolvimento da aprendizagem. Flexibilizações metodológicas e curriculares, ações de recuperação contínua, ajustes na rotina e ações pontuais do monitor são suficientes. O monitor permanece dentro da sala de aula para apoio aos estudantes que necessitarem."
                else: ind, desc = "Apoio contínuo (Muito Substancial)", "Existe comprometimento acentuado na compreensão e/ou execução das atividades pedagógicas. Comportamento desafiador, auto e/ou heterolesivo. Além de adaptações curriculares e metodológicas e das ações do Atendimento Educacional Especializado, há necessidade da atuação de um monitor específico para o apoio nas ações de vida prática e diária, na interação social, comunicação e nas atividades pedagógicas. Há necessidade da atuação direta de um único monitor."

            data_aval2['pontos'] = pontos
            data_aval2['indicacao'] = ind
            
            st.divider()
            st.markdown("### 3. Conclusão e Parecer")
            col_score, col_result = st.columns([1, 2])
            col_score.metric("Pontuação Total", f"{pontos} / 18")
            col_result.success(f"**Resultado Final:** {ind}\n\n{desc}")

            data_aval2['parecer'] = st.text_area("4. Parecer Pedagógico (Contextualize as estratégias e, se houver necessidade, nomeie o monitor responsável)", value=data_aval2.get('parecer', ''), disabled=is_monitor, height=120)
            
            st.divider()
            st.markdown("### 4. Responsáveis pela Avaliação")
            c_r1, c_r2 = st.columns(2)
            data_aval2['resp_diretor'] = c_r1.text_input("Diretor(a) de Escola", value=data_aval2.get('resp_diretor', ''), disabled=is_monitor)
            data_aval2['resp_poli'] = c_r2.text_input("Professor(a) Polivalente/Regente", value=data_aval2.get('resp_poli', ''), disabled=is_monitor)
            
            c_r3, c_r4 = st.columns(2)
            data_aval2['resp_aee'] = c_r3.text_input("Professor(a) AEE", value=data_aval2.get('resp_aee', ''), disabled=is_monitor)
            data_aval2['resp_monitor'] = c_r4.text_input("Monitor(a) / Apoio Escolar (se houver)", value=data_aval2.get('resp_monitor', ''), disabled=is_monitor)

            d_val = data_aval2.get('data_emissao')
            if isinstance(d_val, str): 
                try: d_val = datetime.strptime(d_val, '%Y-%m-%d').date()
                except: d_val = date.today()
            if not isinstance(d_val, date): d_val = date.today()
            
            data_aval2['data_emissao'] = st.date_input("Data de Emissão do Documento", value=d_val, format="DD/MM/YYYY", disabled=is_monitor)

            st.markdown("---")
            c_sv, c_pd = st.columns(2)
            
            if not is_monitor:
                if c_sv.button("💾 Salvar Avaliação 2.0", type="primary", use_container_width=True):
                    if isinstance(data_aval2['data_emissao'], date): 
                        data_aval2['data_emissao'] = data_aval2['data_emissao'].strftime("%Y-%m-%d")
                    save_student("AVALIACAO2", data_aval2.get('nome', 'aluno'), data_aval2, "Final")

            if c_pd.button("👁️ Gerar Documento Oficial PDF", use_container_width=True):
                log_action(data_aval2.get('nome'), "Gerou PDF", "Avaliação de Apoio 2.0")
                
                pdf = OfficialPDF('P', 'mm', 'A4')
                pdf.add_page()
                pdf.set_margins(15, 15, 15)
                pdf.set_auto_page_break(auto=True, margin=20)

                # --- CABEÇALHO ---
                if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 15, 10, 25)
                if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 170, 6, 25)

                pdf.set_xy(0, 15)
                pdf.set_font("Arial", "B", 12)
                pdf.cell(210, 6, clean_pdf_text("PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
                pdf.cell(190, 6, clean_pdf_text("SECRETARIA MUNICIPAL DE EDUCAÇÃO         "), 0, 1, 'C')
                pdf.ln(8)
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, clean_pdf_text("AVALIAÇÃO PEDAGÓGICA PARA APOIO ESCOLAR"), 0, 1, 'C')
                
                pdf.set_font("Arial", "BU", 11)
                pdf.set_text_color(80, 80, 80)
                titulo_etapa = "Educação Infantil - Creche" if versao == "Creche" else "Educação Infantil / Pré-escola e Ensino Fundamental"
                pdf.cell(0, 5, clean_pdf_text(titulo_etapa), 0, 1, 'C')
                pdf.set_text_color(0, 0, 0)
                
                pdf.ln(5)

                # Helper visual para blocos do PDF
                def draw_section_box(pdf_obj, title):
                    pdf_obj.set_fill_color(230, 230, 230)
                    pdf_obj.set_font("Arial", "B", 10)
                    pdf_obj.cell(0, 8, clean_pdf_text(f"  {title}"), border=1, ln=1, align='L', fill=True)
                    pdf_obj.ln(1)

                # --- 1. IDENTIFICAÇÃO ---
                draw_section_box(pdf, "1. IDENTIFICAÇÃO DO ESTUDANTE")
                pdf.set_font("Arial", "B", 10); pdf.cell(35, 7, "Unidade Escolar:", "LT", 0); pdf.set_font("Arial", "", 10); pdf.cell(0, 7, clean_pdf_text("CEIEF Rafael Affonso Leite"), "TR", 1)
                pdf.set_font("Arial", "B", 10); pdf.cell(35, 7, "Estudante:", "L", 0); pdf.set_font("Arial", "", 10); pdf.cell(0, 7, clean_pdf_text(data_aval2.get('nome', '')), "R", 1)
                
                nasc_val = data_aval2.get('nasc', '')
                if isinstance(nasc_val, str) and len(nasc_val) == 10 and nasc_val.count('-') == 2:
                    try: nasc_val = datetime.strptime(nasc_val, "%Y-%m-%d").strftime("%d/%m/%Y")
                    except: pass
                elif isinstance(nasc_val, date):
                    nasc_val = nasc_val.strftime("%d/%m/%Y")
                    
                idade_calc = data_caso.get('idade', data_pei.get('idade', ''))
                
                pdf.set_font("Arial", "B", 10); pdf.cell(35, 7, "Nascimento:", "LB", 0); pdf.set_font("Arial", "", 10); pdf.cell(60, 7, clean_pdf_text(nasc_val), "B", 0)
                pdf.set_font("Arial", "B", 10); pdf.cell(15, 7, "Idade:", "B", 0); pdf.set_font("Arial", "", 10); pdf.cell(0, 7, clean_pdf_text(idade_calc), "RB", 1)
                
                pdf.ln(5)

                # --- 2. DIAGNÓSTICO ---
                draw_section_box(pdf, "2. DIAGNÓSTICO CLÍNICO")
                pdf.set_font("Arial", "", 9)
                pdf.cell(0, 6, clean_pdf_text("De acordo com o diagnóstico do estudante:"), 0, 1)
                
                selected_defs = data_aval2.get('defic_chk', [])
                
                def prt_chk_grid(label1, label2):
                    v1 = "X" if label1 in selected_defs else "  "
                    v2 = "X" if label2 in selected_defs else "  "
                    pdf.set_x(20)
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(8, 5, f"( {v1} )", 0, 0, 'C')
                    pdf.set_font("Arial", "", 9)
                    pdf.cell(82, 5, clean_pdf_text(label1), 0, 0)
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(8, 5, f"( {v2} )", 0, 0, 'C')
                    pdf.set_font("Arial", "", 9)
                    pdf.cell(82, 5, clean_pdf_text(label2), 0, 1)
                
                pdf.ln(2)
                prt_chk_grid("Deficiência auditiva/surdez", "Deficiência física")
                prt_chk_grid("Deficiência intelectual", "Deficiência múltipla")
                prt_chk_grid("Deficiência visual", "Transtorno do Espectro Autista")
                prt_chk_grid("Síndrome de Down", "Altas habilidades/Superdotação")
                
                pdf.ln(2)
                pdf.set_x(20)
                pdf.set_font("Arial", "B", 9)
                pdf.cell(12, 6, "Outra: ", 0, 0)
                pdf.set_font("Arial", "U", 9)
                pdf.cell(0, 6, clean_pdf_text(data_aval2.get('diagnostico_outra', '____________________________________________________')), 0, 1)
                pdf.ln(5)

                # --- 3. PRESSUPOSTOS LEGAIS E NORMATIVOS ---
                draw_section_box(pdf, "3. PRESSUPOSTOS LEGAIS E NORMATIVOS")
                
                textos_legais = [
                    ("1- Lei nº 12.764/2012 institui a Política Nacional de Proteção dos Direitos da Pessoa com Transtorno do Espectro Autista e em seu artigo 3º indica:", 
                     "§1º Em casos de comprovada necessidade, a pessoa com transtorno do espectro autista incluída nas classes comuns de ensino regular, nos termos do inciso IV do art. 2º, terá direito a acompanhante especializado."),
                    ("2- Lei Nº 13.146/2015 institui a Lei Brasileira de Inclusão da Pessoa com Deficiência (Estatuto da Pessoa com Deficiência) que em seu art. 3º, inciso XIII, descreve as ações referentes ao apoio:",
                     "Pessoa que exerce atividades de alimentação, higiene e locomoção do estudante com deficiência e atua em todas as atividades escolares nas quais se fizer necessária, em todos os níveis e modalidades de ensino, em instituições públicas e privadas, excluídas as técnicas ou os procedimentos identificados com profissões legalmente estabelecidas;"),
                    ("3- Decreto nº 12.686/2025 institui a Política Nacional de Educação Especial Inclusiva e a Rede Nacional de Educação Especial Inclusiva e descreve em seu Art.14 – sobre a atuação do profissional de apoio escolar:",
                     "I - na locomoção, no acesso e na participação dos estudantes em todos os espaços e atividades pedagógicas;\nII - na higiene e na alimentação, guardado o respeito ao corpo e à privacidade, ao tempo e às escolhas dos estudantes;\nIII - na interação social e na comunicação, a partir do reconhecimento das diferentes formas de expressão dos estudantes e da pluralidade dos meios e modos de comunicação; e\nIV - na utilização de eventuais tecnologias e recursos auxiliares desenvolvidos pelo AEE, de modo a favorecer o convívio entre pares e a livre expressão dos estudantes nas atividades e nos espaços escolares.\n§ 1º O profissional de apoio escolar atuará em todas as atividades escolares, e deverá reportar-se à equipe pedagógica, sempre que se fizer necessário.\n§ 2º A oferta do profissional de apoio escolar independerá de resultado de diagnóstico, laudo, relatório ou qualquer documento emitido por profissional de saúde."),
                    ("4- Lei Ordinária nº 7.146/2025 institui as Diretrizes Municipais de Educação Especial para Pessoas com Transtorno do Espectro Autista (TEA) em Limeira e aborda em seu Art.10 sobre o acompanhante especializado:",
                     "O Acompanhante Especializado atuará com os estudantes que apresentam necessidades substanciais de apoio ou muito substanciais envolvendo atividades de comunicação, interação social, cuidados pessoais, alimentação, higiene, locomoção e atividades escolares, tendo como premissa o desenvolvimento da autonomia do estudante. Parágrafo único. Para a identificação da necessidade de apoio escolar, além daquele dispensado ao coletivo da turma, e tomada de decisões quanto ao atendimento necessário, a escola deverá realizar com assessoramento técnico, avaliação do estudante o processo de ensino aprendizagem."),
                    ("5- O Decreto Municipal 23, de 26 de janeiro de 2026, reafirma:",
                     "SEÇÃO IV\nDos Serviços Disponibilizados para Apoio Escolar\nArt. 12. A Secretaria Municipal de Educação disponibilizará ao estudante, com necessidade de apoio substancial ou muito substancial, profissional de apoio escolar, que, atuará na sala de aula e demais espaços escolares sob orientação da equipe pedagógica da unidade escolar.\nParágrafo único. A oferta do profissional de apoio escolar será avaliada pela equipe pedagógica, por meio da avaliação pedagógica de necessidade de apoio e independerá de resultado de diagnóstico, laudo, relatório ou qualquer documento emitido por profissional de saúde.\nArt. 13. O profissional de apoio escolar, nos termos do disposto no Decreto Federal nº 12.686, de 20 de outubro de 2025, atuará nas atividades de comunicação e interação social; no auxílio necessário aos estudantes que não consigam realizar com autonomia e independência as atividades de cuidados pessoais, de alimentação, de higiene e de locomoção e, também, na mediação para superação das dificuldades relacionadas às atividades escolares.\nParágrafo único. Os serviços de apoio escolar poderão ser compartilhados entre os estudantes da mesma turma.\nArt. 14. A disponibilização de um profissional de apoio escolar se justifica quando a necessidade específica de um estudante não for atendida no contexto geral dos cuidados disponibilizados ao coletivo dos estudantes.\nArt. 15. O profissional de apoio escolar terá:\nI - formação inicial de, no mínimo, nível médio;\nII - formação continuada, com carga horária de, no mínimo, cento e oitenta horas, nos termos do disposto no art. 15 do Decreto Federal nº 12.686, de 20 de outubro de 2025.\nParágrafo único. O serviço de apoio escolar terá como objetivo a garantia do bem-estar no ambiente escolar e o desenvolvimento da autonomia do estudante.\nArt. 16. A Secretaria Municipal de Educação disponibilizará, aos profissionais da rede municipal de ensino, ações de formação continuada e de formação em serviço nas temáticas da Educação Especial.")
                ]
                
                for main_text, sub_text in textos_legais:
                    pdf.set_x(15)
                    pdf.set_font("Arial", "B", 9)
                    pdf.multi_cell(180, 5, clean_pdf_text(main_text), 0, 'J')
                    if sub_text:
                        pdf.set_x(25)
                        pdf.set_font("Arial", "", 9)
                        pdf.multi_cell(170, 5, clean_pdf_text(sub_text), 0, 'J')
                    pdf.ln(3)

                pdf.ln(2)
                pdf.set_fill_color(245, 245, 245)
                pdf.set_font("Arial", "B", 9)
                pdf.cell(180, 8, clean_pdf_text("Observação normativa: O diagnóstico clínico não gera direito automático ao apoio contínuo."), 1, 1, 'C', True)

                # --- 4. DIMENSÕES AVALIADAS ---
                if pdf.get_y() > 240: pdf.add_page()
                pdf.ln(5)
                draw_section_box(pdf, "4. AVALIAÇÃO PEDAGÓGICA (DIMENSÕES E HABILIDADES)")
                
                pdf.set_font("Arial", "", 9)
                if versao == "Creche":
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(0, 6, clean_pdf_text("1. Princípio Norteador"), 0, 1)
                    pdf.set_font("Arial", "", 9)
                    pdf.multi_cell(0, 5, clean_pdf_text("Na educação infantil, segmento creche, o cuidado integra o currículo. A dependência é característica do desenvolvimento típico. A avaliação considerará descompassos significativos em relação ao esperado para a faixa etária, conforme parâmetros da Organização Mundial da Saúde e da Sociedade Brasileira de Pediatria na definição dos marcos do desenvolvimento infantil."), 0, 'J')
                    pdf.ln(3)
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(0, 6, clean_pdf_text("1.1 Critérios para Indicação de Apoio Contínuo na Creche:"), 0, 1)
                    pdf.set_font("Arial", "", 9)
                    pdf.set_x(20)
                    pdf.multi_cell(0, 5, clean_pdf_text("a. Descompasso significativo e persistente;\nb. Necessidade de cuidado que extrapole o padrão da turma;\nc. Risco permanente à integridade pessoal;\nd. Impossibilidade de participação nas experiências propostas mesmo com adaptações coletivas."), 0, 'L')
                    pdf.ln(2)
                    pdf.set_font("Arial", "B", 9)
                    pdf.multi_cell(0, 5, clean_pdf_text("IMPORTANTE: A necessidade de colo, troca, alimentação assistida ou supervisão constante é inerente a faixa etária de 0 a 3 anos e, por si só, não caracteriza necessidade de apoio escolar além daquele disponibilizado ao coletivo da turma."), 0, 'J')
                else:
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(0, 6, clean_pdf_text("Finalidade da Avaliação Pedagógica de Necessidade de Apoio"), 0, 1)
                    pdf.set_font("Arial", "", 9)
                    pdf.multi_cell(0, 5, clean_pdf_text("Avaliar a necessidade de Profissional de Apoio Escolar com base na funcionalidade do estudante no contexto educacional."), 0, 'J')
                    pdf.ln(3)
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(0, 6, clean_pdf_text("Critérios para Indicação de Apoio Contínuo:"), 0, 1)
                    pdf.set_font("Arial", "", 9)
                    pdf.set_x(20)
                    pdf.multi_cell(0, 5, clean_pdf_text("- Descompasso significativo e persistente;\n- Necessidade de cuidado que extrapole o padrão da turma;\n- Risco permanente à integridade pessoal e de terceiros;\n- Impossibilidade de participação nas experiências propostas mesmo com adaptações coletivas."), 0, 'L')
                
                pdf.ln(4)
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 6, clean_pdf_text("Dimensões Avaliadas (0 a 3 pontos cada)."), 0, 1)
                pdf.ln(2)

                # Renderizando as opções com X e (0, 1, 2, 3) 
                def print_q_card(titulo, id_banco, dict_opcoes):
                    if pdf.get_y() > 250: pdf.add_page()
                    pdf.set_font("Arial", "B", 9)
                    pdf.cell(0, 6, clean_pdf_text(titulo), 0, 1)
                    
                    pdf.set_font("Arial", "", 9)
                    resp = data_aval2.get(id_banco, dict_opcoes[0])
                    for i, op in enumerate(dict_opcoes):
                        chk = "X" if resp == op else " "
                        pdf.set_x(20)
                        pdf.cell(0, 5, clean_pdf_text(f"{i} ( {chk} ) {op}"), 0, 1)
                    pdf.ln(2)

                for label, list_opts in opcoes.items():
                    print_q_card(label, f"val_{label}", list_opts)

                # --- 4.1 BOX DE PONTUAÇÃO (DESTAQUE ANTES DA CONCLUSÃO) ---
                if pdf.get_y() > 250: pdf.add_page()
                pdf.ln(4)
                pdf.set_fill_color(220, 220, 220)
                pdf.set_font("Arial", "B", 12)
                pdf.cell(180, 10, clean_pdf_text(f"PONTUAÇÃO OBTIDA: {pontos} de 18"), 1, 1, 'C', True)
                pdf.ln(5)

                # --- 5. RESULTADO FINAL E CONCLUSÃO ---
                if pdf.get_y() > 220: pdf.add_page()
                draw_section_box(pdf, "5. CONCLUSÃO TÉCNICA E PARECER PEDAGÓGICO")

                pdf.set_fill_color(50, 50, 50)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Arial", "B", 11)
                pdf.cell(0, 10, clean_pdf_text(f" RESULTADO FINAL: {ind.upper()}"), 1, 1, 'C', True)
                
                pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Arial", "", 10)
                pdf.multi_cell(0, 6, clean_pdf_text(f"Atuação/Fundamentação: {desc}"), 1, 'J')
                pdf.ln(4)

                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 6, clean_pdf_text("Parecer Pedagógico Contextualizado:"), 0, 1)
                pdf.set_font("Arial", "", 10)
                parecer_texto = data_aval2.get('parecer', '')
                if not parecer_texto.strip(): parecer_texto = "Nenhum apontamento adicional inserido."
                pdf.multi_cell(0, 6, clean_pdf_text(parecer_texto), 1, 'J')

                pdf.ln(10)
                
                # --- 6. ASSINATURAS E DATA (NOVO LAYOUT) ---
                data_emissao_formatada = data_aval2.get('data_emissao', date.today())
                if isinstance(data_emissao_formatada, str):
                    try: data_emissao_formatada = datetime.strptime(data_emissao_formatada, "%Y-%m-%d").strftime("%d/%m/%Y")
                    except: pass
                else: data_emissao_formatada = data_emissao_formatada.strftime("%d/%m/%Y")
                
                pdf.set_font("Arial", "", 10)
                pdf.cell(0, 6, clean_pdf_text(f"Limeira, {data_emissao_formatada}"), 0, 1, 'R')
                pdf.ln(6)
                
                pdf.set_font("Arial", "B", 9)
                pdf.cell(0, 6, clean_pdf_text("6. Assinaturas dos responsáveis pela avaliação e ciência dos responsáveis:"), 0, 1, 'L')
                pdf.ln(10)
                
                # --- Primeira linha de assinaturas (Polivalente e AEE) ---
                y_sig = pdf.get_y()
                pdf.line(20, y_sig, 90, y_sig)
                pdf.line(120, y_sig, 190, y_sig)
                
                pdf.set_font("Arial", "B", 9)
                pdf.set_xy(20, y_sig + 2)
                pdf.cell(70, 5, clean_pdf_text(data_aval2.get('resp_poli', '')), 0, 0, 'C')
                pdf.set_xy(120, y_sig + 2)
                pdf.cell(70, 5, clean_pdf_text(data_aval2.get('resp_aee', '')), 0, 1, 'C')
                
                pdf.set_font("Arial", "", 9)
                pdf.set_x(20); pdf.cell(70, 4, "Professor(a) Polivalente/Regente", 0, 0, 'C')
                pdf.set_x(120); pdf.cell(70, 4, "Professor(a) AEE", 0, 1, 'C')
                
                # --- Segunda linha de assinaturas (Diretor e Monitor/Se houver) ---
                pdf.ln(15)
                if pdf.get_y() > 270: pdf.add_page(); pdf.ln(10)
                y_sig2 = pdf.get_y()
                
                has_monitor = data_aval2.get('resp_monitor', '').strip() != ''
                
                pdf.line(20, y_sig2, 90, y_sig2)
                if has_monitor:
                    pdf.line(120, y_sig2, 190, y_sig2)
                
                pdf.set_font("Arial", "B", 9)
                pdf.set_xy(20, y_sig2 + 2)
                pdf.cell(70, 5, clean_pdf_text(data_aval2.get('resp_diretor', '')), 0, 0, 'C')
                if has_monitor:
                    pdf.set_xy(120, y_sig2 + 2)
                    pdf.cell(70, 5, clean_pdf_text(data_aval2.get('resp_monitor', '')), 0, 1, 'C')
                else:
                    pdf.ln(5)
                
                pdf.set_font("Arial", "", 9)
                pdf.set_x(20); pdf.cell(70, 4, "Diretor(a) de Escola", 0, 0, 'C')
                if has_monitor:
                    pdf.set_x(120); pdf.cell(70, 4, "Monitor(a) / Apoio Escolar", 0, 1, 'C')
                else:
                    pdf.ln(4)
                
                # --- Terceira linha de assinaturas (Pais) ---
                pdf.ln(15)
                if pdf.get_y() > 270: pdf.add_page(); pdf.ln(10)
                y_sig3 = pdf.get_y()
                pdf.line(55, y_sig3, 155, y_sig3)
                pdf.set_font("Arial", "", 9)
                pdf.set_xy(55, y_sig3 + 2)
                pdf.cell(100, 5, "Assinatura do Responsável pelo Estudante", 0, 1, 'C')

                st.session_state.pdf_bytes_aval2 = get_pdf_bytes(pdf)
                st.rerun()

            if 'pdf_bytes_aval2' in st.session_state:
                st.download_button("📥 Baixar PDF Oficial", st.session_state.pdf_bytes_aval2, f"Avaliacao_Apoio_2_{data_aval2.get('nome','aluno')}.pdf", type="primary", use_container_width=True)

        with tabs[1]:
            st.subheader("🕒 Histórico de Atividades")
            df_hist = safe_read("Historico", ["Data_Hora", "Aluno", "Usuario", "Acao", "Detalhes"])
            if not df_hist.empty and data_aval2.get('nome'):
                student_hist = df_hist[df_hist["Aluno"] == data_aval2.get('nome')]
                if not student_hist.empty:
                    st.dataframe(student_hist.iloc[::-1], use_container_width=True, hide_index=True)
                else: st.info("Sem histórico para este aluno.")
            else: st.info("Histórico vazio.")






























    
 

     # --- RELATÓRIO DIÁRIO ---
    elif doc_mode == "Relatório de Acompanhamento":
        st.markdown("""<div class="header-box"><div class="header-title">Relatório Diário de Acompanhamento</div></div>""", unsafe_allow_html=True)
        st.markdown("""<style>div[data-testid="stFormSubmitButton"] > button {width: 100%; background-color: #dcfce7; color: #166534; border: 1px solid #166534;}</style>""", unsafe_allow_html=True)
        
        # Inicializa se não existir
        if 'data_diario' not in st.session_state: st.session_state.data_diario = {}
        data_diario = st.session_state.data_diario
        if 'logs' not in data_diario: data_diario['logs'] = {}
        
        data_pei = st.session_state.data_pei # Para puxar dados automáticos
        
        tab_fill, tab_gen = st.tabs(["📝 Registro de Atividades", "🖨️ Emissão Mensal"])
        
        with tab_fill:
            with st.form("form_diario_registro"):
                st.subheader("1. Dados Gerais (Configuração)")
                st.caption("Estes dados serão usados no cabeçalho do relatório.")
                
                # Importar dados básicos
                if st.form_submit_button("🔄 Importar Dados do Aluno"):
                    if data_pei:
                        data_diario['nome'] = data_pei.get('nome', '')
                        data_diario['ano_esc'] = data_pei.get('ano_esc', '')
                        data_diario['escola'] = "CEIEF Rafael Affonso Leite"
                        st.success("Dados importados!")
                    else:
                        st.warning("Sem dados PEI para importar.")

                c1, c2 = st.columns(2)
                data_diario['escola'] = c1.text_input("Escola", value=data_diario.get('escola', 'CEIEF Rafael Affonso Leite'))
                data_diario['nome'] = c2.text_input("Estudante", value=data_diario.get('nome', data_pei.get('nome','')), disabled=True)
                
                c3, c4 = st.columns(2)
                data_diario['ano_esc'] = c3.text_input("Ano de Escolaridade", value=data_diario.get('ano_esc', data_pei.get('ano_esc','')))
                data_diario['periodo'] = c4.selectbox("Período", ["Manhã", "Tarde", "Integral"], index=0 if data_diario.get('periodo') == "Manhã" else (1 if data_diario.get('periodo') == "Tarde" else 2))
                
                data_diario['acompanhante'] = st.text_input("Acompanhante (Profissional)", value=data_diario.get('acompanhante', st.session_state.get('usuario_nome','')))
                
                st.divider()
                st.subheader("2. Registro do Dia")
                
                # Seleção da Data para Registro
                col_d_sel, col_info = st.columns([1, 2])
                data_selecionada = col_d_sel.date_input("Selecione a Data", value=date.today(), format="DD/MM/YYYY")
                data_str = data_selecionada.strftime("%Y-%m-%d")
                
                # Recuperar dados existentes para esta data
                log_atual = data_diario['logs'].get(data_str, {})
                
                # Checkbox Falta
                falta_val = log_atual.get('falta', False)
                falta = st.checkbox("Estudante Faltou?", value=falta_val)
                
                # Descrição
                desc_val = log_atual.get('descricao', '')
                descricao = st.text_area("Descrição das atividades realizadas:", value=desc_val, height=150, help="Descreva as atividades ou ocorrências deste dia.")
                
                st.markdown("---")
                # Botão de Salvar
                if st.form_submit_button("💾 Salvar Registro do Dia"):
                    # Atualiza o log no dicionário
                    data_diario['logs'][data_str] = {
                        'falta': falta,
                        'descricao': descricao
                    }
                    # Salva no banco de dados (persistência)
                    save_student("DIARIO", data_diario.get('nome', 'aluno'), data_diario, f"Diário {data_selecionada.strftime('%d/%m')}")
                    st.success(f"Registro de {data_selecionada.strftime('%d/%m/%Y')} salvo com sucesso!")
                    time.sleep(1)
                    st.rerun()

            # Visualização rápida dos últimos registros
            if data_diario['logs']:
                st.divider()
                st.markdown("##### 📅 Registros Recentes")
                # Converter para DF para mostrar
                lista_logs = []
                for d, info in data_diario['logs'].items():
                    lista_logs.append({
                        "Data": datetime.strptime(d, "%Y-%m-%d").date(),
                        "Presença": "Faltou" if info.get('falta') else "Presente",
                        "Resumo Atividade": info.get('descricao', '')[:100] + "..."
                    })
                if lista_logs:
                    df_logs = pd.DataFrame(lista_logs).sort_values("Data", ascending=False)
                    st.dataframe(df_logs, use_container_width=True, hide_index=True)

        with tab_gen:
            st.subheader("Emissão de Relatório Mensal")
            st.caption(f"Código Único do Documento: {data_diario.get('doc_uuid', 'Será gerado na emissão')}")
            
            c_m, c_y = st.columns(2)
            meses = {1:"Janeiro", 2:"Fevereiro", 3:"Março", 4:"Abril", 5:"Maio", 6:"Junho", 7:"Julho", 8:"Agosto", 9:"Setembro", 10:"Outubro", 11:"Novembro", 12:"Dezembro"}
            mes_sel = c_m.selectbox("Mês", list(meses.keys()), format_func=lambda x: meses[x], index=date.today().month - 1)
            ano_sel = c_y.number_input("Ano", min_value=2020, max_value=2030, value=date.today().year)
            
            if st.button("👁️ Gerar PDF Mensal", type="primary"):
                # Filtra logs do mês/ano selecionado
                logs_mensais = {}
                for d_str, info in data_diario['logs'].items():
                    try:
                        d_obj = datetime.strptime(d_str, "%Y-%m-%d").date()
                        if d_obj.month == mes_sel and d_obj.year == ano_sel:
                            logs_mensais[d_str] = info
                    except: pass
                
                if not logs_mensais:
                    st.warning("Não há registros salvos para o período selecionado.")
                else:
                    # Garantir UUID se não tiver
                    if 'doc_uuid' not in data_diario or not data_diario['doc_uuid']:
                        data_diario['doc_uuid'] = str(uuid.uuid4()).upper()
                        save_student("DIARIO", data_diario.get('nome', 'aluno'), data_diario, "Geração UUID")

                    log_action(data_diario.get('nome'), "Gerou PDF", f"Relatório Mensal {mes_sel}/{ano_sel}")
                    
                    # Cria PDF em Retrato ('P')
                    pdf = OfficialPDF('P', 'mm', 'A4')
                    pdf.add_page(); pdf.set_margins(15, 15, 15)
                    
                    # SET SIGNATURE FOOTER (Diario has different signature handling, but let's standardize verification)
                    # For Diário, signatures are usually just the accompanying professional printed
                    signatures_mock = []
                    if data_diario.get('acompanhante'):
                        signatures_mock.append({'name': data_diario.get('acompanhante'), 'role': 'Acompanhante'})
                    pdf.set_signature_footer(signatures_mock, data_diario.get('doc_uuid'))
                    
                    # --- CABEÇALHO ---
                    if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 15, 10, 25)
                    if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 170, 6, 25)

                    # Títulos Centralizados
                    pdf.set_xy(0, 15); pdf.set_font("Arial", "B", 12)
                    pdf.cell(210, 6, clean_pdf_text("PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
                    pdf.cell(180, 6, clean_pdf_text("CEIEF RAFAEL AFFONSO LEITE"), 0, 1, 'C')
                    pdf.ln(8)
                    pdf.set_font("Arial", "B", 16); pdf.cell(0, 10, clean_pdf_text("RELATÓRIO DIÁRIO DE AÇÕES DE ACOMPANHAMENTO ESCOLAR"), 0, 1, 'C')
                    pdf.ln(5)
                    
                    # Dados do Cabeçalho
                    pdf.set_font("Arial", "B", 10)
                    
                    # Linha 1
                    pdf.cell(15, 6, "Escola:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(110, 6, clean_pdf_text(data_diario.get('escola', '')), "B", 0)
                    
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(25, 6, clean_pdf_text("Data (Ref):"), 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(0, 6, f"{meses[mes_sel]}/{ano_sel}", "B", 1)
                    pdf.ln(2)
                    
                    # Linha 2
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(20, 6, "Estudante:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(0, 6, clean_pdf_text(data_diario.get('nome', '')), "B", 1)
                    pdf.ln(2)
                    
                    # Linha 3
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(35, 6, "Ano Escolaridade:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(60, 6, clean_pdf_text(data_diario.get('ano_esc', '')), "B", 0)
                    
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(20, 6, clean_pdf_text("Período:"), 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(0, 6, clean_pdf_text(data_diario.get('periodo', '')), "B", 1)
                    pdf.ln(2)
                    
                    # Linha 4
                    pdf.set_font("Arial", "B", 10)
                    pdf.cell(30, 6, "Acompanhante:", 0, 0)
                    pdf.set_font("Arial", "", 10)
                    pdf.cell(0, 6, clean_pdf_text(data_diario.get('acompanhante', '')), "B", 1)
                    
                    pdf.ln(8)
                    
                    # Tabela
                    pdf.set_font("Arial", "B", 11)
                    pdf.set_fill_color(200, 200, 200)
                    pdf.cell(0, 8, clean_pdf_text("Descrição das atividades realizadas com o estudante"), 1, 1, 'C', True)
                    
                    # Cabeçalho da Tabela
                    pdf.set_font("Arial", "B", 10)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.cell(25, 8, "DATA", 1, 0, 'C', True)
                    pdf.cell(0, 8, clean_pdf_text("ATIVIDADES / OCORRÊNCIAS"), 1, 1, 'C', True)
                   
                    # Conteúdo (Loop)
                    pdf.set_font("Arial", "", 10)
                    
                    # Ordenar dias
                    dias_ordenados = sorted(logs_mensais.keys())
                    
                    for d_str in dias_ordenados:
                        info = logs_mensais[d_str]
                        d_obj = datetime.strptime(d_str, "%Y-%m-%d")
                        d_fmt = d_obj.strftime("%d/%m")
                        
                        texto = info.get('descricao', '')
                        if info.get('falta'):
                            texto = "[ESTUDANTE FALTOU] " + texto
                        
                        pdf.set_x(15)
                        x_start = pdf.get_x()
                        y_start = pdf.get_y()
                        
                        # --- CÁLCULO DE ALTURA CORRIGIDO ---
                        texto_limpo = clean_pdf_text(texto)
                        line_height = 5
                        linhas_totais = 0
                        
                        # Divide o texto pelos "Enters" (\n) e calcula as linhas reais
                        for paragrafo in texto_limpo.split('\n'):
                            largura_paragrafo = pdf.get_string_width(paragrafo)
                            if largura_paragrafo == 0:
                                linhas_totais += 1  # Conta as linhas totalmente em branco
                            else:
                                # Adiciona as quebras automáticas que o FPDF vai fazer por falta de espaço
                                linhas_totais += int(largura_paragrafo / 150) + 1
                                
                        # Calcula a altura final baseada no número real de linhas (multiplicado pela altura da linha)
                        h_row = max(8, (linhas_totais * line_height) + 4) 
                        # -----------------------------------

                        # Check page break
                        if y_start + h_row > 270:
                            pdf.add_page()
                            y_start = pdf.get_y()
                            x_start = pdf.get_x() # Atualiza o X por segurança na nova página
                            
                        # Draw Cells
                        pdf.rect(x_start, y_start, 25, h_row) # Box Data
                        pdf.rect(x_start + 25, y_start, 155, h_row) # Box Desc
                        
                        # Print Data
                        pdf.set_xy(x_start, y_start)
                        pdf.cell(25, h_row, d_fmt, 0, 0, 'C')
                        
                        # Print Desc
                        pdf.set_xy(x_start + 27, y_start + 2)
                        pdf.multi_cell(151, line_height, texto_limpo, 0, 'J')
                        
                        # Move cursor
                        pdf.set_xy(x_start, y_start + h_row)
                        
                    # Assinaturas
                    pdf.ln(10)
                    if pdf.get_y() > 250: pdf.add_page()
                    
                    y = pdf.get_y()
                    pdf.line(15, y+10, 105, y+10)
                    pdf.line(115, y+10, 195, y+10)
                    
                    pdf.set_xy(15, y+11)
                    pdf.set_font("Arial", "", 9)
                    pdf.cell(90, 5, "Assinatura do Acompanhante", 0, 0, 'C')
                    pdf.cell(80, 5, clean_pdf_text("                       Visto da Coordenação/Direção"), 0, 1, 'C')
                    
                    st.session_state.pdf_bytes_diario_mes = get_pdf_bytes(pdf)
                    st.rerun()

            if 'pdf_bytes_diario_mes' in st.session_state:
                file_name_clean = data_diario.get('nome','aluno').replace(" ", "_")
                st.download_button(
                    "📥 BAIXAR RELATÓRIO MENSAL (PDF)", 
                    st.session_state.pdf_bytes_diario_mes, 
                    f"Diario_{file_name_clean}_{mes_sel}_{ano_sel}.pdf", 
                    "application/pdf", 
                    type="primary"
                )

# --- DECLARAÇÃO DE MATRÍCULA (NOVO) ---
    elif doc_mode == "Declaração de Matrícula":
        st.markdown(f"""<div class="header-box"><div class="header-title">Declaração de Matrícula e Atendimento</div></div>""", unsafe_allow_html=True)
        
        data_dec = st.session_state.data_declaracao
        data_pei = st.session_state.data_pei
        data_case = st.session_state.data_case
        data_pdi = st.session_state.data_pdi
        data_aval = st.session_state.get('data_avaliacao', {})
        
        # Helper safe get
        def get_d(d, k, default=""):
            return d.get(k, default) if d.get(k) else default

        with st.form("form_declaracao"):
            st.subheader("Dados da Declaração")
            st.caption("Os dados são pré-carregados dos outros documentos (PEI, PDI, Avaliação, Estudo de Caso) se disponíveis. Verifique e complemente se necessário.")
            
            # --- LÓGICA DE DADOS PADRÃO (AUTOPREENCHIMENTO) ---
            # Se o campo já estiver salvo em data_dec, usa ele. Senão, tenta buscar nos outros docs.
            
            # Nome
            val_nome = data_dec.get('nome') or get_d(data_pei, 'nome') or get_d(data_case, 'nome') or get_d(data_aval, 'nome') or st.session_state.get('aluno_selecionado', '')
            data_dec['nome'] = val_nome

            # Turma/Ano
            val_turma = data_dec.get('turma') or get_d(data_pei, 'ano_esc') or get_d(data_case, 'ano_esc') or get_d(data_aval, 'ano_esc')
            
            # Período
            val_periodo = data_dec.get('periodo') or get_d(data_case, 'periodo', 'Manhã')
            
            # Deficiência
            val_defic = data_dec.get('deficiencia')
            if not val_defic:
                if data_pei.get('defic_txt'): val_defic = data_pei['defic_txt']
                elif data_pei.get('diag_tipo'): val_defic = ", ".join(data_pei['diag_tipo'])
                elif data_aval.get('defic_chk'): val_defic = ", ".join(data_aval['defic_chk'])
            
            # Professores (Prioridade: PEI -> Avaliação -> Vazio)
            val_poli = data_dec.get('prof_poli') or get_d(data_pei, 'prof_poli') or get_d(data_aval, 'resp_sala')
            val_arte = data_dec.get('prof_arte') or get_d(data_pei, 'prof_arte') or get_d(data_aval, 'resp_arte')
            val_ef = data_dec.get('prof_ef') or get_d(data_pei, 'prof_ef') or get_d(data_aval, 'resp_ef')
            val_tec = data_dec.get('prof_tec') or get_d(data_pei, 'prof_tec') # Linguagens e Tecnologias
            val_aee = data_dec.get('prof_aee') or get_d(data_pei, 'prof_aee') or get_d(data_aval, 'resp_ee')
            
            # AEE Detalhes (Prioridade: PDI -> Vazio)
            val_aee_mod = data_dec.get('aee_modalidade') or get_d(data_pdi, 'aee_tipo')
            val_aee_comp = data_dec.get('aee_composicao') or get_d(data_pdi, 'aee_comp')
            val_aee_tempo = data_dec.get('aee_tempo') or get_d(data_pdi, 'aee_tempo', '50 minutos')
            
            # Apoio Escolar (Prioridade: Avaliação de Apoio -> Vazio)
            val_tem_apoio = data_dec.get('tem_apoio')
            val_nome_apoio = data_dec.get('nome_apoio')
            
            if not val_tem_apoio:
                # Inferência automática baseada na Avaliação de Apoio
                nivel = data_aval.get('conclusao_nivel', '')
                apoio_ex = data_aval.get('apoio_existente', '')
                if "Nível 2" in nivel or "Nível 3" in nivel or apoio_ex:
                    val_tem_apoio = 'Sim'
                    if not val_nome_apoio: val_nome_apoio = apoio_ex
                else:
                    val_tem_apoio = 'Não'

            # --- RENDERIZAÇÃO DOS CAMPOS ---
            
            c1, c2 = st.columns([3, 1])
            data_dec['nome'] = c1.text_input("Nome do Estudante", value=val_nome, disabled=True)
            data_dec['turma'] = c2.text_input("Turma/Ano", value=val_turma)
            
            c3, c4 = st.columns([1, 2])
            per_opts = ["Manhã", "Tarde", "Integral"]
            p_idx = per_opts.index(val_periodo) if val_periodo in per_opts else 0
            data_dec['periodo'] = c3.selectbox("Período", per_opts, index=p_idx)
            data_dec['deficiencia'] = c4.text_input("Deficiência / Transtorno", value=val_defic)
            
            st.divider()
            st.markdown("##### Quadro Docente")
            d1, d2 = st.columns(2)
            data_dec['prof_poli'] = d1.text_input("Professor(a) Regente", value=val_poli)
            data_dec['prof_arte'] = d2.text_input("Professor(a) Arte", value=val_arte)
            d3, d4 = st.columns(2)
            data_dec['prof_ef'] = d3.text_input("Professor(a) Ed. Física", value=val_ef)
            data_dec['prof_tec'] = d4.text_input("Professor(a) Linguagens e Tecnologias", value=val_tec)
            
            st.divider()
            st.markdown("##### Atendimento Educacional Especializado (AEE)")
            data_dec['prof_aee'] = st.text_input("Professor(a) Sala de Recursos", value=val_aee)
            
            a1, a2 = st.columns(2)
            data_dec['aee_modalidade'] = a1.text_input("Modalidade", value=val_aee_mod, help="Ex: Sala de Recursos, Colaborativo")
            data_dec['aee_composicao'] = a2.text_input("Forma de Atendimento", value=val_aee_comp, help="Ex: Individual, Grupo")
            
            a3, a4 = st.columns(2)
            data_dec['aee_tempo'] = a3.text_input("Tempo por atendimento", value=val_aee_tempo)
            data_dec['aee_freq'] = a4.text_input("Qtd. Atendimentos Semanais", value=data_dec.get('aee_freq', ''))
            
            st.divider()
            st.markdown("##### Apoio Escolar")
            has_apoio_idx = 0 if val_tem_apoio == 'Sim' else 1
            data_dec['tem_apoio'] = st.radio("Possui Profissional de Apoio?", ["Sim", "Não"], index=has_apoio_idx, horizontal=True)
            
            if data_dec['tem_apoio'] == 'Sim':
                data_dec['nome_apoio'] = st.text_input("Nome do Profissional de Apoio", value=val_nome_apoio)
            else:
                data_dec['nome_apoio'] = "" # Limpa se não tiver

            st.divider()
            
            if not is_monitor:
                if st.form_submit_button("🔄 Atualizar dados (re-importar)"):
                    # Ao submeter, os valores recalculados acima serão usados nos widgets e salvos automaticamente no session_state pelo streamlit
                    # Apenas exibimos uma mensagem
                    st.toast("Dados atualizados com base nos documentos!", icon="🔄")
                
                # Botão Salvar (Para persistir no banco)
                if st.form_submit_button("💾 Salvar Declaração"):
                    save_student("DECLARACAO", data_dec['nome'], data_dec, "Geral")
            else:
                st.info("Modo visualização (Monitor).")

        # Signatures section for Declaration
        st.divider()
        st.subheader("Assinaturas Digitais")
        st.caption(f"Código Único: {data_dec.get('doc_uuid', 'Salvar para gerar')}")
        
        current_signatures = data_dec.get('signatures', [])
        if current_signatures:
            for sig in current_signatures:
                st.success(f"Assinado por {sig['name']} em {sig['date']}")
        
        user_name = st.session_state.get('usuario_nome', '')
        already_signed = any(s['name'] == user_name for s in current_signatures)
        
        if not already_signed and not is_monitor:
            if st.button("🖊️ Assinar Declaração"):
                new_sig = {
                    "name": user_name,
                    "role": "Profissional",
                    "date": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                    "hash": str(uuid.uuid4())
                }
                if 'signatures' not in data_dec: data_dec['signatures'] = []
                data_dec['signatures'].append(new_sig)
                save_student("DECLARACAO", data_dec.get('nome'), data_dec, "Assinatura")
                st.rerun()

        # PDF Button
        if st.button("👁️ GERAR DECLARAÇÃO (PDF)"):
            log_action(data_dec.get('nome'), "Gerou PDF", "Declaração")
            
            pdf = OfficialPDF('P', 'mm', 'A4')
            pdf.add_page(); pdf.set_margins(20, 20, 20)
            pdf.set_signature_footer(data_dec.get('signatures', []), data_dec.get('doc_uuid', ''))
            
            if os.path.exists("logo_prefeitura.png"): pdf.image("logo_prefeitura.png", 20, 10, 25)
            if os.path.exists("logo_escola.png"): pdf.image("logo_escola.png", 165, 10, 25)
            
            pdf.set_xy(0, 20)
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 8, clean_pdf_text("            PREFEITURA MUNICIPAL DE LIMEIRA"), 0, 1, 'C')
            pdf.cell(0, 8, clean_pdf_text("SECRETARIA MUNICIPAL DE EDUCAÇÃO"), 0, 1, 'C')
            
            pdf.ln(20)
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, clean_pdf_text("DECLARAÇÃO DE MATRÍCULA E ATENDIMENTO"), 0, 1, 'C')
            pdf.ln(10)
            
            pdf.set_font("Arial", "", 12)
            texto_inicial = (
                f"Declaramos para os devidos fins que o(a) estudante {data_dec.get('nome', '').upper()}, "
                f"matriculado(a) na turma {data_dec.get('turma', '')}, período {data_dec.get('periodo', '').upper()}, "
                f"desta unidade escolar, frequenta as aulas regularmente."
            )
            pdf.multi_cell(0, 8, clean_pdf_text(texto_inicial))
            pdf.ln(5)
            
            texto_defic = f"O(A) estudante apresenta {data_dec.get('deficiencia', 'não informado')} e recebe acompanhamento pedagógico dos seguintes docentes:"
            pdf.multi_cell(0, 8, clean_pdf_text(texto_defic))
            pdf.ln(2)
            
            pdf.set_x(30)
            pdf.cell(0, 8, clean_pdf_text(f"- Professor(a) Regente: {data_dec.get('prof_poli', '')}"), 0, 1)
            pdf.set_x(30)
            pdf.cell(0, 8, clean_pdf_text(f"- Professor(a) Arte: {data_dec.get('prof_arte', '')}"), 0, 1)
            pdf.set_x(30)
            pdf.cell(0, 8, clean_pdf_text(f"- Professor(a) Ed. Física: {data_dec.get('prof_ef', '')}"), 0, 1)
            if data_dec.get('prof_tec'):
                pdf.set_x(30)
                pdf.cell(0, 8, clean_pdf_text(f"- Professor(a) Linguagens e Tecnologias: {data_dec.get('prof_tec', '')}"), 0, 1)
            
            pdf.ln(5)
            texto_aee = (
                f"No que tange ao Atendimento Educacional Especializado (AEE), o estudante é atendido pelo(a) "
                f"professor(a) {data_dec.get('prof_aee', '')}, na modalidade {data_dec.get('aee_modalidade', '')}, "
                f"de forma {data_dec.get('aee_composicao', '')}, com duração de {data_dec.get('aee_tempo', '')}, "
                f"{data_dec.get('aee_freq', '')} vezes por semana."
            )
            pdf.multi_cell(0, 8, clean_pdf_text(texto_aee))
            
            if data_dec.get('tem_apoio') == 'Sim':
                pdf.ln(5)
                pdf.multi_cell(0, 8, clean_pdf_text(f"Conta com o acompanhamento do Profissional de Apoio Escolar: {data_dec.get('nome_apoio', '')}."))
            
            pdf.ln(20)
            pdf.cell(0, 8, clean_pdf_text(f"Limeira, {datetime.now().strftime('%d/%m/%Y')}."), 0, 1, 'R')
            
            pdf.ln(30)
            pdf.cell(0, 8, "___________________________________________________", 0, 1, 'C')
            pdf.cell(0, 8, "Assinatura do Responsável / Direção", 0, 1, 'C')

            st.session_state.pdf_bytes_dec = get_pdf_bytes(pdf)
            st.rerun()

        if 'pdf_bytes_dec' in st.session_state:
            st.download_button("📥 BAIXAR DECLARAÇÃO", st.session_state.pdf_bytes_dec, f"Declaracao_{data_dec.get('nome','aluno')}.pdf", "application/pdf", type="primary")


elif modulo_atuacao == "🏫 Ensino Regular":
    
    # ==============================================================================
    # CARREGAMENTO GLOBAL DE CONFIGURAÇÕES (O CÉREBRO DO SISTEMA)
    # ==============================================================================
    df_config = safe_read("Config_Ata", ["chave", "valor"])
    
    def get_config(chave, padrao):
        if not df_config.empty and chave in df_config["chave"].values:
            return df_config.loc[df_config["chave"] == chave, "valor"].values[0]
        return padrao

    # --- Textos do Fundamental ---
    texto_base_padrao_ef = "Com base: na Resolução SME nº 07/2024, considerando as orientações da Resolução nº 02/2025 que atualiza o calendário escolar da Rede Municipal em decorrência da portaria nº 729 de 21 de fevereiro de 2025, que dispõe sobre o Calendário Escolar do ano de 2026 das Escolas da Rede Municipal de Ensino de Limeira, e no inciso V do artigo 5º, faz a indicação sobre a realização do Conselho de Ciclo/ Educação Infantil e Educação de Jovens e Adultos; no plano de trabalho para o ano de 2026, produzido no Conselho de Ciclo do 3º trimestre de 2025; na avaliação diagnóstica elaborada em fevereiro de 2026 e nas avaliações realizadas na unidade escolar no primeiro trimestre de 2026. Essa ata possibilita a análise sobre aprendizagem e desempenho dos estudantes e os resultados das estratégias de ensino empregadas."
    texto_base_ata_ef = get_config("texto_base_ata", texto_base_padrao_ef)
    
    propostas_padrao_ef = """1. Recuperação contínua de aprendizagem dos estudantes;
2. Intervenções pontuais e individuais;
3. Organização de recursos pedagógicos e situações didáticas eficientes e coerentes;
4. Encaminhamento à Direção/Serviço Social Escolar para busca ativa de estudantes com baixa frequência;
5. Proposta de compensação de ausências para o próximo trimestre;
6. Informar as famílias dos alunos com desempenho insuficiente e/ou baixa frequência visando a conscientização;
7. Indicar o aluno para Ação Pedagógica Complementar;
8. Propor atividades interdisciplinares objetivando o avanço do processo de aprendizagem;
9. Emitir relatórios solicitando suporte e avaliação de profissionais da saúde;
10. Sistematizar atividades para consolidação dos conteúdos;"""
    propostas_ata_ef = get_config("propostas_ata", propostas_padrao_ef)

    # --- Textos e Conteúdos da Educação Infantil ---
    texto_base_padrao_inf = "Com base na Resolução SME nº 07/24, considerando as orientações da Resolução nº 02/2025 que atualiza o calendário escolar da Rede Municipal em decorrência da portaria nº 729 de 21 de fevereiro de 2025, que dispõe sobre o Calendário Escolar do ano de 2026 das Escolas da Rede Municipal de Ensino de Limeira, especialmente no inciso V do artigo 5º que indica a realização do Conselho de Educação Infantil, na avaliação diagnóstica produzida em fevereiro de 2026 e nas avaliações realizadas na unidade escolar no primeiro trimestre de 2026. Essa ata possibilita a análise sobre aprendizagem e desempenho dos estudantes e os resultados das estratégias de ensino empregadas."
    texto_base_ata_inf = get_config("texto_base_ata_inf", texto_base_padrao_inf)
    
    propostas_padrao_inf = "1. \n2. \n3. \n4. \n5. \n6. "
    propostas_ata_inf = get_config("propostas_ata_inf", propostas_padrao_inf)

    def clean_pdf_text(texto):
        if not isinstance(texto, str):
            return str(texto)
        texto = texto.replace('“', '"').replace('”', '"')
        texto = texto.replace('‘', "'").replace('’', "'")
        texto = texto.replace('–', '-').replace('—', '-')
        texto = texto.replace('\u200b', '').replace('\xa0', ' ')
        return texto
    
    def get_criterios_infantil(etapa):
        # Define os padrões para a 1ª Etapa (como Semente)
        padrao_lv = "Oralidade: (Pronúncia correta das palavras; Participação atenta nas exposições orais escutando com atenção, respondendo e elaborando questões); Leitura: (Compreensão do significado das palavras; Socialização de critérios de escolha e de apreciação estética de leituras; Leitura de sílabas canônica e não canônicas; Localização de Informações explícitas); Análise Linguística: (Escrita de palavras utilizando a direção convencional; Reconhecimento e utilização das letras do alfabeto para a produção escrita; Traçado de modo convencional as letras com o auxílio do(a) professor(a)); Produção: (Produção oral de textos com destino escrito considerando gênero trabalhado (silhueta), (interlocutor) (sentido a partir de uma situação dada))."
        padrao_lm = "Álgebra (Classificação por semelhanças e diferenças); Estatística (Identificação de informações em gráficos de colunas); Geometria (Noções de: direcionalidade, proximidade, interioridade, exterioridade, reconhecimento de figuras geométricas planas e espaciais); Grandezas e medidas (Noções de: tempo, massa, capacidade e comprimento); Números e operações (Sistema de numeração decimal: récita numérica, reconhecimento da representação simbólica, contagem, associação quantidade/número e Ideias das operações: resolução de situações-problema com ideia de juntar - com apoio de material manipulativo)."
        padrao_is = "Sistema biológico: Saúde (Princípios de higiene pessoal: banho diário, cuidado com os dentes, lavagem das mãos); Sujeito histórico: Dados pessoais, Identidade e Escola (Nome e Sobrenome; nomes dos colegas e educadores; regras de convivência no ambiente escolar)."
        padrao_arte = "Identificação e nomeação de formas geométricas básicas; Identificação e nomeação de cores; Identificação de instrumentos e suas famílias; Relação do som com a fonte sonora; Expressão de forma oral, gestual, utilizando a imaginação na representação teatral."
        padrao_ccm = "Vivência de diversas formas de deslocamento nas situações de brincadeira (andar para frente, andar para trás, quadrupejar, saltar com um dos pés, saltitar, correr); Participação em brincadeira(s) cantada(s), realizando os movimentos sugeridos; Participação em brincadeira envolvendo a imitação, utilizando a linguagem corporal; Reconhecimento em si das diversas partes do corpo; Identificação da relação entre seu corpo e o espaço: Experimentação de movimentos estáticos e dinâmicos de equilíbrio."
        padrao_lt = "Exploração de mídias e tecnologias; Interação com diferentes linguagens midiáticas; Uso criativo de recursos tecnológicos de forma orientada e lúdica."
        padrao_libras = "Contato inicial com a Língua Brasileira de Sinais; Percepção visual e espacial; Reconhecimento e imitação de sinais básicos e expressões faciais do cotidiano escolar."
        
        # Se for outra etapa e não tiver configurado ainda, volta vazio para a escola preencher
        if etapa != "1ª Etapa":
            padrao_lv, padrao_lm, padrao_is, padrao_arte, padrao_ccm, padrao_lt, padrao_libras = "", "", "", "", "", "", ""
            
        return {
            "LV": get_config(f"crit_lv_{etapa}", padrao_lv),
            "LM": get_config(f"crit_lm_{etapa}", padrao_lm),
            "IS": get_config(f"crit_is_{etapa}", padrao_is),
            "Arte": get_config(f"crit_arte_{etapa}", padrao_arte),
            "CCM": get_config(f"crit_ccm_{etapa}", padrao_ccm),
            "LT": get_config(f"crit_lt_{etapa}", padrao_lt),
            "LIBRAS": get_config(f"crit_libras_{etapa}", padrao_libras)
        }
    
    # Matriz de Professores Padrão (Semente)
    MATRIZ_SEED = [
        {"Ciclo": "Bilíngue Fund.", "Turma": "2 e 3", "Disciplina": "Polivalente", "Professor": "Flaviany Miranda Montovani"},
        {"Ciclo": "Bilíngue Fund.", "Turma": "2 e 3", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Bilíngue Fund.", "Turma": "2 e 3", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Bilíngue Fund.", "Turma": "2 e 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 1", "Disciplina": "Polivalente", "Professor": "Juliana Aparecida da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 1", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 1", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 1", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 2", "Disciplina": "Polivalente", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 2", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 2", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 2", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 3", "Disciplina": "Polivalente", "Professor": "Marcela Buck de Gaspari"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 3", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 3", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "1º Ano 3", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},    
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 1", "Disciplina": "Polivalente", "Professor": "Iara Cristina Galdino"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 1", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 1", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 2", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},    
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 2", "Disciplina": "Polivalente", "Professor": "Natália dos Santos Lima Fula"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 2", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 2", "Disciplina": "Educação Física", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 3", "Disciplina": "Polivalente", "Professor": "Amanda Mussi"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 3", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},    
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 3", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 3", "Disciplina": "Educação Física", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "2º Ano 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 1", "Disciplina": "Polivalente", "Professor": "Marcia Regina Biserra Branco"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 1", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 1", "Disciplina": "Educação Física", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 1", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},    
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 2", "Disciplina": "Polivalente", "Professor": "Alessandra Rigon Ribeiro"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 2", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 2", "Disciplina": "Educação Física", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 2", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},    
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 3", "Disciplina": "Polivalente", "Professor": "Regiane Faustino"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 3", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 3", "Disciplina": "Educação Física", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "Ciclo I (1º ao 3º ano)", "Turma": "3º Ano 3", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},    
        
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Língua Portuguesa", "Professor": "Eliana Cristina de Carvalho Gabriel"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Matemática", "Professor": "Daiane Luzia de Matos Bueno"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Ciências, Hist. e Geo.", "Professor": "Valdineia Elisabete da Silva Augusto"}, 
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 1", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Língua Portuguesa", "Professor": "Eliana Cristina de Carvalho Gabriel"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Matemática", "Professor": "Daiane Luzia de Matos Bueno"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Ciências, Hist. e Geo.", "Professor": "Valdineia Elisabete da Silva Augusto"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 2", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Língua Portuguesa", "Professor": "Eliana Cristina de Carvalho Gabriel"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Matemática", "Professor": "Daiane Luzia de Matos Bueno"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Ciências, Hist. e Geo.", "Professor": "Valdineia Elisabete da Silva Augusto"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Josiane Modesto da Silva"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "4º Ano 3", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 1", "Disciplina": "Polivalente", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 1", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 1", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 1", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 2", "Disciplina": "Polivalente", "Professor": "Nathalia Teixeira Marcal Ribeiro"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 2", "Disciplina": "Artes", "Professor": "Bruna Thais Berini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 2", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 2", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 3", "Disciplina": "Polivalente", "Professor": "Denise Teixeira Coelho Soffiati"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 3", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 3", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Ciclo II (4º e 5º ano)", "Turma": "5º Ano 3", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},
        
        # --- EDUCAÇÃO INFANTIL ---
        {"Ciclo": "Bilíngue Inf.", "Turma": "1", "Disciplina": "Polivalente", "Professor": "Luciana Martinati Tetzner"},
        {"Ciclo": "Bilíngue Inf.", "Turma": "1", "Disciplina": "Artes", "Professor": "Karen Cristina Fernandes Donatti"},
        {"Ciclo": "Bilíngue Inf.", "Turma": "1", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "Bilíngue Inf.", "Turma": "1", "Disciplina": "Educação Física", "Professor": "Karina Brum Rivabene Gracini"},

        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 1", "Disciplina": "Professor de Ed. Infantil", "Professor": "Luciana Lopes Faber"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 1", "Disciplina": "Artes", "Professor": "Karen Cristina Fernandes Donatti"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 1", "Disciplina": "Educação Física", "Professor": "Michel Luciano de Lima"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Karina Brum Rivabene Gracini"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 1", "Disciplina": "Libras", "Professor": "Flaviany Miranda Mantovani"},

        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 2", "Disciplina": "Professor de Ed. Infantil", "Professor": "Sandra Maria Ribeiro dos Santos Brito"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 2", "Disciplina": "Artes", "Professor": "Karen Cristina Fernandes Donatti"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 2", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 2", "Disciplina": "Libras", "Professor": "Viviane Ayumi Yanase Doi Quesada"},

        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 3", "Disciplina": "Professor de Ed. Infantil", "Professor": "Amanda Mussi"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 3", "Disciplina": "Artes", "Professor": "Karen Cristina Fernandes Donatti"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 3", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "1ª Etapa", "Turma": "1ª Etapa 3", "Disciplina": "Libras", "Professor": "Flaviany Miranda Mantovani"},
        
        
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 1", "Disciplina": "Professor de Ed. Infantil", "Professor": "Keila Maria Ribeiro dos Santos"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 1", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 1", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "1ª Etapa", "Turma": "2ª Etapa 1", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},

        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 2", "Disciplina": "Professor de Ed. Infantil", "Professor": "Sandra Maria Ribeiro dos Santos Brito"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 2", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 2", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "1ª Etapa", "Turma": "2ª Etapa 2", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},

        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 3", "Disciplina": "Professor de Ed. Infantil", "Professor": "Adriana Mello Costa"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 3", "Disciplina": "Artes", "Professor": "Jordana Lima Alvez"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 3", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "2ª Etapa", "Turma": "2ª Etapa 3", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},
        {"Ciclo": "1ª Etapa", "Turma": "2ª Etapa 3", "Disciplina": "Libras", "Professor": "Solange Dalosto Campos"},
        
        {"Ciclo": "Maternal II", "Turma": "Maternal II 1", "Disciplina": "Professor de Ed. Infantil", "Professor": "Fernanda Dumit Graf"},
        {"Ciclo": "Maternal II", "Turma": "Maternal II 1", "Disciplina": "Artes", "Professor": "Bruna Thais Bernini Guedes"},
        {"Ciclo": "Maternal II", "Turma": "Maternal II 1", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Maternal II", "Turma": "Maternal II 1", "Disciplina": "Linguagens e Tecnologias", "Professor": "Karina Brum Rivabene Gracini"},

        {"Ciclo": "Maternal II", "Turma": "Maternal II 2", "Disciplina": "Professor de Ed. Infantil", "Professor": "Lucineide Almeida da Silva"},
        {"Ciclo": "Maternal II", "Turma": "Maternal II 2", "Disciplina": "Artes", "Professor": "Rebeca Grancieri da Cruz"},
        {"Ciclo": "Maternal II", "Turma": "Maternal II 2", "Disciplina": "Educação Física", "Professor": "Fernando Indig Bongiovanni"},
        {"Ciclo": "Maternal II", "Turma": "Maternal II 2", "Disciplina": "Linguagens e Tecnologias", "Professor": "Elaine Cristina Neves Fahl"},

        {"Ciclo": "Maternal I", "Turma": "Maternal I 1", "Disciplina": "Polivalente", "Professor": "Giovanna Guarache de Miranda"},
    ]
    matriz_json = get_config("matriz_professores", "")
    df_matriz = pd.DataFrame(json.loads(matriz_json)) if matriz_json else pd.DataFrame(MATRIZ_SEED)
    
    # Matriz Gestão Semente
    GESTAO_SEED = [
        {"Nome": "Luciana Lopes Faber", "Cargo": "Prof. Coordenador"},
        {"Nome": "Oelen Fernando Pedro", "Cargo": "Prof. Coordenador"},
        {"Nome": "Luciana Martinati Tetzner", "Cargo": "Vice-Diretor"},
        {"Nome": "Noreh Cristina Heldt Aldrigui", "Cargo": "Vice-Diretor"},
        {"Nome": "Marília Motta Camargo dos Reis", "Cargo": "Vice-Diretor"},
        {"Nome": "José Victor Souza Gallo", "Cargo": "Diretor de Escola"}
    ]
    gestao_json = get_config("matriz_gestao", "")
    df_gestao = pd.DataFrame(json.loads(gestao_json)) if gestao_json else pd.DataFrame(GESTAO_SEED)


    # ==============================================================================
    # 1. TELA: NOVA ATA DE CONSELHO
    # ==============================================================================
    if app_mode_regular == "📝 Nova Ata de Conselho":

        # 1. CRIE/DEFINA A VARIÁVEL AQUI ANTES DE QUALQUER COISA
        modalidade_ata = st.radio(
            "Selecione a modalidade da ata:", 
            ["Ensino Fundamental", "Educação Infantil"],
            horizontal=True
        )
        st.markdown(f"""<div class="header-box"><div class="header-title">Conselho de Classe / Termo</div><div class="header-subtitle">{modalidade_ata}</div></div>""", unsafe_allow_html=True)
        
        # ------------------------------------------------------------------------------
        # MÓDULO: ENSINO FUNDAMENTAL
        # ------------------------------------------------------------------------------
        if "Fundamental" in modalidade_ata:
            if 'data_ata_ef' not in st.session_state:
                st.session_state.data_ata_ef = {
                    'abaixo_basico': [{"Estudante": "", "LP": "", "M": "", "H": "", "G": "", "C": "", "A": "", "EF": "", "LT": "", "LIBRAS": ""}],
                    'basico': [{"Estudante": "", "Ações (LP e Mat)": ""}],
                    'obs_especiais': [{"Estudante": "", "Desempenho/Observação": ""}],
                    'encaminhamentos': [{"Estudante": "", "Motivo": ""}],
                    'mat_tardia': [{"Estudante": "", "Data Matrícula": "", "Total Frequência": ""}],
                    'obs_apc': "",
                    'obs_outras': "",
                    'assinaturas': [{"Nome": "", "Cargo/Atuação": ""}]
                }
                
            if 'ata_turma_confirmada' not in st.session_state:
                st.session_state.ata_turma_confirmada = None
                st.session_state.ata_ciclo_confirmado = None
            
            data_ata = st.session_state.data_ata_ef
            
            for key in ['abaixo_basico', 'basico', 'obs_especiais', 'encaminhamentos', 'mat_tardia', 'assinaturas']:
                if isinstance(data_ata.get(key), pd.DataFrame):
                    data_ata[key] = data_ata[key].to_dict('records')
            
        # --- PORTÃO DE ENTRADA (GATE) ---
            # --- PORTÃO DE ENTRADA (GATE) ---
            if not st.session_state.ata_turma_confirmada:
                st.markdown("""
                <div style='background-color: white; padding: 2rem; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); border: 1px solid #e2e8f0;'>
                    <h3 style='color: #1e293b; margin-top:0;'>🚪 Selecione a Turma e o Trimestre</h3>
                    <p style='color: #64748b;'>Se já existir uma ata salva para esta turma e trimestre, ela será carregada automaticamente para você continuar.</p>
                </div>
                <br>
                """, unsafe_allow_html=True)
                
                c_c, c_t, c_tri = st.columns([2, 2, 1])
                ciclo_sel = c_c.selectbox("1. Selecione o Ciclo:", ["Bilíngue Fund.","Ciclo I (1º ao 3º ano)", "Ciclo II (4º e 5º ano)"])
                
                turmas_bd = df_matriz[df_matriz['Ciclo'] == ciclo_sel]['Turma'].unique().tolist()
                turmas_disp = turmas_bd + ["Outra Turma..."]
                
                turma_sel = c_t.selectbox("2. Selecione a Turma:", turmas_disp)
                # AJUSTE 1: Trimestre sem valor padrão (obrigando a escolha)
                trimestre_sel = c_tri.selectbox("3. Trimestre:", ["1º Trimestre", "2º Trimestre", "3º Trimestre"], index=None, placeholder="Selecione...")
                
                if turma_sel == "Outra Turma...":
                    turma_sel = st.text_input("Digite o nome da turma:")
                
                st.write("")
                if st.button("✅ Confirmar e Acessar Formulário", type="primary", use_container_width=True):
                    # BARREIRA DE VALIDAÇÃO
                    if not turma_sel or not trimestre_sel:
                        st.warning("⚠️ Atenção: Por favor, selecione a Turma e o Trimestre antes de continuar.")
                    else:
                        # AJUSTE 2: Limpa o PDF da memória ao entrar em uma ata
                        st.session_state.pop('pdf_bytes_ata', None)
                        
                        id_buscado = f"{turma_sel} - {trimestre_sel} (Ensino Fundamental)"
                        df_atas = safe_read("Atas_Conselho", ["id_ata", "modalidade", "turma", "dados_json"])
                        
                        if not df_atas.empty and id_buscado in df_atas["id_ata"].values:
                            dados_row = df_atas[df_atas["id_ata"] == id_buscado].iloc[0]
                            st.session_state.data_ata_ef = json.loads(dados_row["dados_json"])
                            st.toast("Rascunho anterior carregado com sucesso!", icon="🔄")
                        else:
                            # AJUSTE 3: Assinaturas preenchidas no ato da criação
                            assinaturas_auto = []
                            professores_adicionados = set()
                            df_turma = df_matriz[(df_matriz['Ciclo'] == ciclo_sel) & (df_matriz['Turma'] == turma_sel)]
                            
                            if not df_turma.empty:
                                for _, row in df_turma.iterrows():
                                    materia = row['Disciplina']
                                    nome_prof = row['Professor']
                                    if nome_prof and nome_prof not in professores_adicionados:
                                        cargo_formatado = "Prof. Polivalente" if materia == "Polivalente" else f"Prof. de {materia}"
                                        assinaturas_auto.append({"Nome": nome_prof, "Cargo/Atuação": f"{cargo_formatado} (Atuante na Turma)"})
                                        professores_adicionados.add(nome_prof)
                                        
                                
                                df_ciclo = df_matriz[(df_matriz['Ciclo'] == ciclo_sel) & (df_matriz['Turma'] != turma_sel)]
                                for _, row in df_ciclo.iterrows():
                                    materia = row['Disciplina']
                                    nome_prof = row['Professor']
                                    if nome_prof and nome_prof not in professores_adicionados:
                                        cargo_formatado = "Prof. Polivalente" if materia == "Polivalente" else f"Prof. de {materia}"
                                        assinaturas_auto.append({"Nome": nome_prof, "Cargo/Atuação": f"{cargo_formatado} (Atuante no Ciclo)"})
                                        professores_adicionados.add(nome_prof)
                                        
                                for _, row in df_gestao.iterrows():
                                    if row['Nome']:
                                        assinaturas_auto.append({"Nome": row['Nome'], "Cargo/Atuação": row['Cargo']})
                                        
                            if not assinaturas_auto:
                                assinaturas_auto = [{"Nome": "", "Cargo/Atuação": ""}]

                            st.session_state.data_ata_ef = {
                                'abaixo_basico': [{"Estudante": "", "LP": "", "M": "", "H": "", "G": "", "C": "", "A": "", "EF": "", "LT": "", "LIBRAS": ""}],
                                'basico': [{"Estudante": "", "Ações (LP e Mat)": ""}],
                                'obs_especiais': [{"Estudante": "", "Desempenho/Observação": ""}],
                                'encaminhamentos': [{"Estudante": "", "Motivo": ""}],
                                'mat_tardia': [{"Estudante": "", "Data Matrícula": "", "Total Frequência": ""}],
                                'obs_apc': "", 
                                'obs_outras': "",
                                'assinaturas': assinaturas_auto, # Inserido automaticamente aqui
                                'ciclo': ciclo_sel,
                                'turma': turma_sel,
                                'trimestre': trimestre_sel
                            }
                        
                        st.session_state.ata_ciclo_confirmado = ciclo_sel
                        st.session_state.ata_turma_confirmada = turma_sel
                        st.rerun()
            
            # --- FORMULÁRIO DO FUNDAMENTAL ---
            else:
                c_info, c_btn = st.columns([4, 1])
                c_info.success(f"📌 **Ata em edição:** {st.session_state.ata_ciclo_confirmado} - {st.session_state.ata_turma_confirmada}")
                if c_btn.button("⬅️ Trocar Turma", use_container_width=True):
                    st.session_state.ata_turma_confirmada = None
                    st.session_state.pop('pdf_bytes_ata', None) # Limpa PDF ao voltar
                    st.rerun()
                    
                tabs = st.tabs(["1. Identificação", "2. Síntese", "3. Plano de Ação", "4. Observações", "5. Finalização"])
                
                with tabs[0]:
                    st.subheader("Dados da Unidade e Ciclo")
                    c1, c2, c3 = st.columns([2, 1, 1])
                    data_ata['escola'] = c1.text_input("Unidade Escolar", value=data_ata.get('escola', "CEIEF Rafael Affonso Leite"))
                    
                    tri_opts = ["1º Trimestre", "2º Trimestre", "3º Trimestre"]
                    tri_idx = tri_opts.index(data_ata.get('trimestre', "1º Trimestre")) if data_ata.get('trimestre') in tri_opts else 0
                    data_ata['trimestre'] = c2.selectbox("Trimestre", tri_opts, index=tri_idx)
                    
                    data_ata['ano_letivo'] = c3.text_input("Ano Letivo", value=data_ata.get('ano_letivo', str(date.today().year)))
                    
                    st.markdown("---")
                    c4, c5 = st.columns(2)
                    c4.text_input("Turma/Ano", value=st.session_state.ata_turma_confirmada, disabled=True)
                    c5.text_input("Ciclo", value=st.session_state.ata_ciclo_confirmado, disabled=True)

                with tabs[1]:
                    st.subheader("Síntese Avaliativa da Classe")
                    st.info("Descreva o desempenho alcançado pela classe em cada componente curricular no trimestre atual.")
                    
                    c_lp, c_mat = st.columns(2)
                    data_ata['sin_lp'] = c_lp.text_area("Língua Portuguesa", value=data_ata.get('sin_lp', ''), placeholder="Ex: Descreva o desenvolvimento geral da turma, destacando as aprendizagens consolidadas no nível Adequado e as dificuldades pontuais do nível Básico. Relate também as principais defasagens dos alunos Abaixo do Básico e as intervenções específicas direcionadas aos estudantes da inclusão (AEE).", height=120)
                    data_ata['sin_mat'] = c_mat.text_area("Matemática", value=data_ata.get('sin_mat', ''), height=120)
                    
                    c_h, c_g = st.columns(2)
                    data_ata['sin_hist'] = c_h.text_area("História", value=data_ata.get('sin_hist', ''), height=120)
                    data_ata['sin_geo'] = c_g.text_area("Geografia", value=data_ata.get('sin_geo', ''), height=120)
                    
                    c_c, c_a = st.columns(2)
                    data_ata['sin_cien'] = c_c.text_area("Ciências", value=data_ata.get('sin_cien', ''), height=120)
                    data_ata['sin_arte'] = c_a.text_area("Arte", value=data_ata.get('sin_arte', ''), height=120)
                    
                    c_ef, c_lt = st.columns(2)
                    data_ata['sin_ef'] = c_ef.text_area("Educação Física", value=data_ata.get('sin_ef', ''), height=120)
                    data_ata['sin_lt'] = c_lt.text_area("Linguagens e Tecnologias", value=data_ata.get('sin_lt', ''), height=120)
                    
                    data_ata['sin_libras'] = st.text_area("Libras", value=data_ata.get('sin_libras', ''), height=120)

                with tabs[2]:
                    st.subheader("Plano de Ação (Abaixo do Básico)")
                    st.caption("Insira os **números das propostas de recuperação** nas disciplinas correspondentes (ex: 1, 3, 10).")
                    
                    for i, row in enumerate(data_ata['abaixo_basico']):
                        with st.container():
                            st.markdown(f"**Estudante {i+1}**")
                            c_est, c_del = st.columns([11, 1])
                            row['Estudante'] = c_est.text_input(f"Nome do Estudante", value=row.get('Estudante', ''), key=f"ab_est_{i}", label_visibility="collapsed", placeholder="Nome do Aluno")
                            if c_del.button("🗑️", key=f"del_ab_{i}", help="Excluir Linha"):
                                data_ata['abaixo_basico'].pop(i); st.rerun()
                            
                            cc = st.columns(9)
                            row['LP'] = cc[0].text_input("LP", value=row.get('LP', ''), key=f"ab_lp_{i}")
                            row['M'] = cc[1].text_input("M", value=row.get('M', ''), key=f"ab_m_{i}")
                            row['H'] = cc[2].text_input("H", value=row.get('H', ''), key=f"ab_h_{i}")
                            row['G'] = cc[3].text_input("G", value=row.get('G', ''), key=f"ab_g_{i}")
                            row['C'] = cc[4].text_input("C", value=row.get('C', ''), key=f"ab_c_{i}")
                            row['A'] = cc[5].text_input("A", value=row.get('A', ''), key=f"ab_a_{i}")
                            row['EF'] = cc[6].text_input("EF", value=row.get('EF', ''), key=f"ab_ef_{i}")
                            row['LT'] = cc[7].text_input("LT", value=row.get('LT', ''), key=f"ab_lt_{i}")
                            row['LIBRAS'] = cc[8].text_input("LIB", value=row.get('LIBRAS', ''), key=f"ab_lib_{i}")
                            st.markdown("<hr style='margin: 10px 0;'>", unsafe_allow_html=True)
                    
                    if st.button("➕ Adicionar Novo Estudante", key="add_ab"):
                        data_ata['abaixo_basico'].append({"Estudante": "", "LP": "", "M": "", "H": "", "G": "", "C": "", "A": "", "EF": "", "LT": "", "LIBRAS": ""})
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown("**Propostas de Recuperação da Gestão:**")
                    st.markdown(propostas_ata_ef)
                    
                    st.divider()
                    st.subheader("Plano de Ação (Básico)")
                    st.caption("Cite todos os alunos correspondentes de uma só vez e defina o plano de ação único para cada área.")
                    
                    c_lp, c_mat = st.columns(2)
                    
                    with c_lp:
                        st.markdown("**📖 Língua Portuguesa**")
                        data_ata['basico_lp_estudantes'] = st.text_area(
                            "Alunos", 
                            value=data_ata.get('basico_lp_estudantes', ''), 
                            key="basico_lp_est", 
                            height=80, 
                            placeholder="Ex: João, Maria, Pedro..."
                        )
                        data_ata['basico_lp_acoes'] = st.text_area(
                            "Plano de Ação", 
                            value=data_ata.get('basico_lp_acoes', ''), 
                            key="basico_lp_ac", 
                            height=120, 
                            placeholder="Ações que serão desenvolvidas em LP..."
                        )
                        
                    with c_mat:
                        st.markdown("**🧮 Matemática**")
                        data_ata['basico_mat_estudantes'] = st.text_area(
                            "Alunos", 
                            value=data_ata.get('basico_mat_estudantes', ''), 
                            key="basico_mat_est", 
                            height=80, 
                            placeholder="Ex: Lucas, Ana, Beatriz..."
                        )
                        data_ata['basico_mat_acoes'] = st.text_area(
                            "Plano de Ação", 
                            value=data_ata.get('basico_mat_acoes', ''), 
                            key="basico_mat_ac", 
                            height=120, 
                            placeholder="Ações que serão desenvolvidas em Matemática..."
                        )

                with tabs[3]:
                    st.subheader("3. Observações Gerais")
                    
                    st.markdown("**a) Desempenho de alunos especiais (laudados)**")
                    for i, row in enumerate(data_ata['obs_especiais']):
                        c1, c2, c3 = st.columns([3, 6, 1])
                        row['Estudante'] = c1.text_input("Estudante Especial", value=row.get('Estudante', ''), key=f"obs_est_{i}")
                        row['Desempenho/Observação'] = c2.text_area("Desempenho/Observações - Preenchimento Exclusivo AEE", value=row.get('Desempenho/Observação', ''), key=f"obs_des_{i}", height=68)
                        if c3.button("🗑️", key=f"del_obs_{i}"):
                            data_ata['obs_especiais'].pop(i); st.rerun()
                    if st.button("➕ Adicionar Aluno Especial", key="add_obs"):
                        data_ata['obs_especiais'].append({"Estudante": "", "Desempenho/Observação": ""})
                        st.rerun()
                    
                    st.divider()
                    st.markdown("**b) Alunos encaminhados (Conselho Tutelar ou Serviço Social)**")
                    for i, row in enumerate(data_ata['encaminhamentos']):
                        c1, c2, c3 = st.columns([3, 6, 1])
                        row['Estudante'] = c1.text_input("Estudante Encaminhado", value=row.get('Estudante', ''), key=f"enc_est_{i}")
                        row['Motivo'] = c2.text_area("Motivo do Encaminhamento - Preenchimento pelo Professor Polivalente", value=row.get('Motivo', ''), key=f"enc_mot_{i}", height=68)
                        if c3.button("🗑️", key=f"del_enc_{i}"):
                            data_ata['encaminhamentos'].pop(i); st.rerun()
                    if st.button("➕ Adicionar Encaminhamento", key="add_enc"):
                        data_ata['encaminhamentos'].append({"Estudante": "", "Motivo": ""})
                        st.rerun()

                    st.divider()
                    st.markdown("**c) Estudantes Matriculados Tardiamente**")
                    for i, row in enumerate(data_ata['mat_tardia']):
                        c1, c2, c3, c4 = st.columns([4, 3, 3, 1])
                        row['Estudante'] = c1.text_input("Estudante", value=row.get('Estudante', ''), key=f"mat_est_{i}")
                        row['Data Matrícula'] = c2.text_input("Data da Matrícula", value=row.get('Data Matrícula', ''), key=f"mat_data_{i}")
                        row['Total Frequência'] = c3.text_input("Frequência (Dias)", value=row.get('Total Frequência', ''), key=f"mat_freq_{i}")
                        if c4.button("🗑️", key=f"del_mat_{i}"):
                            data_ata['mat_tardia'].pop(i); st.rerun()
                    if st.button("➕ Adicionar Matrícula Tardia", key="add_mat"):
                        data_ata['mat_tardia'].append({"Estudante": "", "Data Matrícula": "", "Total Frequência": ""})
                        st.rerun()

                    st.divider()
                    
                    # ================= NOVO BLOCO APC =================
                    st.markdown("**d) Alunos matriculados na Ação Pedagógica Complementar (APC)**")
                    data_ata['obs_apc'] = st.text_area("Resumo/Observações sobre os estudantes da APC:", value=data_ata.get('obs_apc', ''), height=100)
                    st.divider()
                    # ==================================================
                    
                    st.divider()
                    st.markdown("**e) Outras Observações**")
                    data_ata['obs_outras'] = st.text_area("Campo livre para quaisquer outras observações da turma:", value=data_ata.get('obs_outras', ''), height=120)

                with tabs[4]:
                    st.subheader("Finalização e Assinaturas")
                    
                    if st.button("🤖 Preencher Assinaturas Automaticamente", type="primary"):
                        ciclo_atual = st.session_state.ata_ciclo_confirmado
                        turma_atual = st.session_state.ata_turma_confirmada
                        
                        df_turma = df_matriz[(df_matriz['Ciclo'] == ciclo_atual) & (df_matriz['Turma'] == turma_atual)]
                        
                        if not df_turma.empty:
                            lista_final = []
                            professores_adicionados = set()
                            
                            for _, row in df_turma.iterrows():
                                materia = row['Disciplina']
                                nome_prof = row['Professor']
                                if nome_prof and nome_prof not in professores_adicionados:
                                    cargo_formatado = "Prof. Polivalente" if materia == "Polivalente" else f"Prof. de {materia}"
                                    lista_final.append({"Nome": nome_prof, "Cargo/Atuação": f"{cargo_formatado} (Atuante na Turma)"})
                                    professores_adicionados.add(nome_prof)
                                    
                            
                            df_ciclo = df_matriz[(df_matriz['Ciclo'] == ciclo_atual) & (df_matriz['Turma'] != turma_atual)]
                            for _, row in df_ciclo.iterrows():
                                materia = row['Disciplina']
                                nome_prof = row['Professor']
                                if nome_prof and nome_prof not in professores_adicionados:
                                    cargo_formatado = "Prof. Polivalente" if materia == "Polivalente" else f"Prof. de {materia}"
                                    lista_final.append({"Nome": nome_prof, "Cargo/Atuação": f"{cargo_formatado} (Atuante no Ciclo)"})
                                    professores_adicionados.add(nome_prof)
                                            
                            for _, row in df_gestao.iterrows():
                                if row['Nome']:
                                    lista_final.append({"Nome": row['Nome'], "Cargo/Atuação": row['Cargo']})
                            
                            data_ata['assinaturas'] = lista_final
                            st.success("✅ Grade de assinaturas preenchida com sucesso para a turma selecionada!")
                            st.rerun()
                        else:
                            st.error("A turma atual não foi encontrada na Matriz de Automação (Acesse a aba 'Configurações' para verificar).")

                    st.divider()
                    st.markdown("**Participantes da Reunião**")
                    st.caption("Você pode usar as setinhas (⬆️/⬇️) para mudar a ordem que os nomes vão aparecer na folha do PDF.")
                    
                    if 'assinaturas' not in data_ata:
                        data_ata['assinaturas'] = [{"Nome": "", "Cargo/Atuação": ""}]
                        
                    for i, row in enumerate(data_ata['assinaturas']):
                        c1, c2, c3 = st.columns([5, 4, 2])
                        row['Nome'] = c1.text_input("Nome", value=row.get('Nome', ''), key=f"sig_nome_{i}", placeholder="Ex: José Victor Souza Gallo", label_visibility="collapsed")
                        row['Cargo/Atuação'] = c2.text_input("Cargo", value=row.get('Cargo/Atuação', ''), key=f"sig_cargo_{i}", placeholder="Ex: Diretor de Escola", label_visibility="collapsed")
                        
                        b1, b2, b3 = c3.columns(3)
                        if b1.button("⬆️", key=f"up_sig_{i}", disabled=(i == 0), help="Mover para cima"):
                            data_ata['assinaturas'][i], data_ata['assinaturas'][i-1] = data_ata['assinaturas'][i-1], data_ata['assinaturas'][i]
                            st.rerun()
                        if b2.button("⬇️", key=f"dw_sig_{i}", disabled=(i == len(data_ata['assinaturas']) - 1), help="Mover para baixo"):
                            data_ata['assinaturas'][i], data_ata['assinaturas'][i+1] = data_ata['assinaturas'][i+1], data_ata['assinaturas'][i]
                            st.rerun()
                        if b3.button("🗑️", key=f"del_sig_{i}", help="Excluir assinatura"):
                            data_ata['assinaturas'].pop(i)
                            st.rerun()
                            
                    if st.button("➕ Adicionar Participante Manualmente", key="add_sig"):
                        data_ata['assinaturas'].append({"Nome": "", "Cargo/Atuação": ""})
                        st.rerun()
                        
                    st.divider()
                    
                    if st.button("💾 Salvar Ata", use_container_width=True, type="secondary"):
                        try:
                            dados_para_salvar = {}
                            for key, value in data_ata.items():
                                if isinstance(value, pd.DataFrame):
                                    dados_para_salvar[key] = value.to_dict(orient='records')
                                else:
                                    dados_para_salvar[key] = value
                            
                            novo_json = json.dumps(dados_para_salvar, ensure_ascii=False)
                            id_ata = f"{data_ata.get('turma', 'SemTurma')} - {data_ata.get('trimestre', 'SemTri')} ({modalidade_ata})"
                            
                            # Cria o registro isolado
                            novo_registro = {
                                "id_ata": id_ata, 
                                "modalidade": modalidade_ata, 
                                "turma": data_ata.get('turma', ''), 
                                "dados_json": novo_json
                            }
                            
                            # SALVAMENTO CIRÚRGICO (UPSERT): Insere a ata nova ou atualiza a existente. 
                            # Esqueça o safe_update! Ele não vai mais apagar a tabela.
                            supabase.table("Atas_Conselho").upsert(novo_registro).execute()
                            
                            st.success(f"✅ Ata salva com segurança!")
                        except Exception as e:
                            st.error(f"Erro ao salvar: {e}")


                    if st.button("👁️ GERAR ATA COMPLETA (PDF)", type="primary", use_container_width=True):
                        try:
                            pdf = OfficialPDF('P', 'mm', 'A4')
                            pdf.doc_type = "Ata" 
                            pdf.set_margins(15, 35, 15)
                            pdf.set_auto_page_break(auto=True, margin=20)
                            pdf.add_page()
                            
                            def calc_lines(txt, w):
                                # A MÁGICA VEM AQUI: Limpa o texto antes da biblioteca tentar medir!
                                txt = clean_pdf_text(txt)
                                
                                if not txt: return 1
                                lines = 0
                                for par in txt.split('\n'):
                                    words = par.split(' ')
                                    curr_w = 0
                                    for word in words:
                                        word_w = pdf.get_string_width(word + ' ') # O erro acontecia nesta linha!
                                        if curr_w + word_w > w:
                                            lines += 1; curr_w = word_w
                                        else:
                                            curr_w += word_w
                                    lines += 1
                                return lines
                            
                            # --- CABEÇALHO ---
                            pdf.set_font("Arial", "B", 10)
                            escola_nome = data_ata.get('escola', 'CEIEF Rafael Affonso Leite').upper()
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(escola_nome), 0, 1, 'C')
                            
                            trimestre = data_ata.get('trimestre', '').upper()
                            ano = data_ata.get('ano_letivo', '')
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(f"ATA DESCRITIVA DO CONSELHO DE CICLO/TERMO - {trimestre} DE {ano}"), 0, 1, 'C')
                            
                            turma = data_ata.get('turma', '')
                            ciclo = data_ata.get('ciclo', '')
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(f"Turma: {turma} | Ciclo: {ciclo}"), 0, 1, 'C')
                            pdf.ln(5)
                            
                            # ==============================================================================
                            # INÍCIO DA SUPER CAIXA
                            # ==============================================================================
                            
                            # --- 1. SÍNTESE AVALIATIVA ---
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_fill_color(220, 220, 220)
                            pdf.set_x(15)
                            pdf.cell(180, 6, "SÍNTESE AVALIATIVA", "LTR", 1, 'C', True)
                            
                            pdf.set_x(15)
                            pdf.cell(180, 2, "", "LR", 1) 
                            
                            pdf.set_font("Arial", "", 10)
                            pdf.set_x(15)
                            pdf.multi_cell(180, 5, clean_pdf_text(texto_base_ata_ef), "LR", 'J')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 4, "", "LR", 1)
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("1- Síntese avaliativa da classe:"), "LR", 1, 'L')
                            
                            texto_sint = "A partir dos diferentes instrumentos avaliativos e da análise dos resultados, o desempenho alcançado pela classe em cada componente curricular no trimestre atual se apresenta da seguinte forma:"
                            pdf.set_font("Arial", "", 10)
                            pdf.set_x(15)
                            pdf.multi_cell(180, 5, clean_pdf_text(texto_sint), "LR", 'J')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 4, "", "LR", 1)
                            
                            disciplinas = [
                                ("Língua Portuguesa", data_ata.get('sin_lp', '')),
                                ("Matemática", data_ata.get('sin_mat', '')),
                                ("História", data_ata.get('sin_hist', '')),
                                ("Geografia", data_ata.get('sin_geo', '')),
                                ("Ciências", data_ata.get('sin_cien', '')),
                                ("Arte", data_ata.get('sin_arte', '')),
                                ("Educação Física", data_ata.get('sin_ef', '')),
                                ("Linguagens e Tecnologias", data_ata.get('sin_lt', '')),
                                ("Libras", data_ata.get('sin_libras', ''))
                            ]
                            
                            for i, (nome, texto) in enumerate(disciplinas):
                                if texto.strip() != "":
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {nome}:"), "LR", 1, 'L')
                                    
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {texto}"), "LR", 'J')
                                
                                if i < len(disciplinas) - 1:
                                    pdf.set_x(15)
                                    pdf.cell(180, 4, "", "LR", 1) 
                            
                            # --- 2. PLANO DE AÇÃO ---
                            pdf.set_x(15)
                            pdf.cell(180, 5, "", "LR", 1) 
                            
                            if pdf.get_y() > 230: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("2- Plano de Ação para os estudantes de acordo com desempenho:"), "LR", 1, 'L')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 3, "", "LR", 1)
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("-Estudantes com desempenho Abaixo do Básico:"), "LR", 1, 'L')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 2, "", "LR", 1)
                            
                            pdf.set_font("Arial", "B", 10)
                            col_w = [54, 14, 14, 14, 14, 14, 14, 14, 14, 14]
                            headers = ["Estudante", "LP", "M", "H", "G", "C", "A", "EF", "LT", "LIB"]
                            pdf.set_x(15)
                            for i, h in enumerate(headers):
                                pdf.cell(col_w[i], 6, h, 1, 0, 'C')
                            pdf.ln()
                            
                            pdf.set_font("Arial", "", 4)
                            lista_abaixo = data_ata.get('abaixo_basico', [])
                                
                            def truncate_str(texto, max_w):
                                # Limpa antes de usar no laço while
                                texto = clean_pdf_text(texto)
                                
                                while pdf.get_string_width(texto) > max_w - 2: 
                                    texto = texto[:-1]
                                return texto

                            for row in lista_abaixo:
                                estudante = str(row.get('Estudante', '')).strip()
                                if estudante: 
                                    pdf.set_x(15)
                                    estudante_seguro = truncate_str(estudante, col_w[0])
                                    pdf.cell(col_w[0], 6, clean_pdf_text(estudante_seguro), 1, 0, 'L')
                                    pdf.cell(col_w[1], 6, clean_pdf_text(str(row.get('LP', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[2], 6, clean_pdf_text(str(row.get('M', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[3], 6, clean_pdf_text(str(row.get('H', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[4], 6, clean_pdf_text(str(row.get('G', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[5], 6, clean_pdf_text(str(row.get('C', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[6], 6, clean_pdf_text(str(row.get('A', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[7], 6, clean_pdf_text(str(row.get('EF', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[8], 6, clean_pdf_text(str(row.get('LT', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[9], 6, clean_pdf_text(str(row.get('LIBRAS', ''))), 1, 1, 'C')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 3, "", "LR", 1) 
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("*Propostas de Recuperação:"), "LR", 1, 'L')
                            pdf.set_font("Arial", "", 10)
                            
                            for prop in propostas_ata_ef.split('\n'):
                                if prop.strip():
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(prop.strip()), "LR", 'J')

                            pdf.set_x(15)
                            pdf.cell(180, 5, "", "LR", 1) 
                            
                            if pdf.get_y() > 230: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("-Estudantes com desempenho Básico:"), "LR", 1, 'L')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 2, "", "LR", 1) 

                            # --- FUNÇÃO AUXILIAR PARA DESENHAR AS DISCIPLINAS ---
                            def desenhar_bloco_basico(titulo, estudantes, acoes):
                                if not estudantes.strip() and not acoes.strip(): return
                                
                                # Pega a posição atual diretamente do gerador de PDF
                                y = pdf.get_y() 
                                
                                if y > 250:
                                    pdf.set_x(15); pdf.cell(180, 1, "", "LRB", 1) # Fecha borda
                                    pdf.add_page()
                                    pdf.set_x(15); pdf.cell(180, 2, "", "LTR", 1) # Abre borda
                                    y = pdf.get_y() # Atualiza Y após a nova página

                                # Título da disciplina
                                pdf.set_fill_color(240, 240, 240)
                                pdf.rect(15, y, 180, 6, 'FD')
                                pdf.set_xy(15, y)
                                pdf.set_font("Arial", "B", 10)
                                pdf.cell(180, 6, clean_pdf_text(titulo), 0, 1, 'C')
                                
                                y = pdf.get_y() # Atualiza Y após inserir o título

                                # Calcula a altura necessária baseado no texto longo
                                linhas_est = calc_lines(estudantes, 58)
                                linhas_acao = calc_lines(acoes, 118)
                                h_row = max(10, max(linhas_est, linhas_acao) * 5 + 4)

                                if y + h_row > 275:
                                    pdf.set_x(15); pdf.cell(180, 1, "", "LRB", 1)
                                    pdf.add_page()
                                    pdf.set_x(15); pdf.cell(180, 2, "", "LTR", 1)
                                    y = pdf.get_y() # Atualiza Y após a nova página

                                # Desenha os blocos
                                pdf.rect(15, y, 60, h_row)
                                pdf.rect(75, y, 120, h_row)

                                pdf.set_font("Arial", "", 10)
                                pdf.set_xy(15, y + 2)
                                pdf.multi_cell(60, 5, clean_pdf_text(estudantes), 0, 'L')
                                pdf.set_xy(75, y + 2)
                                pdf.multi_cell(120, 5, clean_pdf_text(acoes), 0, 'J')

                                pdf.set_xy(15, y + h_row)
                                pdf.cell(180, 2, "", "LR", 1) # Espaçamento pós bloco

                            # --- CHAMADA PARA RENDERIZAR OS BLOCOS ---
                            lp_est = data_ata.get('basico_lp_estudantes', '')
                            lp_ac  = data_ata.get('basico_lp_acoes', '')
                            mat_est = data_ata.get('basico_mat_estudantes', '')
                            mat_ac  = data_ata.get('basico_mat_acoes', '')

                            desenhar_bloco_basico("Plano de Ação - Língua Portuguesa", lp_est, lp_ac)
                            desenhar_bloco_basico("Plano de Ação - Matemática", mat_est, mat_ac)

                            # --- 3. OBSERVAÇÕES GERAIS ---
                            pdf.set_x(15)
                            pdf.cell(180, 5, "", "LR", 1) 
                            
                            if pdf.get_y() > 230: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 6, clean_pdf_text("3- Observações Gerais:"), "LR", 1, 'L')
                            pdf.set_font("Arial", "", 10)
                            
                            prefix_code = 97 
                            
                            lista_esp = data_ata.get('obs_especiais', [])
                            esp_valid = [r for r in lista_esp if str(r.get('Estudante', '')).strip()]
                            if esp_valid:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Desempenho de alunos especiais (laudados):"), "LR", 1, 'L')
                                prefix_code += 1
                                for row in esp_valid:
                                    est = str(row.get('Estudante', '')).strip()
                                    obs = str(row.get('Desempenho/Observação', '')).strip()
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {est}:"), "LR", 1, 'L')
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {obs}"), "LR", 'J')
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)
                            
                            lista_enc = data_ata.get('encaminhamentos', [])
                            enc_valid = [r for r in lista_enc if str(r.get('Estudante', '')).strip()]
                            if enc_valid:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Alunos encaminhados (Conselho Tutelar/Serv. Social):"), "LR", 1, 'L')
                                prefix_code += 1
                                for row in enc_valid:
                                    est = str(row.get('Estudante', '')).strip()
                                    mot = str(row.get('Motivo', '')).strip()
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {est}:"), "LR", 1, 'L')
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {mot}"), "LR", 'J')
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)

                            lista_tardia = data_ata.get('mat_tardia', [])
                            tardia_valid = [r for r in lista_tardia if str(r.get('Estudante', '')).strip()]
                            if len(tardia_valid) > 0:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Estudantes Matriculados Tardiamente:"), "LR", 1, 'L')
                                prefix_code += 1
                                for row in tardia_valid:
                                    est = str(row.get('Estudante', '')).strip()
                                    mat = str(row.get('Data Matrícula', '')).strip()
                                    freq = str(row.get('Total Frequência', '')).strip()
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {est}:"), "LR", 1, 'L')
                                    texto_tardio = f"Matriculado(a) nesta sala em {mat}. Portanto, o período correspondente à sua matrícula compreende um total de {freq} dias letivos."
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {texto_tardio}"), "LR", 'J')
                            else:
                                pdf.set_font("Arial", "", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, f" {chr(prefix_code)}) Sem matrículas tardias registradas no período.", "LR", 1)
                                prefix_code += 1

                            pdf.set_x(15)
                            pdf.cell(180, 2, "", "LR", 1)

                            # ================= NOVO BLOCO APC NO PDF =================
                            obs_apc = data_ata.get('obs_apc', '').strip()
                            if obs_apc:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Alunos matriculados na Ação Pedagógica Complementar (APC):"), "LR", 1, 'L')
                                prefix_code += 1
                                
                                pdf.set_font("Arial", "", 10)
                                pdf.set_x(15)
                                pdf.multi_cell(180, 5, clean_pdf_text(f"  {obs_apc}"), "LR", 'J')
                                
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)
                            # =========================================================
                            
                            obs_outras = data_ata.get('obs_outras', '').strip()
                            if obs_outras:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Outras Observações:"), "LR", 1, 'L')
                                pdf.set_font("Arial", "", 10)
                                pdf.set_x(15)
                                pdf.multi_cell(180, 5, clean_pdf_text(f"  {obs_outras}"), "LR", 'J')
                                
                            pdf.set_x(15)
                            pdf.cell(180, 3, "", "LRB", 1) 

                            # --- ASSINATURAS ---
                            pdf.ln(5)
                            if pdf.get_y() > 220: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_fill_color(220, 220, 220)
                            pdf.set_x(15)
                            pdf.cell(180, 6, "ASSINATURA DOS PARTICIPANTES NA REUNIÃO DO CONSELHO DE CICLO", 1, 1, 'C', True)
                            
                            lista_assinaturas = data_ata.get('assinaturas', [])
                            sigs_validas = [s for s in lista_assinaturas if str(s.get('Nome', '')).strip()]
                            
                            if not sigs_validas:
                                pdf.set_x(15)
                                pdf.cell(180, 20, "Nenhuma assinatura cadastrada para esta ata.", 1, 1, 'C')
                            else:
                                cols = 4
                                cell_w = 180 / cols
                                cell_h = 24 
                                
                                x_start = 15
                                y = pdf.get_y()
                                
                                for i, sig in enumerate(sigs_validas):
                                    col = i % cols
                                    if col == 0 and i > 0:
                                        y += cell_h
                                    
                                    if y + cell_h > 275:
                                        pdf.add_page()
                                        y = pdf.get_y()
                                    
                                    x = x_start + (col * cell_w)
                                    pdf.rect(x, y, cell_w, cell_h)
                                    
                                    pdf.line(x + 4, y + 10, x + cell_w - 4, y + 10)
                                    
                                    nome = clean_pdf_text(str(sig.get('Nome', '')).strip())
                                    cargo_full = clean_pdf_text(str(sig.get('Cargo/Atuação', '')).strip())
                                    
                                    cargo = cargo_full
                                    atuacao = ""
                                    if "(" in cargo_full:
                                        parts = cargo_full.split("(")
                                        cargo = parts[0].strip()
                                        atuacao = parts[1].replace(")", "").strip()
                                    
                                    font_size = 7
                                    pdf.set_font("Arial", "B", font_size)
                                    while pdf.get_string_width(nome) > cell_w - 2 and font_size > 4.5:
                                        font_size -= 0.5
                                        pdf.set_font("Arial", "B", font_size)
                                    
                                    pdf.set_xy(x + 1, y + 11) 
                                    pdf.multi_cell(cell_w - 2, 3, clean_pdf_text(nome), 0, 'C')
                                    
                                    pdf.set_font("Arial", "", 6)
                                    curr_y = pdf.get_y()
                                    pdf.set_xy(x + 1, curr_y)
                                    pdf.multi_cell(cell_w - 2, 3, clean_pdf_text(cargo), 0, 'C')
                                    
                                    if atuacao:
                                        curr_y = pdf.get_y()
                                        pdf.set_xy(x + 1, curr_y)
                                        pdf.multi_cell(cell_w - 2, 3, clean_pdf_text(atuacao), 0, 'C')
                                
                                pdf.set_y(y + cell_h)

                            st.session_state.pdf_bytes_ata = get_pdf_bytes(pdf)
                            st.success("✅ PDF gerado com sucesso!")
                        except Exception as e:
                            st.error(f"Erro ao desenhar o PDF: {e}")

                if 'pdf_bytes_ata' in st.session_state:
                    turma_limpa = str(data_ata.get('turma', 'Turma')).replace('/', '-').replace('\\', '-')
                    trimestre_limpo = str(data_ata.get('trimestre', 'Trimestre')).replace('/', '-')
                    st.download_button("📥 BAIXAR ATA", st.session_state.pdf_bytes_ata, f"Ata_{turma_limpa}_{trimestre_limpo}.pdf", "application/pdf", type="primary")

        # ------------------------------------------------------------------------------
        # MÓDULO: EDUCAÇÃO INFANTIL
        # ------------------------------------------------------------------------------
        elif "Infantil" in modalidade_ata:
            if 'data_ata_inf' not in st.session_state:
                st.session_state.data_ata_inf = {
                    'abaixo_basico': [{"Estudante": "", "LV": "", "LM": "", "IS": "", "A": "", "CCM": "", "LT": "", "LIBRAS": ""}],
                    'obs_especiais': [{"Estudante": "", "Desempenho/Observação": ""}],
                    'encaminhamentos': [{"Estudante": "", "Motivo": ""}],
                    'mat_tardia': [{"Estudante": "", "Data Matrícula": "", "Total Frequência": ""}],
                    'obs_outras': "",
                    'assinaturas': [{"Nome": "", "Cargo/Atuação": ""}]
                }
                
            if 'ata_turma_confirmada_inf' not in st.session_state:
                st.session_state.ata_turma_confirmada_inf = None
                st.session_state.ata_ciclo_confirmado_inf = None
            
            data_inf = st.session_state.data_ata_inf
            
            for key in ['abaixo_basico', 'obs_especiais', 'encaminhamentos', 'mat_tardia', 'assinaturas']:
                if isinstance(data_inf.get(key), pd.DataFrame):
                    data_inf[key] = data_inf[key].to_dict('records')
            
# --- PORTÃO DE ENTRADA (GATE) ---
            # --- PORTÃO DE ENTRADA (GATE) ---
            if not st.session_state.ata_turma_confirmada_inf:
                st.markdown("""
                <div style='background-color: #fdf4ff; padding: 2rem; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); border: 1px solid #f5d0fe;'>
                    <h3 style='color: #86198f; margin-top:0;'>🧸 Selecione a Turma e o Trimestre</h3>
                    <p style='color: #a21caf;'>Se já existir uma ata salva para esta turma e trimestre, ela será carregada automaticamente para você continuar.</p>
                </div>
                <br>
                """, unsafe_allow_html=True)
                
                c_c, c_t, c_tri = st.columns([2, 2, 1])
                ciclo_sel = c_c.selectbox("1. Selecione a Fase/Etapa:", ["Bilíngue Inf.", "Maternal I", "Maternal II", "1ª Etapa", "2ª Etapa", "Educação Infantil"])
                
                if ciclo_sel == "Educação Infantil":
                    turmas_bd = df_matriz[df_matriz['Ciclo'].str.contains("Infantil|Etapa|Maternal", na=False)]['Turma'].unique().tolist()
                else:
                    turmas_bd = df_matriz[df_matriz['Ciclo'] == ciclo_sel]['Turma'].unique().tolist()
                
                turmas_disp = turmas_bd + ["Outra Turma..."]
                
                turma_sel = c_t.selectbox("2. Selecione a Turma:", turmas_disp)
                # AJUSTE 1: Index nulo
                trimestre_sel = c_tri.selectbox("3. Trimestre:", ["1º Trimestre", "2º Trimestre", "3º Trimestre"], key="tri_gate_inf", index=None, placeholder="Selecione...")
                
                if turma_sel == "Outra Turma...":
                    turma_sel = st.text_input("Digite o nome da turma:")
                
                st.write("")
                if st.button("✅ Confirmar e Acessar Formulário do Infantil", type="primary", use_container_width=True):
                    if not turma_sel or not trimestre_sel:
                        st.warning("⚠️ Atenção: Por favor, selecione a Turma e o Trimestre antes de continuar.")
                    else:
                        st.session_state.pop('pdf_bytes_ata', None) # Limpa cache PDF
                        
                        id_buscado = f"{turma_sel} - {trimestre_sel} (Infantil)"
                        df_atas = safe_read("Atas_Conselho", ["id_ata", "modalidade", "turma", "dados_json"])
                        
                        if not df_atas.empty and id_buscado in df_atas["id_ata"].values:
                            dados_row = df_atas[df_atas["id_ata"] == id_buscado].iloc[0]
                            st.session_state.data_ata_inf = json.loads(dados_row["dados_json"])
                            st.toast("Rascunho anterior carregado com sucesso!", icon="🔄")
                        else:
                            # Preenchimento automático Infantil
                            assinaturas_auto = []
                            professores_adicionados = set()
                            df_turma = df_matriz[(df_matriz['Turma'] == turma_sel)]
                            
                            if not df_turma.empty:
                                for _, row in df_turma.iterrows():
                                    materia = row['Disciplina']
                                    nome_prof = row['Professor']
                                    if nome_prof and nome_prof not in professores_adicionados:
                                        assinaturas_auto.append({"Nome": nome_prof, "Cargo/Atuação": f"{materia} (Atuante na Turma)"})
                                        professores_adicionados.add(nome_prof)
                                        
                                for _, row in df_gestao.iterrows():
                                    if row['Nome']:
                                        assinaturas_auto.append({"Nome": row['Nome'], "Cargo/Atuação": row['Cargo']})
                                        
                            if not assinaturas_auto:
                                assinaturas_auto = [{"Nome": "", "Cargo/Atuação": ""}]

                            st.session_state.data_ata_inf = {
                                'abaixo_basico': [{"Estudante": "", "LV": "", "LM": "", "IS": "", "A": "", "CCM": "", "LT": "", "LIBRAS": ""}],
                                'obs_especiais': [{"Estudante": "", "Desempenho/Observação": ""}],
                                'encaminhamentos': [{"Estudante": "", "Motivo": ""}],
                                'mat_tardia': [{"Estudante": "", "Data Matrícula": "", "Total Frequência": ""}],
                                'obs_outras': "",
                                'assinaturas': assinaturas_auto, # Assinaturas incluídas!
                                'ciclo': ciclo_sel,
                                'turma': turma_sel,
                                'trimestre': trimestre_sel
                            }
                        
                        st.session_state.ata_ciclo_confirmado_inf = ciclo_sel
                        st.session_state.ata_turma_confirmada_inf = turma_sel
                        st.rerun()
                        
            # --- FORMULÁRIO DA EDUCAÇÃO INFANTIL ---
            else:
                c_info, c_btn = st.columns([4, 1])
                c_info.success(f"🧸 **Ata do Infantil em edição:** {st.session_state.ata_ciclo_confirmado_inf} - {st.session_state.ata_turma_confirmada_inf}")
                if c_btn.button("⬅️ Trocar Turma", use_container_width=True, key="btn_trocar_inf"):
                    st.session_state.ata_turma_confirmada_inf = None
                    st.session_state.pop('pdf_bytes_ata', None) # Limpa cache PDF ao voltar
                    st.rerun()
                
                criterios_etapa = get_criterios_infantil(st.session_state.ata_ciclo_confirmado_inf)
                    
                tabs = st.tabs(["1. Identificação", "2. Síntese Avaliativa", "3. Plano de Ação", "4. Observações", "5. Finalização"])
                
                with tabs[0]:
                    st.subheader("Dados da Unidade Escolar")
                    c1, c2, c3 = st.columns([2, 1, 1])
                    data_inf['escola'] = c1.text_input("Unidade Escolar", value=data_inf.get('escola', "CEIEF Rafael Affonso Leite"), key="inf_esc")
                    
                    tri_opts = ["1º Trimestre", "2º Trimestre", "3º Trimestre"]
                    tri_idx = tri_opts.index(data_inf.get('trimestre', "1º Trimestre")) if data_inf.get('trimestre') in tri_opts else 0
                    data_inf['trimestre'] = c2.selectbox("Trimestre", tri_opts, index=tri_idx, key="inf_tri")
                    
                    data_inf['ano_letivo'] = c3.text_input("Ano Letivo", value=data_inf.get('ano_letivo', str(date.today().year)), key="inf_ano")
                    
                    st.markdown("---")
                    c4, c5 = st.columns(2)
                    c4.text_input("Turma", value=st.session_state.ata_turma_confirmada_inf, disabled=True, key="inf_turma")
                    c5.text_input("Etapa/Fase", value=st.session_state.ata_ciclo_confirmado_inf, disabled=True, key="inf_ciclo")

                with tabs[1]:
                    st.subheader("Síntese Avaliativa da Classe - Campos de Experiência e Disciplinas")
                    st.info("Abaixo de cada campo, digite a síntese correspondente à turma. Os conteúdos avaliados (configurados para esta etapa) já estão listados como referência e constarão no PDF gerado.")
                    
                    st.markdown("**Linguagem Verbal:** descrever o desenvolvimento da classe referente a:")
                    if criterios_etapa["LV"]: st.caption(f"_{criterios_etapa['LV']}_")
                    data_inf['sin_lv'] = st.text_area("Síntese - Linguagem Verbal", value=data_inf.get('sin_lv', ''), height=100, label_visibility="collapsed", key="txt_lv")
                    
                    st.markdown("**Linguagem Matemática:** descrever o desenvolvimento da classe referente a:")
                    if criterios_etapa["LM"]: st.caption(f"_{criterios_etapa['LM']}_")
                    data_inf['sin_lm'] = st.text_area("Síntese - Linguagem Matemática", value=data_inf.get('sin_lm', ''), height=100, label_visibility="collapsed", key="txt_lm")
                    
                    st.markdown("**Indivíduo e Sociedade:** descrever o desenvolvimento da classe em relação a:")
                    if criterios_etapa["IS"]: st.caption(f"_{criterios_etapa['IS']}_")
                    data_inf['sin_is'] = st.text_area("Síntese - Indivíduo e Sociedade", value=data_inf.get('sin_is', ''), height=100, label_visibility="collapsed", key="txt_is")
                    
                    st.markdown("**Arte:** descrever o desenvolvimento da classe referente a:")
                    if criterios_etapa["Arte"]: st.caption(f"_{criterios_etapa['Arte']}_")
                    data_inf['sin_arte'] = st.text_area("Síntese - Arte", value=data_inf.get('sin_arte', ''), height=100, label_visibility="collapsed", key="txt_arte")
                    
                    st.markdown("**Cultura Corporal e Movimento:** descrever o desenvolvimento da classe em relação a:")
                    if criterios_etapa["CCM"]: st.caption(f"_{criterios_etapa['CCM']}_")
                    data_inf['sin_ccm'] = st.text_area("Síntese - Cultura Corporal", value=data_inf.get('sin_ccm', ''), height=100, label_visibility="collapsed", key="txt_ccm")

                    st.markdown("**Linguagens e Tecnologias:** descrever o desenvolvimento da classe em relação a:")
                    if criterios_etapa["LT"]: st.caption(f"_{criterios_etapa['LT']}_")
                    data_inf['sin_lt'] = st.text_area("Síntese - Linguagens e Tecnologias", value=data_inf.get('sin_lt', ''), height=100, label_visibility="collapsed", key="txt_lt_inf")

                    st.markdown("**Libras:** descrever o desenvolvimento da classe em relação a:")
                    if criterios_etapa["LIBRAS"]: st.caption(f"_{criterios_etapa['LIBRAS']}_")
                    data_inf['sin_libras'] = st.text_area("Síntese - Libras", value=data_inf.get('sin_libras', ''), height=100, label_visibility="collapsed", key="txt_libras_inf")

                with tabs[2]:
                    st.subheader("Plano de Ação")
                    st.caption("Estudantes com desenvolvimento abaixo do esperado em relação aos conteúdos do currículo.")
                    
                    for i, row in enumerate(data_inf['abaixo_basico']):
                        with st.container():
                            st.markdown(f"**Estudante {i+1}**")
                            c_est, c_del = st.columns([11, 1])
                            row['Estudante'] = c_est.text_input(f"Nome", value=row.get('Estudante', ''), key=f"inf_ab_est_{i}", label_visibility="collapsed")
                            if c_del.button("🗑️", key=f"inf_del_ab_{i}"):
                                data_inf['abaixo_basico'].pop(i); st.rerun()
                            
                            cc = st.columns(7)
                            row['LV'] = cc[0].text_input("LV", value=row.get('LV', ''), key=f"inf_ab_lv_{i}")
                            row['LM'] = cc[1].text_input("LM", value=row.get('LM', ''), key=f"inf_ab_lm_{i}")
                            row['IS'] = cc[2].text_input("IS", value=row.get('IS', ''), key=f"inf_ab_is_{i}")
                            row['A'] = cc[3].text_input("A", value=row.get('A', ''), key=f"inf_ab_a_{i}")
                            row['CCM'] = cc[4].text_input("CCM", value=row.get('CCM', ''), key=f"inf_ab_ccm_{i}")
                            row['LT'] = cc[5].text_input("LT", value=row.get('LT', ''), key=f"inf_ab_lt_{i}")
                            row['LIBRAS'] = cc[6].text_input("LIB", value=row.get('LIBRAS', ''), key=f"inf_ab_lib_{i}")
                            st.markdown("<hr style='margin: 10px 0;'>", unsafe_allow_html=True)
                    
                    if st.button("➕ Adicionar Estudante", key="inf_add_ab"):
                        data_inf['abaixo_basico'].append({"Estudante": "", "LV": "", "LM": "", "IS": "", "A": "", "CCM": "", "LT": "", "LIBRAS": ""})
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown("**Propostas para Intervenção/Recuperação de aprendizagem:**")
                    st.markdown(propostas_ata_inf)

                with tabs[3]:
                    st.subheader("3. Observações Gerais")
                    
                    st.markdown("**a) Desempenho de alunos especiais (laudados)**")
                    for i, row in enumerate(data_inf['obs_especiais']):
                        c1, c2, c3 = st.columns([3, 6, 1])
                        row['Estudante'] = c1.text_input("Estudante Especial", value=row.get('Estudante', ''), key=f"inf_obs_est_{i}")
                        row['Desempenho/Observação'] = c2.text_area("Desempenho/Observações", value=row.get('Desempenho/Observação', ''), key=f"inf_obs_des_{i}", height=68)
                        if c3.button("🗑️", key=f"inf_del_obs_{i}"):
                            data_inf['obs_especiais'].pop(i); st.rerun()
                    if st.button("➕ Adicionar Aluno Especial", key="inf_add_obs"):
                        data_inf['obs_especiais'].append({"Estudante": "", "Desempenho/Observação": ""})
                        st.rerun()
                    
                    st.divider()
                    st.markdown("**b) Alunos encaminhados (Conselho Tutelar ou Serviço Social)**")
                    for i, row in enumerate(data_inf['encaminhamentos']):
                        c1, c2, c3 = st.columns([3, 6, 1])
                        row['Estudante'] = c1.text_input("Estudante Encaminhado", value=row.get('Estudante', ''), key=f"inf_enc_est_{i}")
                        row['Motivo'] = c2.text_area("Motivo do Encaminhamento", value=row.get('Motivo', ''), key=f"inf_enc_mot_{i}", height=68)
                        if c3.button("🗑️", key=f"inf_del_enc_{i}"):
                            data_inf['encaminhamentos'].pop(i); st.rerun()
                    if st.button("➕ Adicionar Encaminhamento", key="inf_add_enc"):
                        data_inf['encaminhamentos'].append({"Estudante": "", "Motivo": ""})
                        st.rerun()

                    st.divider()
                    st.markdown("**c) Estudantes Matriculados Tardiamente**")
                    for i, row in enumerate(data_inf['mat_tardia']):
                        c1, c2, c3, c4 = st.columns([4, 3, 3, 1])
                        row['Estudante'] = c1.text_input("Estudante", value=row.get('Estudante', ''), key=f"inf_mat_est_{i}")
                        row['Data Matrícula'] = c2.text_input("Data da Matrícula", value=row.get('Data Matrícula', ''), key=f"inf_mat_data_{i}")
                        row['Total Frequência'] = c3.text_input("Frequência (Dias)", value=row.get('Total Frequência', ''), key=f"inf_mat_freq_{i}")
                        if c4.button("🗑️", key=f"inf_del_mat_{i}"):
                            data_inf['mat_tardia'].pop(i); st.rerun()
                    if st.button("➕ Adicionar Matrícula Tardia", key="inf_add_mat"):
                        data_inf['mat_tardia'].append({"Estudante": "", "Data Matrícula": "", "Total Frequência": ""})
                        st.rerun()
                        
                    st.divider()
                    st.markdown("**d) Outras Observações**")
                    data_inf['obs_outras'] = st.text_area("Campo livre para quaisquer outras observações da turma:", value=data_inf.get('obs_outras', ''), height=120, key="inf_obs_outras")

                with tabs[4]:
                    st.subheader("Finalização e Assinaturas")
                    
                    if st.button("🤖 Preencher Assinaturas Automaticamente", type="primary", key="inf_btn_auto"):
                        turma_atual = st.session_state.ata_turma_confirmada_inf
                        df_turma = df_matriz[(df_matriz['Turma'] == turma_atual)]
                        
                        if not df_turma.empty:
                            lista_final = []
                            professores_adicionados = set()
                            
                            for _, row in df_turma.iterrows():
                                materia = row['Disciplina']
                                nome_prof = row['Professor']
                                if nome_prof and nome_prof not in professores_adicionados:
                                    lista_final.append({"Nome": nome_prof, "Cargo/Atuação": f"{materia} (Atuante na Turma)"})
                                    professores_adicionados.add(nome_prof)
                                    
                            for _, row in df_gestao.iterrows():
                                if row['Nome']:
                                    lista_final.append({"Nome": row['Nome'], "Cargo/Atuação": row['Cargo']})
                            
                            data_inf['assinaturas'] = lista_final
                            st.success("✅ Grade preenchida com sucesso para a turma selecionada!")
                            st.rerun()
                        else:
                            st.error("Turma não encontrada na Matriz de Automação (Aba Configurações).")

                    st.divider()
                    st.markdown("**Participantes da Reunião**")
                    
                    if 'assinaturas' not in data_inf:
                        data_inf['assinaturas'] = [{"Nome": "", "Cargo/Atuação": ""}]
                        
                    for i, row in enumerate(data_inf['assinaturas']):
                        c1, c2, c3 = st.columns([5, 4, 2])
                        row['Nome'] = c1.text_input("Nome", value=row.get('Nome', ''), key=f"inf_sig_n_{i}", label_visibility="collapsed")
                        row['Cargo/Atuação'] = c2.text_input("Cargo", value=row.get('Cargo/Atuação', ''), key=f"inf_sig_c_{i}", label_visibility="collapsed")
                        
                        b1, b2, b3 = c3.columns(3)
                        if b1.button("⬆️", key=f"inf_up_{i}", disabled=(i == 0)):
                            data_inf['assinaturas'][i], data_inf['assinaturas'][i-1] = data_inf['assinaturas'][i-1], data_inf['assinaturas'][i]
                            st.rerun()
                        if b2.button("⬇️", key=f"inf_dw_{i}", disabled=(i == len(data_inf['assinaturas']) - 1)):
                            data_inf['assinaturas'][i], data_inf['assinaturas'][i+1] = data_inf['assinaturas'][i+1], data_inf['assinaturas'][i]
                            st.rerun()
                        if b3.button("🗑️", key=f"inf_del_{i}"):
                            data_inf['assinaturas'].pop(i)
                            st.rerun()
                            
                    if st.button("➕ Adicionar Participante Manualmente", key="inf_add_sig"):
                        data_inf['assinaturas'].append({"Nome": "", "Cargo/Atuação": ""})
                        st.rerun()
                        
                    st.divider()
                    
                    if st.button("💾 Salvar Ata do Infantil", use_container_width=True, type="secondary", key="inf_save"):
                        try:
                            dados_para_salvar = {}
                            for key, value in data_inf.items():
                                if isinstance(value, pd.DataFrame):
                                    dados_para_salvar[key] = value.to_dict(orient='records')
                                else:
                                    dados_para_salvar[key] = value
                            
                            novo_json = json.dumps(dados_para_salvar, ensure_ascii=False)
                            id_ata = f"{data_inf.get('turma', 'SemTurma')} - {data_inf.get('trimestre', 'SemTri')} (Infantil)"
                            
                            # Cria apenas o registro desta ata específica
                            novo_registro = {
                                "id_ata": id_ata, 
                                "modalidade": "Educação Infantil", 
                                "turma": data_inf.get('turma', ''), 
                                "dados_json": novo_json
                            }
                            
                            # Faz a atualização direta na linha do Supabase, sem mexer no resto da tabela
                            supabase.table("Atas_Conselho").upsert(novo_registro).execute()
                            
                            st.success(f"✅ Ata do Infantil salva com sucesso e segurança!")
                        except Exception as e:
                            st.error(f"Erro ao salvar: {e}")

                    if st.button("👁️ GERAR ATA DO INFANTIL (PDF)", type="primary", use_container_width=True, key="inf_pdf"):
                        try:
                            pdf = OfficialPDF('P', 'mm', 'A4')
                            pdf.doc_type = "Ata" 
                            pdf.set_margins(15, 35, 15)
                            pdf.set_auto_page_break(auto=True, margin=20)
                            pdf.add_page()
                            
                            def calc_lines(txt, w):
                                # A MÁGICA VEM AQUI: Limpa o texto antes da biblioteca tentar medir!
                                txt = clean_pdf_text(txt)
                                
                                if not txt: return 1
                                lines = 0
                                for par in txt.split('\n'):
                                    words = par.split(' ')
                                    curr_w = 0
                                    for word in words:
                                        word_w = pdf.get_string_width(word + ' ') # O erro acontecia nesta linha!
                                        if curr_w + word_w > w:
                                            lines += 1; curr_w = word_w
                                        else:
                                            curr_w += word_w
                                    lines += 1
                                return lines
                            
                            # --- CABEÇALHO ---
                            pdf.set_font("Arial", "B", 10)
                            escola_nome = data_inf.get('escola', 'CEIEF Rafael Affonso Leite').upper()
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(escola_nome), 0, 1, 'C')
                            
                            trimestre = data_inf.get('trimestre', '').upper()
                            ano = data_inf.get('ano_letivo', '')
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(f"REGISTRO E CONTROLE DO ACOMPANHAMENTO ESCOLAR"), 0, 1, 'C')
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(f"EDUCAÇÃO INFANTIL - CONSELHO DE CLASSE/TERMO - {trimestre} DE {ano}"), 0, 1, 'C')
                            
                            turma = data_inf.get('turma', '')
                            ciclo = data_inf.get('ciclo', '')
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text(f"Turma: {turma} | Etapa: {ciclo}"), 0, 1, 'C')
                            pdf.ln(5)
                            
                            # --- SÍNTESE AVALIATIVA ---
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_fill_color(220, 220, 220)
                            pdf.set_x(15)
                            pdf.cell(180, 6, "SÍNTESE AVALIATIVA", "LTR", 1, 'C', True)
                            
                            pdf.set_x(15)
                            pdf.cell(180, 2, "", "LR", 1) 
                            
                            pdf.set_font("Arial", "", 10)
                            pdf.set_x(15)
                            pdf.multi_cell(180, 5, clean_pdf_text(texto_base_ata_inf), "LR", 'J')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 4, "", "LR", 1)
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("1- Síntese avaliativa da classe:"), "LR", 1, 'L')
                            
                            texto_sint = "descrever o desenvolvimento da classe em cada componente curricular, considerando os aspectos avaliados no trimestre:"
                            pdf.set_font("Arial", "", 10)
                            pdf.set_x(15)
                            pdf.multi_cell(180, 5, clean_pdf_text(texto_sint), "LR", 'J')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 4, "", "LR", 1)
                            
                            criterios_etapa = get_criterios_infantil(st.session_state.ata_ciclo_confirmado_inf)
                            
                            disciplinas_inf = [
                                ("Linguagem Verbal", criterios_etapa["LV"], data_inf.get('sin_lv', '')),
                                ("Linguagem Matemática", criterios_etapa["LM"], data_inf.get('sin_lm', '')),
                                ("Indivíduo e Sociedade", criterios_etapa["IS"], data_inf.get('sin_is', '')),
                                ("Arte", criterios_etapa["Arte"], data_inf.get('sin_arte', '')),
                                ("Cultura Corporal e Movimento", criterios_etapa["CCM"], data_inf.get('sin_ccm', '')),
                                ("Linguagens e Tecnologias", criterios_etapa["LT"], data_inf.get('sin_lt', '')),
                                ("Libras", criterios_etapa["LIBRAS"], data_inf.get('sin_libras', ''))
                            ]
                            
                            for i, (nome, crit, texto) in enumerate(disciplinas_inf):
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                
                                h_box = calc_lines(crit, 168) * 5 if crit else 5
                                h_total = 5 + 2 + h_box + 2
                                if pdf.get_y() + h_total > 275:
                                    pdf.set_x(15)
                                    pdf.cell(180, 1, "", "LRB", 1)
                                    pdf.add_page()
                                    pdf.set_x(15)
                                    pdf.cell(180, 2, "", "LTR", 1)
                                
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {nome}: descrever o desenvolvimento da classe referente a:"), "LR", 1, 'L')
                                
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)
                                
                                y_start = pdf.get_y()
                                pdf.set_font("Arial", "", 9)
                                pdf.set_x(20)
                                pdf.multi_cell(170, 5, clean_pdf_text(crit if crit else "Critérios não configurados para esta etapa."), 1, 'J')
                                y_end = pdf.get_y()
                                
                                pdf.line(15, y_start, 15, y_end)
                                pdf.line(195, y_start, 195, y_end)
                                
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)
                                
                                if texto.strip():
                                    pdf.set_font("Arial", "", 10)
                                    h_texto = calc_lines(texto, 178) * 5
                                    if pdf.get_y() + h_texto > 275:
                                        pdf.set_x(15)
                                        pdf.cell(180, 1, "", "LRB", 1)
                                        pdf.add_page()
                                        pdf.set_x(15)
                                        pdf.cell(180, 2, "", "LTR", 1)
                                        
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {texto}"), "LR", 'J')
                                
                                if i < len(disciplinas_inf) - 1:
                                    pdf.set_x(15)
                                    pdf.cell(180, 4, "", "LR", 1) 
                            
                            # --- 2. PLANO DE AÇÃO ---
                            pdf.set_x(15)
                            pdf.cell(180, 5, "", "LR", 1) 
                            
                            if pdf.get_y() > 230: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.multi_cell(180, 5, clean_pdf_text("2- Plano de Ação para os estudantes com desenvolvimento abaixo do esperado em relação aos conteúdos do currículo:"), "LR", 'L')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 3, "", "LR", 1)
                            
                            pdf.set_font("Arial", "B", 9)
                            col_w = [54, 18, 18, 18, 18, 18, 18, 18]
                            headers = ["Estudante", "LV", "LM", "IS", "A", "CCM", "LT", "LIB"]
                            pdf.set_x(15)
                            for i, h in enumerate(headers):
                                pdf.cell(col_w[i], 6, h, 1, 0, 'C')
                            pdf.ln()
                            
                            pdf.set_font("Arial", "", 10)
                            lista_abaixo = data_inf.get('abaixo_basico', [])
                                
                            def truncate_str(texto, max_w):
                                # Limpa antes de usar no laço while
                                texto = clean_pdf_text(texto)
                                
                                while pdf.get_string_width(texto) > max_w - 2: 
                                    texto = texto[:-1]
                                return texto

                            for row in lista_abaixo:
                                estudante = str(row.get('Estudante', '')).strip()
                                if estudante: 
                                    pdf.set_x(15)
                                    estudante_seguro = truncate_str(estudante, col_w[0])
                                    pdf.cell(col_w[0], 6, clean_pdf_text(estudante_seguro), 1, 0, 'L')
                                    pdf.cell(col_w[1], 6, clean_pdf_text(str(row.get('LV', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[2], 6, clean_pdf_text(str(row.get('LM', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[3], 6, clean_pdf_text(str(row.get('IS', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[4], 6, clean_pdf_text(str(row.get('A', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[5], 6, clean_pdf_text(str(row.get('CCM', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[6], 6, clean_pdf_text(str(row.get('LT', ''))), 1, 0, 'C')
                                    pdf.cell(col_w[7], 6, clean_pdf_text(str(row.get('LIBRAS', ''))), 1, 1, 'C')
                            
                            pdf.set_x(15)
                            pdf.cell(180, 3, "", "LR", 1) 
                            
                            pdf.set_font("Arial", "B", 6.5)
                            pdf.set_x(15)
                            pdf.cell(180, 5, clean_pdf_text("*Propostas para Intervenção/Recuperação de aprendizagem:"), "LR", 1, 'L')
                            pdf.set_font("Arial", "", 10)
                            
                            for prop in propostas_ata_inf.split('\n'):
                                if prop.strip():
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(prop.strip()), "LR", 'J')

                            pdf.set_x(15)
                            pdf.cell(180, 5, "", "LR", 1) 

                            # --- 3. OBSERVAÇÕES GERAIS ---
                            if pdf.get_y() > 230: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_x(15)
                            pdf.cell(180, 6, clean_pdf_text("3- Observações Gerais:"), "LR", 1, 'L')
                            pdf.set_font("Arial", "", 10)
                            
                            prefix_code = 97 
                            
                            lista_esp = data_inf.get('obs_especiais', [])
                            esp_valid = [r for r in lista_esp if str(r.get('Estudante', '')).strip()]
                            if esp_valid:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Desempenho de alunos especiais (laudados):"), "LR", 1, 'L')
                                prefix_code += 1
                                for row in esp_valid:
                                    est = str(row.get('Estudante', '')).strip()
                                    obs = str(row.get('Desempenho/Observação', '')).strip()
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {est}:"), "LR", 1, 'L')
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {obs}"), "LR", 'J')
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)
                            
                            lista_enc = data_inf.get('encaminhamentos', [])
                            enc_valid = [r for r in lista_enc if str(r.get('Estudante', '')).strip()]
                            if enc_valid:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Alunos encaminhados (Conselho Tutelar/Serv. Social):"), "LR", 1, 'L')
                                prefix_code += 1
                                for row in enc_valid:
                                    est = str(row.get('Estudante', '')).strip()
                                    mot = str(row.get('Motivo', '')).strip()
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {est}:"), "LR", 1, 'L')
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {mot}"), "LR", 'J')
                                pdf.set_x(15)
                                pdf.cell(180, 2, "", "LR", 1)

                            lista_tardia = data_inf.get('mat_tardia', [])
                            tardia_valid = [r for r in lista_tardia if str(r.get('Estudante', '')).strip()]
                            if len(tardia_valid) > 0:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Estudantes Matriculados Tardiamente:"), "LR", 1, 'L')
                                prefix_code += 1
                                for row in tardia_valid:
                                    est = str(row.get('Estudante', '')).strip()
                                    mat = str(row.get('Data Matrícula', '')).strip()
                                    freq = str(row.get('Total Frequência', '')).strip()
                                    pdf.set_font("Arial", "B", 10)
                                    pdf.set_x(15)
                                    pdf.cell(180, 5, clean_pdf_text(f"  {chr(149)}  {est}:"), "LR", 1, 'L')
                                    texto_tardio = f"Matriculado(a) nesta sala em {mat}. Portanto, o período correspondente à sua matrícula compreende um total de {freq} dias letivos."
                                    pdf.set_font("Arial", "", 10)
                                    pdf.set_x(15)
                                    pdf.multi_cell(180, 5, clean_pdf_text(f"  {texto_tardio}"), "LR", 'J')
                            else:
                                pdf.set_font("Arial", "", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, f" {chr(prefix_code)}) Sem matrículas tardias registradas no período.", "LR", 1)
                                prefix_code += 1

                            pdf.set_x(15)
                            pdf.cell(180, 2, "", "LR", 1)
                            
                            obs_outras = data_inf.get('obs_outras', '').strip()
                            if obs_outras:
                                pdf.set_font("Arial", "B", 10)
                                pdf.set_x(15)
                                pdf.cell(180, 5, clean_pdf_text(f"{chr(prefix_code)}) Outras Observações:"), "LR", 1, 'L')
                                pdf.set_font("Arial", "", 10)
                                pdf.set_x(15)
                                pdf.multi_cell(180, 5, clean_pdf_text(f"  {obs_outras}"), "LR", 'J')
                                
                            pdf.set_x(15)
                            pdf.cell(180, 3, "", "LRB", 1) 

                            # --- ASSINATURAS ---
                            pdf.ln(5)
                            if pdf.get_y() > 220: pdf.add_page()
                            
                            pdf.set_font("Arial", "B", 10)
                            pdf.set_fill_color(220, 220, 220)
                            pdf.set_x(15)
                            pdf.cell(180, 6, "ASSINATURA DOS PARTICIPANTES NA REUNIÃO DO CONSELHO DE EDUCAÇÃO INFANTIL", 1, 1, 'C', True)
                            
                            lista_assinaturas = data_inf.get('assinaturas', [])
                            sigs_validas = [s for s in lista_assinaturas if str(s.get('Nome', '')).strip()]
                            
                            if not sigs_validas:
                                pdf.set_x(15)
                                pdf.cell(180, 20, "Nenhuma assinatura cadastrada para esta ata.", 1, 1, 'C')
                            else:
                                cols = 4
                                cell_w = 180 / cols
                                cell_h = 24 
                                
                                x_start = 15
                                y = pdf.get_y()
                                
                                for i, sig in enumerate(sigs_validas):
                                    col = i % cols
                                    if col == 0 and i > 0:
                                        y += cell_h
                                    
                                    if y + cell_h > 275:
                                        pdf.add_page()
                                        y = pdf.get_y()
                                    
                                    x = x_start + (col * cell_w)
                                    pdf.rect(x, y, cell_w, cell_h)
                                    
                                    pdf.line(x + 4, y + 10, x + cell_w - 4, y + 10)
                                    
                                    nome = clean_pdf_text(str(sig.get('Nome', '')).strip())
                                    cargo_full = clean_pdf_text(str(sig.get('Cargo/Atuação', '')).strip())    
                                    
                                    cargo = cargo_full
                                    atuacao = ""
                                    if "(" in cargo_full:
                                        parts = cargo_full.split("(")
                                        cargo = parts[0].strip()
                                        atuacao = parts[1].replace(")", "").strip()
                                    
                                    font_size = 7
                                    pdf.set_font("Arial", "B", font_size)
                                    while pdf.get_string_width(nome) > cell_w - 2 and font_size > 4.5:
                                        font_size -= 0.5
                                        pdf.set_font("Arial", "B", font_size)
                                    
                                    pdf.set_xy(x + 1, y + 11) 
                                    pdf.multi_cell(cell_w - 2, 3, clean_pdf_text(nome), 0, 'C')
                                    
                                    pdf.set_font("Arial", "", 6)
                                    curr_y = pdf.get_y()
                                    pdf.set_xy(x + 1, curr_y)
                                    pdf.multi_cell(cell_w - 2, 3, clean_pdf_text(cargo), 0, 'C')
                                    
                                    if atuacao:
                                        curr_y = pdf.get_y()
                                        pdf.set_xy(x + 1, curr_y)
                                        pdf.multi_cell(cell_w - 2, 3, clean_pdf_text(atuacao), 0, 'C')
                                
                                pdf.set_y(y + cell_h)

                            st.session_state.pdf_bytes_ata = get_pdf_bytes(pdf)
                            st.success("✅ PDF do Infantil gerado com sucesso!")
                        except Exception as e:
                            st.error(f"Erro ao desenhar o PDF: {e}")

                if 'pdf_bytes_ata' in st.session_state:
                    turma_limpa = str(data_inf.get('turma', 'Turma')).replace('/', '-').replace('\\', '-')
                    trimestre_limpo = str(data_inf.get('trimestre', 'Trimestre')).replace('/', '-')
                    st.download_button("📥 BAIXAR ATA DO INFANTIL", st.session_state.pdf_bytes_ata, f"Ata_Infantil_{turma_limpa}_{trimestre_limpo}.pdf", "application/pdf", type="primary")

    # ==============================================================================
    # 2. TELA: HISTÓRICO DE ATAS
    # ==============================================================================
    if app_mode_regular == "📂 Histórico de Atas":
        st.markdown('<div class="header-box"><div class="header-title">Histórico de Atas</div></div>', unsafe_allow_html=True)
        df_atas = safe_read("Atas_Conselho", ["id_ata", "modalidade", "turma", "dados_json"])
        
        if df_atas.empty:
            st.info("Nenhuma ata salva ainda.")
        else:
            df_display = df_atas[["id_ata", "modalidade", "turma"]].copy()
            df_display.columns = ["ID / Título da Ata", "Modalidade", "Turma"]
            st.dataframe(df_display, use_container_width=True, hide_index=True)
            
            st.divider()
            st.subheader("🔄 Carregar Ata")
            c_sel, c_btn = st.columns([3, 1])
            ata_selecionada = c_sel.selectbox("Selecione a Ata:", df_atas["id_ata"].tolist(), label_visibility="collapsed")
            
            if c_btn.button("Carregar Dados", type="primary", use_container_width=True):
                dados_row = df_atas[df_atas["id_ata"] == ata_selecionada].iloc[0]
                try:
                    dados_json = json.loads(dados_row["dados_json"])
                    for key in ['abaixo_basico', 'basico', 'mat_tardia', 'obs_especiais', 'encaminhamentos', 'assinaturas']:
                        if key in dados_json and isinstance(dados_json[key], list):
                            dados_json[key] = pd.DataFrame(dados_json[key])
                            
                    if "Fundamental" in dados_row["modalidade"]:
                        st.session_state.data_ata_ef = dados_json
                        st.session_state.ata_turma_confirmada = dados_json.get('turma', '')
                        st.session_state.ata_ciclo_confirmado = dados_json.get('ciclo', '')
                        
                    elif "Infantil" in dados_row["modalidade"]:
                        st.session_state.data_ata_inf = dados_json
                        st.session_state.ata_turma_confirmada_inf = dados_json.get('turma', '')
                        st.session_state.ata_ciclo_confirmado_inf = dados_json.get('ciclo', '')
                        
                    st.success(f"Ata '{ata_selecionada}' carregada! Vá para 'Nova Ata' e certifique-se de escolher a modalidade correta no menu lateral.")
                except Exception as e:
                    st.error(f"Erro ao carregar: {e}")
            
            st.divider()
            if not is_monitor:
                with st.expander("⚠️ Excluir Ata"):
                    c_del_sel, c_del_btn = st.columns([3, 1])
                    ata_excluir = c_del_sel.selectbox("Ata:", df_atas["id_ata"].tolist(), key="del_ata_sel", label_visibility="collapsed")
                    if c_del_btn.button("🗑️ Excluir", type="secondary", use_container_width=True):
                        safe_update("Atas_Conselho", df_atas[df_atas["id_ata"] != ata_excluir])
                        st.success("Excluída!")
                        time.sleep(1); st.rerun()

    # ==============================================================================
    # 3. TELA: CONFIGURAÇÕES (O CÉREBRO DA AUTOMAÇÃO)
    # ==============================================================================
    if app_mode_regular == "⚙️ Configurações":
        st.markdown('<div class="header-box"><div class="header-title">Configurações do Sistema</div><div class="header-subtitle">Textos Base e Matriz de Professores</div></div>', unsafe_allow_html=True)
        
        t_conf = st.tabs(["📝 Textos Fund.", "🧸 Textos Inf.", "👨‍🏫 Matriz de Professores", "👔 Matriz da Gestão"])
        
        with t_conf[0]:
            st.info("💡 As edições salvas aqui serão utilizadas automaticamente nas novas Atas do Ensino Fundamental.")
            novo_texto_base = st.text_area("Texto Base da Síntese Avaliativa (Legislações)", value=texto_base_ata_ef, height=180)
            novas_propostas = st.text_area("Propostas de Recuperação", value=propostas_ata_ef, height=250)
            
            if st.button("💾 Salvar Textos Fundamental", type="primary", use_container_width=True):
                if not df_config.empty and "texto_base_ata" in df_config["chave"].values: df_config.loc[df_config["chave"] == "texto_base_ata", "valor"] = novo_texto_base
                else: df_config = pd.concat([df_config, pd.DataFrame([{"chave": "texto_base_ata", "valor": novo_texto_base}])], ignore_index=True)
                
                if not df_config.empty and "propostas_ata" in df_config["chave"].values: df_config.loc[df_config["chave"] == "propostas_ata", "valor"] = novas_propostas
                else: df_config = pd.concat([df_config, pd.DataFrame([{"chave": "propostas_ata", "valor": novas_propostas}])], ignore_index=True)
                
                safe_update("Config_Ata", df_config)
                st.success("✅ Textos atualizados!")

        with t_conf[1]:
            st.info("💡 Escolha a Etapa/Maternal e defina os conteúdos cobrados para aquela idade.")
            
            etapa_edit = st.selectbox("Selecione a Etapa para editar os conteúdos:", ["1ª Etapa", "2ª Etapa", "Maternal II", "Maternal I"])
            criterios_tela = get_criterios_infantil(etapa_edit)
            
            novo_texto_base_inf = st.text_area("Texto Base da Síntese Avaliativa (Infantil)", value=texto_base_ata_inf, height=120)
            novas_propostas_inf = st.text_area("Propostas de Intervenção (Infantil)", value=propostas_ata_inf, height=120)
            
            st.markdown(f"**Conteúdos Avaliados - {etapa_edit}**")
            novo_crit_lv = st.text_area("Linguagem Verbal", value=criterios_tela["LV"])
            novo_crit_lm = st.text_area("Linguagem Matemática", value=criterios_tela["LM"])
            novo_crit_is = st.text_area("Indivíduo e Sociedade", value=criterios_tela["IS"])
            novo_crit_arte = st.text_area("Arte", value=criterios_tela["Arte"])
            novo_crit_ccm = st.text_area("Cultura Corporal e Movimento", value=criterios_tela["CCM"])
            novo_crit_lt = st.text_area("Linguagens e Tecnologias", value=criterios_tela["LT"])
            novo_crit_libras = st.text_area("Libras", value=criterios_tela["LIBRAS"])
            
            if st.button("💾 Salvar Textos e Conteúdos do Infantil", type="primary", use_container_width=True):
                if not df_config.empty and "texto_base_ata_inf" in df_config["chave"].values: df_config.loc[df_config["chave"] == "texto_base_ata_inf", "valor"] = novo_texto_base_inf
                else: df_config = pd.concat([df_config, pd.DataFrame([{"chave": "texto_base_ata_inf", "valor": novo_texto_base_inf}])], ignore_index=True)
                
                if not df_config.empty and "propostas_ata_inf" in df_config["chave"].values: df_config.loc[df_config["chave"] == "propostas_ata_inf", "valor"] = novas_propostas_inf
                else: df_config = pd.concat([df_config, pd.DataFrame([{"chave": "propostas_ata_inf", "valor": novas_propostas_inf}])], ignore_index=True)
                
                chaves_crit = [
                    (f"crit_lv_{etapa_edit}", novo_crit_lv), 
                    (f"crit_lm_{etapa_edit}", novo_crit_lm), 
                    (f"crit_is_{etapa_edit}", novo_crit_is), 
                    (f"crit_arte_{etapa_edit}", novo_crit_arte), 
                    (f"crit_ccm_{etapa_edit}", novo_crit_ccm),
                    (f"crit_lt_{etapa_edit}", novo_crit_lt),
                    (f"crit_libras_{etapa_edit}", novo_crit_libras)
                ]
                
                for k, v in chaves_crit:
                    if not df_config.empty and k in df_config["chave"].values: df_config.loc[df_config["chave"] == k, "valor"] = v
                    else: df_config = pd.concat([df_config, pd.DataFrame([{"chave": k, "valor": v}])], ignore_index=True)

                safe_update("Config_Ata", df_config)
                st.success(f"✅ Textos e Conteúdos salvos com sucesso para a {etapa_edit}!")

        with t_conf[2]:
            st.info("💡 Edite a tabela para alterar a atribuição de aulas.")
            
            config_col_matriz = {
                "Ciclo": st.column_config.SelectboxColumn("Ciclo", options=["Ciclo I (1º ao 3º ano)", "Ciclo II (4º e 5º ano)", "1ª Etapa", "2ª Etapa", "Maternal II", "Educação Infantil"]),
                "Turma": st.column_config.TextColumn("Turma (Ex: 5º Ano 1)"),
                "Disciplina": st.column_config.TextColumn("Disciplina (Ex: Polivalente, Matemática)"),
                "Professor": st.column_config.TextColumn("Nome do Professor")
            }
            
            df_matriz_editada = st.data_editor(df_matriz, column_config=config_col_matriz, num_rows="dynamic", use_container_width=True, hide_index=True)
            
            if st.button("💾 Salvar Matriz de Professores", type="primary", use_container_width=True):
                novo_matriz_json = df_matriz_editada.to_json(orient='records')
                if not df_config.empty and "matriz_professores" in df_config["chave"].values:
                    df_config.loc[df_config["chave"] == "matriz_professores", "valor"] = novo_matriz_json
                else:
                    df_config = pd.concat([df_config, pd.DataFrame([{"chave": "matriz_professores", "valor": novo_matriz_json}])], ignore_index=True)
                safe_update("Config_Ata", df_config)
                st.success("✅ Matriz de professores salva com sucesso!")

        with t_conf[3]:
            st.info("💡 Edite a tabela para atualizar quem assina como Equipe Gestora nas Atas.")
            
            config_col_gestao = {
                "Nome": st.column_config.TextColumn("Nome do Profissional"),
                "Cargo": st.column_config.SelectboxColumn("Cargo", options=["Prof. Coordenador", "Vice-Diretor", "Diretor de Escola"])
            }
            
            df_gestao_editada = st.data_editor(df_gestao, column_config=config_col_gestao, num_rows="dynamic", use_container_width=True, hide_index=True)
            
            if st.button("💾 Salvar Matriz de Gestão", type="primary", use_container_width=True):
                novo_gestao_json = df_gestao_editada.to_json(orient='records')
                if not df_config.empty and "matriz_gestao" in df_config["chave"].values:
                    df_config.loc[df_config["chave"] == "matriz_gestao", "valor"] = novo_gestao_json
                else:
                    df_config = pd.concat([df_config, pd.DataFrame([{"chave": "matriz_gestao", "valor": novo_gestao_json}])], ignore_index=True)
                safe_update("Config_Ata", df_config)
                st.success("✅ Matriz da gestão salva com sucesso!")

    

# ==============================================================================
# MÓDULO 4: AGENDAMENTO SALA DE INFORMÁTICA (MIGRADO PARA SUPABASE)
# ==============================================================================

    elif app_mode_regular == "💻 Agendamento Informática":
        st.markdown('<div class="header-box"><div class="header-title">💻 Agendamento - Sala de Informática</div></div>', unsafe_allow_html=True)
        st.markdown("Reserve a sala de computadores para a sua turma do Ensino Regular.")
        st.divider()

        # --- GRADE FIXA ANUAL ATUALIZADA ---
        grade_fixa = {
            0: [ # Segunda-feira
                {"Horario": "07:00 - 07:50", "Professor": "Prof. Fernando", "Turma": "1º ANO 1: Linguagens e Tecnologias"},
                {"Horario": "11:10 - 12:00", "Professor": "Prof. Bruna", "Turma": "5º ANO 1: Linguagens e Tecnologias"},
                {"Horario": "12:30 - 13:20", "Professor": "Prof. Elaine", "Turma": "Etapa 2-2: Linguagens e Tecnologias"},
                {"Horario": "13:20 - 14:10", "Professor": "Prof. Elaine", "Turma": "Etapa 2-3: Linguagens e Tecnologias"},
                {"Horario": "14:10 - 15:00", "Professor": "Prof. Fernando", "Turma": "1º ANO 2: Linguagens e Tecnologias"},
                {"Horario": "15:00 - 15:50", "Professor": "Prof. Fernando", "Turma": "2º ANO 1: Linguagens e Tecnologias"}
            ],
            1: [ # Terça-feira
                {"Horario": "07:00 - 07:50", "Professor": "Prof. Josiane", "Turma": "4º ANO 2: Linguagens e Tecnologias"},
                {"Horario": "09:30 - 10:20", "Professor": "Prof. Josiane", "Turma": "4º ANO 1: Linguagens e Tecnologias"},
                {"Horario": "11:10 - 12:00", "Professor": "Prof. Bruna", "Turma": "5º ANO 3: Linguagens e Tecnologias"},
                {"Horario": "12:30 - 13:20", "Professor": "Prof. Elaine", "Turma": "Etapa 2-1: Linguagens e Tecnologias"},
                {"Horario": "13:20 - 14:10", "Professor": "Prof. Elaine", "Turma": "Mat. 2-2: Linguagens e Tecnologias"}
            ],
            2: [ # Quarta-feira
                {"Horario": "07:50 - 08:40", "Professor": "Prof. Karina", "Turma": "Mat. 2-1: Linguagens e Tecnologias"},
                {"Horario": "08:40 - 09:30", "Professor": "Prof. Karina", "Turma": "Etapa 1-1: Linguagens e Tecnologias"},
                {"Horario": "15:50 - 16:40", "Professor": "Prof. Fernando", "Turma": "2º ANO 2: Linguagens e Tecnologias"}
            ],
            3: [ # Quinta-feira
                {"Horario": "09:30 - 10:20", "Professor": "Prof. Bruna", "Turma": "5º ANO 2: Linguagens e Tecnologias"},
                {"Horario": "10:20 - 11:10", "Professor": "Prof. Josiane", "Turma": "4º ANO 3: Linguagens e Tecnologias"},
                {"Horario": "12:30 - 13:20", "Professor": "Prof. Elaine", "Turma": "3º ANO 2: Linguagens e Tecnologias"},
                {"Horario": "13:20 - 14:10", "Professor": "Prof. Elaine", "Turma": "3º ANO 1: Linguagens e Tecnologias"},
                {"Horario": "14:10 - 15:00", "Professor": "Prof. Elaine", "Turma": "3º ANO 3: Linguagens e Tecnologias"},
                {"Horario": "15:00 - 15:50", "Professor": "Prof. Fernando", "Turma": "1º ANO 3: Linguagens e Tecnologias"}
            ],
            4: [ # Sexta-feira
                {"Horario": "09:30 - 10:20", "Professor": "Prof. Ramon", "Turma": "Etapa 1-2: Linguagens e Tecnologias"},
                {"Horario": "10:20 - 11:10", "Professor": "Prof. Ramon", "Turma": "Etapa 1-3: Linguagens e Tecnologias"},
                {"Horario": "15:00 - 15:50", "Professor": "Prof. Fernando", "Turma": "2º ANO 3: Linguagens e Tecnologias"}
            ],
            5: [], 6: []
        }

        # --- LEITURA DO BANCO (SUPABASE) ---
        # Substituímos o conn.read pelo safe_read configurado no passo anterior
        df_agendamentos = safe_read("Agendamentos", ["Data", "Horario", "Professor", "Turma"])

        col_form, col_view = st.columns([1, 1.2], gap="large")

        with col_view:
            st.subheader("📅 Grade do Dia")
            data_selecionada = st.date_input("Escolha a data para visualizar/agendar:", format="DD/MM/YYYY")
            data_str = data_selecionada.strftime("%d/%m/%Y")
            dia_semana_idx = data_selecionada.weekday() 
            
            # 1. Carregar Agendamentos Fixos
            lista_fixos = grade_fixa.get(dia_semana_idx, [])
            df_fixos = pd.DataFrame(lista_fixos)
            if not df_fixos.empty:
                df_fixos["Tipo"] = "Fixo (Anual)"
            
            # 2. Carregar Agendamentos Avulsos do Supabase para a data específica
            # Verificamos se há dados para evitar erro de filtro em DF vazio
            if not df_agendamentos.empty:
                df_dinamico = df_agendamentos[df_agendamentos["Data"] == data_str].copy()
            else:
                df_dinamico = pd.DataFrame()

            if not df_dinamico.empty:
                df_dinamico["Tipo"] = "Reserva Avulsa"
                df_dinamico = df_dinamico[["Horario", "Professor", "Turma", "Tipo"]]
            
            # 3. Juntar tudo para a tabela
            frames = []
            if not df_fixos.empty: frames.append(df_fixos)
            if not df_dinamico.empty: frames.append(df_dinamico)
            
            if len(frames) > 0:
                df_dia_completo = pd.concat(frames, ignore_index=True)
                df_dia_completo = df_dia_completo.sort_values(by="Horario")
                
                st.dataframe(df_dia_completo[["Horario", "Professor", "Turma", "Tipo"]], use_container_width=True, hide_index=True)
                horarios_ocupados = df_dia_completo["Horario"].tolist()
            else:
                st.info(f"A sala de informática está totalmente livre no dia {data_str}.")
                horarios_ocupados = []

        with col_form:
            st.subheader("Novo Agendamento Avulso")
            horarios_escola = [
                "07:00 - 07:50", "07:50 - 08:40", "08:40 - 09:30", "09:30 - 10:20", "10:20 - 11:10", "11:10 - 12:00",
                "12:30 - 13:20", "13:20 - 14:10", "14:10 - 15:00", "15:00 - 15:50", "15:50 - 16:40", "16:40 - 17:30"
            ]
            
            horarios_disponiveis = [h for h in horarios_escola if h not in horarios_ocupados]

            with st.form("form_agendamento", clear_on_submit=True):
                # Pegamos o nome do professor logado se disponível para facilitar
                default_prof = st.session_state.get('usuario_nome', "")
                professor = st.text_input("Nome do Professor(a)", value=default_prof, placeholder="Ex: Prof. Silva")
                turma = st.text_input("Turma", placeholder="Ex: 6º Ano A")
                
                if not horarios_disponiveis:
                    st.error("Todos os horários estão lotados para este dia!")
                    horario_escolhido = None
                else:
                    horario_escolhido = st.selectbox("Horários Disponíveis", horarios_disponiveis)
                
                submit_agendamento = st.form_submit_button("💾 Confirmar Reserva", use_container_width=True)

                if submit_agendamento:
                    if not professor or not turma:
                        st.error("Por favor, preencha o Nome e a Turma.")
                    elif not horario_escolhido:
                        st.error("Selecione um horário válido.")
                    else:
                        novo_registro = pd.DataFrame([{"Data": data_str, "Horario": horario_escolhido, "Professor": professor, "Turma": turma}])
                        df_atualizado = pd.concat([df_agendamentos, novo_registro], ignore_index=True)
                        
                        # Substituímos o conn.update pelo safe_update para persistir no Supabase
                        if safe_update("Agendamentos", df_atualizado):
                            st.success(f"✅ Sala reservada com sucesso para {turma} às {horario_escolhido}!")
                            st.rerun()
                        else:
                            st.error("Erro ao salvar o agendamento no Supabase.")
# ==============================================================================
# NOVO MÓDULO: CARÔMETRO INTERATIVO (VERSÃO CORRIGIDA E BLINDADA)
# ==============================================================================
elif app_mode and "Carômetro" in app_mode:

    st.markdown('<div class="header-box"><div class="header-title">🖼️ Carômetro de Estudantes</div></div>', unsafe_allow_html=True)
    st.markdown("Visualize todos os alunos. A foto atualizada aqui reflete em todos os documentos.")
    st.divider()
    
    df_full = load_db()
    
    if df_full is None or df_full.empty:
        st.warning("⚠️ Nenhum aluno cadastrado.")
    else:
        # 1. LIMPEZA E AGREGAÇÃO (Garante que alunos únicos apareçam)
        df_temp = df_full.copy()
        df_temp['nome'] = df_temp['nome'].str.strip()
        nomes_unicos = sorted(df_temp['nome'].unique())
        
        # 2. CONFIGURAÇÃO DA GRADE
        cols = st.columns(5)
        idx_col = 0
        
        st.markdown("""
            <style>
            .caro-foto-frame { height: 160px; width: 100%; display: flex; align-items: center; justify-content: center; overflow: hidden; border-radius: 8px; background-color: #f8fafc; margin: 10px 0; border: 1px dashed #cbd5e1; }
            .caro-nome { font-weight: 800; color: #1e3a8a; font-size: 11px; height: 35px; display: flex; align-items: center; justify-content: center; text-align: center; text-transform: uppercase; line-height: 1.1; overflow: hidden; }
            .caro-prof { font-size: 10px; color: #64748b; line-height: 1.2; margin-bottom: 8px; text-align: center; height: 30px; overflow: hidden; }
            .stFileUploader section { padding: 0 !important; }
            </style>
        """, unsafe_allow_html=True)

        for nome_aluno in nomes_unicos:
            # Busca registros desse aluno específico no DataFrame carregado
            registros_aluno = df_temp[df_temp['nome'] == nome_aluno]
            
            foto_url = None
            foto_b64 = None
            prof_encontrado = "Não informado"
            
            for _, r in registros_aluno.iterrows():
                try:
                    if r["dados_json"]:
                        d = json.loads(r["dados_json"])
                        if not foto_url: foto_url = d.get("foto_url")
                        if not foto_b64: foto_b64 = d.get("foto_base64")
                        
                        if prof_encontrado == "Não informado":
                            prof_encontrado = d.get("prof_aee") or d.get("resp_ee") or d.get("prof_poli") or "Não informado"
                except Exception:
                    pass

            with cols[idx_col]:
                with st.container(border=True):
                    st.markdown(f'<div class="caro-nome">{nome_aluno}</div>', unsafe_allow_html=True)
                    
                    # --- LÓGICA DE EXIBIÇÃO CORRIGIDA ---
                    if foto_url and str(foto_url).startswith("http"):
                        # Se for link da nuvem, carrega a URL pura
                        img_html = f"<img src='{foto_url}' style='width: 100%; height: 100%; object-fit: cover;'>"
                    elif foto_b64:
                        # Limpa qualquer prefixo residual que possa ter ficado no base64
                        clean_b64 = foto_b64.replace("data:image/jpeg;base64,", "").replace("data:image/png;base64,", "")
                        # Monta a string base64 corretamente
                        img_html = f"<img src='data:image/jpeg;base64,{clean_b64}' style='width: 100%; height: 100%; object-fit: cover;'>"
                    else:
                        img_html = "<div style='font-size: 40px; opacity: 0.2;'>👤</div>"
                    
                    st.markdown(f'<div class="caro-foto-frame">{img_html}</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="caro-prof"><b>Prof(a) AEE:</b><br>{prof_encontrado}</div>', unsafe_allow_html=True)

                    # --- UPLOAD DE FOTOS (ATUALIZADO PARA BUCKETS) ---
                    key_up = f"caro_up_{nome_aluno.replace(' ', '_')}"
                    new_file = st.file_uploader("Trocar", type=["jpg", "png", "jpeg"], key=key_up, label_visibility="collapsed")
                    
                    if new_file:
                        try:
                            import uuid
                            
                            # 1. Faz Upload da nova foto para o Supabase Storage
                            nome_arquivo = f"{uuid.uuid4()}.jpg"
                            supabase.storage.from_("fotos_alunos").upload(
                                file=new_file.getvalue(),
                                path=nome_arquivo,
                                file_options={"content-type": new_file.type}
                            )
                            # 2. Pega o Link Público
                            nova_url = supabase.storage.from_("fotos_alunos").get_public_url(nome_arquivo)
                            
                            # 3. Atualização Cirúrgica no Banco (Sem usar safe_update)
                            # Puxamos apenas as linhas desse aluno
                            res_aluno = supabase.table("Alunos").select("id, dados_json").eq("nome", nome_aluno).execute()
                            registros_aluno_db = res_aluno.data if hasattr(res_aluno, 'data') else res_aluno
                            
                            for row_db in registros_aluno_db:
                                d_s = json.loads(row_db['dados_json'])
                                d_s['foto_url'] = nova_url
                                
                                # Limpa o base64 antigo se existir, para manter o banco leve
                                if 'foto_base64' in d_s:
                                    del d_s['foto_base64']
                                
                                # Atualiza APENAS esta linha no Supabase
                                supabase.table("Alunos").update({"dados_json": json.dumps(d_s, ensure_ascii=False)}).eq("id", row_db["id"]).execute()
                            
                            st.success("✅ Foto atualizada!")
                            time.sleep(1)
                            st.rerun()
                            
                        except Exception as e_up:
                            st.error(f"Erro ao salvar imagem: {e_up}")

            idx_col = (idx_col + 1) % 5

    # ==============================================================================
# MÓDULO: PLANEJAMENTO CURRICULAR E SEMANAL (SISTEMA PLANEJAR COMPLETO)
# ==============================================================================
if app_mode_regular == "📖 Planejamento Curricular":
    
    # --- CONFIGURAÇÕES DE EMAIL ---
    EMAIL_REMETENTE = "coord.rafaelaffonsoleite@gmail.com" 
    SENHA_APP_GOOGLE = "olsi hriz zocu oiyt" 
    EMAIL_COORDENACAO = "coord.rafaelaffonsoleite@gmail.com" 
    
    from dados_curriculo import CURRICULO_DB

    # --- GESTÃO DE ESTADO ISOLADA PARA O MÓDULO ---
    if 'plan_step' not in st.session_state: st.session_state.plan_step = 1
    if 'plan_conteudos' not in st.session_state: st.session_state.plan_conteudos = []
    if 'plan_config' not in st.session_state: st.session_state.plan_config = {}

    def set_plan_step(s): st.session_state.plan_step = s

    def clean(t): 
        return t.encode('latin-1', 'replace').decode('latin-1') if t else ""
        
    def get_brazil_time():
        fuso_br = timezone(timedelta(hours=-3))
        return datetime.now(fuso_br)

    def enviar_email_automatico(pdf_bytes, dados, nome_arquivo):
        if "xxxx" in SENHA_APP_GOOGLE: return False, "Configuração de e-mail pendente."
        try:
            msg = MIMEMultipart()
            msg['From'] = EMAIL_REMETENTE
            msg['To'] = EMAIL_COORDENACAO
            recipients = [EMAIL_COORDENACAO]
            if dados.get('email_prof') and "@" in dados['email_prof']:
                msg['Cc'] = dados['email_prof']
                recipients.append(dados['email_prof'])
            msg['Subject'] = f"Planejamento Entregue: {dados['professor']} - {dados['mes']}"
            corpo = f"Olá,\n\nUm novo planejamento pedagógico foi gerado.\n\nDocente: {dados['professor']}\nAno/Turma: {dados['ano']} - {', '.join(dados['turmas'])}\nPeríodo: {dados['periodo']} ({dados['quinzena']})"
            msg.attach(MIMEText(corpo, 'plain'))
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(pdf_bytes)
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename= {nome_arquivo}.pdf")
            msg.attach(part)
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(EMAIL_REMETENTE, SENHA_APP_GOOGLE)
            server.sendmail(EMAIL_REMETENTE, recipients, msg.as_string())
            server.quit()
            return True, "E-mail enviado com sucesso!"
        except Exception as e:
            return False, f"Falha no envio do e-mail: {str(e)}"

    st.markdown('<div class="header-box"><div class="header-title">📖 Sistema Planejar Integrado</div></div>', unsafe_allow_html=True)
    st.write("")
    

# =========================================================================
    # DEFINIÇÃO DE PERMISSÕES DA EQUIPA GESTORA
    # =========================================================================
    MATRICULAS_GESTAO = ['8257601', '8844051', '8084912', '8829405', '8011512', '8258411', '7047682', '88286861']

    matricula_atual = st.session_state.get('usuario_matricula', '') 
    is_gestor = matricula_atual in MATRICULAS_GESTAO

    # =========================================================================
    # FUNÇÃO: GERAR PDF DO HISTÓRICO COM VALIDAÇÃO
    # =========================================================================
    def gerar_pdf_historico(plano):
        # Função auxiliar para evitar erros de acentuação no PDF
        def cln(text):
            return str(text).encode('latin-1', 'replace').decode('latin-1')
            
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        
        # Logotipo (se existir)
        logo_e = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo_escola.jpg"
        if os.path.exists(logo_e): pdf.image(logo_e, 175, 8, 25)
        
        # Cabeçalho
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, cln('CEIEF RAFAEL AFFONSO LEITE'), 0, 1, 'C')
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 5, cln('Planejamento de Unidade de Ensino - Arquivo Oficial'), 0, 1, 'C')
        pdf.ln(10)
        
        # Bloco de Informações Iniciais
        pdf.set_fill_color(245, 247, 250)
        pdf.set_font("Arial", 'B', 9)
        pdf.cell(0, 7, cln(f"DOCENTE: {plano.get('Professor', '')}"), 1, 1, 'L', True)
        pdf.cell(0, 7, cln(f"TURMA(S): {plano.get('Turma', '')}"), 1, 1, 'L', True)
        pdf.cell(0, 7, cln(f"PERÍODO: {plano.get('Periodo', 'Não informado')}"), 1, 1, 'L', True)
        pdf.cell(0, 7, cln(f"DATA DO REGISTRO: {plano.get('Data', '')}"), 1, 1, 'L', True)
        pdf.ln(5)
        
        # Função interna para adicionar as secções de texto
        def add_secao(titulo, texto):
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 8, cln(titulo), 0, 1)
            pdf.set_font("Arial", '', 9)
            pdf.multi_cell(0, 5, cln(texto))
            pdf.ln(3)

        # Adicionar os conteúdos do planeamento
        add_secao("Matriz Curricular e Objetivos:", plano.get('Objetivos', ''))
        add_secao("Estratégias e Situação Didática:", plano.get('Estrategias', ''))
        add_secao("Recursos e Materiais:", plano.get('Recursos', ''))
        add_secao("Critérios de Avaliação:", plano.get('Avaliacao', ''))
        
        # ========================================================
        # BLOCO DE VALIDAÇÃO DA GESTÃO NO PDF
        # ========================================================
        pdf.ln(5)
        pdf.set_fill_color(220, 230, 240)
        pdf.set_font("Arial", 'B', 10)
        pdf.cell(0, 8, cln(" PARECER PEDAGÓGICO (EQUIPE GESTORA) "), 1, 1, 'C', True)
        
        status_plano = plano.get('Status', 'Aguardando')
        obs_plano = plano.get('Observacoes', '')
        if not obs_plano or obs_plano.strip() == "":
            obs_plano = "Sem apontamentos registrados pela coordenação."
            
        pdf.set_font("Arial", 'B', 9)
        pdf.cell(0, 7, cln(f" Status Atual: {status_plano}"), 'LR', 1, 'L')
        pdf.set_font("Arial", '', 9)
        pdf.multi_cell(0, 5, cln(f" Observações e Devolutivas:\n {obs_plano}"), 'LRB', 'L')
        
        # Rodapé
        pdf.set_auto_page_break(False)
        pdf.set_y(-15)
        pdf.set_font('Arial', 'I', 7)
        try:
            agora_str = get_brazil_time().strftime("%d/%m/%Y %H:%M:%S")
        except:
            from datetime import datetime
            agora_str = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        pdf.cell(0, 10, cln(f'Emitido via Sistema Planejar Integra em: {agora_str} (Cópia de Arquivo)'), 0, 0, 'C')
        
        return bytes(pdf.output())

    # =========================================================================
    # CRIAÇÃO DOS SEPARADORES (TABS) DINÂMICOS
    # =========================================================================
    if is_gestor:
        tab_realizar, tab_historico, tab_validacao = st.tabs(["📝 Realizar Planejamento", "📂 Meus Planejamentos", "✅ Validação Pedagógica"])
    else:
        tab_realizar, tab_historico = st.tabs(["📝 Realizar Planejamento", "📂 Meus Planejamentos"])

    # -------------------------------------------------------------------------
    # ABA 1: FAZER O PLANEJAMENTO (PASSOS 1, 2 E 3)
    # -------------------------------------------------------------------------
    with tab_realizar:
        progresso = {1: 33, 2: 66, 3: 100}
        st.progress(progresso[st.session_state.plan_step])
        st.write("")

        # --- PASSO 1: IDENTIFICAÇÃO ---
        if st.session_state.plan_step == 1:
            st.markdown("### 📋 Identificação do Planejamento")
            
            c1, c2 = st.columns(2)
            with c1:
                default_nome = st.session_state.get('usuario_nome', '')
                if not default_nome: default_nome = st.session_state.plan_config.get('professor', '')
                professor = st.text_input("PROFESSOR(A) RESPONSÁVEL", value=default_nome)
            with c2:
                email_prof = st.text_input("E-MAIL DO PROFESSOR (Para receber cópia)", value=st.session_state.plan_config.get('email_prof', ''))

            c3, c4 = st.columns(2)
            with c3:
                anos = list(CURRICULO_DB.keys())
                if "Maternal I" in anos: anos.remove("Maternal I"); anos.insert(0, "Maternal I")
                saved_ano = st.session_state.plan_config.get('ano')
                idx_ano = anos.index(saved_ano) if saved_ano in anos else 0
                ano = st.selectbox("ANO DE ESCOLARIDADE", anos, index=idx_ano)
                
                if "Maternal" in ano: opts = [f"{ano} - Turma 1", f"{ano} - Turma 2"]
                else:
                    qtd = {"Etapa I": 3, "Etapa II": 3, "1º Ano": 3, "2º Ano": 3, "3º Ano": 3, "4º Ano": 3, "5º Ano": 3}
                    max_t = qtd.get(ano, 3)
                    opts = [f"{prefix}{i}" for i in range(1, max_t + 1) for prefix in ([f"{ano} - Turma " if "Etapa" in ano else f"{ano} "])]
                
                valid_defaults = [t for t in st.session_state.plan_config.get('turmas', []) if t in opts]
                turmas = st.multiselect("TURMAS VINCULADAS", opts, default=valid_defaults)
            
            with c4:
                meses = {2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"}
                saved_mes = st.session_state.plan_config.get('mes')
                idx_mes = list(meses.values()).index(saved_mes) if saved_mes in list(meses.values()) else 0
                mes_nome = st.selectbox("MÊS DE REFERÊNCIA", list(meses.values()), index=idx_mes)
                mes_num = [k for k, v in meses.items() if v == mes_nome][0]
                
                if mes_num == 2:
                    quinzena_label = "Mês Inteiro"
                    periodo_texto = "01/02/2026 a 28/02/2026"
                    trimestre_doc = "1º Trimestre"
                    st.info("Nota: Fevereiro é Planejamento Mensal.")
                else:
                    q_sel = st.radio("PERÍODO DE EXECUÇÃO", ["1ª Quinzena (01-15)", "2ª Quinzena (16-Fim)"], horizontal=True)
                    quinzena_label = q_sel.split(" (")[0]
                    tri = "1º Trimestre" if mes_num <= 4 else "2º Trimestre" if mes_num <= 8 else "3º Trimestre"
                    ultimo = calendar.monthrange(2026, mes_num)[1]
                    periodo_texto = f"01/{mes_num:02d}/2026 a 15/{mes_num:02d}/2026" if "1ª" in q_sel else f"16/{mes_num:02d}/2026 a {ultimo}/{mes_num:02d}/2026"
                    trimestre_doc = tri
            
            if st.button("Avançar para Matriz Curricular ➔", type="primary", use_container_width=True):
                if not professor or not turmas or not email_prof:
                    st.error("ERRO: Preencha todos os campos obrigatórios, incluindo o e-mail.")
                else:
                    if st.session_state.plan_config.get('ano') != ano: st.session_state.plan_conteudos = []
                    st.session_state.plan_config = {
                        'professor': professor, 'email_prof': email_prof, 'ano': ano, 
                        'turmas': turmas, 'mes': mes_nome, 'periodo': periodo_texto, 
                        'trimestre': trimestre_doc, 'quinzena': quinzena_label
                    }
                    set_plan_step(2); st.rerun()

        # --- PASSO 2: MATRIZ ---
        elif st.session_state.plan_step == 2:
            ano_sel = st.session_state.plan_config['ano']
            st.markdown(f"### 📖 Matriz Curricular: **{ano_sel}**")
            
            dados = CURRICULO_DB.get(ano_sel, {})
            
            lista_libras = []
            if ano_sel in ["Etapa I", "Etapa II"]:
                lista_libras = LIBRAS_INFANTIL
            elif ano_sel in ["1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano"]:
                lista_libras = LIBRAS_FUNDAMENTAL

            titulos = []
            chaves = []
            
            if ano_sel == "Maternal I":
                titulos = ["🗣️ Linguagem Verbal", "🔢 Linguagem Matemática", "👥 Indivíduo e Sociedade"]
                chaves = ["LINGUAGEM VERBAL", "LINGUAGEM MATEMÁTICA", "INDIVÍDUO E SOCIEDADE"]
            elif ano_sel == "Maternal II":
                titulos = ["💻 Cultura Digital", "📱 Mundo Digital", "🧩 Pensamento Computacional", "🗣️ Inglês: Oralidade"]
                chaves = ["CULTURA DIGITAL", "MUNDO DIGITAL", "PENSAMENTO COMPUTACIONAL", "INGLÊS: ORALIDADE"]
            else:
                op_tec = [k for k in dados.keys() if "INGLÊS" not in k.upper()]
                op_ing = [k for k in dados.keys() if "INGLÊS" in k.upper()]
                
                if op_tec:
                    titulos.append("💻 Tecnologia & Cultura Digital")
                    chaves.append(op_tec)
                if op_ing:
                    titulos.append("🗣️ Língua Inglesa")
                    chaves.append(op_ing)

            if lista_libras:
                titulos.append("🤟 LIBRAS (Nivelamento)")
                chaves.append("ABA_LIBRAS")

            abas = st.tabs(titulos)

            for idx, aba in enumerate(abas):
                with aba:
                    categoria = chaves[idx]
                    
                    if categoria == "ABA_LIBRAS":
                        c1, c2 = st.columns(2)
                        opcoes_g = sorted(list(set([it['geral'] for it in lista_libras])))
                        g_sel = c1.selectbox("EIXO / TÓPICO (Libras)", opcoes_g, key=f"lib_g_{idx}")
                        
                        itens_filtrados = [it for it in lista_libras if it['geral'] == g_sel]
                        opcoes_e = [it['especifico'] for it in itens_filtrados]
                        e_sel = c2.selectbox("CONTEÚDO / PRÁTICA", opcoes_e, key=f"lib_e_{idx}")
                        
                        sel = next((it for it in itens_filtrados if it['especifico'] == e_sel), None)
                        if sel:
                            st.info(f"**Objetivo:** {sel['objetivo']}")
                            if st.button("Adicionar à Lista ➕", key=f"btn_lib_{idx}"):
                                st.session_state.plan_conteudos.append({
                                    'tipo': 'Libras', 
                                    'eixo': sel['eixo'], 
                                    'geral': g_sel, 
                                    'especifico': e_sel, 
                                    'objetivo': sel['objetivo']
                                })
                                st.toast("Conteúdo de Libras adicionado!")
                                
                    elif ano_sel in ["Maternal I", "Maternal II"]:
                        if categoria in dados:
                            c1, c2 = st.columns(2)
                            opcoes_g = sorted(list(set([it['geral'] for it in dados[categoria]])))
                            g_sel = c1.selectbox("CONTEÚDO GERAL", opcoes_g, key=f"inf_g_{idx}")
                            
                            itens_filtrados = [it for it in dados[categoria] if it['geral'] == g_sel]
                            opcoes_e = [it['especifico'] for it in itens_filtrados]
                            e_sel = c2.selectbox("CONTEÚDO ESPECÍFICO", opcoes_e, key=f"inf_e_{idx}")
                            
                            sel = next((it for it in itens_filtrados if it['especifico'] == e_sel), None)
                            if sel:
                                st.info(f"**Objetivo:** {sel['objetivo']}")
                                if st.button("Adicionar à Lista ➕", key=f"btn_inf_{idx}"):
                                    st.session_state.plan_conteudos.append({
                                        'tipo': categoria, 
                                        'eixo': sel['eixo'], 
                                        'geral': g_sel, 
                                        'especifico': e_sel, 
                                        'objetivo': sel['objetivo']
                                    })
                                    st.toast("Item adicionado!")
                                    
                    else:
                        filtros = categoria 
                        if filtros:
                            c1, c2 = st.columns(2)
                            g = c1.selectbox("EIXO / TÓPICO", filtros, key=f"f_g_{idx}")
                            e = c2.selectbox("CONTEÚDO / PRÁTICA", [it['especifico'] for it in dados[g]], key=f"f_e_{idx}")
                            
                            sel = next((it for it in dados[g] if it['especifico'] == e), None)
                            if sel:
                                st.info(f"**Objetivo:** {sel['objetivo']}")
                                if st.button("Adicionar à Lista ➕", key=f"btn_f_{idx}"):
                                    label_tipo = "Tecnologia" if idx == 0 else "Inglês"
                                    st.session_state.plan_conteudos.append({
                                        'tipo': label_tipo, 
                                        'eixo': sel['eixo'], 
                                        'geral': g, 
                                        'especifico': e, 
                                        'objetivo': sel['objetivo']
                                    })
                                    st.toast("Item adicionado!")

            if st.session_state.plan_conteudos:
                st.markdown("#### Conteúdos Selecionados")
                for i, it in enumerate(st.session_state.plan_conteudos):
                    col_t, col_b = st.columns([0.90, 0.10])
                    col_t.success(f"**[{it['tipo']}]** {it['geral']}: {it['especifico']}")
                    if col_b.button("Remover", key=f"del_{i}"): 
                        st.session_state.plan_conteudos.pop(i)
                        st.rerun()

            c1, c2 = st.columns(2)
            c1.button("⬅ Voltar", on_click=set_plan_step, args=(1,))
            if c2.button("Avançar para Detalhamento ➔", type="primary", use_container_width=True):
                if not st.session_state.plan_conteudos: st.error("Selecione pelo menos um conteúdo.")
                else: set_plan_step(3); st.rerun()

        # --- PASSO 3: DETALHAMENTO E EXPORTAÇÃO ---
        elif st.session_state.plan_step == 3:
            st.markdown("### ✍️ Detalhamento Pedagógico")
            
            obj_esp = st.text_area("Objetivos Específicos", height=100, value=st.session_state.plan_config.get('obj_esp', ''))
            c1, c2 = st.columns(2)
            with c1: sit = st.text_area("Situação didática", height=220, value=st.session_state.plan_config.get('sit', ''))
            with c2: rec = st.text_area("Recursos e Materiais", height=220, value=st.session_state.plan_config.get('rec', 'Descritos na situação didática'))
            recup = st.text_area("Recuperação Contínua", height=100, value=st.session_state.plan_config.get('recup', ''))
            
            st.session_state.plan_config.update({'obj_esp': obj_esp, 'sit': sit, 'rec': rec, 'recup': recup})

            def gerar_pdf(dados, conteudos):
                pdf = FPDF(); pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=30)
                logo_e = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo_escola.jpg"
                if os.path.exists(logo_e): pdf.image(logo_e, 175, 8, 25)
                pdf.set_font('Arial', 'B', 14); pdf.cell(0, 10, clean('CEIEF RAFAEL AFFONSO LEITE'), 0, 1, 'C')
                pdf.set_font('Arial', '', 10); pdf.cell(0, 5, clean('Planejamento de Unidade de Ensino'), 0, 1, 'C'); pdf.ln(10)
                pdf.set_fill_color(245, 247, 250); pdf.set_font("Arial", 'B', 9)
                pdf.cell(0, 7, clean(f"DOCENTE: {dados['professor']}"), 1, 1, 'L', True)
                pdf.cell(0, 7, clean(f"ANO: {dados['ano']} | TURMAS: {', '.join(dados['turmas'])}"), 1, 1, 'L', True)
                pdf.cell(0, 7, clean(f"MES: {dados['mes']} | PERIODO: {dados['quinzena']} | TRIMESTRE: {dados['trimestre']}"), 1, 1, 'L', True)
                pdf.cell(0, 7, clean(f"INTERVALO: {dados['periodo']}"), 1, 1, 'L', True); pdf.ln(5)
                pdf.set_font("Arial", 'B', 10); pdf.cell(0, 8, clean("MATRIZ CURRICULAR SELECIONADA"), 0, 1)
                pdf.set_fill_color(230, 230, 230); pdf.set_font("Arial", 'B', 8)
                col_w = [45, 75, 70]
                pdf.cell(col_w[0], 7, clean("Eixo / Tema"), 1, 0, 'C', True); pdf.cell(col_w[1], 7, clean("Habilidade Especifica"), 1, 0, 'C', True); pdf.cell(col_w[2], 7, clean("Objetivo do Ano"), 1, 1, 'C', True)
                pdf.set_font("Arial", '', 8)
                for it in conteudos:
                    x, y = pdf.get_x(), pdf.get_y()
                    pdf.multi_cell(col_w[0], 5, clean(f"{it['eixo']}\n({it['geral']})"), 0, 'L')
                    y1 = pdf.get_y(); pdf.set_xy(x + col_w[0], y)
                    pdf.multi_cell(col_w[1], 5, clean(it['especifico']), 0, 'L')
                    y2 = pdf.get_y(); pdf.set_xy(x + col_w[0] + col_w[1], y)
                    pdf.multi_cell(col_w[2], 5, clean(it['objetivo']), 0, 'L')
                    y3 = pdf.get_y(); max_y = max(y1, y2, y3); h_row = max_y - y
                    pdf.set_xy(x, y); pdf.cell(col_w[0], h_row, "", 1, 0); pdf.cell(col_w[1], h_row, "", 1, 0); pdf.cell(col_w[2], h_row, "", 1, 1)
                    pdf.set_y(max_y)
                pdf.ln(5); pdf.set_font("Arial", 'B', 10); pdf.cell(0, 8, clean("DETALHAMENTO PEDAGOGICO"), 0, 1)
                for l, v in [("Objetivos Especificos", dados['obj_esp']), ("Situação didática", dados['sit']), ("Recursos e Materiais", dados['rec']), ("Recuperação Contínua", dados['recup'])]:
                    pdf.set_font("Arial", 'B', 9); pdf.cell(0, 5, clean(l + ":"), 0, 1); pdf.set_font("Arial", '', 9); pdf.multi_cell(0, 5, clean(v)); pdf.ln(2)
                pdf.set_font("Arial", 'B', 9); pdf.cell(0, 5, clean("Avaliação:"), 0, 1)
                pdf.line(pdf.get_x(), pdf.get_y()+5, 200, pdf.get_y()+5); pdf.line(pdf.get_x(), pdf.get_y()+12, 200, pdf.get_y()+12); pdf.ln(15)
                pdf.set_auto_page_break(False); pdf.set_y(-15); pdf.set_font('Arial', 'I', 7)
                pdf.cell(0, 10, clean(f'Emitido via Sistema Planejar Integrado em: {get_brazil_time().strftime("%d/%m/%Y %H:%M:%S")} (GMT-3)'), 0, 0, 'C')
                pdf.set_auto_page_break(True, margin=30)
                return bytes(pdf.output())

            def gerar_docx(dados, conteudos):
                doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Arial'; font.size = Pt(10)
                table_h = doc.add_table(rows=1, cols=2); table_h.autofit = False; table_h.columns[0].width = Cm(14); table_h.columns[1].width = Cm(4)
                p = table_h.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.add_run("CEIEF RAFAEL AFFONSO LEITE\n").bold = True; p.add_run("Planejamento Digital de Linguagens e Tecnologias")
                logo_e = "logo_escola.png" if os.path.exists("logo_escola.png") else "logo_escola.jpg"
                if os.path.exists(logo_e): table_h.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; table_h.cell(0,1).paragraphs[0].add_run().add_picture(logo_e, width=Cm(3.0))
                doc.add_paragraph(); p_info = doc.add_paragraph(); p_info.add_run(f"DOCENTE: {dados['professor']}\n").bold = True; p_info.add_run(f"ANO: {dados['ano']} | TURMAS: {', '.join(dados['turmas'])}\n"); p_info.add_run(f"MES: {dados['mes']} | PERIODO: {dados['quinzena']} | TRIMESTRE: {dados['trimestre']}\n"); p_info.add_run(f"INTERVALO: {dados['periodo']}")
                doc.add_heading("Matriz Curricular Selecionada", 2); table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
                hdr = table.rows[0].cells; hdr[0].text = 'Eixo / Tema'; hdr[1].text = 'Habilidade Especifica'; hdr[2].text = 'Objetivo do Ano'
                for cell in hdr: cell.paragraphs[0].runs[0].bold = True
                for it in conteudos:
                    row = table.add_row().cells; row[0].text = f"{it['eixo']}\n({it['geral']})"; row[1].text = it['especifico']; row[2].text = it['objetivo']
                doc.add_heading("Detalhamento Pedagogico", 2)
                for l, v in [("Objetivos Especificos", dados['obj_esp']), ("Situação didática", dados['sit']), ("Recursos e Materiais", dados['rec']), ("Recuperação Contínua", dados['recup'])]:
                    p = doc.add_paragraph(); p.add_run(l + ": ").bold = True; p.add_run(v)
                p_aval = doc.add_paragraph(); p_aval.add_run("Avaliação: ").bold = True
                doc.add_paragraph("_" * 80); doc.add_paragraph("_" * 80)
                doc.add_paragraph(f"\nEmitido eletronicamente em: {get_brazil_time().strftime('%d/%m/%Y %H:%M:%S')} (GMT-3)")
                f = BytesIO(); doc.save(f); f.seek(0); return f

            c1, c2 = st.columns(2)
            if c1.button("⬅ Matriz"): set_plan_step(2); st.rerun()
            if c2.button("GERAR PLANEJAMENTO FINAL 🚀", type="primary", use_container_width=True):
                if not all([obj_esp, sit, rec, recup]): st.error("Erro: Preencha todos os campos.")
                else:
                    with st.spinner("Salvando no Supabase, gerando documentos e enviando e-mail..."):
                        f_data = st.session_state.plan_config
                        
                        # ==========================================================
                        # 1. GUARDAR NO SUPABASE
                        # ==========================================================
                        lista_obj_texto = " | ".join([f"({c['tipo']}) {c['especifico']}" for c in st.session_state.plan_conteudos])
                        turmas_juntas = f"{f_data['ano']} ({', '.join(f_data['turmas'])})"
                        
                        texto_periodo_banco = f"{f_data['mes']} - {f_data['quinzena']} ({f_data['periodo']})"
                        
                        novo_reg = {
                            "Data": get_brazil_time().strftime("%d/%m/%Y"),
                            "Periodo": texto_periodo_banco, 
                            "Professor": f_data['professor'],
                            "Turma": turmas_juntas,
                            "Componente": "Vários (Ver Objetivos)",
                            "Objetivos": lista_obj_texto,
                            "Estrategias": f_data['sit'],
                            "Recursos": f_data['rec'],
                            "Avaliacao": "Ver detalhamento no PDF gerado",
                            "Status": "Aguardando", 
                            "Observacoes": ""       
                        }
                        
                        try:
                            supabase.table("Planejamento").insert(novo_reg).execute()
                            salvou_banco = True
                        except Exception as e:
                            st.error(f"Erro ao inserir REGISTRO no Supabase: {e}")
                            salvou_banco = False
                        
                        # ==========================================================
                        # 2. GERAR PDF E WORD
                        # ==========================================================
                        w_file = gerar_docx(f_data, st.session_state.plan_conteudos)
                        p_file = gerar_pdf(f_data, st.session_state.plan_conteudos)
                        nome_arq = f"Plan_{f_data['mes']}_{f_data['ano'].replace(' ','')}"
                        
                        # ==========================================================
                        # 3. ENVIAR E-MAIL
                        # ==========================================================
                        if f_data.get('email_prof'):
                            sucesso_email, msg_email = enviar_email_automatico(p_file, f_data, nome_arq)
                            if sucesso_email: st.success(f"📧 {msg_email}")
                            else: st.warning(f"⚠️ Arquivos gerados, mas o e-mail falhou: {msg_email}")
                        else:
                            st.info("ℹ️ E-mail não enviado (endereço do professor não informado).")

                        if salvou_banco:
                            st.success("✅ Planejamento guardado com sucesso na base de dados Supabase!")
                        else:
                            st.error("⚠️ Erro ao guardar na base de dados.")

                        cd1, cd2 = st.columns(2)
                        cd1.download_button("📄 Download WORD", w_file, f"{nome_arq}.docx", use_container_width=True)
                        cd2.download_button("📕 Download PDF", p_file, f"{nome_arq}.pdf", use_container_width=True)

    # -------------------------------------------------------------------------
    # ABA 2: CONSULTAR O HISTÓRICO NO SUPABASE
    # -------------------------------------------------------------------------
    with tab_historico:
        st.subheader("🗄️ Meus Planejamentos Guardados")
        
        df_todos = safe_read("Planejamento", ["id", "Data", "Periodo", "Professor", "Turma", "Componente", "Objetivos", "Estrategias", "Recursos", "Avaliacao", "Status", "Observacoes"])
        
        if not df_todos.empty:
            if 'Periodo' not in df_todos.columns: df_todos['Periodo'] = 'Não informado'
            else: df_todos['Periodo'] = df_todos['Periodo'].fillna('Não informado')
            
            if 'Status' not in df_todos.columns: df_todos['Status'] = 'Aguardando'
            else: df_todos['Status'] = df_todos['Status'].fillna('Aguardando')
            
            if 'Observacoes' not in df_todos.columns: df_todos['Observacoes'] = ''
            
            meu_nome = st.session_state.get('usuario_nome', '')
            df_meu_hist = df_todos[df_todos["Professor"] == meu_nome].sort_index(ascending=False)
            
            if df_meu_hist.empty:
                st.info("Ainda não tem planeamentos guardados no sistema.")
            else:
                opcoes_planos = [f"{row['Periodo']} | {row['Turma']} (Criado em {row['Data']})" for idx, row in df_meu_hist.iterrows()]
                escolha = st.selectbox("Escolha um planejamento para visualizar:", opcoes_planos)
                
                idx_escolhido = df_meu_hist.index[opcoes_planos.index(escolha)]
                plano = df_meu_hist.loc[idx_escolhido]
                id_plano = plano['id'] 
                
                st.markdown("---")
                st.markdown(f"### 📄 Planejamento: {plano['Periodo']}")
                
                status_atual = plano.get('Status', 'Aguardando')
                if status_atual == 'Validado':
                    st.success("✅ **STATUS:** Este planejamento foi validado pela Equipe Gestora.")
                elif status_atual == 'Correção':
                    st.error("⚠️ **STATUS:** Este planejamento requer adequações. Leia os apontamentos abaixo e edite.")
                else:
                    st.warning("⏳ **STATUS:** Planejamento aguardando análise da Equipe Gestora.")
                
                obs_gestao = plano.get('Observacoes', '')
                if obs_gestao and obs_gestao.strip() != "":
                    st.info(f"**🗣️ Apontamentos da Gestão:**\n\n{obs_gestao}")
                
                col_a, col_b = st.columns(2)
                with col_a:
                    st.write(f"**🏫 Turma:** {plano['Turma']}")
                    st.write(f"**📅 Salvo em:** {plano['Data']}")
                with col_b:
                    st.write(f"**👤 Professor:** {plano['Professor']}")
                
                st.info(f"**🎯 Objetivos e Conteúdos:**\n\n{plano['Objetivos']}")
                
                with st.expander("📝 Estratégias e Metodologia", expanded=True):
                    st.write(plano['Estrategias'])
                
                with st.expander("🛠️ Recursos e Avaliação"):
                    st.write(f"**Materiais:** {plano['Recursos']}")
                    st.write(f"**Avaliação:** {plano['Avaliacao']}")
                
                st.markdown("---")
                st.markdown("### ⚙️ Ações do Planejamento")
                
                # --- NOVO: BOTÃO DE DOWNLOAD DO PDF DO HISTÓRICO ---
                try:
                    pdf_hist_bytes = gerar_pdf_historico(plano)
                    st.download_button(
                        label="📥 Baixar Documento PDF Oficial",
                        data=pdf_hist_bytes,
                        file_name=f"Planejamento_{plano['Professor'].split()[0]}_{plano['Data'].replace('/','-')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as e:
                    st.warning(f"Não foi possível gerar o PDF deste REGISTRO. Erro interno: {e}")
                
                # --- EDIÇÃO ---
                with st.expander("✏️ Editar este planejamento"):
                    with st.form(key=f"form_edit_{id_plano}"):
                        st.markdown("Altere os campos abaixo e salve para atualizar o banco de dados:")
                        edit_estrategias = st.text_area("Estratégias e Metodologia", value=plano['Estrategias'], height=150)
                        edit_recursos = st.text_area("Recursos e Materiais", value=plano['Recursos'], height=100)
                        edit_avaliacao = st.text_area("Avaliação", value=plano['Avaliacao'], height=100)
                        
                        if st.form_submit_button("💾 Salvar Alterações", type="primary"):
                            try:
                                supabase.table("Planejamento").update({
                                    "Estrategias": edit_estrategias,
                                    "Recursos": edit_recursos,
                                    "Avaliacao": edit_avaliacao
                                }).eq("id", id_plano).execute()
                                
                                st.success("Planejamento atualizado com sucesso!")
                                st.rerun() 
                            except Exception as e:
                                st.error(f"Erro ao atualizar no banco de dados: {e}")

                # --- EXCLUSÃO ---
                if st.button("🗑️ Apagar este planejamento", type="secondary", use_container_width=True):
                    try:
                        supabase.table("Planejamento").delete().eq("id", id_plano).execute()
                        st.success("Planejamento apagado com sucesso!")
                        st.rerun() 
                    except Exception as e:
                        st.error(f"Erro ao apagar no banco de dados: {e}")
                        
        else:
            st.write("O banco de dados de planejamentos está vazio.")

    # -------------------------------------------------------------------------
    # ABA 3: VALIDAÇÃO PEDAGÓGICA (APENAS PARA A EQUIPA GESTORA)
    # -------------------------------------------------------------------------
    if is_gestor:
        with tab_validacao:
            st.subheader("✅ Gestão e Validação de Planejamentos")
            
            df_geral = safe_read("Planejamento", ["id", "Data", "Periodo", "Professor", "Turma", "Componente", "Objetivos", "Estrategias", "Recursos", "Avaliacao", "Status", "Observacoes"])
            
            if not df_geral.empty:
                if 'Periodo' not in df_geral.columns: df_geral['Periodo'] = 'Não informado'
                else: df_geral['Periodo'] = df_geral['Periodo'].fillna('Não informado')
                
                if 'Status' not in df_geral.columns: df_geral['Status'] = 'Aguardando'
                else: df_geral['Status'] = df_geral['Status'].fillna('Aguardando')
                
                if 'Observacoes' not in df_geral.columns: df_geral['Observacoes'] = ''
                
                status_filtro = st.radio("Filtrar por Status:", ["Todos", "Aguardando", "Correção", "Validado"], horizontal=True)
                
                df_filtrado = df_geral if status_filtro == "Todos" else df_geral[df_geral['Status'] == status_filtro]
                
                if df_filtrado.empty:
                    st.info(f"Nenhum planejamento encontrado com o status '{status_filtro}'.")
                else:
                    opcoes_val = [f"Prof. {row['Professor']} | {row['Periodo']} | {row['Turma']} ({row['Status']})" for idx, row in df_filtrado.iterrows()]
                    escolha_val = st.selectbox("Selecione o planejamento para análise:", opcoes_val)
                    
                    idx_val = df_filtrado.index[opcoes_val.index(escolha_val)]
                    plano_val = df_filtrado.loc[idx_val]
                    id_plano_val = plano_val['id']
                    
                    st.markdown("---")
                    st.markdown(f"### Análise: Prof. {plano_val['Professor']} ({plano_val['Turma']})")
                    st.caption(f"**Período:** {plano_val['Periodo']} | **Salvo no sistema em:** {plano_val['Data']}")
                    
                    with st.expander("🔍 Visualizar Conteúdo do Planejamento", expanded=True):
                        st.info(f"**Objetivos/Habilidades:**\n{plano_val['Objetivos']}")
                        st.write(f"**Situação Didática:**\n{plano_val['Estrategias']}")
                        st.write(f"**Recursos e Avaliação:**\n{plano_val['Recursos']} | {plano_val['Avaliacao']}")
                    
                    # --- NOVO: BOTÃO DE DOWNLOAD DO PDF PARA A GESTÃO ---
                    try:
                        pdf_val_bytes = gerar_pdf_historico(plano_val)
                        st.download_button(
                            label="📥 Baixar Cópia em PDF (Com Parecer Atual)",
                            data=pdf_val_bytes,
                            file_name=f"Avaliacao_{plano_val['Professor'].split()[0]}_{plano_val['Data'].replace('/','-')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.warning("Não foi possível gerar o PDF de pré-visualização.")

                    with st.form(key=f"form_val_{id_plano_val}"):
                        st.markdown("#### 📋 Parecer da Coordenação/Direção")
                        
                        lista_status = ["Aguardando", "Validado", "Correção"]
                        idx_status_atual = lista_status.index(plano_val['Status']) if plano_val['Status'] in lista_status else 0
                        
                        novo_status = st.radio("Mudar Status para:", lista_status, index=idx_status_atual, horizontal=True)
                        
                        nova_obs = st.text_area("Apontamentos / Devolutiva:", 
                                              value=plano_val.get('Observacoes', ''), 
                                              height=120,
                                              help="Se solicitar correção, explique aqui o que precisa ser alterado.")
                        
                        if st.form_submit_button("Salvar Validação", type="primary"):
                            try:
                                supabase.table("Planejamento").update({
                                    "Status": novo_status,
                                    "Observacoes": nova_obs
                                }).eq("id", id_plano_val).execute()
                                
                                st.success("Parecer pedagógico salvo com sucesso!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erro ao salvar validação: {e}")
            else:
                st.write("Nenhum planejamento foi enviado pelos professores ainda.")

        # ==============================================================================
        # FERRAMENTA DE BACKUP VERSIONADO (COFRE NO GOOGLE SHEETS E LOCAL)
        # ==============================================================================
        
if app_mode_regular == "💾 Cofre de Segurança":

        st.divider()
        st.markdown("💾 Cofre de Segurança")
        st.caption("Crie cópias de segurança do seu banco de dados na nuvem (Google Sheets) ou baixe para o seu computador.")

        # --- OPÇÃO 1: BACKUP NA NUVEM (NOVA ABA NO GOOGLE SHEETS) ---
        if st.button("🔄 Gerar Backup de Segurança (Nuvem)", type="primary", use_container_width=True):
            with st.spinner("Construindo aba e copiando dados no Google Sheets..."):
                try:
                    df_backup = load_db()
                    
                    if df_backup is not None and not df_backup.empty:
                        agora = pd.Timestamp.now().strftime("%d_%m_%Y_%H%M")
                        nome_aba_backup = f"BKP_{agora}"
                        
                        from streamlit_gsheets import GSheetsConnection
                        conn_backup = st.connection("gsheets", type=GSheetsConnection)
                        
                        # 1. A SOLUÇÃO: Forçar a criação da aba vazia antes de jogar os dados
                        try:
                            # Puxa o link da sua planilha que já está salvo nos segredos do sistema
                            url_planilha = st.secrets["connections"]["gsheets"]["spreadsheet"]
                            planilha_mestra = conn_backup.client.open_by_url(url_planilha)
                            
                            # Manda o Google criar a aba com espaço de sobra (ex: 2000 linhas)
                            planilha_mestra.add_worksheet(title=nome_aba_backup, rows="2000", cols="40")
                        except Exception as erro_aba:
                            # Se der um pequeno erro na criação, ele ignora e tenta salvar mesmo assim
                            pass 
                        
                        # 2. Agora sim, atualiza a aba recém-criada com a tabela inteira!
                        conn_backup.update(worksheet=nome_aba_backup, data=df_backup)
                        
                        st.success(f"✅ Backup blindado com sucesso! A aba '{nome_aba_backup}' foi criada na sua planilha original.")
                    else:
                        st.error("🛑 O banco do Supabase está vazio. Operação cancelada.")
                except Exception as e:
                    st.error(f"Erro na comunicação com a nuvem: {e}")

        # --- OPÇÃO 2: BACKUP FÍSICO (DOWNLOAD IMEDIATO) ---
        st.markdown("<br>", unsafe_allow_html=True)
        st.caption("Prefere ter o arquivo físico no seu computador para segurança máxima? Baixe agora (não precisa de internet):")
        
        try:
            df_fisico = load_db()
            if df_fisico is not None and not df_fisico.empty:
                # Converte o banco de dados para formato CSV (Excel)
                csv = df_fisico.to_csv(index=False).encode('utf-8')
                agora_str = pd.Timestamp.now().strftime("%d_%m_%Y")
                
                st.download_button(
                    label="📥 Baixar Backup Físico (Excel/CSV)",
                    data=csv,
                    file_name=f"Backup_Alunos_{agora_str}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
        except Exception:
            pass # Apenas esconde o botão se o banco estiver indisponível no momento


# ==============================================================================
# MÓDULO EXCLUSIVO: ABA CARÔMETRO (ENSINO REGULAR)
# ==============================================================================
elif app_mode_regular == "🖼️ Carômetro Escolar":
    st.markdown('<div class="header-box"><div class="header-title">🖼️ Carômetro Escolar</div></div>', unsafe_allow_html=True)
    st.markdown("Visualize as turmas. A foto atualizada aqui reflete no sistema.")
    st.divider()

    turmas_escola = [
        "B1-1", "B2-1", "B2-2", "M1-1", "M2-1", "M2-2",
        "ETAPA 1-1", "ETAPA 1-2", "ETAPA 1-3",
        "ETAPA 2-1", "ETAPA 2-2", "ETAPA 2-3",
        "1º ANO 1", "1º ANO 2", "1º ANO 3", "2º ANO 1", "2º ANO 2", "2º ANO 3",
        "3º ANO 1", "3º ANO 2", "3º ANO 3", "4º ANO 1", "4º ANO 2", "4º ANO 3",
        "5º ANO 1", "5º ANO 2", "5º ANO 3",
        "BILÍNGUE 1", "BILÍNGUE 2", "BILÍNGUE 3"
    ]

    turma_sel = st.selectbox("Selecione a Turma:", ["-- Escolha --"] + turmas_escola)

    if turma_sel != "-- Escolha --":
        # Carrega os alunos do banco de dados do Carômetro
        df_caro = load_carometro_db()
        
        if not df_caro.empty:
            df_turma = df_caro[df_caro['turma'] == turma_sel].sort_values(by="nome")
        else:
            df_turma = pd.DataFrame()

        if df_turma.empty:
            st.warning(f"⚠️ Nenhum aluno encontrado na turma {turma_sel}.")
        else:
            # 1. CONFIGURAÇÃO DA GRADE (Idêntico ao AEE)
            cols = st.columns(5)
            idx_col = 0
            
            # CSS EXATAMENTE IGUAL AO DO AEE
            st.markdown("""
                <style>
                .caro-foto-frame { height: 160px; width: 100%; display: flex; align-items: center; justify-content: center; overflow: hidden; border-radius: 8px; background-color: #f8fafc; margin: 10px 0; border: 1px dashed #cbd5e1; }
                .caro-nome { font-weight: 800; color: #1e3a8a; font-size: 11px; height: 35px; display: flex; align-items: center; justify-content: center; text-align: center; text-transform: uppercase; line-height: 1.1; overflow: hidden; }
                .caro-prof { font-size: 10px; color: #64748b; line-height: 1.2; margin-bottom: 8px; text-align: center; height: 20px; overflow: hidden; }
                .stFileUploader section { padding: 0 !important; }
                </style>
            """, unsafe_allow_html=True)

            for row in df_turma.itertuples():
                with cols[idx_col]:
                    with st.container(border=True):
                        st.markdown(f'<div class="caro-nome">{row.nome}</div>', unsafe_allow_html=True)
                        
                        # Verifica se tem foto
                        if row.foto_base64:
                            img_html = f"<img src='data:image/jpeg;base64,{row.foto_base64}' style='width: 100%; height: 100%; object-fit: cover;'>"
                        else:
                            img_html = "<div style='font-size: 40px; opacity: 0.2;'>👤</div>"
                            
                        st.markdown(f'<div class="caro-foto-frame">{img_html}</div>', unsafe_allow_html=True)
                        st.markdown(f'<div class="caro-prof"><b>Turma:</b> {row.turma}</div>', unsafe_allow_html=True)
                        
                        # Botão de upload idêntico ao AEE (Label "collapsed" para ficar clean)
                        key_up = f"caro_up_reg_{row.id}"
                        new_file = st.file_uploader("Trocar", type=["jpg", "png", "jpeg"], key=key_up, label_visibility="collapsed")
                        
                        if new_file:
                            try:
                                from PIL import Image
                                import io, base64, time
                                
                                img = Image.open(new_file)
                                if img.mode != 'RGB': img = img.convert('RGB')
                                img.thumbnail((400, 500))
                                buf = io.BytesIO()
                                img.save(buf, format="JPEG", quality=85)
                                nova_foto_b64 = base64.b64encode(buf.getvalue()).decode()
                                
                                # Salva a nova foto no banco de dados usando a função do carômetro
                                if save_carometro_entry(row.nome, row.turma, nova_foto_b64):
                                    st.success("✅ Foto atualizada!")
                                    time.sleep(1)
                                    st.rerun()
                                else:
                                    st.error("Erro ao atualizar banco.")
                            except Exception as e_up:
                                st.error(f"Erro na imagem: {e_up}")

                idx_col = (idx_col + 1) % 5


 

















import uuid
import time
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from datetime import datetime
from fpdf import FPDF

# ==============================================================================
# MÓDULO: ADMINISTRATIVO (ALMOXARIFADO ESCOLAR - NÃO PEDAGÓGICO)
# ==============================================================================
if app_mode_adm == "📦 Almoxarifado Escolar":
    st.markdown('<div class="header-box"><div class="header-title">📦 Gestão de Almoxarifado Escolar</div></div>', unsafe_allow_html=True)
    
    if st.button("⬅️ Voltar ao Menu Inicial", key="voltar_adm"):
        st.session_state.modulo_atuacao = None
        st.rerun()

    # Leitura de dados via Supabase (utilizando a tua função safe_read personalizada)
    df_estoque = safe_read("Almoxarifado_Estoque", ["id", "item", "quantidade", "categoria"])
    df_pedidos = safe_read("Almoxarifado_Pedidos", ["id", "data", "professor", "item", "quantidade", "status"])
    
    # ---------------------------------------------------------
    # 🔒 CONTROLO DE ACESSO E SEGURANÇA
    # ---------------------------------------------------------
    # Lista oficial de gestores autorizados
    MATRICULAS_GESTAO = ['8257601', '8844051', '8084912', '8829405', '8011512', '8258411', '7047682', '88286861']
    
    # Normalização da matrícula para evitar erros de tipo (int/str)
    matricula_atual = str(st.session_state.get('usuario_matricula', '')).strip() 
    eh_gestao = matricula_atual in MATRICULAS_GESTAO

    # =========================================================
    # VISUALIZAÇÃO DO COMPROVATIVO DE EXPEDIÇÃO (MODAL)
    # =========================================================
    if eh_gestao and 'comprovante_almox' in st.session_state and st.session_state.comprovante_almox:
        st.success("✅ Saída de materiais processada com sucesso no inventário!")
        st.markdown("### 🖨️ Comprovativo de Entrega")
        
        cupom_html = f"""
        <div id="cupom_impressao" style="width: 280px; padding: 15px; border: 1px dashed #000; font-family: 'Courier New', Courier, monospace; font-size: 13px; background: #fff; color: #000; margin: 0 auto;">
            <div style="text-align: center; margin-bottom: 10px;">
                <strong>CEIEF RAFAEL AFFONSO LEITE</strong><br>
                ALMOXARIFADO ESCOLAR<br>
                NÃO PEDAGÓGICO<br>
                COMPROVATIVO DE ENTREGA<br>
                --------------------------------
            </div>
            <strong>Data:</strong> {st.session_state.comprovante_almox['data']}<br>
            <strong>Requisitante:</strong> {st.session_state.comprovante_almox['professor']}<br>
            --------------------------------<br>
            <strong>MATERIAIS ENTREGUES:</strong><br>
            <ul style="padding-left: 15px; margin-top: 5px; margin-bottom: 5px;">
                {st.session_state.comprovante_almox['itens_html']}
            </ul>
            --------------------------------<br>
            <div style="text-align: center; margin-top: 30px;">
                ________________________________<br>
                <span style="font-size: 11px;">Assinatura do Requisitante</span><br><br>
                ________________________________<br>
                <span style="font-size: 11px;">Responsável Almoxarifado</span>
            </div>
        </div>
        
        <script>
        function imprimirCupom() {{
            var conteudo = document.getElementById('cupom_impressao').outerHTML;
            var tela_impressao = window.open('', '', 'height=600,width=400');
            tela_impressao.document.write('<html><head><title>Talão Almoxarifado</title></head>');
            tela_impressao.document.write('<body style="margin: 0; padding: 10px;">');
            tela_impressao.document.write(conteudo);
            tela_impressao.document.write('</body></html>');
            tela_impressao.document.close();
            tela_impressao.focus();
            setTimeout(function() {{ tela_impressao.print(); }}, 500);
        }}
        </script>
        
        <button onclick="imprimirCupom()" style="display: block; width: 100%; max-width: 330px; margin: 15px auto; padding: 12px; background-color: #2e7d32; color: white; border: none; border-radius: 6px; font-weight: bold; font-size: 16px; cursor: pointer;">🖨️ Imprimir na Bobina</button>
        """
        components.html(cupom_html, height=550)
        
        if st.button("🔄 Concluir e Retornar ao Painel", use_container_width=True):
            st.session_state.comprovante_almox = None
            st.rerun()

    else:
        # Definição das Abas conforme o nível de acesso
        if eh_gestao:
            tab_req, tab_retro, tab_baixa, tab_estoque = st.tabs([
                "🙋 Nova Solicitação", "📝 Registro Manual", "📦 Expedição", "📈 Inventário"
            ])
        else:
            tab_req, = st.tabs(["🙋 Nova Solicitação"])

        # --- ABA 1: SOLICITAÇÃO (TODOS OS UTILIZADORES) ---
        with tab_req:
            st.subheader("Requisição de Materiais")
            if not df_estoque.empty:
                if 'reset_carrinho' not in st.session_state:
                    st.session_state.reset_carrinho = 0

                selecionados = st.multiselect(
                    "Selecione os materiais no catálogo:", 
                    df_estoque['item'].tolist(),
                    key=f"carrinho_{st.session_state.reset_carrinho}"
                )
                
                if selecionados:
                    with st.form("form_solicitacao"):
                        dict_qtds = {}
                        cols = st.columns(2)
                        for i, item in enumerate(selecionados):
                            target = cols[i % 2]
                            dict_qtds[item] = target.number_input(f"📦 {item}", min_value=1, step=1)
                        
                        if st.form_submit_button("Submeter Solicitação", type="primary", use_container_width=True):
                            agora = datetime.now().strftime("%d/%m/%Y %H:%M")
                            user_nome = st.session_state.get('usuario_nome', 'Usuário').upper()
                            for item, qtd in dict_qtds.items():
                                supabase.table("Almoxarifado_Pedidos").insert({
                                    "id": str(uuid.uuid4()), "data": agora, "professor": user_nome,
                                    "item": item, "quantidade": qtd, "status": "Pendente"
                                }).execute()
                            st.success("✅ Solicitação enviada!")
                            st.session_state.reset_carrinho += 1
                            time.sleep(1); st.rerun()
            
            st.divider()
            st.markdown("##### 🕒 Histórico de Solicitações")
            meus_pedidos = df_pedidos[df_pedidos['professor'] == st.session_state.get('usuario_nome', '').upper()]
            if not meus_pedidos.empty:
                st.dataframe(meus_pedidos.iloc[::-1][["data", "item", "quantidade", "status"]], use_container_width=True, hide_index=True)

        # =========================================================
        # 🔒 FERRAMENTAS EXCLUSIVAS DA GESTÃO
        # =========================================================
        if eh_gestao:
            # --- ABA 2: REGISTRO MANUAL (ENTREGA JÁ REALIZADA) ---
            with tab_retro:
                st.subheader("Registro de Entregas Realizadas")
                df_prof = safe_read("Professores", ["nome"])
                df_mon = safe_read("Monitores", ["nome"])
                lista_profs = sorted(list(set([str(n).upper() for n in (df_prof['nome'].tolist() + df_mon['nome'].tolist() + df_pedidos['professor'].tolist()) if pd.notnull(n)])))

                with st.form("REGISTRO_manual", clear_on_submit=True):
                    col_n, col_d = st.columns([2, 1])
                    requisitante = col_n.selectbox("Profissional:", ["-- Selecione --"] + lista_profs)
                    nome_extra = col_n.text_input("Ou digite o nome completo (Maiúsculas):")
                    data_e = col_d.date_input("Data da Entrega:", datetime.now())
                    itens_m = st.multiselect("Materiais Fornecidos:", df_estoque['item'].tolist())
                    
                    st.divider()
                    dict_manual = {}
                    if itens_m:
                        c_m = st.columns(2)
                        for i, it in enumerate(itens_m):
                            dict_manual[it] = c_m[i % 2].number_input(f"{it}", min_value=1, step=1)

                    if st.form_submit_button("🚀 Efetivar REGISTRO e Gerar Talão", type="primary", use_container_width=True):
                        final_nome = nome_extra.strip().upper() if nome_extra.strip() else requisitante
                        if final_nome != "-- Selecione --" and itens_m:
                            html_itens = ""
                            for it, q in dict_manual.items():
                                supabase.table("Almoxarifado_Pedidos").insert({"id": str(uuid.uuid4()), "data": data_e.strftime("%d/%m/%Y"), "professor": final_nome, "item": it, "quantidade": q, "status": "Entregue"}).execute()
                                stock_at = df_estoque[df_estoque['item'] == it].iloc[0]['quantidade']
                                supabase.table("Almoxarifado_Estoque").update({"quantidade": max(0, int(stock_at) - q)}).eq("item", it).execute()
                                html_itens += f"<li>{q}x {it}</li>"
                            st.session_state.comprovante_almox = {'data': datetime.now().strftime("%d/%m/%Y %H:%M"), 'professor': final_nome, 'itens_html': html_itens}
                            st.rerun()

            # --- ABA 3: EXPEDIÇÃO (SAÍDA PARCIAL E INTELIGENTE) ---
            with tab_baixa:
                st.subheader("Processamento de Expedição")
                pendentes = df_pedidos[df_pedidos['status'] == 'Pendente']
                if pendentes.empty:
                    st.info("Nenhuma solicitação pendente.")
                else:
                    for (dt, prof), grupo in pendentes.groupby(['data', 'professor']):
                        with st.expander(f"📦 Solicitação: {prof} | {dt}"):
                            with st.form(f"f_{dt}_{prof}".replace(" ","")):
                                picks = {}
                                for _, row in grupo.iterrows():
                                    c1, c2, c3 = st.columns([0.5, 3, 1.5])
                                    at_stock = df_estoque[df_estoque['item'] == row['item']].iloc[0]['quantidade']
                                    chk = c1.checkbox("", value=True, key=f"c_{row['id']}")
                                    c2.markdown(f"**{row['item']}** \nSolicitado: {row['quantidade']} | Stock: {at_stock}")
                                    lib = c3.number_input("Qtd a Entregar", min_value=1, value=int(row['quantidade']), key=f"q_{row['id']}")
                                    picks[row['id']] = {'entregar': chk, 'qtd': lib, 'item': row['item'], 'original': int(row['quantidade'])}
                                
                                if st.form_submit_button("Confirmar Expedição e Imprimir", type="primary"):
                                    html_talao = ""
                                    for rid, info in picks.items():
                                        if info['entregar']:
                                            supabase.table("Almoxarifado_Pedidos").update({"status": "Entregue", "quantidade": info['qtd']}).eq("id", rid).execute()
                                            if info['qtd'] < info['original']:
                                                supabase.table("Almoxarifado_Pedidos").insert({"id": str(uuid.uuid4()), "data": dt, "professor": prof, "item": info['item'], "quantidade": info['original'] - info['qtd'], "status": "Pendente"}).execute()
                                            stk = df_estoque[df_estoque['item'] == info['item']].iloc[0]['quantidade']
                                            supabase.table("Almoxarifado_Estoque").update({"quantidade": max(0, int(stk) - info['qtd'])}).eq("item", info['item']).execute()
                                            html_talao += f"<li>{info['qtd']}x {info['item']}</li>"
                                    if html_talao:
                                        st.session_state.comprovante_almox = {'data': datetime.now().strftime("%d/%m/%Y %H:%M"), 'professor': prof, 'itens_html': html_talao}
                                        st.rerun()

            # --- ABA 4: INVENTÁRIO (RELATÓRIO PDF PROFISSIONAL) ---
            with tab_estoque:
                col_tab, col_in = st.columns([3, 2])
                with col_tab:
                    st.subheader("Posição do Inventário")
                    if not df_estoque.empty:
                        def export_pdf(df, user):
                            pdf = FPDF()
                            pdf.add_page()
                            pdf.set_font("Arial", "B", 14)
                            pdf.cell(0, 8, "CEIEF RAFAEL AFFONSO LEITE", ln=True, align="C")
                            pdf.set_font("Arial", "B", 10)
                            pdf.cell(0, 6, "ALMOXARIFADO ESCOLAR - NAO PEDAGOGICO", ln=True, align="C")
                            pdf.cell(0, 6, "RELATORIO DE INVENTARIO", ln=True, align="C")
                            pdf.ln(5)
                            pdf.set_font("Arial", "", 9)
                            pdf.cell(0, 5, f"Emissao: {datetime.now().strftime('%d/%m/%Y %H:%M')} | Por: {user}", ln=True)
                            pdf.ln(5)
                            pdf.set_fill_color(200, 220, 255)
                            pdf.set_font("Arial", "B", 9)
                            pdf.cell(100, 8, "Material", 1, 0, "L", True)
                            pdf.cell(50, 8, "Categoria", 1, 0, "L", True)
                            pdf.cell(40, 8, "Saldo", 1, 1, "C", True)
                            fill = False
                            for _, r in df.sort_values("item").iterrows():
                                pdf.set_fill_color(245, 245, 245)
                                pdf.set_font("Arial", "", 9)
                                pdf.cell(100, 7, str(r['item'])[:55].encode('latin-1','replace').decode('latin-1'), 1, 0, "L", fill)
                                pdf.cell(50, 7, str(r['categoria'])[:25].encode('latin-1','replace').decode('latin-1'), 1, 0, "L", fill)
                                pdf.cell(40, 7, str(r['quantidade']), 1, 1, "C", fill)
                                fill = not fill
                            
                            res = pdf.output()
                            return bytes(res) if not isinstance(res, str) else res.encode("latin-1")

                        st.download_button("📄 Exportar PDF Oficial", export_pdf(df_estoque, st.session_state.get('usuario_nome','')), f"Inventario_{datetime.now().strftime('%d_%m')}.pdf", "application/pdf", use_container_width=True)
                    st.dataframe(df_estoque[["item", "quantidade", "categoria"]].sort_values("item"), use_container_width=True, hide_index=True)

                with col_in:
                    st.subheader("📥 Entrada de Materiais")
                    if 'reset_in' not in st.session_state: st.session_state.reset_in = 0
                    chegaram = st.multiselect("Itens Recebidos:", df_estoque['item'].tolist(), key=f"in_{st.session_state.reset_in}")
                    if chegaram:
                        with st.form("f_entrada"):
                            ins = {i: st.number_input(f"Qtd: {i}", min_value=1) for i in chegaram}
                            if st.form_submit_button("Confirmar Entrada"):
                                for it, q in ins.items():
                                    at = df_estoque[df_estoque['item'] == it].iloc[0]['quantidade']
                                    supabase.table("Almoxarifado_Estoque").update({"quantidade": int(at) + q}).eq("item", it).execute()
                                st.success("Inventário Atualizado!"); st.session_state.reset_in += 1; time.sleep(1); st.rerun()


































import uuid
import time
import math  # <--- ADICIONE ESTA LINHA AQUI
import requests
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
import base64
from datetime import datetime, timedelta

# ==============================================================================
# MÓDULO: GESTÃO DE ACERVO E BIBLIOTECONOMIA (SALA DE LEITURA)
# ==============================================================================
if st.session_state.get("modulo_atuacao") in ["📚  Sala de Leitura", "📚 Sala de Leitura"]:
    st.markdown('<div class="header-box"><div class="header-title">📚 Gestão de Acervo e Leitores</div></div>', unsafe_allow_html=True)
    
    if st.button("⬅️ Retornar ao Menu Inicial", key="voltar_bib"):
        st.session_state.modulo_atuacao = None
        st.rerun()

    # --- INTEGRAÇÃO CARÔMETRO & SEGURANÇA ---
    MATRICULAS_GESTAO = ['8257601', '8844051', '8084912', '8829405', '8011512', '8258411', '7047682', '88286861']
    eh_gestao = str(st.session_state.get('usuario_matricula', '')).strip() in MATRICULAS_GESTAO
    usuario_nome = st.session_state.get('usuario_nome', 'Visitante').upper()

    df_carometro = load_carometro_db() 
    if not df_carometro.empty:
        df_carometro['display_leitor'] = df_carometro['nome'].str.upper() + " (" + df_carometro['turma'].str.upper() + ")"
        lista_leitores = sorted(df_carometro['display_leitor'].tolist())
    else:
        lista_leitores = []

    # Leitura das Tabelas
    df_acervo = safe_read("Biblioteca_Acervo", ["id", "isbn", "titulo", "autor", "editora", "genero", "serie", "capa_url"])
    df_exemplares = safe_read("Biblioteca_Exemplares", ["id", "id_acervo", "tombo", "disponivel", "localizacao"])
    df_emp = safe_read("Biblioteca_Emprestimos", ["id", "id_exemplar", "leitor", "data_saida", "data_prevista", "status"])

    # =========================================================
    # FUNÇÕES AUXILIARES VISUAIS - SALA DE LEITURA
    # =========================================================
    def gerar_capa_dinamica(id_livro, titulo, autor):
        """
        Gera o HTML de uma capa de livro colorida com base no ID da obra.
        Cada livro terá uma cor consistente baseada no seu ID UUID.
        """
        import zlib
        import pandas as pd
        
        # Paleta de degradês modernos e vibrantes
        bg_palettes = [
            "linear-gradient(135deg, #0d9488 0%, #064e3b 100%)", # Teal-Green
            "linear-gradient(135deg, #2563eb 0%, #172554 100%)", # Blue
            "linear-gradient(135deg, #7c3aed 0%, #4c1d95 100%)", # Purple
            "linear-gradient(135deg, #ea580c 0%, #7c2d12 100%)", # Orange
            "linear-gradient(135deg, #db2777 0%, #831843 100%)", # Pink
            "linear-gradient(135deg, #ca8a04 0%, #713f12 100%)", # Yellow/Brown
            "linear-gradient(135deg, #16a34a 0%, #14532d 100%)", # Green
            "linear-gradient(135deg, #dc2626 0%, #7f1d1d 100%)", # Red
        ]
        
        # Usa o ID do livro (UUID) para gerar um hash estável
        hash_val = zlib.adler32(str(id_livro).encode('utf-8'))
        bg_style = bg_palettes[hash_val % len(bg_palettes)]
        
        # Normalização dos textos (Título maior, autor menor)
        tit_norm = str(titulo).strip().upper()
        if len(tit_norm) > 45:
            tit_norm = tit_norm[:42] + "..."
            
        aut_norm = str(autor).strip().upper() if pd.notnull(autor) else "AUTOR NÃO INFORMADO"
        if len(aut_norm) > 30:
            aut_norm = aut_norm[:27] + "..."
            
        # O SEGREDO DO BUG DO </div> ESTÁ AQUI: Tudo numa linha só!
        return f'<div class="dynamic-cover" style="background: {bg_style};"><div class="dynamic-cover-title">{tit_norm}</div><div class="dynamic-cover-divider"></div><div class="dynamic-cover-author">{aut_norm}</div><div class="dynamic-cover-branding">INTEGRA BIBLIO</div></div>'

    # --- ESTILO VISUAL (CSS) ---
    st.markdown("""
    <style>
        /* Container Principal do Card */
        .book-card { 
            background: #f8fafc; 
            border-radius: 12px; 
            padding: 15px; 
            border: 1px solid #e2e8f0; 
            text-align: center; 
            height: 100%; 
            transition: all 0.3s ease; 
            display: flex;
            flex-direction: column;
            justify-content: space-between;
        }
        .book-card:hover { 
            transform: translateY(-8px); 
            box-shadow: 0 10px 20px rgba(0,0,0,0.1); 
            border-color: #cbd5e1;
        }
        
        /* Capa Gerada Dinamicamente (HTML) */
        .dynamic-cover {
            height: 200px;
            border-radius: 8px;
            box-shadow: 2px 2px 10px rgba(0,0,0,0.2);
            padding: 15px;
            color: white;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            border: 1px solid rgba(255,255,255,0.1);
            position: relative;
            overflow: hidden;
        }
        
        .dynamic-cover-title {
            font-weight: 800;
            font-size: 14px;
            text-transform: uppercase;
            margin-bottom: 8px;
            line-height: 1.1;
            text-shadow: 1px 1px 2px rgba(0,0,0,0.5);
            text-align: center;
            font-family: 'Open Sans', sans-serif;
            max-height: 65px;
            overflow: hidden;
        }
        
        .dynamic-cover-divider {
            border-top: 2px solid rgba(255,255,255,0.4);
            width: 80%;
            margin: 10px 0;
        }
        
        .dynamic-cover-author {
            font-size: 10px;
            font-weight: 400;
            font-style: italic;
            opacity: 0.9;
            text-align: center;
            text-transform: uppercase;
        }
        
        .dynamic-cover-branding {
            position: absolute;
            bottom: 5px;
            right: 8px;
            font-size: 8px;
            font-weight: 900;
            opacity: 0.3;
            letter-spacing: 1px;
        }
        
        /* Título Externo (para quando tem imagem real) */
        .external-book-title {
            font-weight: bold;
            font-size: 14px;
            margin-top: 10px;
            color: #1e293b;
            height: 40px;
            overflow: hidden;
            line-height: 1.2;
        }
        
        /* Badges de Conquistas */
        .badge-box { background: linear-gradient(135deg, #fbbf24 0%, #d97706 100%); color: white; padding: 15px; border-radius: 10px; text-align: center; }
        .badge-title { font-weight: 800; font-size: 18px; }
    </style>
    """, unsafe_allow_html=True)

    # --- NAVEGAÇÃO ---
    if eh_gestao:
        tabs = st.tabs(["🔍 Consulta", "👤 Prontuário", "🔄 Circulação", "🗃️ Incorporação (ISBN/Excel)", "📊 Auditoria"])
        tab_vitrine, tab_perfil, tab_circ, tab_cat, tab_dash = tabs
    else:
        tab_vitrine, tab_perfil = st.tabs(["🔍 Consulta", "👤 Meu Prontuário"])

    # --- 1. PORTAL DO LEITOR (VITRINE INTERATIVA COM PAGINAÇÃO) ---
    with tab_vitrine:
        import math 
        
        st.subheader("📚 Portal do Leitor - Acervo da Escola")
        
        # Resumo da Biblioteca (Métricas compactas) no topo
        total_titulos = len(df_acervo)
        total_exemplares = len(df_exemplares) if not df_exemplares.empty else 0
        exemplares_disponiveis = len(df_exemplares[df_exemplares['disponivel'] == True]) if not df_exemplares.empty else 0
        
        c_m1, c_m2, c_m3 = st.columns(3)
        c_m1.metric(label="Títulos Cadastrados", value=total_titulos)
        c_m2.metric(label="Total de Exemplares Físicos", value=total_exemplares)
        c_m3.metric(label="Exemplares Disponíveis Agora", value=exemplares_disponiveis)
        st.markdown("---")
        
        # Barra de pesquisa e controlos
        c_pesq, c_ordem, c_itens = st.columns([2, 1, 1])
        termo_busca = c_pesq.text_input("🔍 Pesquisar Título, Autor ou Gênero:", placeholder="Ex: Chapeuzinho...")
        ordem = c_ordem.selectbox("Ordenar por:", ["Mais Recentes", "Título (A-Z)"])
        itens_por_pagina = c_itens.selectbox("Exibir:", [12, 24, 48, 100], index=0)

        if not df_acervo.empty:
            
            # 1. Aplica o filtro de busca
            if termo_busca:
                res_vitrine = df_acervo[
                    df_acervo['titulo'].str.contains(termo_busca, case=False) | 
                    df_acervo['autor'].str.contains(termo_busca, case=False) |
                    df_acervo['genero'].astype(str).str.contains(termo_busca, case=False)
                ]
            else:
                res_vitrine = df_acervo.copy()

            # 2. Ordenação
            if ordem == "Título (A-Z)":
                res_vitrine = res_vitrine.sort_values(by="titulo")
            else:
                res_vitrine = res_vitrine.iloc[::-1]

            if not res_vitrine.empty:
                # --- LÓGICA DE PAGINAÇÃO ---
                total_paginas = max(1, math.ceil(len(res_vitrine) / itens_por_pagina))
                
                if total_paginas > 1:
                    c_vazio, c_pag = st.columns([4, 1])
                    pagina_atual = c_pag.number_input(f"Página (1 a {total_paginas})", min_value=1, max_value=total_paginas, value=1)
                else:
                    pagina_atual = 1
                    
                inicio = (pagina_atual - 1) * itens_por_pagina
                fim = inicio + itens_por_pagina
                res_pagina = res_vitrine.iloc[inicio:fim]

                # --- DESENHO DOS LIVROS (GRID) ---
                cols_vitrine = st.columns(4)
                for idx, (_, livro) in enumerate(res_pagina.iterrows()):
                    with cols_vitrine[idx % 4]:
                        tem_capa_real = pd.notnull(livro['capa_url']) and livro['capa_url'] != "" and str(livro['capa_url']).startswith('http')
                        
                        if tem_capa_real:
                            capa_path = livro['capa_url']
                            markdown_capa = f'<div style="text-align: center; margin-bottom: 10px;"><img src="{capa_path}" style="width: 100%; border-radius: 8px; box-shadow: 2px 2px 10px rgba(0,0,0,0.1); height: 200px; object-fit: cover;"><div class="external-book-title">{livro["titulo"]}</div></div>'
                        else:
                            capa_dinamica_html = gerar_capa_dinamica(livro['id'], livro['titulo'], livro['autor'])
                            markdown_capa = f'<div style="text-align: center; margin-bottom: 10px;">{capa_dinamica_html}</div>'
                        
                        st.markdown(markdown_capa, unsafe_allow_html=True)
                        
                        with st.popover("📖 Ver Detalhes", use_container_width=True):
                            st.markdown(f"### {livro['titulo']}")
                            c_capa, c_info = st.columns([1, 2])
                            
                            with c_capa:
                                if tem_capa_real:
                                    st.image(livro['capa_url'], use_container_width=True)
                                else:
                                    st.markdown(capa_dinamica_html, unsafe_allow_html=True)
                            
                            with c_info:
                                st.markdown(f"**✍️ Autor:** {livro['autor']}")
                                st.markdown(f"**🏢 Editora:** {livro['editora']}")
                                st.markdown(f"**🏷️ Gênero:** {livro['genero']}")
                                st.markdown(f"**📚 Série:** {livro['serie']}")
                            
                            st.markdown("---")
                            st.markdown("#### 📍 Situação dos Exemplares")
                            
                            meus_exemplares = df_exemplares[df_exemplares['id_acervo'] == livro['id']]
                            
                            if not meus_exemplares.empty:
                                for _, ex in meus_exemplares.iterrows():
                                    status_cor = "🟢 Disponível" if ex['disponivel'] else "🔴 Emprestado"
                                    loc = ex['localizacao'] if ex['localizacao'] else "Não definida"
                                    st.info(f"**Tombo:** {ex['tombo']} | **Local:** {loc} | **Status:** {status_cor}")
                            else:
                                st.warning("Nenhum exemplar físico localizado para esta obra.")
                
                # Rodapé informativo
                if total_paginas > 1:
                    st.markdown("---")
                    st.caption(f"Página {pagina_atual} de {total_paginas}. Total: {len(res_vitrine)} livros encontrados.")
            else:
                st.warning("Nenhum livro encontrado com o termo pesquisado.")
                
        else:
            st.info("O acervo está vazio. Vá em 'Incorporação' para adicionar livros.")

    # --- 2. PRONTUÁRIO DO LEITOR & RANKING DE CONQUISTAS ---
    with tab_perfil:
        st.subheader("👤 Consulta de Prontuário por Aluno")
        
        # 1. Seleção do Aluno
        leitor_selecionado = st.selectbox(
            "Selecione o aluno para visualizar o histórico:",
            ["-- Selecione um Aluno --"] + lista_leitores,
            key="sel_prontuario"
        )

        if leitor_selecionado != "-- Selecione um Aluno --":
            col_p1, col_p2 = st.columns([2, 1])
            
            # Filtros de dados para o aluno selecionado
            emprestimos_aluno = df_emp[df_emp['leitor'] == leitor_selecionado]
            ativos_aluno = emprestimos_aluno[emprestimos_aluno['status'] == 'Ativo']
            historico_aluno = emprestimos_aluno[emprestimos_aluno['status'] == 'Devolvido']
            total_lidos = len(historico_aluno)
            
            with col_p1:
                st.markdown(f"### 📖 Histórico de {leitor_selecionado.split('(')[0]}")
                
                # Seção de Livros com o Aluno agora
                st.markdown("#### 📤 Em posse do aluno:")
                if ativos_aluno.empty:
                    st.info("O aluno não possui livros pendentes de devolução.")
                else:
                    for _, r in ativos_aluno.iterrows():
                        # Cruzamento para pegar o título do livro
                        ex_info = df_exemplares[df_exemplares['id'] == r['id_exemplar']]
                        titulo_livro = "Obra não identificada"
                        if not ex_info.empty:
                            ac_info = df_acervo[df_acervo['id'] == ex_info.iloc[0]['id_acervo']]
                            titulo_livro = ac_info.iloc[0]['titulo'] if not ac_info.empty else "Título não localizado"
                        
                        st.warning(f"**{titulo_livro}** (Tombo: {r['id_exemplar']})  \n📅 Devolução prevista: **{r['data_prevista']}**")

                # Seção de Histórico Completo
                st.markdown("---")
                st.markdown("#### 📚 Obras já lidas:")
                if historico_aluno.empty:
                    st.write("Este aluno ainda não completou nenhuma leitura no sistema.")
                else:
                    # Criar uma lista simplificada para o histórico
                    lista_h = []
                    for _, h in historico_aluno.iterrows():
                        ex_h = df_exemplares[df_exemplares['id'] == h['id_exemplar']]
                        tit_h = "Obra Desconhecida"
                        if not ex_h.empty:
                            ac_h = df_acervo[df_acervo['id'] == ex_h.iloc[0]['id_acervo']]
                            tit_h = ac_h.iloc[0]['titulo'] if not ac_h.empty else "Título não localizado"
                        lista_h.append({"Obra": tit_h, "Data da Retirada": h['data_saida']})
                    
                    st.table(pd.DataFrame(lista_h))

            with col_p2:
                st.markdown("### 🏆 Conquistas")
                nivel = "Iniciante" if total_lidos < 5 else "Prata" if total_lidos < 15 else "Mestre"
                icone = "🌱" if total_lidos < 5 else "🥈" if total_lidos < 15 else "💎"
                
                st.markdown(f"""
                    <div class="badge-box">
                        <div style="font-size:50px">{icone}</div>
                        <div class="badge-title">Leitor {nivel}</div>
                        <div style="font-size:20px; font-weight:bold;">{total_lidos}</div>
                        <div>livros lidos</div>
                    </div>
                """, unsafe_allow_html=True)
                
                if total_lidos >= 1:
                    st.success("Parabéns pelo engajamento!")

        st.markdown("---")
        
        # 2. RANKING GERAL DE LEITURA (Gamificação)
        st.subheader("🥇 Ranking de Leitores da Escola")
        if not df_emp.empty:
            # Contagem de livros devolvidos por leitor
            ranking = df_emp[df_emp['status'] == 'Devolvido']['leitor'].value_counts().reset_index()
            ranking.columns = ['Aluno', 'Livros Lidos']
            
            # Pega os Top 10
            top_10 = ranking.head(10)
            
            col_r1, col_r2 = st.columns([2, 1])
            with col_r1:
                st.write("Confira os alunos que mais leram livros este ano:")
                st.dataframe(top_10, use_container_width=True, hide_index=True)
            
            with col_r2:
                if not top_10.empty:
                    st.info(f"🌟 O maior leitor é:  \n**{top_10.iloc[0]['Aluno']}**")
        else:
            st.info("O ranking será gerado assim que as primeiras devoluções forem registradas.")

    # --- 3. ABA DE CIRCULAÇÃO (MODELO COMPLETO LIMEIRA - 10CM) ---
    if eh_gestao:
        with tab_circ:
            st.subheader("🚀 Balcão de Atendimento - Sala de Leitura")
            c_out, c_in = st.columns(2)
            
            # Lista de frases de incentivo
            frases_leitura = [
                "Ler é ganhar asas para o mundo.",
                "Um livro é um brinquedo feito com letras. Ler é brincar!",
                "Quem lê, viaja sem sair do lugar.",
                "A leitura é uma porta aberta para o conhecimento.",
                "Livros são amigos que nos ensinam grandes lições.",
                "No mundo dos livros, você pode ser o que quiser!",
                "Um livro por vez, uma aventura por dia.",
                "Ler alimenta a imaginação e fortalece o saber.",
                "A cada página lida, uma nova descoberta.",
                "Cultive o hábito da leitura e colha sabedoria."
            ]
            import random

            # ==========================================
            # COLUNA ESQUERDA: EMPRÉSTIMO (MODELO INSTITUCIONAL COMPLETO)
            # ==========================================
            with c_out:
                st.markdown("### 📤 Saída de Livro")
                with st.form("form_emprestimo_institucional"):
                    leitor_sel = st.selectbox("Selecionar Aluno:", ["-- Selecione --"] + lista_leitores)
                    tombo_out = st.text_input("Bipar ou Digitar Tombo:")
                    atendente = st.text_input("Atendente:", value="Responsável Sala de Leitura")
                    
                    if st.form_submit_button("Efetivar Empréstimo", type="primary"):
                        if leitor_sel != "-- Selecione --" and tombo_out:
                            t_limpo = str(tombo_out).strip().replace('.0', '')
                            res_ex = supabase.table("Biblioteca_Exemplares").select("*").eq("tombo", t_limpo).execute()
                            
                            if res_ex.data:
                                exemplar = res_ex.data[0]
                                if exemplar['disponivel']:
                                    res_ob = supabase.table("Biblioteca_Acervo").select("*").eq("id", exemplar['id_acervo']).execute()
                                    if res_ob.data:
                                        obra = res_ob.data[0]
                                        info_aluno = df_carometro[df_carometro['display_leitor'] == leitor_sel].iloc[0] if not df_carometro[df_carometro['display_leitor'] == leitor_sel].empty else {}
                                        
                                        venc = (datetime.now() + timedelta(days=7)).strftime("%d/%m/%Y")
                                        hj = datetime.now().strftime("%d/%m/%Y %H:%M")

                                        # Updates no Banco
                                        supabase.table("Biblioteca_Emprestimos").insert({
                                            "id": str(uuid.uuid4()), "id_exemplar": exemplar['id'], 
                                            "leitor": leitor_sel, "data_saida": hj, "data_prevista": venc, "status": "Ativo"
                                        }).execute()
                                        supabase.table("Biblioteca_Exemplares").update({"disponivel": False}).eq("id", exemplar['id']).execute()
                                        
                                        st.session_state.comprovante = {
                                            "aluno": leitor_sel,
                                            "turma": info_aluno.get('turma', info_aluno.get('Série', 'N/A')),
                                            "livro": obra.get('titulo', 'N/A'),
                                            "autor": obra.get('autor', 'N/I'),
                                            "editora": obra.get('editora', 'N/I'),
                                            "genero": obra.get('genero', 'Geral'),
                                            "serie_livro": obra.get('SÉRIE', 'Fundamental'),
                                            "tombo": t_limpo,
                                            "atendente": atendente,
                                            "vencimento": venc,
                                            "data": hj,
                                            "frase": random.choice(frases_leitura)
                                        }
                                        st.success("Empréstimo Registrado!"); time.sleep(0.5); st.rerun()
                                    else: st.error("Obra não encontrada.")
                                else: st.warning("Livro já emprestado.")
                            else: st.error("Tombo não localizado.")

                if 'comprovante' in st.session_state:
                    cp = st.session_state.comprovante
                    html_final = f"""
                    <div id="cupom_doc" style="width:88mm; padding:3mm; box-sizing:border-box; font-family:'Courier New', Courier, monospace; font-size:12px; border:1px solid #000; background:#fff; color:#000; margin:0 auto;">
                        <div style="text-align: center; border-bottom: 2px solid #000; padding-bottom: 8px;">
                            <b style="font-size:14px;">PREFEITURA MUNICIPAL DE LIMEIRA</b><br>
                            <b style="font-size:15px;">CEIEF RAFAEL AFFONSO LEITE</b><br>
                            <span style="font-size:12px; font-weight:bold;">SALA DE LEITURA</span><br>
                            <span style="font-size:10px;">Rua Antonio Alves de Oliveira, 250 - Jd. Pres. Dutra - Limeira - São Paulo</span><br>
                            <span style="font-size:10px;">Contato: (19) 3495-5390</span>
                        </div>

                        <div style="margin-top:12px;">
                            <b>ALUNO:</b> {cp['aluno']}<br>
                            <b>TURMA:</b> {cp['turma']} | <b>STATUS:</b> Ativo<br>
                            <b>ATENDENTE:</b> {cp['atendente']}
                        </div>
                        <hr style="border:0; border-top:1px dashed #000;">

                        <div style="font-size:11px; margin-top:10px;">
                            <b>OBRA:</b> {cp['livro']}<br>
                            <b>AUTOR:</b> {cp['autor']}<br>
                            <b>EDITORA:</b> {cp['editora']}<br>
                            <b>GÊNERO:</b> {cp['genero']} | <b>COLEÇÃO:</b> {cp['serie_livro']}<br>
                            <b>TOMBO:</b> {cp['tombo']} | <b>ESTADO:</b> [ ]Ótimo [ ]Bom
                        </div>

                        <div style="margin-top:12px; background:#000; color:#fff; text-align:center; padding:8px;">
                            <span style="font-size:11px;">DATA DE DEVOLUÇÃO:</span><br>
                            <b style="font-size:24px;">{cp['vencimento']}</b>
                        </div>

                        <div style="margin-top:12px; font-size:10px; text-align:justify; border:1px solid #ccc; padding:6px;">
                            <center><b style="font-size:11px;">TERMO DE COMPROMISSO E ZELO</b></center>
                            O aluno/responsável declara-se ciente da guarda e conservação desta obra. 
                            Proibido: riscar, dobrar páginas ou manusear próximo a líquidos. 
                            Em caso de perda ou dano, deverá repor por exemplar idêntico.
                        </div>

                        <div style="margin-top:25px; text-align:center;">
                            ____________________________<br>Assinatura do Aluno<br><br>
                            ____________________________<br>Visto Sala de Leitura
                        </div>

                        <div style="margin-top:20px; text-align:center; font-size:11px; border-top:1px solid #000; padding-top:8px;">
                            <i>"{cp['frase']}"</i><br>
                            <span style="font-size:9px;">Sistema Integra - {cp['data']}</span>
                        </div>
                    </div>
                    <script>setTimeout(function(){{ window.print(); }}, 1000);</script>
                    """
                    st.components.v1.html(html_final, height=900) 
                    if st.button("Finalizar Atendimento"):
                        del st.session_state.comprovante; st.rerun()

            # ==========================================
            # COLUNA DIREITA: DEVOLUÇÃO (RECEBIMENTO)
            # ==========================================
            with c_in:
                st.markdown("### 📥 Devolução de Livro")
                
                tombo_in = st.text_input("Bipar ou Digitar Tombo para Devolução:", key="input_dev")
                
                if st.button("Confirmar Devolução", type="primary", key="btn_dev"):
                    if tombo_in:
                        try:
                            t_limpo = str(tombo_in).strip().replace('.0', '')
                            
                            # 1. Busca o exemplar
                            res_ex = supabase.table("Biblioteca_Exemplares").select("id").eq("tombo", t_limpo).execute()
                            
                            if res_ex.data:
                                id_ex = res_ex.data[0]['id']
                                
                                # 2. Busca se existe um empréstimo ATIVO para esse exemplar
                                res_emp = supabase.table("Biblioteca_Emprestimos").select("*").eq("id_exemplar", id_ex).eq("status", "Ativo").execute()
                                
                                if res_emp.data:
                                    id_emp = res_emp.data[0]['id']
                                    
                                    # 3. Atualiza o empréstimo para 'Devolvido' e o exemplar para 'Disponível'
                                    supabase.table("Biblioteca_Emprestimos").update({"status": "Devolvido"}).eq("id", id_emp).execute()
                                    supabase.table("Biblioteca_Exemplares").update({"disponivel": True}).eq("id", id_ex).execute()
                                    
                                    st.success(f"✅ O livro do Tombo {t_limpo} foi devolvido com sucesso!")
                                    time.sleep(2) 
                                    st.rerun()
                                else:
                                    st.warning(f"O Tombo {t_limpo} já consta como disponível ou não possui empréstimo ativo.")
                            else:
                                st.error("Tombo não encontrado no sistema.")
                        
                        except Exception as e:
                            st.error(f"❌ Erro de comunicação com o banco: {e}")
                    else:
                        st.error("Por favor, digite ou bipe um número de tombo.")

                st.markdown("---")
                st.info("💡 Dica: Ao confirmar a devolução, o livro volta a aparecer imediatamente como 'Disponível' na consulta do acervo.")
                    
        # =========================================================
        # 4. INCORPORAÇÃO TÉCNICA (CENTRAL DE ACERVO)
        # =========================================================
        with tab_cat:
            st.subheader("Adicionar Novos Livros ao Sistema")
            
            # Sub-navegação interna para organizar os métodos
            metodo = st.radio("Escolha o método de cadastro:", 
                             ["Pelo Título (Busca Capa)", "Pelo ISBN (Automático)", "Cadastro Manual", "Importar Planilha"],
                             horizontal=True)

            # --- MÉTODO 1: BUSCA POR TÍTULO (ÓTIMO PARA TER A FOTO) ---
            if metodo == "Pelo Título (Busca Capa)":
                with st.form("form_titulo"):
                    st.markdown("#### 🔍 Localizar Obra e Capa")
                    c_t, c_a = st.columns(2)
                    t_busca = c_t.text_input("Título do Livro:")
                    a_busca = c_a.text_input("Autor (opcional):")
                    if st.form_submit_button("Pesquisar na Web"):
                        query = f"intitle:{t_busca}" + (f"+inauthor:{a_busca}" if a_busca else "")
                        try:
                            r = requests.get(f"https://www.googleapis.com/books/v1/volumes?q={query}&maxResults=1").json()
                            if "items" in r:
                                info = r["items"][0]["volumeInfo"]
                                st.session_state.temp_add = {
                                    "titulo": info.get("title", t_busca).upper(),
                                    "autor": ", ".join(info.get("authors", [a_busca])).upper(),
                                    "editora": info.get("publisher", "").upper(),
                                    "capa": info.get("imageLinks", {}).get("thumbnail", "").replace("http:", "https:"),
                                    "genero": info.get("categories", [""])[0],
                                    "resumo": info.get("description", "")
                                }
                                st.success("Sugestão de capa e dados localizada!")
                            else: st.error("Não encontramos capas automáticas para este título.")
                        except: st.error("Erro na busca online.")

                if 'temp_add' in st.session_state and st.session_state.temp_add is not None:
                    with st.form("confirm_add_titulo"):
                        v = st.session_state.temp_add
                        c_img, c_form = st.columns([1, 3])
                        if v['capa']: c_img.image(v['capa'], width=150)
                        
                        tit_f = c_form.text_input("Título:", value=v['titulo'])
                        aut_f = c_form.text_input("Autor:", value=v['autor'])
                        edit_f = c_form.text_input("Editora:", value=v['editora'])
                        
                        st.markdown("---")
                        st.markdown("#### 📍 Dados da sua Planilha")
                        col1, col2, col3 = st.columns(3)
                        tombo_f = col1.text_input("TOMBO (Número da Etiqueta):")
                        local_f = col2.text_input("LOCAL (Ex: C-21):")
                        serie_f = col3.text_input("SÉRIE (Ex: Fundamental):")
                        gen_f = st.text_input("GÊNERO:", value=v['genero'])

                        if st.form_submit_button("💾 Salvar no Sistema", type="primary"):
                            if tombo_f and tit_f:
                                id_a = str(uuid.uuid4())
                                supabase.table("Biblioteca_Acervo").insert({
                                    "id": id_a, "titulo": tit_f, "autor": aut_f, "editora": edit_f,
                                    "capa_url": v['capa'], "genero": gen_f, "serie": serie_f
                                }).execute()
                                supabase.table("Biblioteca_Exemplares").insert({
                                    "id": str(uuid.uuid4()), "id_acervo": id_a, "tombo": tombo_f, 
                                    "localizacao": local_f, "disponivel": True
                                }).execute()
                                st.success("Livro cadastrado com sucesso!"); st.session_state.temp_add = None; time.sleep(1); st.rerun()
                            else: st.error("Título e Tombo são obrigatórios.")

            # --- MÉTODO 2: ISBN (O MAIS RÁPIDO PARA LIVROS NOVOS) ---
            elif metodo == "Pelo ISBN (Automático)":
                st.info("Bipe o código de barras do livro aqui.")

            # --- MÉTODO 3: CADASTRO MANUAL (TOTAL LIBERDADE) ---
            elif metodo == "Cadastro Manual":
                with st.form("form_manual"):
                    st.markdown("#### 📝 Preencher Dados Manuais")
                    col_a, col_b = st.columns(2)
                    m_tit = col_a.text_input("TÍTULO:")
                    m_aut = col_b.text_input("AUTOR:")
                    m_edit = col_a.text_input("EDITORA:")
                    m_gen = col_b.text_input("GÊNERO:")
                    
                    st.markdown("---")
                    col_c, col_d, col_e = st.columns(3)
                    m_tombo = col_c.text_input("TOMBO:")
                    m_local = col_d.text_input("LOCAL:")
                    m_serie = col_e.text_input("SÉRIE:")
                    
                    m_capa = st.text_input("URL da Foto da Capa (opcional):", placeholder="https://...")

                    if st.form_submit_button("💾 Cadastrar Manualmente"):
                        if m_tit and m_tombo:
                            id_a = str(uuid.uuid4())
                            supabase.table("Biblioteca_Acervo").insert({
                                "id": id_a, "titulo": m_tit.upper(), "autor": m_aut.upper(), 
                                "editora": m_edit.upper(), "genero": m_gen.upper(), "serie": m_serie.upper(), "capa_url": m_capa
                            }).execute()
                            supabase.table("Biblioteca_Exemplares").insert({
                                "id": str(uuid.uuid4()), "id_acervo": id_a, "tombo": m_tombo, 
                                "localizacao": m_local.upper(), "disponivel": True
                            }).execute()
                            st.success("Cadastro manual realizado!"); time.sleep(1); st.rerun()
                        else: st.error("Preencha ao menos Título e Tombo.")

            # --- MÉTODO 4: IMPORTAÇÃO PLANILHA ---
            elif metodo == "Importar Planilha":
                st.info("Envie sua planilha (.csv ou .xlsx). O sistema dividirá os dados automaticamente entre Acervo e Exemplares.")
                
                arquivo_upload = st.file_uploader("Selecione o arquivo:", type=["csv", "xlsx"])

                if arquivo_upload:
                    if st.button("🚀 Iniciar Importação em Lote", type="primary"):
                        with st.spinner("Lendo arquivo e preparando lotes... Isso pode levar alguns segundos."):
                            try:
                                if arquivo_upload.name.endswith('.csv'):
                                    df = pd.read_csv(arquivo_upload, sep=';', encoding='latin1')
                                else:
                                    df = pd.read_excel(arquivo_upload)
                                
                                df.columns = df.columns.str.strip().str.lower()
                                df = df.where(pd.notnull(df), None)

                                registros_acervo = []
                                registros_exemplares = []

                                for index, row in df.iterrows():
                                    if row.get("tombo") is None or row.get("titulo") is None:
                                        continue
                                    
                                    id_acervo_gerado = str(uuid.uuid4())
                                    id_exemplar_gerado = str(uuid.uuid4())

                                    registros_acervo.append({
                                        "id": id_acervo_gerado,
                                        "titulo": str(row.get("titulo", "")).strip().upper(),
                                        "autor": str(row.get("autor", "")).strip().upper() if row.get("autor") else "",
                                        "editora": str(row.get("editora", "")).strip().upper() if row.get("editora") else "",
                                        "genero": str(row.get("genero", "")).strip().upper() if row.get("genero") else "",
                                        "serie": str(row.get("serie", "")).strip().upper() if row.get("serie") else "",
                                        "capa_url": "" 
                                    })

                                    registros_exemplares.append({
                                        "id": id_exemplar_gerado,
                                        "id_acervo": id_acervo_gerado,
                                        "tombo": str(int(row["tombo"])) if isinstance(row["tombo"], float) else str(row["tombo"]),
                                        "localizacao": str(row.get("local", "")).strip().upper() if row.get("local") else "",
                                        "disponivel": True 
                                    })

                                tamanho_lote = 500
                                total_lotes = math.ceil(len(registros_acervo) / tamanho_lote)
                                
                                barra_progresso = st.progress(0)
                                st.write(f"Preparando para inserir {len(registros_acervo)} livros...")

                                for i in range(total_lotes):
                                    inicio = i * tamanho_lote
                                    fim = inicio + tamanho_lote
                                    
                                    if registros_acervo[inicio:fim]:
                                        supabase.table("Biblioteca_Acervo").insert(registros_acervo[inicio:fim]).execute()
                                    
                                    if registros_exemplares[inicio:fim]:
                                        supabase.table("Biblioteca_Exemplares").insert(registros_exemplares[inicio:fim]).execute()
                                    
                                    barra_progresso.progress((i + 1) / total_lotes)

                                st.success(f"✅ Sucesso! {len(registros_acervo)} livros foram importados preservando os tombos e a separação correta.")
                                st.balloons()

                            except Exception as e:
                                st.error(f"❌ Erro ao processar a importação: {e}")

        # =========================================================
        # 5. ABA DE AUDITORIA E RELATÓRIOS (DASHBOARD)
        # =========================================================
        with tab_dash:
            st.subheader("📊 Painel de Auditoria e Relatórios")
            
            # 1. Métricas Globais (Cards Superiores)
            c1, c2, c3, c4 = st.columns(4)
            
            total_acervo = len(df_acervo) if not df_acervo.empty else 0
            total_exemplares = len(df_exemplares) if not df_exemplares.empty else 0
            emp_ativos = len(df_emp[df_emp['status'] == 'Ativo']) if not df_emp.empty else 0
            disponiveis = len(df_exemplares[df_exemplares['disponivel'] == True]) if not df_exemplares.empty else 0
            
            c1.metric("Obras Cadastradas", total_acervo)
            c2.metric("Total Físico (Exemplares)", total_exemplares)
            c3.metric("Empréstimos Ativos", emp_ativos, delta="-ocupados", delta_color="inverse")
            c4.metric("Disponíveis na Estante", disponiveis, delta="+livres")
            
            st.markdown("---")
            
            # 2. Relatório de Empréstimos Ativos (Com cruzamento de dados)
            st.markdown("#### 🚨 Acompanhamento de Empréstimos e Atrasos")
            
            if not df_emp.empty:
                ativos = df_emp[df_emp['status'] == 'Ativo'].copy()
                
                if not ativos.empty:
                    dados_relatorio = []
                    hoje = datetime.now()
                    
                    for _, emp in ativos.iterrows():
                        ex_info = df_exemplares[df_exemplares['id'] == emp['id_exemplar']]
                        if not ex_info.empty:
                            tombo_val = ex_info.iloc[0]['tombo']
                            id_ac = ex_info.iloc[0]['id_acervo']
                            
                            ac_info = df_acervo[df_acervo['id'] == id_ac]
                            titulo_val = ac_info.iloc[0]['titulo'] if not ac_info.empty else "Título não localizado"
                            
                            data_prev_str = emp['data_prevista']
                            dias_atraso = 0
                            is_atrasado = False
                            status_texto = "✅ No Prazo"
                            
                            try:
                                data_prev = datetime.strptime(data_prev_str, "%d/%m/%Y")
                                diferenca = (hoje.date() - data_prev.date()).days
                                
                                if diferenca > 0:
                                    is_atrasado = True
                                    dias_atraso = diferenca
                                    status_texto = f"🚨 Atrasado ({dias_atraso} dias)"
                            except Exception:
                                status_texto = "❓ Erro na Data"

                            dados_relatorio.append({
                                "Leitor / Aluno": emp['leitor'],
                                "Título da Obra": titulo_val,
                                "Tombo": tombo_val,
                                "Data de Retirada": emp['data_saida'],
                                "Devolução Prevista": data_prev_str,
                                "Status": status_texto,
                                "_is_atrasado": is_atrasado,
                                "_dias_atraso": dias_atraso
                            })
                            
                    if dados_relatorio:
                        df_relatorio = pd.DataFrame(dados_relatorio)
                        df_relatorio = df_relatorio.sort_values(by=['_is_atrasado', '_dias_atraso'], ascending=[False, False])
                        
                        df_exibicao = df_relatorio.drop(columns=['_is_atrasado', '_dias_atraso'])
                        
                        pesquisa_aluno = st.text_input("🔍 Buscar Aluno:", placeholder="Digite o nome ou turma...")
                        if pesquisa_aluno:
                            df_exibicao = df_exibicao[df_exibicao['Leitor / Aluno'].str.contains(pesquisa_aluno, case=False, na=False)]
                            
                        st.dataframe(df_exibicao, use_container_width=True, hide_index=True)
                        
                        st.markdown("---")
                        
                        # 3. IMPRESSÃO DE AVISOS DE COBRANÇA
                        st.markdown("#### 🖨️ Emissão de Avisos de Cobrança (Atrasados)")
                        
                        df_atrasados = df_relatorio[df_relatorio['_is_atrasado'] == True]
                        
                        if not df_atrasados.empty:
                            st.warning(f"Existem **{len(df_atrasados)}** empréstimos atrasados.")
                            
                            opcoes_cobranca = [
                                f"{row['Leitor / Aluno']} | Livro: {row['Título da Obra']} | Tombo: {row['Tombo']}"
                                for _, row in df_atrasados.iterrows()
                            ]
                            
                            idx_selecionado = st.selectbox("Selecione o empréstimo para gerar o bilhete:", range(len(opcoes_cobranca)), format_func=lambda x: opcoes_cobranca[x])
                            
                            if st.button("🖨️ Imprimir Aviso de Cobrança", type="primary"):
                                aluno_cobrar = df_atrasados.iloc[idx_selecionado]
                                
                                html_cobranca = f"""
                                <div id="cobranca_doc" style="width:88mm; padding:3mm; box-sizing:border-box; font-family:'Courier New', Courier, monospace; font-size:12px; border:1px solid #000; background:#fff; color:#000; margin:0 auto;">
                                    <div style="text-align: center; border-bottom: 2px solid #000; padding-bottom: 8px;">
                                        <b style="font-size:14px;">PREFEITURA MUNICIPAL DE LIMEIRA</b><br>
                                        <b style="font-size:15px;">CEIEF RAFAEL AFFONSO LEITE</b><br>
                                        <span style="font-size:12px; font-weight:bold;">SALA DE LEITURA</span><br>
                                    </div>

                                    <div style="margin-top:15px; text-align:center; border:2px dashed #000; padding:5px;">
                                        <b style="font-size:16px;">⚠️ AVISO DE ATRASO ⚠️</b>
                                    </div>

                                    <div style="margin-top:15px;">
                                        <b>ALUNO:</b> {aluno_cobrar['Leitor / Aluno']}<br>
                                        <b>TOMBO:</b> {aluno_cobrar['Tombo']}<br>
                                        <b>OBRA:</b> {aluno_cobrar['Título da Obra']}
                                    </div>

                                    <div style="margin-top:15px; background:#000; color:#fff; text-align:center; padding:10px;">
                                        <span style="font-size:12px;">VENCIDO DESDE:</span><br>
                                        <b style="font-size:22px;">{aluno_cobrar['Devolução Prevista']}</b><br>
                                        <span style="font-size:14px; color:#ffcccc;">({aluno_cobrar['_dias_atraso']} dias de atraso)</span>
                                    </div>

                                    <div style="margin-top:15px; font-size:11px; text-align:justify; border:1px solid #ccc; padding:8px;">
                                        <b>Sr(a) Responsável,</b><br><br>
                                        Consta em nossos registros que o livro acima não foi devolvido no prazo estipulado. Solicitamos a devolução <b>imediata</b> para que outros alunos também possam realizar a leitura.<br><br>
                                        Em caso de perda ou dano, favor procurar a Direção da escola.
                                    </div>

                                    <div style="margin-top:35px; text-align:center;">
                                        ________________________________<br>Ciência do Responsável<br><br>
                                        ________________________________<br>Data
                                    </div>
                                </div>
                                <script>setTimeout(function(){{ window.print(); }}, 1000);</script>
                                """
                                st.components.v1.html(html_cobranca, height=750)
                        else:
                            st.success("Excelente! Não há nenhum livro atrasado no momento.")

                    else:
                        st.warning("Não foi possível cruzar os dados dos exemplares.")
                else:
                    st.success("🎉 Todos os livros da escola estão devolvidos no momento!")
            else:
                st.info("Nenhum histórico de circulação encontrado.")
                
            st.markdown("---")
            
            # 4. Ferramentas de Exportação
            st.markdown("#### 🗃️ Exportação de Dados")
            col_down1, col_down2 = st.columns(2)
            
            with col_down1:
                if not df_acervo.empty:
                    csv_acervo = df_acervo.to_csv(index=False, sep=';').encode('latin1', errors='replace')
                    st.download_button(
                        label="📥 Descarregar Acervo Completo",
                        data=csv_acervo,
                        file_name=f"acervo_escola_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv",
                        use_container_width=True
                    )
            
            with col_down2:
                if not df_emp.empty:
                    csv_emp = df_emp.to_csv(index=False, sep=';').encode('latin1', errors='replace')
                    st.download_button(
                        label="📥 Descarregar Histórico de Circulação",
                        data=csv_emp,
                        file_name=f"circulacao_escola_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv",
                        use_container_width=True
                    )





























elif app_mode_adm == "🖨️ Emissão de Boletins":
    import hashlib
    import uuid
    import io
    import zipfile
    from datetime import datetime
    import pandas as pd
    import re
    from fpdf import FPDF
    
    st.markdown('<div class="header-box"><div class="header-title">🖨️ Emissão de Boletins Oficiais</div></div>', unsafe_allow_html=True)
    
    if st.button("⬅️ Voltar ao Menu Inicial", key="voltar_boletins"):
        st.session_state.modulo_atuacao = None
        st.rerun()

    # --- FUNÇÃO DE FORMATAÇÃO DE NOMES (ABREVIAÇÃO INTELIGENTE) ---
    def formatar_nome_prof(nome_sujo):
        if not nome_sujo or "Docente" in nome_sujo:
            return "Docente"
        
        # Limpeza de códigos de matrícula e caracteres especiais
        nome_limpo = re.sub(r'\d{3}\.?\d{3}-\d{2}\s*-\s*', '', nome_sujo).strip()
        nome_limpo = nome_limpo.title()
        
        # Correção de OCR (Nomes grudados na extração)
        nome_limpo = nome_limpo.replace("Julianaaparecida", "Juliana Aparecida")
        nome_limpo = nome_limpo.replace("Limaalvez", "Lima Alvez")
        
        partes = nome_limpo.split()
        if len(partes) <= 2:
            return nome_limpo
            
        primeiro = partes[0]
        ultimo = partes[-1]
        # Mantém a primeira letra dos nomes do meio
        meio = [p[0] + "." for p in partes[1:-1] if len(p) > 2]
        return f"{primeiro} {' '.join(meio)} {ultimo}"

    # --- CLASSE DO PDF ---
    class BoletimPDF(FPDF):
        def header(self):
            try:
                self.image("logo_prefeitura.png", 10, 8, 25)
                self.image("logo_escola.png", 175, 8, 25)
            except Exception:
                pass 

            self.set_font("helvetica", "B", 12)
            self.cell(0, 6, "PREFEITURA MUNICIPAL DE LIMEIRA/SP", border=0, new_x="LMARGIN", new_y="NEXT", align="C")
            self.cell(0, 6, "SECRETARIA MUNICIPAL DE EDUCAÇÃO", border=0, new_x="LMARGIN", new_y="NEXT", align="C")
            self.set_font("helvetica", "B", 14)
            self.cell(0, 8, "CEIEF RAFAEL AFFONSO LEITE", border=0, new_x="LMARGIN", new_y="NEXT", align="C")
            self.set_font("helvetica", "", 10)
            self.cell(0, 6, "ACOMPANHAMENTO ESCOLAR - ANO LETIVO 2026", border=0, new_x="LMARGIN", new_y="NEXT", align="C")
            self.ln(10)

    # --- GERAÇÃO DO BOLETIM INDIVIDUAL ---
    def gerar_boletim_pdf(aluno, diretor, coord, p_poli, p_arte, p_edf, p_tec, dias_tri1):
        pdf = BoletimPDF()
        pdf.add_page()
        
        # Identificação
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("helvetica", "B", 11)
        pdf.cell(140, 8, f" ALUNO(A): {aluno['Nome']}", border=1, fill=True)
        pdf.cell(50, 8, f" SITUAÇÃO: {aluno['Situação']}", border=1, new_x="LMARGIN", new_y="NEXT", fill=True)
        
        pdf.set_font("helvetica", "", 10)
        pdf.cell(95, 8, f" SÉRIE/TURMA: {aluno['Turma']}", border=1)
        pdf.cell(95, 8, f" PERÍODO: {aluno['Periodo']}", border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)
        
        # Tabela de Frequência Trimestral
        pdf.set_fill_color(230, 230, 230)
        pdf.set_font("helvetica", "B", 9)
        pdf.cell(190, 7, " CONTROLE DE FREQUÊNCIA", border=1, new_x="LMARGIN", new_y="NEXT", align="C", fill=True)
        pdf.cell(55, 7, "Discriminação", border=1, align="C")
        pdf.cell(33, 7, "1º Trimestre", border=1, align="C")
        pdf.cell(33, 7, "2º Trimestre", border=1, align="C")
        pdf.cell(33, 7, "3º Trimestre", border=1, align="C")
        pdf.cell(36, 7, "Total Anual", border=1, new_x="LMARGIN", new_y="NEXT", align="C")
        
        pdf.set_font("helvetica", "", 9)
        labels = [("Dias Letivos", dias_tri1), ("Ausências", aluno['Faltas_Total']), ("Frequência (%)", f"{aluno['Freq_Perc']}%")]
        for desc, val in labels:
            pdf.cell(55, 7, f" {desc}", border=1)
            pdf.cell(33, 7, str(val), border=1, align="C")
            pdf.cell(33, 7, "-", border=1, align="C")
            pdf.cell(33, 7, "-", border=1, align="C")
            pdf.cell(36, 7, "-", border=1, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)

        # Tabela de Aproveitamento
        pdf.set_fill_color(200, 220, 255) 
        pdf.set_font("helvetica", "B", 9)
        pdf.cell(60, 8, "Componente Curricular", border=1, align="C", fill=True)
        pdf.cell(50, 8, "Professor(a)", border=1, align="C", fill=True)
        pdf.cell(20, 8, "1º Tri", border=1, align="C", fill=True)
        pdf.cell(20, 8, "2º Tri", border=1, align="C", fill=True)
        pdf.cell(20, 8, "3º Tri", border=1, align="C", fill=True)
        pdf.cell(20, 8, "Final", border=1, new_x="LMARGIN", new_y="NEXT", align="C", fill=True)
        
        pdf.set_font("helvetica", "", 8)
        disciplinas = [
            ("Língua Portuguesa", p_poli, aluno['LP']), ("Matemática", p_poli, aluno['MAT']),
            ("Ciências", p_poli, aluno['CIE']), ("História", p_poli, aluno['HIST']),
            ("Geografia", p_poli, aluno['GEOG']), ("Arte", p_arte, aluno['ART']),
            ("Educação Física", p_edf, aluno['EF']), ("Linguagens e Tecnologias", p_tec, aluno['TEC'])
        ]
        
        for disc, prof, nota in disciplinas:
            pdf.cell(60, 7, f" {disc}", border=1)
            pdf.cell(50, 7, f" {prof}", border=1)
            pdf.cell(20, 7, nota, border=1, align="C")
            pdf.cell(20, 7, "-", border=1, align="C")
            pdf.cell(20, 7, "-", border=1, align="C")
            pdf.cell(20, 7, "-", border=1, new_x="LMARGIN", new_y="NEXT", align="C")
        
        pdf.ln(8)
        
        pdf.set_font("helvetica", "I", 7)
        pdf.multi_cell(0, 3.5, "Legenda: AD (Adequado), AV (Avançado), B (Básico), AB (Abaixo do Básico), NA (Não Avaliado).")
        pdf.ln(5)

        # Autenticação e Assinatura Eletrônica
        data_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        h = hashlib.sha256(f"{aluno['Nome']}{diretor}{data_at}".encode()).hexdigest()
        cod = str(uuid.uuid4()).split('-')[0].upper() + "-" + str(uuid.uuid4()).split('-')[1].upper()

        pdf.set_fill_color(248, 248, 248)
        pdf.set_font("helvetica", "B", 8)
        pdf.cell(0, 6, " AUTENTICAÇÃO E ASSINATURA ELETRÔNICA", border="LTR", new_x="LMARGIN", new_y="NEXT", fill=True)
        pdf.set_font("helvetica", "", 8)
        
        texto_assinatura = (
            "As notas contidas neste boletim foram inseridas no sistema pelos respectivos professores titulares. "
            "O documento foi assinado digitalmente e emitido sob a responsabilidade da equipe gestora:\n\n"
            f"DIRETOR(A): {diretor}\n"
            f"PROFESSOR(A) COORDENADOR(A): {coord}\n\n"
            f"Código de Autenticação: {cod}  |  Hash SHA-256: {h}\n"
            f"Data/Hora de Emissão: {data_at}\n"
            "A validade deste documento pode ser consultada no site https://integra.streamlit.app/."
        )
        pdf.multi_cell(0, 4.5, texto_assinatura, border="LBR", fill=True)
        
        # Retorna o PDF gerado e os códigos para salvar no Supabase
        return pdf.output(dest="S"), cod, h

    # --- PROCESSAMENTO DO ARQUIVO PDF ---
    st.divider()
    arquivo_pdf = st.file_uploader("📄 Selecione a Ata Escolar (PDF)", type=["pdf"])

    if arquivo_pdf is not None:
        with st.spinner("Extraindo e sincronizando dados da Ata Escolar..."):
            import pdfplumber
            
            alunos_extraidos = []
            texto_total = ""
            dias_letivos = "57"
            
            try:
                # Extrai todo o texto respeitando o layout
                with pdfplumber.open(arquivo_pdf) as pdf:
                    for p in pdf.pages:
                        txt = p.extract_text(layout=True)
                        if txt: texto_total += txt + "\n"
                
                linhas = [l.strip() for l in texto_total.split('\n') if l.strip()]
                
                # Valores padrão
                p_poli, p_arte, p_edf, p_tec = "Juliana A. Da Silva", "Jordana L. Alvez", "Michel L. D. Lima", "Fernando I. Bongiovanni"
                diretor, coord = "José Victor Souza Gallo", "Oelen Fernando Pedro"

                # Varrer o texto para identificar equipe e alunos
                for i, l in enumerate(linhas):
                    if "DIAS LETIVOS:" in l:
                        m = re.search(r"DIAS LETIVOS:\s*(\d+)", l)
                        if m: dias_letivos = m.group(1)
                    
                    # Captura Equipe
                    if "Professor(a)" in l and "Arte" in l:
                        partes_assinatura = re.split(r'\s{2,}', linhas[i-1])
                        if len(partes_assinatura) >= 3:
                            p_poli = formatar_nome_prof(partes_assinatura[0])
                            p_arte = formatar_nome_prof(partes_assinatura[1])
                            p_edf = formatar_nome_prof(partes_assinatura[2])
                    
                    if "Tecnologias" in l and "Coordenador" in l:
                        partes_assinatura2 = re.split(r'\s{2,}', linhas[i-1])
                        if len(partes_assinatura2) >= 3:
                            p_tec = formatar_nome_prof(partes_assinatura2[0])
                            coord = formatar_nome_prof(partes_assinatura2[1])
                            diretor = formatar_nome_prof(partes_assinatura2[2])

                    # Captura Alunos
                    if re.match(r"^\d{2}\s", l) and " ATIVO " in l:
                        pre, pos = l.split(" ATIVO ")
                        d = pos.strip().split()
                        
                        if len(d) >= 13:
                            alunos_extraidos.append({
                                "Nome": pre[2:].strip(),
                                "Situação": "ATIVO",
                                "Turma": "1º Ano 01",
                                "Periodo": "Manhã",
                                "Faltas_Total": d[0],
                                "Freq_Perc": d[4],
                                "LP": d[5],
                                "MAT": d[6],
                                "CIE": d[7],
                                "HIST": d[8],
                                "GEOG": d[9],
                                "ART": d[10],
                                "EF": d[11],
                                "TEC": d[12]
                            })

                df_final = pd.DataFrame(alunos_extraidos)
                
                if not df_final.empty:
                    st.success(f"✅ Ata processada. {len(df_final)} boletins prontos para emissão.")
                    
                    with st.expander("🔍 Conferência da Equipe e Alunos lidos"):
                        st.write(f"**Polivalente:** {p_poli} | **Tecnologias:** {p_tec}")
                        st.write(f"**Gestor:** {diretor} | **Coord:** {coord}")
                        st.dataframe(df_final[["Nome", "Freq_Perc", "LP", "MAT"]], hide_index=True)

                    if st.button("🚀 Gerar, Assinar e Salvar Boletins (ZIP)", type="primary"):
                        z_io = io.BytesIO()
                        
                        with st.spinner("Gerando PDFs e salvando assinaturas no Supabase..."):
                            with zipfile.ZipFile(z_io, "a", zipfile.ZIP_DEFLATED) as zf:
                                for _, r in df_final.iterrows():
                                    
                                    # Chama a função que devolve o PDF e os dados de validação
                                    pdf_out, cod_gerado, hash_gerado = gerar_boletim_pdf(r, diretor, coord, p_poli, p_arte, p_edf, p_tec, dias_letivos)
                                    
                                    # Insere o arquivo no ZIP final
                                    zf.writestr(f"Boletim_{r['Nome'].replace(' ','_')}.pdf", pdf_out)
                                    
                                    # --- SALVAMENTO NO SUPABASE ---
                                    try:
                                        dados_validacao = {
                                            "codigo_curto": cod_gerado,
                                            "hash_sha256": hash_gerado,
                                            "nome_aluno": r['Nome'],
                                            "trimestre": "1º Trimestre",
                                            "ano_letivo": "2026",
                                            "emitido_por": f"{coord} e {diretor}"
                                        }
                                        # Envia para a tabela (verifique se 'supabase' está instanciado no seu app_pei.py)
                                        supabase.table("validacao_documentos").insert(dados_validacao).execute()
                                    except Exception as e_db:
                                        st.warning(f"Erro ao salvar registro de validação no banco para {r['Nome']}: {e_db}")
                        
                        st.success("🎉 Lote de boletins gerado e chaves de validação registradas com sucesso!")
                        st.download_button(
                            label="📥 Baixar Arquivo ZIP com Boletins",
                            data=z_io.getvalue(),
                            file_name="Boletins_Assinados_1Tri.zip",
                            mime="application/zip",
                            type="primary"
                        )
                else:
                    st.warning("A leitura foi concluída, mas nenhum aluno 'ATIVO' foi identificado na tabela.")
                    
            except Exception as e:
                st.error(f"Erro crítico no processamento da Ata: {e}")













































# =====================================================================
# MÓDULO: ÁLBUM PREMIUM E ARENA DE JOGOS (CLOUD SAFE + ANIMAÇÃO DE PACOTE)
# Tabelas: estudantes, figurinhas, inventario_album, banca_trocas, ranking_jogos
# =====================================================================
import os
import random
import time
import json
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

# --- Configurações Supabase (Mantenha as suas credenciais originais aqui) ---
# supabase: Client = st.session_state.get('supabase_client') 
# (Assegure-se que a conexão Supabase está ativa antes de chamar estas funções)

# =====================================================================
# FUNÇÕES DE BANCO DE DADOS (SUPABASE)
# =====================================================================

def salvar_ranking(ra, nome, jogo, pontuacao, detalhes=""):
    try:
        supabase.table("ranking_jogos").insert({
            "estudante_ra": ra,
            "nome_estudante": nome,
            "jogo": jogo,
            "pontuacao": pontuacao,
            "detalhes": detalhes
        }).execute()
    except Exception as e:
        pass # Ignora erro silenciosamente caso a tabela ainda não exista

def puxar_dados_album_estudante(ra):
    res_est = supabase.table("estudantes").select("nome, turma, pacotes_disponiveis").eq("ra", ra).execute()
    if not res_est.data: return {"pacotes": 0, "coladas": [], "repetidas": [], "catalogo_ids": [], "turma": "Sem Turma", "total": 0}
    
    pacotes = res_est.data[0]['pacotes_disponiveis']
    turma_aluno = res_est.data[0]['turma']
    
    # Puxar catálogo completo (turma + lendárias)
    res_turma = supabase.table("figurinhas").select("id").eq("turma", turma_aluno).execute()
    res_lendarias = supabase.table("figurinhas").select("id").eq("tipo", "lendaria").execute()
    catalogo_ids = sorted(list(set([f['id'] for f in res_turma.data] + [f['id'] for f in res_lendarias.data])))
    
    # Puxar inventário
    res_inv = supabase.table("inventario_album").select("figurinha_id, quantidade").eq("estudante_ra", ra).execute()
    
    coladas = []; repetidas = []
    for item in res_inv.data:
        f_id = item['figurinha_id']; qtd = item.get('quantidade', 1)
        if f_id in catalogo_ids:
            if qtd >= 1: coladas.append(f_id)
            if qtd > 1:
                for _ in range(qtd - 1): repetidas.append(f_id)
    
    # Puxar detalhes das figurinhas do catálogo para a animação
    if catalogo_ids:
        res_detalhes = supabase.table("figurinhas").select("id, nome, foto_path, tipo").in_("id", catalogo_ids).execute()
        mapa_detalhes = {f['id']: f for f in res_detalhes.data}
    else:
        mapa_detalhes = {}

    return {
        "pacotes": pacotes, 
        "coladas": coladas, 
        "repetidas": repetidas, 
        "catalogo_ids": catalogo_ids, 
        "mapa_detalhes": mapa_detalhes,
        "turma": turma_aluno, 
        "total": len(catalogo_ids)
    }

def processar_abertura_pacote_supa(ra, catalogo_ids_permitidos):
    if not catalogo_ids_permitidos: return None
    res_est = supabase.table("estudantes").select("pacotes_disponiveis").eq("ra", ra).execute()
    if not res_est.data: return None
    disponiveis = res_est.data[0]['pacotes_disponiveis']
    if disponiveis <= 0: return None
    
    # 1. Debitar pacote
    supabase.table("estudantes").update({"pacotes_disponiveis": disponiveis - 1}).eq("ra", ra).execute()
    
    # 2. Sortear 5 figurinhas
    sorteio = random.choices(catalogo_ids_permitidos, k=5)
    
    # 3. Atualizar Inventário (Lógica Upsert manual para Cloud)
    res_inv = supabase.table("inventario_album").select("figurinha_id, quantidade").eq("estudante_ra", ra).in_("figurinha_id", sorteio).execute()
    qtd_map = {item['figurinha_id']: item['quantidade'] for item in res_inv.data}
    
    sorteio_counts = {}
    for f in sorteio: sorteio_counts[f] = sorteio_counts.get(f, 0) + 1
    
    for f_id, count in sorteio_counts.items():
        if f_id in qtd_map:
            supabase.table("inventario_album").update({"quantidade": qtd_map[f_id] + count}).eq("estudante_ra", ra).eq("figurinha_id", f_id).execute()
        else:
            supabase.table("inventario_album").insert({"estudante_ra": ra, "figurinha_id": f_id, "quantidade": count}).execute()
            
    return sorteio

# =====================================================================
# ESTILOS CSS GERAIS (ÁLBUM PREMIUM)
# =====================================================================
def injetar_css_album_premium():
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Oswald:wght@400;600;700&family=Poppins:wght@400;600&display=swap');
        
        /* Container Principal */
        .album-premium-container { font-family: 'Poppins', sans-serif; background-color: #f4f7f6; padding: 20px; border-radius: 15px; }
        
        /* Cabeçalho */
        .header-premium { background: linear-gradient(135deg, #004d23 0%, #009c3b 50%, #004d23 100%); padding: 25px; border-radius: 20px; text-align: center; border: 4px solid #d4af37; box-shadow: 0px 10px 30px rgba(0, 156, 59, 0.4); margin-bottom: 25px; }
        .header-premium h1 { font-family: 'Oswald', sans-serif; margin: 0; color: #ffdf00; font-size: 3rem; text-transform: uppercase; text-shadow: 3px 3px 0px #002776; }
        
        /* Slots do Álbum */
        .fig-wrapper { width: 100%; aspect-ratio: 3 / 4; margin-bottom: 20px; perspective: 1000px; }
        .slot-vazio { width: 100%; height: 100%; border: 2px dashed rgba(0, 156, 59, 0.3); background-color: rgba(255,255,255,0.8); display: flex; flex-direction: column; align-items: center; justify-content: center; border-radius: 8px; color: rgba(0, 156, 59, 0.4); }
        .slot-vazio .numero { font-family: 'Oswald', sans-serif; font-size: 2.5rem; font-weight: 700; }
        
        .slot-preenchido { width: 100%; height: 100%; background-color: #ffffff; padding: 4%; border-radius: 8px; box-shadow: 0px 5px 15px rgba(0,0,0,0.2); display: flex; flex-direction: column; border: 1px solid #e0e0e0; transition: transform 0.3s ease; position: relative; }
        .slot-preenchido:hover { transform: scale(1.05) translateY(-5px); z-index: 5; }
        
        .foto-area { background-color: #e0e0e0; flex-grow: 1; border-radius: 4px; border: 2px solid #002776; overflow: hidden; display: flex; align-items: center; justify-content: center; }
        .foto-img { width: 100%; height: 100%; object-fit: cover; }
        
        .foto-rodape { height: 25%; background: #009c3b; margin-top: 5px; border-radius: 4px; display: flex; flex-direction: column; align-items: center; justify-content: center; color: #ffffff; font-family: 'Oswald', sans-serif; text-align: center; padding: 2px; }
        .rodape-num { font-size: 0.7rem; color: #ffdf00; }
        .rodape-nome { font-size: 0.75rem; font-weight: 700; text-transform: uppercase; line-height: 1.1; }
        
        /* Tipo Lendária */
        .lendaria { background: linear-gradient(135deg, #ffd700, #ffb300); border: none; box-shadow: 0px 5px 20px rgba(255, 215, 0, 0.5); }
        .lendaria .foto-area { border: 2px solid #b8860b; }
        .lendaria .foto-rodape { background: linear-gradient(to bottom, #d4af37, #b8860b); }
        </style>
    """, unsafe_allow_html=True)

# =====================================================================
# COMPONENTE DE ANIMAÇÃO: RASGAR PACOTINHO (HTML/CSS/JS)
# =====================================================================
def renderizar_animacao_abertura(figurinhas_sorteadas_detalhes):
    import json
    import streamlit.components.v1 as components
    
    figs_json = json.dumps(figurinhas_sorteadas_detalhes)
    
    # Mantivemos a URL original, mas adicionamos uma camada de segurança no CSS
    url_pacote_fechado = "https://bkqhbwnphnnueyyhdqbn.supabase.co/storage/v1/object/public/fotos_alunos/8A4EB18F-22F4-4076-997B-4C285995DE5F-%281%29.jpg"
    url_costas_figurinha = "https://bkqhbwnphnnueyyhdqbn.supabase.co/storage/v1/object/public/fotos_alunos/23A39E31-E7F4-444E-BC12-508181FC50E7.jpg"

    html_content = f"""
    <!DOCTYPE html>
    <html lang="pt-br">
    <head>
    <meta charset="UTF-8">
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Oswald:wght@600&family=Poppins:wght@700&display=swap');
        
        body {{ margin: 0; padding: 0; display: flex; justify-content: center; align-items: center; background: transparent; overflow: hidden; font-family: 'Poppins', sans-serif; }}
        #cena-animacao {{ position: relative; width: 100%; height: 500px; display: flex; justify-content: center; align-items: center; overflow: hidden; }}

        .pacotinho-wrapper {{ position: relative; width: 220px; height: 320px; cursor: pointer; transition: transform 0.3s ease; z-index: 10; }}
        .pacotinho-wrapper:hover {{ transform: scale(1.05) rotate(2deg); }}
        
        /* A MÁGICA DE DEBUG ESTÁ AQUI: Adicionamos cor de fundo, fallback de erro e texto centralizado */
        .pacotinho-imagem {{ 
            width: 100%; height: 100%; 
            background-color: #d4af37; /* Dourado de fallback */
            background-image: url('{url_pacote_fechado}'); 
            background-size: cover; background-position: center; 
            border-radius: 15px; box-shadow: 0 15px 35px rgba(0,0,0,0.4); 
            position: relative; overflow: hidden; border: 3px solid #ffdf00; 
            display: flex; justify-content: center; align-items: center; text-align: center; color: #fff; font-size: 24px; text-shadow: 2px 2px 4px #000;
        }}
        
        /* Se a imagem quebrar, esse texto aparece no fundo dourado */
        .pacotinho-imagem::before {{ content: 'PACOTE MISTERIOSO'; z-index: 1; }}

        .pacotinho-imagem::after {{ content: ''; position: absolute; top: -50%; left: -50%; width: 200%; height: 200%; background: linear-gradient(45deg, transparent, rgba(255,255,255,0.6), transparent); transform: rotate(45deg); animation: brilhoPacote 3s infinite; z-index: 2; }}

        .wrapper-rasgando {{ animation: tremer 0.5s ease-in-out; }}
        .rasgo-superior, .rasgo-inferior {{ position: absolute; width: 100%; height: 50%; background-color: #d4af37; background-image: url('{url_pacote_fechado}'); background-size: 220px 320px; left: 0; transition: all 0.6s cubic-bezier(0.68, -0.55, 0.27, 1.55); opacity: 1; border-radius: 15px; border: 3px solid #ffdf00; box-sizing: border-box; }}
        .rasgo-superior {{ top: 0; background-position: top center; border-bottom: none; border-radius: 15px 15px 0 0; }}
        .rasgo-inferior {{ bottom: 0; background-position: bottom center; border-top: none; border-radius: 0 0 15px 15px; }}
        
        .cena-aberta .rasgo-superior {{ transform: translateY(-150px) rotate(-15deg); opacity: 0; }}
        .cena-aberta .rasgo-inferior {{ transform: translateY(150px) rotate(15deg); opacity: 0; }}
        .explosao-luz {{ position: absolute; width: 300px; height: 300px; background: radial-gradient(circle, rgba(255,223,0,1) 0%, rgba(255,255,255,0) 70%); border-radius: 50%; opacity: 0; transform: scale(0.1); transition: all 0.5s ease-out; z-index: 5; filter: blur(10px); }}
        .cena-aberta .explosao-luz {{ opacity: 1; transform: scale(1); }}

        .area-figurinhas {{ position: absolute; width: 100%; height: 100%; top: 0; left: 0; display: flex; justify-content: center; align-items: center; z-index: 20; pointer-events: none; }}
        .mini-figurinha {{ position: absolute; width: 110px; height: 150px; background-color: white; border-radius: 8px; box-shadow: 0 10px 20px rgba(0,0,0,0.3); border: 2px solid #002776; overflow: hidden; display: flex; flex-direction: column; opacity: 0; transform: translate(0, 0) scale(0.2) rotate(0deg); transition: all 0.8s cubic-bezier(0.175, 0.885, 0.32, 1.275); backface-visibility: hidden; background-image: url('{url_costas_figurinha}'); background-size: cover; }}
        
        .revelar-fig {{ background-image: none !important; animation: virarFigurinha 0.6s forwards; pointer-events: auto; cursor: pointer; }}

        .cena-aberta #fig-0 {{ opacity: 1; transform: translate(-220px, 0px) scale(1) rotate(-10deg); }}
        .cena-aberta #fig-1 {{ opacity: 1; transform: translate(-110px, -30px) scale(1) rotate(-5deg); transition-delay: 0.1s; }}
        .cena-aberta #fig-2 {{ opacity: 1; transform: translate(0px, -50px) scale(1) rotate(0deg); transition-delay: 0.2s; }}
        .cena-aberta #fig-3 {{ opacity: 1; transform: translate(110px, -30px) scale(1) rotate(5deg); transition-delay: 0.3s; }}
        .cena-aberta #fig-4 {{ opacity: 1; transform: translate(220px, 0px) scale(1) rotate(10deg); transition-delay: 0.4s; }}

        .fig-conteudo {{ display: none; width: 100%; height: 100%; flex-direction: column; padding: 3px; box-sizing: border-box; }}
        .revelar-fig .fig-conteudo {{ display: flex; }}
        .fig-img-area {{ flex-grow: 1; background: #e0e0e0; border-radius: 4px; border: 1px solid #002776; overflow: hidden; }}
        .fig-img-area img {{ width: 100%; height: 100%; object-fit: cover; }}
        .fig-txt-area {{ height: 35px; background: #009c3b; color: white; margin-top: 3px; border-radius: 4px; font-family: 'Oswald', sans-serif; display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center; }}
        .fig-txt-area .num {{ font-size: 10px; color: #ffdf00; }}
        .fig-txt-area .nome {{ font-size: 11px; font-weight: 600; text-transform: uppercase; line-height: 1; max-width: 95%; overflow: hidden; }}

        #btn-concluir {{ position: absolute; bottom: 20px; padding: 12px 30px; background: #002776; color: white; border: none; border-radius: 50px; font-size: 18px; font-weight: bold; cursor: pointer; display: none; z-index: 100; box-shadow: 0 5px 15px rgba(0,0,0,0.3); transition: background 0.3s; }}
        #instrucao {{ position: absolute; top: 20px; color: #004d23; font-size: 20px; font-weight: bold; text-align: center; width: 100%; z-index: 30; text-shadow: 0 2px 4px white; }}

        @keyframes brilhoPacote {{ 0% {{ left: -50%; top: -50%; }} 100% {{ left: 100%; top: 100%; }} }}
        @keyframes tremer {{ 0% {{ transform: rotate(0deg); }} 20% {{ transform: rotate(-5deg); }} 40% {{ transform: rotate(5deg); }} 60% {{ transform: rotate(-3deg); }} 80% {{ transform: rotate(3deg); }} 100% {{ transform: rotate(0deg); }} }}
        @keyframes virarFigurinha {{ 0% {{ transform: scale(1) rotateY(0deg); }} 50% {{ transform: scale(1.1) rotateY(90deg); }} 100% {{ transform: scale(1) rotateY(180deg); }} }}
    </style>
    </head>
    <body>

    <div id="cena-animacao">
        <div id="instrucao">👋 Clique no pacotinho para rasgar!</div>
        <div class="explosao-luz"></div>
        <div class="pacotinho-wrapper" id="pacote-clicavel" onclick="rasgarPacote()">
            <div class="pacotinho-imagem" id="pacote-inteiro"></div>
            <div class="rasgo-superior" id="metade-sup" style="display:none;"></div>
            <div class="rasgo-inferior" id="metade-inf" style="display:none;"></div>
        </div>
        <div class="area-figurinhas" id="container-figurinhas"></div>
        <button id="btn-concluir" onclick="finalizarAbertura()">Guardar na Coleção</button>
    </div>

    <script>
        const dadosFigurinhas = {figs_json};
        let pacoteRasgado = false;
        let figsReveladas = 0;

        const container = document.getElementById('container-figurinhas');
        dadosFigurinhas.forEach((fig, index) => {{
            const div = document.createElement('div');
            div.className = 'mini-figurinha';
            div.id = `fig-${{index}}`;
            div.onclick = function() {{ revelarFigurinha(index); }};
            div.innerHTML = `
                <div class="fig-conteudo">
                    <div class="fig-img-area">
                        <img src="${{fig.foto_path}}" alt="${{fig.nome}}" onerror="this.src='https://i.imgur.com/w9O7B4G.png'">
                    </div>
                    <div class="fig-txt-area">
                        <span class="num">Nº ${{fig.id}}</span>
                        <span class="nome">${{fig.nome}}</span>
                    </div>
                </div>
            `;
            container.appendChild(div);
        }});

        function rasgarPacote() {{
            if (pacoteRasgado) return;
            pacoteRasgado = true;

            const wrapper = document.getElementById('pacote-clicavel');
            const inteiro = document.getElementById('pacote-inteiro');
            const sup = document.getElementById('metade-sup');
            const inf = document.getElementById('metade-inf');
            const instrucao = document.getElementById('instrucao');
            const cena = document.getElementById('cena-animacao');

            wrapper.classList.add('wrapper-rasgando');
            instrucao.innerText = "💥💥💥";
            
            setTimeout(() => {{
                inteiro.style.display = 'none';
                sup.style.display = 'block';
                inf.style.display = 'block';
                cena.classList.add('cena-aberta');
                setTimeout(() => {{ wrapper.style.display = 'none'; }}, 600);
                instrucao.innerHTML = "✨ Toca nas cartas para revelar! ✨";
                instrucao.style.color = "#d4af37";
            }}, 500);
        }}

        function revelarFigurinha(index) {{
            const fig = document.getElementById(`fig-${{index}}`);
            if (fig.classList.contains('revelar-fig')) return;
            fig.classList.add('revelar-fig');
            figsReveladas++;
            if (figsReveladas === 5) {{
                setTimeout(() => {{
                    document.getElementById('btn-concluir').style.display = 'block';
                    document.getElementById('instrucao').innerText = "🎉 Excelente sorte! 🎉";
                }}, 800);
            }}
        }}

        function finalizarAbertura() {{
            window.parent.postMessage({{type: 'streamlit:wrapper_event', data: {{action: 'recarregar_album'}}}}, '*');
        }}
    </script>
    </body>
    </html>
    """
    components.html(html_content, height=520)

# =====================================================================
# RENDERIZAÇÃO DO MÓDULO DO ÁLBUM
# =====================================================================
def render_modulo_album():
    injetar_css_album_premium()
    
    # RA do estudante logado (Supomos que esteja no session_state)
    estudante_ra = st.session_state.get('usuario_ra', 'RA-TESTE')
    estudante_nome = st.session_state.get('usuario_nome', 'Estudante')
    
    # Carregar dados do DB
    dados_db = puxar_dados_album_estudante(estudante_ra)
    
    st.markdown('<div class="album-premium-container">', unsafe_allow_html=True)
    st.markdown(f'<div class="header-premium"><h1>⚽ SUPER ÁLBUM DOS CRAQUES ⚽</h1><p style="font-size: 1.2rem; color: #ffffff; font-weight: 600;">O álbum de figurinhas da turma</p></div>', unsafe_allow_html=True)
    
    aba_album, aba_pacotes, aba_trocas, aba_jogos, aba_ranking = st.tabs(["📖 Meu Álbum", "📦 Abrir Pacotinhos", "🤝 Banca de Trocas", "🏟️ Arena de Jogos", "🏆 Ranking Geral"])
    
    # --- ABA: MEU ÁLBUM ---
    with aba_album:
        figurinhas_por_pagina = 10
        total_figuras = dados_db['total']
        
        if total_figuras == 0:
            st.warning("Nenhuma figurinha cadastrada para a sua turma.")
        else:
            if 'pag_album' not in st.session_state: st.session_state['pag_album'] = 0
            total_paginas = (total_figuras // figurinhas_por_pagina) + (1 if total_figuras % figurinhas_por_pagina > 0 else 0)
            
            c_prev, c_page, c_next = st.columns([1, 2, 1])
            with c_prev:
                if st.button("⬅️ ANTERIOR", use_container_width=True) and st.session_state['pag_album'] > 0:
                    st.session_state['pag_album'] -= 1; st.rerun()
            with c_page:
                coladas = len(dados_db['coladas'])
                progresso = coladas / total_figuras if total_figuras > 0 else 0
                st.markdown(f"<div style='text-align: center; color: #004d23; font-family: Oswald; font-size: 18px;'>PÁGINA {st.session_state['pag_album'] + 1} DE {total_paginas}</div>", unsafe_allow_html=True)
                st.progress(progresso, text=f"Progresso: {coladas} coladas de {total_figuras}")
            with c_next:
                if st.button("PRÓXIMA ➡️", use_container_width=True) and st.session_state['pag_album'] < (total_paginas - 1):
                    st.session_state['pag_album'] += 1; st.rerun()
            
            # Renderizar Grid
            inicio_idx = st.session_state['pag_album'] * figurinhas_por_pagina
            ids_da_pagina = dados_db['catalogo_ids'][inicio_idx : inicio_idx + figurinhas_por_pagina]
            mapa = dados_db['mapa_detalhes']
            
            for linha in range(0, len(ids_da_pagina), 5):
                cols = st.columns(5)
                for i, col in enumerate(cols):
                    if linha + i < len(ids_da_pagina):
                        f_id = ids_da_pagina[linha + i]
                        info = mapa.get(f_id, {})
                        
                        f_nome = info.get('nome', '???')
                        f_tipo = info.get('tipo', 'comum')
                        f_foto = info.get('foto_path', '')
                        
                        classe_lendaria = "lendaria" if f_tipo == "lendaria" else ""
                        esta_colada = f_id in dados_db['coladas']
                        
                        html_fig = ""
                        if esta_colada:
                            # Slot Preenchido
                            foto_html = f'<img src="{f_foto}" class="foto-img" onerror="this.src=\'https://i.imgur.com/w9O7B4G.png\'">' if f_foto else '<span style="font-size: 2rem;">📸</span>'
                            html_fig = f"""
                            <div class="fig-wrapper"><div class="slot-preenchido {classe_lendaria}">
                                <div class="foto-area">{foto_html}</div>
                                <div class="foto-rodape"><span class="rodape-num">Nº {f_id}</span><span class="rodape-nome">{f_nome}</span></div>
                            </div></div>
                            """
                        else:
                            # Slot Vazio
                            html_fig = f"""
                            <div class="fig-wrapper"><div class="slot-vazio">
                                <div class="numero">{f_id}</div>
                                <div style="font-size: 0.7rem; font-weight: bold; text-transform: uppercase;">Faltando</div>
                            </div></div>
                            """
                        col.markdown(html_fig, unsafe_allow_html=True)

    # --- ABA: ABRIR PACOTES (COM ANIMAÇÃO) ---
    with aba_pacotes:
        st.markdown(f"<h2 style='font-family: Oswald; color: #004d23; text-align: center;'>📦 Salão de Abertura</h2>", unsafe_allow_html=True)
        st.markdown(f"<h3 style='text-align: center; color: #009c3b;'>Possui <b style='font-size: 2rem; color: #002776;'>{dados_db['pacotes']}</b> pacotinhos fechados.</h3>", unsafe_allow_html=True)
        
        if 'processando_abertura' not in st.session_state:
            st.session_state['processando_abertura'] = False
        if 'figurinhas_para_animar' not in st.session_state:
            st.session_state['figurinhas_para_animar'] = None

        area_acao = st.empty()

        if dados_db['pacotes'] > 0:
            if not st.session_state['processando_abertura']:
                with area_acao.container():
                    st.write("---")
                    st.markdown("<p style='text-align: center; color: #666;'>Prepare o seu dedo! Clique no botão abaixo para trazer o seu pacotinho para a mesa.</p>", unsafe_allow_html=True)
                    _, col_btn, _ = st.columns([1, 2, 1])
                    if col_btn.button("✨ TRAZER PACOTINHO PARA A MESA! ✨", type="primary", use_container_width=True):
                        
                        sorteio_ids = processar_abertura_pacote_supa(estudante_ra, dados_db['catalogo_ids'])
                        if sorteio_ids:
                            mapa_geral = dados_db['mapa_detalhes']
                            detalhes_sorteio = []
                            
                            for f_id in sorteio_ids:
                                info = mapa_geral.get(f_id, {'id': f_id, 'nome': 'Desconhecido', 'foto_path': '', 'tipo': 'comum'})
                                
                                # CORRECÇÃO DA URL DA FIGURINHA: Garante o endereço completo para o iframe
                                foto = info.get('foto_path', '')
                                if foto and not foto.startswith('http'):
                                    # Concatena o domínio público do bucket do Supabase
                                    foto = f"https://bkqhbwnphnnueyyhdqbn.supabase.co/storage/v1/object/public/fotos_alunos/{foto}"
                                
                                info_completa = info.copy()
                                info_completa['foto_path'] = foto
                                detalhes_sorteio.append(info_completa)
                            
                            st.session_state['figurinhas_para_animar'] = detalhes_sorteio
                            st.session_state['processando_abertura'] = True
                            st.balloons()
                            st.rerun()
            else:
                area_acao.empty() 
                st.write("---")
                
                # Executa a animação com as URLs corrigidas e absolutas
                renderizar_animacao_abertura(st.session_state['figurinhas_para_animar'])
                
                if st.button("Finalizar e ver Álbum", use_container_width=True):
                    st.session_state['processando_abertura'] = False
                    st.session_state['figurinhas_para_animar'] = None
                    st.rerun()
        else:
            area_acao.error("🚨 Você não tem pacotinhos disponíveis.")

    with aba_trocas:
        st.markdown("### 🤝 Mercado de Transferências")
        res_cat = supabase.table("figurinhas").select("id, nome, tipo").in_("id", dados_db['catalogo_ids']).execute()
        catalogo_map = {f['id']: f for f in res_cat.data}
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### 📥 Suas Repetidas")
            if dados_db['repetidas']:
                contagem = {x: dados_db['repetidas'].count(x) for x in set(dados_db['repetidas'])}
                for f_id, qtd in contagem.items():
                    info = catalogo_map.get(f_id, {})
                    st.info(f"🟢 **Nº {f_id} - {info.get('nome', 'N/A')}** {'✨(LENDÁRIA)' if info.get('tipo') == 'lendaria' else ''} — Tem {qtd} extra(s)")
            else: st.write("Não possui cartas repetidas.")
        with col2:
            st.markdown("#### 📢 Oferta de Troca")
            if dados_db['repetidas']:
                fig_oferecida = st.selectbox("Dar Figurinha:", list(set(dados_db['repetidas'])))
                faltantes = [x for x in dados_db['catalogo_ids'] if x not in dados_db['coladas']]
                if faltantes:
                    fig_desejada = st.selectbox("Receber Figurinha:", faltantes)
                    if st.button("Registrar no Mural", type="primary", use_container_width=True):
                        supabase.table("banca_trocas").insert({"estudante_ra": estudante_ra, "id_oferecida": fig_oferecida, "id_desejada": fig_desejada, "status": "ativo"}).execute()
                        st.success("Proposta gravada!"); time.sleep(1); st.rerun()
                else: st.success("Álbum completo!")
            else: st.warning("Abra pacotes primeiro.")

        st.write("---"); st.markdown("#### 📢 Mural de Anúncios da Escola")
        res_mural = supabase.table("banca_trocas").select("*").eq("status", "ativo").order("data_criacao", desc=True).execute()
        if res_mural.data:
            res_estudantes = supabase.table("estudantes").select("ra, nome").execute()
            estudantes_map = {e['ra']: e['nome'] for e in res_estudantes.data}
            for anuncio in res_mural.data:
                item_id = anuncio['id']; dono_ra = anuncio['estudante_ra']; dono_nome = estudantes_map.get(dono_ra, "Aluno")
                id_of = anuncio['id_oferecida']; id_des = anuncio['id_desejada']
                if id_of in dados_db['catalogo_ids'] or id_des in dados_db['catalogo_ids']:
                    info_of = catalogo_map.get(id_of, {}); tipo_of = info_of.get('tipo', 'comum'); info_des = catalogo_map.get(id_des, {})
                    escala_ouro = tipo_of == "lendaria"; border_color = "#d4af37" if escala_ouro else "#009c3b"; bg_card = "#fffdf0" if escala_ouro else "#ffffff"
                    st.markdown(f"""<div style='background-color: {bg_card}; padding: 15px; border-radius: 10px; border-left: 6px solid {border_color}; box-shadow: 2px 2px 6px rgba(0,0,0,0.05); margin-bottom: 10px;'><strong style='color: #002776;'>{dono_nome}</strong> {"<span style='color:#b8860b;'>[LENDÁRIA]</span>" if escala_ouro else ""}<br><span style='font-size:0.9rem; color:#444;'>Dá a <b>Nº {id_of} ({info_of.get('nome', '')})</b> em troca da <b>Nº {id_des} ({info_des.get('nome', '')})</b></span></div>""", unsafe_allow_html=True)
                    if dono_ra != estudante_ra:
                        if st.button(f"Aceitar Troca com {dono_nome} (Nº {id_of})##{item_id}", use_container_width=True):
                            res_check = supabase.table("inventario_album").select("quantidade").eq("estudante_ra", estudante_ra).eq("figurinha_id", id_des).execute()
                            if res_check.data and res_check.data[0]['quantidade'] > 1:
                                supabase.table("banca_trocas").update({"status": "concluido"}).eq("id", item_id).execute()
                                q_atual = res_check.data[0]['quantidade']
                                supabase.table("inventario_album").update({"quantidade": q_atual - 1}).eq("estudante_ra", estudante_ra).eq("figurinha_id", id_des).execute()
                                res_q2 = supabase.table("inventario_album").select("quantidade").eq("estudante_ra", estudante_ra).eq("figurinha_id", id_of).execute()
                                if res_q2.data: supabase.table("inventario_album").update({"quantidade": res_q2.data[0]['quantidade'] + 1}).eq("estudante_ra", estudante_ra).eq("figurinha_id", id_of).execute()
                                else: supabase.table("inventario_album").insert({"estudante_ra": estudante_ra, "figurinha_id": id_of, "quantidade": 1}).execute()
                                res_d1 = supabase.table("inventario_album").select("quantidade").eq("estudante_ra", dono_ra).eq("figurinha_id", id_of).execute()
                                if res_d1.data: supabase.table("inventario_album").update({"quantidade": res_d1.data[0]['quantidade'] - 1}).eq("estudante_ra", dono_ra).eq("figurinha_id", id_of).execute()
                                res_d2 = supabase.table("inventario_album").select("quantidade").eq("estudante_ra", dono_ra).eq("figurinha_id", id_des).execute()
                                if res_d2.data: supabase.table("inventario_album").update({"quantidade": res_d2.data[0]['quantidade'] + 1}).eq("estudante_ra", dono_ra).eq("figurinha_id", id_des).execute()
                                else: supabase.table("inventario_album").insert({"estudante_ra": dono_ra, "figurinha_id": id_des, "quantidade": 1}).execute()
                                st.success("Troca realizada com sucesso!"); time.sleep(1); st.rerun()
                            else: st.error(f"Você não possui a figurinha Nº {id_des} sobrando no inventário.")
        else: st.info("Nenhum anúncio de troca em aberto.")

    # =========================================================================
    # ARENA DE JOGOS: HTML PURO COM CÓDIGOS DE VALIDAÇÃO (100% CLOUD SAFE)
    # =========================================================================
    with aba_jogos:
        st.markdown("### 🏟️ Arena de Jogos")
        st.write("Jogue e anote o **Código de Segurança** no final para salvar a sua pontuação!")
        st.write("---")
        
        jogo1, jogo2, jogo3, jogo4, jogo5 = st.tabs(["🧠 Memória", "🏃 Pula Craque", "🧤 Goleiro", "🎯 Falta", "⚽ Embaixadinha"])

        # ---------------- 1. JOGO DA MEMÓRIA INTELIGENTE ----------------
        with jogo1:
            html_memoria = """
            <!DOCTYPE html><html><head><style>@import url('https://fonts.googleapis.com/css2?family=Oswald:wght@600&display=swap');body{font-family:'Oswald',sans-serif;text-align:center;margin:0;background:transparent;}#stats{font-size:20px;color:#004d23;margin-bottom:15px;}.grid{display:flex;flex-wrap:wrap;gap:8px;justify-content:center;max-width:400px;margin:0 auto;}.card{width:65px;height:85px;perspective:1000px;cursor:pointer;}.card-inner{width:100%;height:100%;transition:transform 0.5s;transform-style:preserve-3d;position:relative;border-radius:8px;box-shadow:0 4px 6px rgba(0,0,0,0.2);}.card.open .card-inner{transform:rotateY(180deg);}.card.match .card-inner{transform:rotateY(180deg) scale(1.05);box-shadow:0 0 10px #d4af37;}.card-front,.card-back{width:100%;height:100%;position:absolute;backface-visibility:hidden;border-radius:8px;display:flex;align-items:center;justify-content:center;font-size:2rem;}.card-front{background:linear-gradient(135deg,#004d23,#009c3b);border:2px solid #ffdf00;}.card-back{background:white;transform:rotateY(180deg);border:2px solid #009c3b;}button{margin-top:20px;padding:10px 20px;font-size:20px;background:#d4af37;color:#002776;border:none;border-radius:8px;cursor:pointer;font-family:'Oswald';font-weight:bold;}#gameOver{display:none;font-size:22px;color:#002776;margin-top:20px;background:#fffdf0;padding:15px;border-radius:10px;border:3px solid #d4af37;}</style>
            </head><body>
            <div id="stats">Nível: <span id="lv">1</span>/4 | Erros: <span id="err">0</span> | Tempo: <span id="tm">0</span>s</div>
            <div class="grid" id="board"></div>
            <div id="gameOver"></div>
            <button id="btn" onclick="startLevel()">Iniciar Nível 1 ➡️</button>
            <script>
                const allEmojis = ['⚽','🏆','🏟️','🧤','👟','🇧🇷','🥅','⏱️','🥇','🎉'];
                let level=1; let maxLevel=4; let errors=0; let time=0; let timerId; let openCards=[]; let matched=0; let totalCards=0; let isPlaying=false;
                const board = document.getElementById('board'); const lvEl = document.getElementById('lv'); const errEl = document.getElementById('err'); const tmEl = document.getElementById('tm'); const btn = document.getElementById('btn'); const over = document.getElementById('gameOver');
                function startLevel() {
                    isPlaying=true; btn.style.display='none'; board.innerHTML=''; openCards=[]; matched=0; lvEl.innerText=level;
                    let numPairs = 2 + (level*2); totalCards = numPairs * 2;
                    let emojis = allEmojis.slice(0, numPairs); let cards = [...emojis, ...emojis].sort(() => Math.random() - 0.5);
                    cards.forEach(e => { let c = document.createElement('div'); c.className='card'; c.innerHTML = `<div class="card-inner"><div class="card-front">⚽</div><div class="card-back">${e}</div></div>`; c.onclick = () => flipCard(c, e); board.appendChild(c); });
                    timerId=setInterval(()=>{time++;tmEl.innerText=time;},1000);
                }
                function flipCard(c, e) {
                    if(!isPlaying) return;
                    if(openCards.length<2 && !c.classList.contains('open') && !c.classList.contains('match')) {
                        c.classList.add('open'); openCards.push({el: c, emoji: e});
                        if(openCards.length===2) {
                            setTimeout(() => {
                                if(openCards[0].emoji === openCards[1].emoji) {
                                    openCards[0].el.classList.add('match'); openCards[1].el.classList.add('match'); matched+=2;
                                    if(matched===totalCards) {
                                        clearInterval(timerId); isPlaying=false;
                                        if(level < maxLevel) { level++; btn.innerText="Iniciar Nível " + level + " ➡️"; btn.style.display='inline-block'; }
                                        else { 
                                            over.style.display='block'; board.style.display='none';
                                            let pontos = Math.max(0, 1000 - (time * 2) - (errors * 10));
                                            let segCode = (pontos * 3) + 45;
                                            over.innerHTML = `🏆 PARABÉNS! 🏆<br>Pontos de Memória: <b>${pontos}</b><br><span style="color:red;">CÓDIGO SECRETO: ${segCode}</span><br><br><span style="font-size:16px;">Digite esses números abaixo para salvar!</span>`;
                                        }
                                    }
                                } else { openCards[0].el.classList.remove('open'); openCards[1].el.classList.remove('open'); errors++; errEl.innerText=errors; }
                                openCards=[];
                            }, 700);
                        }
                    }
                }
            </script></body></html>
            """
            components.html(html_memoria, height=520)
            c_m1, c_m2, c_m3 = st.columns([1, 1, 1])
            pts_m = c_m1.number_input("🧠 Pontos no Jogo", min_value=0, max_value=9999, key="p_mem")
            cod_m = c_m2.number_input("🔒 Código Secreto", min_value=0, max_value=99999, key="c_mem")
            if c_m3.button("Salvar Recorde", use_container_width=True, key="btn_mem"):
                if pts_m > 0 and cod_m == (pts_m * 3) + 45:
                    salvar_ranking(estudante_ra, estudante_nome, "memoria", pts_m, "4 Níveis Concluídos")
                    st.success("✅ Salvo com sucesso no Mural!")
                else: st.error("❌ Código Inválido ou Pontuação incorreta.")

        # ---------------- 2. SUPER PULA CRAQUE ----------------
        with jogo2:
            html_pula = """
            <!DOCTYPE html><html><head><style>@import url('https://fonts.googleapis.com/css2?family=Oswald:wght@600&display=swap');body{margin:0;font-family:'Oswald',sans-serif;user-select:none;overflow:hidden;background:transparent;}#game{width:100%;height:350px;background:linear-gradient(to bottom, #87CEEB 0%, #E0F7FA 70%, #4CAF50 70%, #2E7D32 100%);position:relative;border-radius:12px;overflow:hidden;box-shadow:inset 0 0 20px rgba(0,0,0,0.2);cursor:pointer;}.cloud{position:absolute;font-size:40px;color:rgba(255,255,255,0.7);white-space:nowrap;}#player{position:absolute;bottom:105px;left:50px;font-size:55px;z-index:10; filter: drop-shadow(2px 4px 4px rgba(0,0,0,0.3)); display:flex; align-items:center; justify-content:center;} .flip { transform: scaleX(-1); display: inline-block; margin-right: -10px;} #obstacle{position:absolute;bottom:105px;left:800px;font-size:45px;z-index:9; filter: drop-shadow(2px 4px 4px rgba(0,0,0,0.3));}#coin{position:absolute;bottom:200px;left:900px;font-size:35px;z-index:8;}#scoreBoard{position:absolute;top:15px;left:20px;font-size:28px;color:#002776;text-shadow:1px 1px 0px #fff;z-index:20;}#msg{display:none;flex-direction:column;justify-content:center;align-items:center;position:absolute;top:0;left:0;width:100%;height:100%;background:rgba(0,0,0,0.85);color:white;font-size:30px;z-index:30;text-shadow:2px 2px 5px #000;text-align:center;padding:20px;box-sizing:border-box;}button{margin-top:15px;padding:12px 25px;font-size:22px;background:#ffdf00;color:#002776;border:none;border-radius:8px;cursor:pointer;font-family:'Oswald';font-weight:bold;box-shadow:0 4px 10px rgba(0,0,0,0.5); z-index:40; position:relative;}</style>
            </head><body>
            <div id="game"><div class="cloud" style="top:20px; left:100px; font-size:60px;">☁️</div><div class="cloud" style="top:50px; left:400px; font-size:40px;">☁️</div>
            <div id="scoreBoard">Pontos: <span id="score">0</span></div>
            <div id="player"><span class="flip">🏃‍♂️</span><span id="ballRoll">⚽</span></div>
            <div id="obstacle">🚧</div><div id="coin">🏆</div>
            <div id="msg"><div id="gameOverText" style="margin-bottom:10px;">SUPER PULA CRAQUE</div><button id="btnPlay" onclick="startGame(event)">JOGAR</button></div></div>
            <script>
            const p=document.getElementById('player');const o=document.getElementById('obstacle');const c=document.getElementById('coin');
            const sEl=document.getElementById('score');const msg=document.getElementById('msg'); const btn = document.getElementById('btnPlay');
            const clouds=document.querySelectorAll('.cloud'); let isPlaying=false; let score=0; let frame=0;
            let py=105; let vy=0; let gravity=0.5; let isJumping=false; let canDouble=false;
            let ox=800; let speed=4; let cx=1000; let cy=200; let coinActive=true; let animId; let ballRot=0;
            
            msg.style.display='flex';
            
            function jump(e){
                if(e && e.target && e.target.closest && e.target.closest('button')) return;
                if(e && e.cancelable) e.preventDefault();
                if(!isPlaying) return;
                if(!isJumping){
                    vy=12;isJumping=true;canDouble=true;p.innerHTML="<span class='flip'>🤸‍♂️</span><span id='ballRoll'>⚽</span>";
                }else if(canDouble){
                    vy=10;canDouble=false;p.innerHTML="<span class='flip'>🦸‍♂️</span><span id='ballRoll'>⚽</span>";
                }
            }
            document.getElementById('game').addEventListener('mousedown', jump);
            document.getElementById('game').addEventListener('touchstart', jump, {passive:false});
            
            function startGame(e){
                if(e)e.stopPropagation();
                isPlaying=true;msg.style.display='none';score=0;sEl.innerText=0;speed=4;ox=800;cx=1200;coinActive=true;py=105;vy=0;isJumping=false;canDouble=false;p.innerHTML="<span class='flip'>🏃‍♂️</span><span id='ballRoll'>⚽</span>";
                if(animId)cancelAnimationFrame(animId);
                loop();
            }
            
            function loop(){
                if(!isPlaying)return;
                vy-=gravity;py+=vy;
                if(py<=105){
                    py=105;vy=0;
                    if(isJumping){
                        isJumping=false;canDouble=false;
                        p.innerHTML=(frame%10<5)?"<span class='flip'>🏃‍♂️</span><span id='ballRoll'>⚽</span>":"<span class='flip'>🚶‍♂️</span><span id='ballRoll'>⚽</span>";
                    }
                }
                if(!isJumping) p.innerHTML=(frame%16<8)?"<span class='flip'>🏃‍♂️</span><span id='ballRoll'>⚽</span>":"<span class='flip'>🚶‍♂️</span><span id='ballRoll'>⚽</span>"; 
                p.style.bottom=py+'px';
                
                let b=document.getElementById('ballRoll'); if(b){ ballRot+=speed*3; b.style.transform=`rotate(${ballRot}deg)`; }
                
                ox-=speed;
                if(ox<-60){
                    ox=600+Math.random()*400;
                    let em=['🟥','🟨','🚧','🦵'];
                    o.innerText=em[Math.floor(Math.random()*em.length)];
                    speed+=0.04;
                } 
                o.style.left=ox+'px';
                
                cx-=speed*0.8;
                if(cx<-50){
                    cx=800+Math.random()*800;cy=180+Math.random()*80;coinActive=true;c.style.display='block';
                } 
                c.style.left=cx+'px'; c.style.bottom=cy+'px';
                
                clouds[0].style.left=(parseInt(clouds[0].style.left||100)-1)+'px';if(parseInt(clouds[0].style.left)<-100)clouds[0].style.left='800px';
                clouds[1].style.left=(parseInt(clouds[1].style.left||400)-0.5)+'px';if(parseInt(clouds[1].style.left)<-100)clouds[1].style.left='800px';
                
                let pRect={x:50,y:py,w:40,h:50}; let oRect={x:ox+15,y:105,w:30,h:40}; let cRect={x:cx,y:cy,w:30,h:30};
                
                if(pRect.x<oRect.x+oRect.w && pRect.x+pRect.w>oRect.x && pRect.y<oRect.y+oRect.h && pRect.y+pRect.h>oRect.y){gameOver();return;}
                if(coinActive && pRect.x<cRect.x+cRect.w && pRect.x+pRect.w>cRect.x && pRect.y<cRect.y+cRect.h && pRect.y+pRect.h>cRect.y){
                    score+=50;coinActive=false;c.style.display='none';sEl.innerText=score;
                }
                if(frame%10===0){score+=1;sEl.innerText=score;} 
                frame++; animId=requestAnimationFrame(loop);
            }
            function gameOver(){
                isPlaying=false;
                p.innerHTML="<span class='flip'>😵</span>💥";
                msg.style.display='flex';
                let segCode = (score * 5) + 10;
                msg.innerHTML=`<div style="font-size:35px; color:#ffdf00; margin-bottom:10px;">FIM DE JOGO!</div><div>Pontos: ${score}</div><div style='color:#ff4444; font-size:26px; margin: 15px 0;'>CÓDIGO SECRETO: <b>${segCode}</b></div><button onclick="startGame(event)">JOGAR NOVAMENTE</button>`;
            }
            </script></body></html>
            """
            components.html(html_pula, height=380)
            c_p1, c_p2, c_p3 = st.columns([1, 1, 1])
            pts_p = c_p1.number_input("🏃 Pontos Feitos", min_value=0, max_value=99999, key="p_pula")
            cod_p = c_p2.number_input("🔒 Código Secreto", min_value=0, max_value=999999, key="c_pula")
            if c_p3.button("Salvar Recorde", use_container_width=True, key="btn_pula"):
                if pts_p > 0 and cod_p == (pts_p * 5) + 10:
                    salvar_ranking(estudante_ra, estudante_nome, "pula_craque", pts_p, "Tropeçou na corrida")
                    st.success("✅ Código Válido! Salvo com sucesso no Mural!")
                else: st.error("❌ Código de Segurança Inválido ou Pontuação incorreta.")

        # ---------------- 3. GOLEIRO ----------------
        with jogo3:
            html_goleiro = """
            <!DOCTYPE html><html><head><style>
            @import url('https://fonts.googleapis.com/css2?family=Oswald:wght@600&display=swap'); 
            body { font-family: 'Oswald', sans-serif; text-align: center; margin: 0; background-color: transparent; } 
            #gameContainer { position:relative; width:100%; height:450px; overflow:hidden;} 
            .grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 15px; max-width: 350px; margin: 20px auto; } 
            .hole { height: 90px; background: rgba(0, 100, 0, 0.6); border-radius: 10px; display: flex; justify-content: center; align-items: center; font-size: 50px; cursor: pointer; box-shadow: inset 0 5px 15px rgba(0,0,0,0.5); border: 2px solid #004d23; transition: background 0.2s;} 
            .hole:active { background: rgba(0, 150, 0, 0.8); } 
            .hole.active::after { content: '⚽'; animation: pop 0.2s ease-out; } 
            @keyframes pop { from { transform: scale(0); } to { transform: scale(1); } } 
            .stats { font-size: 24px; color: #004d23; margin-top: 10px; } 
            button { font-family: 'Oswald', sans-serif; background: #d4af37; color: #002776; border: none; padding: 12px 25px; font-size: 20px; border-radius: 8px; cursor: pointer; font-weight: bold; box-shadow: 0 4px 6px rgba(0,0,0,0.2); transition: transform 0.1s;} 
            button:active { transform: scale(0.95); } 
            #endMsg { display:none; background:#fffdf0; border: 4px solid #d4af37; padding: 25px; border-radius:15px; font-size: 26px; color: #002776; margin: 40px auto; width: 85%; max-width: 350px; box-shadow: 0 10px 25px rgba(0,0,0,0.4); text-align:center;} 
            </style></head><body>
            <div id="gameContainer">
                <div class="stats" id="statsBar">DEFESAS: <span id="score">0</span> | TEMPO: <span id="time">15</span>s</div>
                <button id="startBtn" onclick="startGame()">INICIAR TREINO (15s)</button>
                <div class="grid" id="grid">
                    <div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div><div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div><div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div>
                    <div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div><div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div><div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div>
                    <div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div><div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div><div class="hole" onmousedown="defend(this)" ontouchstart="defend(this)"></div>
                </div>
                <div id="endMsg"></div>
            </div>
            <script>
                let score=0; let lastHole; let timeUp=false; let timeLeft=15; 
                const holes=document.querySelectorAll('.hole'); const scoreBoard=document.getElementById('score'); 
                const timeBoard=document.getElementById('time'); const btn=document.getElementById('startBtn'); 
                const msg=document.getElementById('endMsg'); const grid=document.getElementById('grid');
                const statsBar=document.getElementById('statsBar');
                
                function randomHole() { const idx=Math.floor(Math.random()*holes.length); const hole=holes[idx]; if (hole===lastHole) return randomHole(); lastHole=hole; return hole; }
                
                function showBall() { 
                    const time=Math.random()*(800-400)+400; 
                    const hole=randomHole(); hole.classList.add('active'); 
                    setTimeout(() => { hole.classList.remove('active'); if (!timeUp) showBall(); }, time); 
                }
                
                function startGame() { 
                    msg.style.display='none'; grid.style.display='grid'; statsBar.style.display='block';
                    scoreBoard.textContent=0; timeBoard.textContent=15; score=0; timeUp=false; timeLeft=15; btn.style.display='none'; 
                    showBall(); 
                    const countdown=setInterval(() => { 
                        timeLeft--; timeBoard.textContent=timeLeft; 
                        if(timeLeft<=0) { 
                            clearInterval(countdown); timeUp=true; 
                            let segCode = (score * 8) + 22; 
                            grid.style.display='none'; statsBar.style.display='none'; btn.style.display='none';
                            msg.style.display='block'; 
                            msg.innerHTML=`<div style="font-size:32px; color:#009c3b; margin-bottom:10px;">FIM DO TEMPO!</div><div>Defesas: <b>${score}</b></div><div style='color:#ff4444; font-size:26px; margin: 15px 0;'>CÓDIGO SECRETO:<br><b>${segCode}</b></div><button onclick="startGame()">JOGAR DE NOVO</button>`; 
                        } 
                    }, 1000); 
                }
                function defend(hole) { if(!hole.classList.contains('active')) return; score++; scoreBoard.textContent=score; hole.classList.remove('active'); }
            </script></body></html>
            """
            components.html(html_goleiro, height=450)
            c_g1, c_g2, c_g3 = st.columns([1, 1, 1])
            pts_g = c_g1.number_input("🧤 Número de Defesas", min_value=0, max_value=999, key="p_gol")
            cod_g = c_g2.number_input("🔒 Código Secreto", min_value=0, max_value=9999, key="c_gol")
            if c_g3.button("Salvar Recorde", use_container_width=True, key="btn_gol"):
                if pts_g > 0 and cod_g == (pts_g * 8) + 22:
                    salvar_ranking(estudante_ra, estudante_nome, "goleiro", pts_g, "Em 15 segundos")
                    st.success("✅ Salvo com sucesso no Mural!")
                else: st.error("❌ Código de Segurança Inválido ou Pontuação incorreta.")

        # ---------------- 4. FALTA PERFEITA (COM ALVO DINÂMICO CORRIGIDO) ----------------
        with jogo4:
            html_falta = """
            <!DOCTYPE html><html><head><style>@import url('https://fonts.googleapis.com/css2?family=Oswald:wght@600&display=swap'); body{margin:0;font-family:'Oswald',sans-serif;user-select:none; overflow:hidden;} #game{width:100%;height:300px;background:linear-gradient(#87CEEB 70%, #4CAF50 30%);position:relative;border-radius:10px;overflow:hidden;box-shadow:inset 0 0 10px rgba(0,0,0,0.3);} #goal{position:absolute;bottom:90px;left:50%;transform:translateX(-50%);width:200px;height:100px;border:5px solid white;border-bottom:none;} #targetZone{position:absolute;bottom:90px;left:50%;transform:translateX(-50%);width:60px;height:100px;background:rgba(0,255,0,0.3);} #target{position:absolute;bottom:130px;left:0;font-size:30px;} #ball{position:absolute;bottom:20px;left:50%;transform:translateX(-50%);font-size:40px;transition:bottom 0.3s ease-out;} button{position:absolute;top:20px;left:50%;transform:translateX(-50%);padding:10px 30px;font-size:20px;font-weight:bold;background:#ffdf00;border:none;border-radius:8px;cursor:pointer;font-family:'Oswald';z-index:20;} #msg{position:absolute;top:0;left:0;width:100%;height:100%;background:rgba(0,0,0,0.8);display:none;flex-direction:column;justify-content:center;align-items:center;font-size:30px;color:white;text-shadow:2px 2px 4px #000;font-weight:bold;z-index:15;} #kickFeedback{position:absolute;top:80px;width:100%;text-align:center;font-size:30px;color:white;text-shadow:2px 2px 4px #000;font-weight:bold;z-index:5;} </style>
            </head><body>
            <div id="game"><div id="goal"></div><div id="targetZone"></div><div id="target">🎯</div><div id="ball">⚽</div><button id="btn" onclick="shoot()">CHUTE 1/5</button><div id="kickFeedback"></div><div id="msg"></div></div>
            <script>
            let tx=0; let d=1; let speed=4; let moving=true; let shots=0; let tScore=0;
            const tg=document.getElementById('target'); const bl=document.getElementById('ball'); const btn=document.getElementById('btn'); const msg=document.getElementById('msg'); const kfb=document.getElementById('kickFeedback');
            const gameDiv=document.getElementById('game');
            function anim(){
                if(!moving)return; 
                let w = gameDiv.clientWidth || window.innerWidth || 350; 
                tx+=speed*d; 
                if(tx>w-40||tx<0)d*=-1; 
                tg.style.left=tx+'px'; 
                requestAnimationFrame(anim); 
            } anim();
            function shoot(){ 
                if(!moving)return; moving=false; bl.style.bottom='130px'; 
                setTimeout(()=>{ 
                    let w=gameDiv.clientWidth || window.innerWidth || 350; 
                    let center=w/2; let tCenter=tx+15; shots++;
                    if(Math.abs(tCenter-center)<40){ kfb.innerText="🎉 GOLAÇO!!! (+100)"; kfb.style.color="#00FF00"; tScore+=100; } else{ kfb.innerText="❌ NA TRAVE!"; kfb.style.color="red"; }
                    if(shots<5){ btn.innerText="CHUTE "+(shots+1)+"/5"; setTimeout(()=>{kfb.innerText=''; bl.style.bottom='20px'; moving=true; anim();}, 1500); } 
                    else{ let segCode = (tScore * 2) + 99; btn.innerText="JOGAR DE NOVO"; btn.setAttribute("onclick", "location.reload()"); msg.style.display="flex"; msg.innerHTML=`FIM DOS 5 CHUTES!<br><span style='color:#ffdf00;'>Total: ${tScore}</span><br><span style='color:red;'>CÓDIGO: ${segCode}</span>`; }
                },300); 
            }
            </script></body></html>
            """
            components.html(html_falta, height=350)
            c_f1, c_f2, c_f3 = st.columns([1, 1, 1])
            pts_f = c_f1.number_input("🎯 Pontos nos Chutes", min_value=0, max_value=9999, key="p_falta")
            cod_f = c_f2.number_input("🔒 Código Secreto", min_value=0, max_value=99999, key="c_falta")
            if c_f3.button("Salvar Recorde", use_container_width=True, key="btn_falta"):
                if pts_f > 0 and cod_f == (pts_f * 2) + 99:
                    salvar_ranking(estudante_ra, estudante_nome, "falta", pts_f, "Em 5 Chutes")
                    st.success("✅ Salvo com sucesso no Mural!")
                else: st.error("❌ Código Inválido ou Pontuação incorreta.")

        # ---------------- 5. EMBAIXADINHA ----------------
        with jogo5:
            html_embaixadinha = """
            <!DOCTYPE html><html><head><style>@import url('https://fonts.googleapis.com/css2?family=Oswald:wght@600&display=swap'); body { margin: 0; padding: 0; overflow: hidden; font-family: 'Oswald', sans-serif; user-select:none;} #gameArea { width: 100%; height: 350px; background: linear-gradient(to bottom, #87CEEB 0%, #87CEEB 70%, #228B22 70%, #228B22 100%); position: relative; border-radius: 12px; box-shadow: inset 0 0 20px rgba(0,0,0,0.2); cursor: pointer; } #score { position: absolute; top: 15px; left: 20px; font-size: 28px; color: white; text-shadow: 2px 2px 4px rgba(0,0,0,0.8); z-index: 10; } #ball { font-size: 70px; position: absolute; left: 50%; transform: translateX(-50%); cursor: pointer; z-index: 5; } #startBtn { position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); padding: 15px 30px; background: #ffdf00; color: #002776; border: none; border-radius: 8px; font-size: 24px; font-family: 'Oswald', sans-serif; cursor: pointer; box-shadow: 0 5px 15px rgba(0,0,0,0.3); z-index: 20; font-weight: bold;} #startBtn:active { transform: translate(-50%, -45%); box-shadow: 0 2px 5px rgba(0,0,0,0.3); } #msg { display:none; position:absolute; top:0; left:0; width:100%; height:100%; background:rgba(0,0,0,0.85); flex-direction:column; justify-content:center; align-items:center; color:white; font-size:32px; z-index:15; text-align:center; } </style>
            </head><body>
            <div id="gameArea"><div id="score">Embaixadinhas: 0</div><div id="ball" style="top: 50px;">⚽</div><button id="startBtn">COMEÇAR</button><div id="msg"></div></div>
            <script>
            const ball=document.getElementById('ball'); const scoreEl=document.getElementById('score'); const startBtn=document.getElementById('startBtn'); const msg=document.getElementById('msg'); let y=50; let vy=0; let gravity=0.4; let isPlaying=false; let score=0; let animId; let rotation=0;
            function update() { if(!isPlaying) return; vy+=gravity; y+=vy; rotation+=vy*2; ball.style.transform=`translateX(-50%) rotate(${rotation}deg)`; if(y>270) { isPlaying=false; ball.style.top='270px'; let segCode = (score * 6) + 33; msg.style.display='flex'; msg.innerHTML=`A BOLA CAIU!<br><span style='color:#ffdf00;'>Total: ${score}</span><br><span style='color:red;'>CÓDIGO: ${segCode}</span><br><button onclick="location.reload()" style="margin-top:10px;">JOGAR DE NOVO</button>`; return; } if(y<0) { y=0; vy=0; } ball.style.top=y+'px'; animId=requestAnimationFrame(update); }
            function kick(e) { if(e){e.preventDefault(); e.stopPropagation();} if(!isPlaying) return; vy=-8; score++; scoreEl.innerText='Embaixadinhas: '+score; }
            ball.addEventListener('mousedown', kick); ball.addEventListener('touchstart', kick, {passive: false});
            startBtn.onclick=function() { isPlaying=true; score=0; scoreEl.innerText='Embaixadinhas: 0'; y=50; vy=-5; rotation=0; startBtn.style.display='none'; msg.style.display='none'; cancelAnimationFrame(animId); update(); }
            </script></body></html>
            """
            components.html(html_embaixadinha, height=380)
            c_e1, c_e2, c_e3 = st.columns([1, 1, 1])
            pts_e = c_e1.number_input("⚽ Toques na Bola", min_value=0, max_value=9999, key="p_emb")
            cod_e = c_e2.number_input("🔒 Código Secreto", min_value=0, max_value=99999, key="c_emb")
            if c_e3.button("Salvar Recorde", use_container_width=True, key="btn_emb"):
                if pts_e > 0 and cod_e == (pts_e * 6) + 33:
                    salvar_ranking(estudante_ra, estudante_nome, "embaixadinha", pts_e, "Deixou a bola cair")
                    st.success("✅ Salvo com sucesso no Mural!")
                else: st.error("❌ Código Inválido ou Pontuação incorreta.")

    # =========================================================================
    # RANKING GERAL OFICIAL
    # =========================================================================
    with aba_ranking:
        st.markdown("<h2 style='text-align: center; font-family: Oswald; color: #d4af37; text-shadow: 1px 1px 2px #000;'>🏆 MURAL DE RECORDES 🏆</h2>", unsafe_allow_html=True)
        jogo_rank = st.selectbox("Escolha o jogo para ver o Ranking Oficial da Escola:", ["🏃 Pula Craque", "🧠 Memória", "🧤 Goleiro", "🎯 Falta", "⚽ Embaixadinha"])
        map_db = {"🏃 Pula Craque": "pula_craque", "🧠 Memória": "memoria", "🧤 Goleiro": "goleiro", "🎯 Falta": "falta", "⚽ Embaixadinha": "embaixadinha"}
        
        try:
            res_rank = supabase.table("ranking_jogos").select("*").eq("jogo", map_db[jogo_rank]).order("pontuacao", desc=True).limit(10).execute()
            if res_rank.data:
                df_rank = pd.DataFrame(res_rank.data)
                df_rank.index = df_rank.index + 1
                df_rank['Aluno'] = df_rank['nome_estudante']
                df_rank['Pontuação'] = df_rank['pontuacao']
                df_rank['Detalhes'] = df_rank['detalhes']
                st.dataframe(df_rank[['Aluno', 'Pontuação', 'Detalhes']], use_container_width=True)
            else:
                st.info("Nenhum recorde registrado ainda. Jogue para ser o Número 1!")
        except Exception as e:
            st.error("Crie a tabela 'ranking_jogos' no Supabase para visualizar o placar!")

    st.markdown('</div>', unsafe_allow_html=True)


# ==============================================================================
# GATILHOS FINAIS DO SISTEMA (FICAM NA ÚLTIMA LINHA DO SEU ARQUIVO)
# ==============================================================================
if st.session_state.get('authenticated'):
    
    # GATILHO 1: ALUNO LOGADO
    if st.session_state.get('modulo_atuacao') == "Álbum do Estudante":
        st.markdown("""
            <style>
                [data-testid="stSidebar"] { display: none !important; }
                [data-testid="collapsedControl"] { display: none !important; }
                .block-container { padding-top: 2rem !important; }
            </style>
        """, unsafe_allow_html=True)
        c_nome, c_sair = st.columns([3, 1])
        with c_nome:
            st.markdown(f"<h3 style='color: #004d23; font-family: Poppins; margin-bottom: 0;'>🎒 Olá, <b>{st.session_state.get('usuario_nome')}</b>!</h3>", unsafe_allow_html=True)
            st.markdown(f"<span style='color: #666; font-size: 0.9rem;'>📌 R.A.: {st.session_state.get('usuario_ra')}</span>", unsafe_allow_html=True)
        with c_sair:
            st.markdown("<div style='margin-top: 10px;'></div>", unsafe_allow_html=True)
            if st.button("🚪 Sair da Conta", type="primary", use_container_width=True):
                st.session_state.authenticated = False
                st.session_state.user_role = None
                st.session_state.usuario_ra = None
                st.session_state.usuario_nome = None
                st.session_state.modulo_atuacao = None
                st.rerun()
        st.write("---")
        render_modulo_album()
        st.stop()

    # GATILHO 2: DIRETOR
    elif st.session_state.get('modulo_atuacao') == "🏫 Ensino Regular":
        if 'app_mode_regular' in locals() and app_mode_regular == "🏆 Álbum de Figurinhas":
            render_modulo_album()







# ==============================================================================
# IMPORTANTE: Adicione as 3 linhas abaixo na primeira linha do seu arquivo,
# junto com os outros "import", caso ainda não as tenha.
# ==============================================================================
import cv2
import numpy as np
from PIL import Image
import io
import base64
import uuid
import time
from datetime import datetime
import pytesseract
from PIL import Image
import re  # Biblioteca para filtrar textos
from PIL import Image, ImageOps

# ... (Seu código existente até chegar no bloco do Módulo Administrativo) ...

# ==============================================================================
# MÓDULO: ADMINISTRATIVO (PATRIMÔNIO E INVENTÁRIO) - VERSÃO DESKTOP/WEB
# ==============================================================================
if app_mode_adm == "🏷️ Patrimônio e Inventário":

    st.markdown('<div class="header-box"><div class="header-title">🏷️ Gestão de Patrimônio</div></div>', unsafe_allow_html=True)
    st.write("") # Espaçamento
    
    col_voltar, _ = st.columns([2, 8])
    with col_voltar:
        if st.button("⬅️ Voltar ao Menu Inicial", key="voltar_patrimonio", use_container_width=True):
            st.session_state.modulo_atuacao = None
            st.rerun()

    # Leitura da tabela no Supabase
    df_patrimonio = safe_read("Patrimonio", ["id", "codigo", "nome", "localizacao", "estado", "conferido", "ultima_conferencia", "foto_base64", "observacao"])

    # Lista de locais organizada
    LOCAIS_ESCOLA = sorted([
        "ALMOXARIFADO", "ARQUIVO MORTO", "BANHEIROS ADM", "BANHEIROS ALUNOS",
        "BANHEIROS FUNC.", "BANHEIROS PROF", "BIBLIOTECA", "CORREDOR FUNDAMENTAL",
        "CORREDOR INFANTIL", "COZINHA DOS PROFESSORES", "COZINHA FUND.", "HALL DE ENTRADA",
        "INFORMÁTICA", "LAVANDERIA", "QUADRA", "REFEITÓRIO BERÇÁRIO", "REFEITÓRIO FUND",
        "REFEITÓRIO INFANTIL", "SALA 09", "SALA 1", "SALA 10", "SALA 11", "SALA 12",
        "SALA 13", "SALA 14", "SALA 15", "SALA 16", "SALA 17", "SALA 18", "SALA 2",
        "SALA 3", "SALA 4", "SALA 5", "SALA 6", "SALA 7", "SALA 8", "SALA BILINGUE",
        "SALA DA COORDENAÇÃO", "SALA DA DIREÇÃO", "SALA DE CAFÉ", "SALA DE ED. FISICA",
        "SALA DE RECURSOS 1", "SALA DOS PROFESSORES", "COZINHA CRECHE", "LACTARIO", "COZINHA FUNDAMENTAL", "BANHEIRO FUNDAMENTAL", "BANHEIRO INFANTIL", "BANHEIRO QUADRA", "HALL COORDENAÇÃO", "SALA RECURSOS 2", "SECRETARIA"
    ])

    if "local_trabalho_atual" not in st.session_state:
        st.session_state.local_trabalho_atual = LOCAIS_ESCOLA[0]

    # ==================================================================
    # --- SISTEMA DE NAVEGAÇÃO À PROVA DE BUGS (Substitui st.tabs) ---
    # ==================================================================
    menu_opcoes = [
        "📊 Visão Geral", 
        "✅ Conferência e Atualização", 
        "➕ Cadastro de Bens", 
        "🖼️ Galeria por Ambiente",
        "🖨️ Etiquetas e Relatórios"
    ]
    
    aba_selecionada = st.radio(
        "Navegação:", 
        menu_opcoes, 
        horizontal=True, 
        label_visibility="collapsed"
    )
    st.markdown("---") # Linha separadora para simular o design de abas

    # ==================================================================
    # --- TELA 0: VISÃO GERAL ---
    # ==================================================================
    if aba_selecionada == "📊 Visão Geral":
        st.header("Visão Geral do Inventário")
        
        if not df_patrimonio.empty:
            # 1. Cálculo das Métricas
            total_bens = len(df_patrimonio)
            
            # Conta quantos itens estão com status 'conferido' igual a True ou "True"
            # (Garante que funcione independentemente de como o Supabase/Pandas tipa o boolean)
            conferidos = len(df_patrimonio[df_patrimonio['conferido'].astype(bool) == True])
            pendentes = total_bens - conferidos
            
            # 2. Exibição dos Indicadores (Cards)
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(label="📦 Total de Bens", value=total_bens)
            with col2:
                st.metric(label="✅ Conferidos", value=conferidos)
            with col3:
                st.metric(label="⚠️ Pendentes", value=pendentes)
                
            st.markdown("---") # Linha divisória antes da tabela

            # 3. Exibição da Tabela
            # Esconde a coluna da imagem base64 na visualização da tabela para não poluir a tela
            df_display = df_patrimonio.drop(columns=['foto_base64'], errors='ignore')
            st.dataframe(df_display, use_container_width=True, hide_index=True)
        else:
            st.info("Nenhum bem cadastrado no banco de dados.")

    # ==================================================================
    # --- TELA 1: CONFERÊNCIA E ATUALIZAÇÃO ---
    # ==================================================================
    elif aba_selecionada == "✅ Conferência e Atualização":
        st.header("Conferência")
        
        col_busca, col_form = st.columns([4, 6], gap="large")
        
        with col_busca:
            st.markdown("#### 1. Configuração de Ambiente")
            st.session_state.local_trabalho_atual = st.selectbox(
                "Ambiente atual de trabalho:", 
                LOCAIS_ESCOLA,
                index=LOCAIS_ESCOLA.index(st.session_state.local_trabalho_atual),
                help="O local selecionado aqui será sugerido automaticamente ao conferir ou cadastrar bens."
            )
            
            st.markdown("---")
            st.markdown("#### 2. Localizar Bem")
            
            metodo_busca = st.selectbox("Método de Leitura:", [
                "Digitar / Leitor USB",
                "Câmera (QR Code)", 
                "Câmera (Ler Plaquinha - OCR)"
            ])
            
            codigo_busca = ""
            
            if metodo_busca == "Digitar / Leitor USB":
                st.info("💡 Conecte o leitor USB, clique no campo abaixo e bipe a etiqueta.")
                codigo_busca = st.text_input("Código do Bem (Tombo):", placeholder="Ex: 123456", key="busca_patrimonio")
                
            elif metodo_busca == "Câmera (QR Code)":
                st.info("💡 Aponte a webcam para o QR Code.")
                foto_qr = st.camera_input("📷 Webcam")
                
                if foto_qr:
                    try:
                        file_bytes = np.asarray(bytearray(foto_qr.read()), dtype=np.uint8)
                        img = cv2.imdecode(file_bytes, 1)
                        detector = cv2.QRCodeDetector()
                        data, bbox, _ = detector.detectAndDecode(img)
                        
                        if data:
                            codigo_busca = data
                            st.success(f"✅ QR Code lido: **{codigo_busca}**")
                        else:
                            st.error("❌ Não foi possível ler o QR Code.")
                    except Exception as e:
                        st.error(f"Erro ao processar a imagem: {e}")
                        
            elif metodo_busca == "Câmera (Ler Plaquinha - OCR)":
                st.info("💡 Centralize o número da placa na webcam.")
                foto_placa = st.camera_input("📷 Webcam")
                
                if foto_placa:
                    with st.spinner("Analisando imagem..."):
                        try:
                            img_placa = Image.open(foto_placa)
                            config_tesseract = r'--psm 6'
                            texto_bruto = pytesseract.image_to_string(img_placa, config=config_tesseract)
                            
                            numeros_encontrados = re.findall(r'\d+', texto_bruto)
                            
                            if numeros_encontrados:
                                texto_extraido = max(numeros_encontrados, key=len)
                                st.success("✅ Número detectado!")
                                codigo_busca = st.text_input("🔍 Confirme o código lido:", value=texto_extraido, key="confirma_ocr")
                            else:
                                st.error("❌ Nenhum número detectado.")
                        except Exception as e:
                            st.error(f"Erro no leitor de placa: {e}")

        with col_form:
            st.markdown("#### 3. Painel de Atualização")
            if codigo_busca:
                bem_encontrado = df_patrimonio[df_patrimonio['codigo'] == codigo_busca]
                
                if not bem_encontrado.empty:
                    bem = bem_encontrado.iloc[0]
                    st.success(f"Bem localizado: **{bem['nome']}** (Cód: {bem['codigo']})")
                    
                    with st.form(f"form_conf_{bem['id']}"):
                        
                        col_f1, col_f2 = st.columns(2)
                        
                        idx_loc_atual = LOCAIS_ESCOLA.index(st.session_state.local_trabalho_atual)
                        with col_f1:
                            nova_loc = st.selectbox("Localização Atual do Bem", LOCAIS_ESCOLA, index=idx_loc_atual)
                        
                        with col_f2:
                            opcoes_estado = ["Ótimo", "Novo", "Bom", "Regular", "Ruim", "Inservível/Sucata", "Em Manutenção"]
                            idx_est = opcoes_estado.index(bem.get('estado', 'Bom')) if bem.get('estado') in opcoes_estado else 1
                            novo_estado = st.selectbox("Estado de Conservação", opcoes_estado, index=idx_est)
                        
                        obs = st.text_area("Observações (Avarias, Transferências, etc.)", value=bem.get('observacao', ''), height=100)
                        
                        st.markdown("**Registro Fotográfico**")
                        col_img_1, col_img_2 = st.columns([1, 1])
                        
                        with col_img_1:
                            if bem.get('foto_base64') and pd.notnull(bem['foto_base64']) and bem['foto_base64'] != "":
                                try:
                                    st.image(base64.b64decode(bem['foto_base64']), caption="Foto Atual", use_container_width=True)
                                except:
                                    st.warning("Erro ao carregar imagem salva.")
                            else:
                                st.info("📸 Sem foto registrada.")
                                
                        with col_img_2:
                            metodo_foto = st.radio("Como inserir a foto?", ["📸 Câmera Direta", "📁 Galeria/Arquivo"], horizontal=True, key=f"metodo_foto_{bem['id']}")
                            
                            nova_foto = None
                            if metodo_foto == "📸 Câmera Direta":
                                nova_foto = st.camera_input("Tire a foto", key=f"cam_{bem['id']}")
                            else:
                                nova_foto = st.file_uploader("Substituir / Inserir Foto", type=["jpg", "png", "jpeg"], key=f"up_{bem['id']}")
                        
                        st.markdown("---")
                        conferido = st.checkbox("✅ Marcar como CONFERIDO E VISTADO nesta data", value=True)
                        
                        col_submit, _ = st.columns([4, 6])
                        with col_submit:
                            submit_btn = st.form_submit_button("💾 Salvar Conferência", type="primary", use_container_width=True)
                        
                        if submit_btn:
                            dados_update = {
                                "localizacao": nova_loc,
                                "estado": novo_estado,
                                "observacao": obs,
                                "conferido": conferido,
                                "ultima_conferencia": datetime.now().strftime("%d/%m/%Y %H:%M")
                            }
                            
                            if nova_foto:
                                try:
                                    img = Image.open(nova_foto)
                                    if img.mode != 'RGB': img = img.convert('RGB')
                                    img.thumbnail((500, 500))
                                    buf = io.BytesIO()
                                    img.save(buf, format="JPEG", quality=80)
                                    dados_update['foto_base64'] = base64.b64encode(buf.getvalue()).decode()
                                except Exception as e:
                                    st.error(f"Erro ao processar a foto: {e}")
                                    
                            try:
                                supabase.table("Patrimonio").update(dados_update).eq("id", bem['id']).execute()
                                st.success("✅ Inventário atualizado com sucesso!")
                                time.sleep(1)
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erro ao salvar no banco: {e}")
                else:
                    st.warning("⚠️ Código não encontrado no sistema. Vá para a aba 'Cadastro de Bens'.")
            else:
                st.info("Aguardando inserção ou leitura do código.")

    # ==================================================================
    # --- TELA 2: CADASTRO DE BENS ---
    # ==================================================================
    elif aba_selecionada == "➕ Cadastro de Bens":
        st.header("Cadastro de Bens")
        st.markdown("Preencha o formulário abaixo para registrar um novo bem no sistema.")
        
        with st.form("form_novo_patrimonio", clear_on_submit=True):
            col1, col2 = st.columns(2)
            
            with col1:
                n_codigo = st.text_input("Código / Tombo *", placeholder="Ex: 98765")
                idx_loc_atual = LOCAIS_ESCOLA.index(st.session_state.local_trabalho_atual)
                n_local = st.selectbox("Localização", LOCAIS_ESCOLA, index=idx_loc_atual)
                
            with col2:
                n_nome = st.text_input("Nome/Descrição do Bem *", placeholder="Ex: Mesa de Escritório MDF")
                n_estado = st.selectbox("Estado de Conservação", ["Novo", "Bom", "Regular", "Ruim", "Inservível/Sucata"])
                
            n_foto = st.file_uploader("Foto Inicial do Bem (Opcional)", type=["jpg", "png", "jpeg"])
            
            st.markdown("---")
            col_submit, _ = st.columns([3, 7])
            with col_submit:
                btn_cadastrar = st.form_submit_button("➕ Cadastrar Bem", type="primary", use_container_width=True)
            
            if btn_cadastrar:
                if n_codigo and n_nome:
                    foto_b64 = ""
                    if n_foto:
                        try:
                            img = Image.open(n_foto)
                            if img.mode != 'RGB': img = img.convert('RGB')
                            img.thumbnail((500, 500))
                            buf = io.BytesIO()
                            img.save(buf, format="JPEG", quality=80)
                            foto_b64 = base64.b64encode(buf.getvalue()).decode()
                        except Exception as e:
                            st.error(f"Erro ao processar imagem: {e}")

                    novo_item = {
                        "id": str(uuid.uuid4()),
                        "codigo": n_codigo,
                        "nome": n_nome,
                        "localizacao": n_local,
                        "estado": n_estado,
                        "conferido": False,
                        "foto_base64": foto_b64,
                        "ultima_conferencia": "Nunca conferido",
                        "observacao": ""
                    }
                    try:
                        supabase.table("Patrimonio").insert(novo_item).execute()
                        st.success("✅ Bem cadastrado com sucesso!")
                        time.sleep(1)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao cadastrar: Verifique se este tombo já não existe. Erro: {e}")
                else:
                    st.error("Preencha os campos obrigatórios (Código e Nome).")

    # ==================================================================
    # --- TELA: GALERIA DE BENS POR AMBIENTE (CARÔMETRO PATRIMONIAL) ---
    # ==================================================================
    elif aba_selecionada == "🖼️ Galeria por Ambiente":
        st.header("Galeria de Bens por Ambiente")
        st.markdown("Selecione um ambiente para visualizar visualmente todos os bens, estado de conservação e observações.")
        
        if df_patrimonio.empty:
            st.warning("⚠️ Nenhum bem cadastrado no banco de dados.")
        else:
            # Puxa apenas os locais que realmente possuem bens cadastrados
            locais_com_bens = sorted([loc for loc in df_patrimonio['localizacao'].unique() if pd.notnull(loc) and str(loc).strip() != ""])
            
            if not locais_com_bens:
                st.info("Nenhum ambiente possui bens cadastrados com localização.")
            else:
                col_filtro, _ = st.columns([4, 6])
                with col_filtro:
                    local_escolhido = st.selectbox("Selecione o Ambiente:", ["-- Escolha o Ambiente --"] + locais_com_bens)

                if local_escolhido != "-- Escolha o Ambiente --":
                    df_filtrado = df_patrimonio[df_patrimonio['localizacao'] == local_escolhido]
                    
                    st.markdown(f"**Total de itens neste ambiente:** {len(df_filtrado)}")
                    st.divider()

                    # Estilização adaptada do seu carômetro
                    st.markdown("""
                        <style>
                        .pat-foto-frame { height: 160px; width: 100%; display: flex; align-items: center; justify-content: center; overflow: hidden; border-radius: 8px; background-color: #f8fafc; margin: 10px 0; border: 1px dashed #cbd5e1; }
                        .pat-nome { font-weight: 800; color: #1e3a8a; font-size: 11px; min-height: 35px; display: flex; align-items: center; justify-content: center; text-align: center; text-transform: uppercase; line-height: 1.1; }
                        .pat-info { font-size: 10px; color: #64748b; line-height: 1.2; text-align: center; min-height: 45px; overflow-y: auto; padding: 2px; }
                        </style>
                    """, unsafe_allow_html=True)

                    cols = st.columns(5)
                    idx_col = 0

                    for _, bem in df_filtrado.iterrows():
                        codigo_bem = bem.get("codigo", "-")
                        nome_bem = bem.get("nome", "Sem descrição")
                        estado = bem.get("estado", "Não informado")
                        obs = bem.get("observacao", "")
                        foto_b64 = bem.get("foto_base64", None)

                        with cols[idx_col]:
                            with st.container(border=True):
                                # Título (Código + Nome)
                                st.markdown(f'<div class="pat-nome">{codigo_bem}<br>{nome_bem}</div>', unsafe_allow_html=True)
                                
                                # Processamento da Imagem
                                if foto_b64 and pd.notnull(foto_b64) and str(foto_b64).strip() != "":
                                    # Limpa possíveis cabeçalhos de data URI para evitar quebra
                                    clean_b64 = str(foto_b64).replace("data:image/jpeg;base64,", "").replace("data:image/png;base64,", "")
                                    img_html = f"<img src='data:image/jpeg;base64,{clean_b64}' style='width: 100%; height: 100%; object-fit: cover;'>"
                                else:
                                    img_html = "<div style='font-size: 40px; opacity: 0.2;'>📦</div>"
                                
                                st.markdown(f'<div class="pat-foto-frame">{img_html}</div>', unsafe_allow_html=True)
                                
                                # Informações Adicionais
                                obs_formatada = f"<br><b style='color:#b91c1c;'>Obs:</b> {obs}" if pd.notnull(obs) and obs.strip() != "" else ""
                                st.markdown(f'<div class="pat-info"><b>Estado:</b> {estado}{obs_formatada}</div>', unsafe_allow_html=True)

                        idx_col = (idx_col + 1) % 5
                        
    # ==================================================================
    # --- TELA 3: IMPRESSÕES E RELATÓRIOS ---
    # ==================================================================
    elif aba_selecionada == "🖨️ Etiquetas e Relatórios":
        st.header("Etiquetas e Relatórios")
        
        # Mudei a sub-aba para radio também para garantir que não vai bugar!
        sub_aba = st.radio("Selecione o que deseja gerar:", [
            "🏷️ Geração de Etiquetas (QR Code)", 
            "📑 Relatórios Oficiais (Padrão CONAM)",
            "📸 Relatório Fotográfico por Setor"
        ], horizontal=True)
        st.write("")
        
        if not df_patrimonio.empty:
            
            if sub_aba == "🏷️ Geração de Etiquetas (QR Code)":
                col_eti_1, col_eti_2 = st.columns([1, 1], gap="large")
                
                with col_eti_1:
                    st.markdown("#### Seleção de Itens")
                    modo_impressao = st.radio("Selecione os dados para gerar a folha de etiquetas:", ["📍 Por Ambiente", "✍️ Seleção Manual"])
                    df_para_imprimir = pd.DataFrame()
                    
                    if modo_impressao == "📍 Por Ambiente":
                        locais_disponiveis = [loc for loc in df_patrimonio['localizacao'].unique() if pd.notnull(loc) and str(loc).strip() != ""]
                        if locais_disponiveis:
                            local_sel = st.selectbox("Selecione o Ambiente:", ["-- Escolha o Ambiente --"] + sorted(locais_disponiveis))
                            if local_sel != "-- Escolha o Ambiente --":
                                df_para_imprimir = df_patrimonio[df_patrimonio['localizacao'] == local_sel]
                                st.info(f"Foram encontrados **{len(df_para_imprimir)}** bens neste ambiente.")
                        else:
                            st.warning("Nenhum ambiente foi registrado ainda.")
                    else:
                        opcoes_bens = df_patrimonio.apply(lambda row: f"{row['codigo']} - {row['nome']}", axis=1).tolist()
                        bens_selecionados = st.multiselect("Pesquise e selecione os bens desejados:", opcoes_bens)
                        if bens_selecionados:
                            codigos_selecionados = [item.split(" - ")[0] for item in bens_selecionados]
                            df_para_imprimir = df_patrimonio[df_patrimonio['codigo'].isin(codigos_selecionados)]

                with col_eti_2:
                    st.markdown("#### Geração do PDF")
                    if not df_para_imprimir.empty:
                        def gerar_pdf_etiquetas(df_bens):
                            pdf = FPDF('P', 'mm', 'A4')
                            pdf.set_auto_page_break(auto=False)
                            pdf.add_page()
                            
                            cols, rows = 4, 5
                            margin_x, margin_y = 10, 15
                            label_w, label_h = 45, 52
                            space_x, space_y = 3, 3
                            
                            idx = 0
                            progress_text = "Gerando QR Codes e montando o PDF..."
                            my_bar = st.progress(0, text=progress_text)
                            total_itens = len(df_bens)
                            
                            for i, (_, bem) in enumerate(df_bens.iterrows()):
                                if idx >= cols * rows:
                                    pdf.add_page()
                                    idx = 0
                                    
                                col = idx % cols
                                row = idx // cols
                                
                                x = margin_x + col * (label_w + space_x)
                                y = margin_y + row * (label_h + space_y)
                                
                                pdf.set_draw_color(200, 200, 200)
                                pdf.rect(x, y, label_w, label_h)
                                
                                pdf.set_xy(x, y + 3)
                                pdf.set_font("Arial", "B", 6)
                                pdf.set_text_color(0, 0, 0)
                                pdf.cell(label_w, 3, "CEIEF RAFAEL AFFONSO LEITE", align='C')
                                
                                pdf.set_xy(x, y + 6)
                                pdf.set_font("Arial", "", 5)
                                pdf.cell(label_w, 3, "INVENTÁRIO / PATRIMÔNIO", align='C')
                                
                                qr_size = 26
                                qr_x = x + (label_w - qr_size) / 2
                                qr_y = y + 11
                                
                                qr_url = f"https://api.qrserver.com/v1/create-qr-code/?size=150x150&data={bem['codigo']}"
                                
                                try:
                                    req = urllib.request.Request(qr_url, headers={'User-Agent': 'Mozilla/5.0'})
                                    with urllib.request.urlopen(req) as response:
                                        img_data = response.read()
                                    
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                                        tmp_file.write(img_data)
                                        tmp_path = tmp_file.name
                                    
                                    pdf.image(tmp_path, x=qr_x, y=qr_y, w=qr_size, h=qr_size)
                                    os.unlink(tmp_path)
                                except Exception:
                                    pdf.set_xy(x, qr_y + 10)
                                    pdf.set_font("Arial", "B", 8)
                                    pdf.cell(label_w, 5, "[ERRO QR]", align='C')
                                
                                pdf.set_xy(x, qr_y + qr_size + 2)
                                pdf.set_font("Arial", "B", 9)
                                pdf.cell(label_w, 5, str(bem['codigo']), align='C')
                                
                                nome = clean_pdf_text(str(bem['nome']))
                                if len(nome) > 35: nome = nome[:32] + "..."
                                    
                                pdf.set_xy(x + 2, qr_y + qr_size + 7)
                                pdf.set_font("Arial", "", 6)
                                pdf.multi_cell(label_w - 4, 3, nome, align='C')
                                
                                idx += 1
                                my_bar.progress((i + 1) / total_itens, text=progress_text)
                            
                            my_bar.empty()
                            return get_pdf_bytes(pdf)

                        st.write(f"Você selecionou **{len(df_para_imprimir)}** itens para impressão.")
                        if st.button("⚙️ Gerar Arquivo PDF", type="primary", use_container_width=True):
                            with st.spinner("Gerando etiquetas..."):
                                st.session_state.pdf_etiquetas = gerar_pdf_etiquetas(df_para_imprimir)
                        
                        if 'pdf_etiquetas' in st.session_state:
                            st.success("PDF gerado com sucesso!")
                            st.download_button(
                                label="📥 Baixar Etiquetas (PDF)", 
                                data=st.session_state.pdf_etiquetas, 
                                file_name=f"Etiquetas_Patrimonio_{datetime.now().strftime('%d%m%Y_%H%M')}.pdf", 
                                mime="application/pdf", 
                                use_container_width=True
                            )
                    else:
                        st.write("Aguardando seleção de itens.")

            elif sub_aba == "📑 Relatórios Oficiais (Padrão CONAM)":
                col_rel_1, col_rel_2 = st.columns([1, 1], gap="large")
                
                with col_rel_1:
                    st.markdown("#### Configuração do Relatório Matricial")
                    st.write("Listagem oficial para assinatura de responsabilidade patrimonial.")
                    
                    locais_rel = [loc for loc in df_patrimonio['localizacao'].unique() if pd.notnull(loc) and str(loc).strip() != ""]
                    local_rel = st.selectbox("Filtrar por Setor/Ambiente:", ["Todos os Ambientes"] + sorted(locais_rel))
                    
                    resp_padrao = st.session_state.get('usuario_nome', '').upper()
                    resp_rel = st.text_input("Responsável pelo Setor (Assinatura):", value=resp_padrao)

                    nivel_detalhe_obs = st.selectbox(
                        "Conteúdo da coluna OBSERVAÇÃO:",
                        [
                            "Somente Estado", 
                            "Estado + Local", 
                            "Estado + Observações do Bem",
                            "Estado + Local + Observações do Bem"
                        ]
                    )

                    class RelatorioConamPDF(FPDF):
                        def __init__(self, setor_nome, responsavel):
                            super().__init__('L', 'mm', 'A4')
                            self.set_margins(10, 10, 10)
                            self.set_auto_page_break(auto=False)
                            self.setor_nome = setor_nome
                            self.responsavel = responsavel
                            self.data_hoje = datetime.now().strftime("%d/%m/%Y")

                        def header(self):
                            self.set_font('Courier', '', 9)
                            self.set_text_color(0, 0, 0)
                            self.set_draw_color(0, 0, 0)
                            
                            self.cell(138, 4, "CN-SIP", 0, 0, 'L')
                            self.cell(138, 4, "CONAM", 0, 1, 'R')

                            self.cell(0, 4, "Prefeitura Municipal de Limeira", 0, 1, 'C')
                            self.cell(0, 4, "Relatorio de Moveis por Setor 01671 SMED - CEIEF RAFAEL AFFONSO LEITE", 0, 1, 'C')
                            self.cell(0, 4, f"Selecao : Bens INCORPORADOS             de             a {self.data_hoje}", 0, 1, 'C')

                            self.cell(138, 4, f"DATA {self.data_hoje}", 0, 0, 'L')
                            self.cell(138, 4, f"PAGINA      {self.page_no()}", 0, 1, 'R')

                            self.line(10, self.get_y(), 287, self.get_y())
                            self.ln(2)

                            self.set_x(70)
                            setor_str = f"01671 SMED - CEIEF RAFAEL AFFONSO LEITE - {self.setor_nome}" if self.setor_nome != "Todos os Ambientes" else "01671 SMED - CEIEF RAFAEL AFFONSO LEITE"
                            self.cell(0, 4, f"Setor       : {setor_str}", 0, 1, 'L')
                            
                            self.set_x(70)
                            self.cell(0, 4, f"Responsavel : {self.responsavel.upper()}", 0, 1, 'L')
                            self.ln(2)

                            self.line(10, self.get_y(), 287, self.get_y())
                            self.ln(1)

                            self.set_x(12)
                            self.cell(30, 4, "No. CHAPA", 0, 0, 'L')
                            self.cell(83, 4, "DESCRICAO", 0, 0, 'L')
                            self.set_x(155)
                            self.cell(0, 4, "OBSERVACAO", 0, 1, 'L')

                            self.line(10, self.get_y(), 287, self.get_y())
                            self.ln(1)
                            self.line(10, self.get_y(), 287, self.get_y())

                with col_rel_2:
                    st.markdown("#### Gerar Listagem")
                    if st.button("⚙️ Processar Relatório CONAM", type="primary", use_container_width=True):
                        df_rel = df_patrimonio.copy()
                        
                        df_rel['codigo_num'] = pd.to_numeric(df_rel['codigo'], errors='coerce')
                        df_rel = df_rel.sort_values(by=['codigo_num', 'codigo'])
                        
                        if local_rel != "Todos os Ambientes":
                            df_rel = df_rel[df_rel['localizacao'] == local_rel]
                        
                        if df_rel.empty:
                            st.warning("Nenhum bem encontrado para os filtros informados.")
                        else:
                            with st.spinner("Desenhando relatório matricial..."):
                                pdf = RelatorioConamPDF(local_rel, resp_rel)
                                pdf.add_page()
                                
                                pdf.set_font('Courier', '', 8)
                                
                                for i, row in df_rel.iterrows():
                                    # NOVA LÓGICA DE QUEBRA DE PÁGINA: baseada no espaço restante e não no número de linhas
                                    if pdf.get_y() > 182:
                                        pdf.line(10, pdf.get_y(), 287, pdf.get_y())
                                        pdf.add_page()
                                        pdf.set_font('Courier', '', 8)
                                        
                                    y = pdf.get_y()
                                    
                                    # 1. Puxa os dados
                                    chapa = str(row['codigo']).zfill(10)
                                    desc = clean_pdf_text(str(row['nome'])).upper()
                                    
                                    estado = str(row.get('estado', '')).strip()
                                    if estado in ['None', 'nan']: estado = ""
                                    
                                    loc = str(row.get('localizacao', '')).strip()
                                    if loc in ['None', 'nan']: loc = ""
                                    
                                    obs_banco = str(row.get('observacao', '')).strip()
                                    if obs_banco in ['None', 'nan']: obs_banco = ""
                                    
                                    # 2. Monta as partes conforme a escolha do SelectBox
                                    obs_parts = []
                                    if estado: obs_parts.append(estado)
                                    if nivel_detalhe_obs in ["Estado + Local", "Estado + Local + Observações do Bem"] and loc:
                                        obs_parts.append(loc)
                                    if nivel_detalhe_obs in ["Estado + Observações do Bem", "Estado + Local + Observações do Bem"] and obs_banco:
                                        obs_parts.append(obs_banco)
                                    
                                    # Tira as reticências e limite de caracteres! Deixa o texto inteiro.
                                    obs_text = clean_pdf_text(" - ".join(obs_parts)).upper()
                                    
                                    # 3. Escreve os textos primeiro para o sistema calcular a altura da linha
                                    pdf.set_xy(11, y + 0.2)
                                    pdf.cell(25, 3, chapa, 0, 0)
                                    
                                    # multi_cell permite que a descrição quebre em várias linhas se for enorme
                                    pdf.set_xy(38, y + 0.2)
                                    pdf.multi_cell(85, 3, desc, 0, 'L')
                                    y_desc = pdf.get_y()
                                    
                                    # multi_cell permite que a observação quebre linha, usando um espaço bem mais largo (158mm)
                                    pdf.set_xy(127, y + 0.2)
                                    pdf.multi_cell(158, 3, obs_text, 0, 'L')
                                    y_obs = pdf.get_y()
                                    
                                    # 4. Calcula a nova altura dinâmica desta linha da tabela
                                    row_height = max(y_desc, y_obs) - y + 0.4
                                    if row_height < 3.4: 
                                        row_height = 3.4 # Garante altura mínima
                                    
                                    # 5. Desenha as divisórias verticais acompanhando a altura dinâmica do texto
                                    pdf.line(10, y, 10, y + row_height)
                                    pdf.line(125, y, 125, y + row_height)
                                    # pdf.line(195, y, 195, y + row_height) -> APAGUEI ESSA LINHA PARA DAR MAIS ESPAÇO PARA O TEXTO
                                    pdf.line(287, y, 287, y + row_height)
                                    
                                    # Desce a "caneta" para a próxima linha
                                    pdf.set_y(y + row_height)
                                    
                                # Desenha a linha horizontal final de fechamento da tabela
                                pdf.line(10, pdf.get_y(), 287, pdf.get_y())
                                
                                if pdf.get_y() > 185:
                                    pdf.add_page()
                                    
                                total_moveis = len(df_rel)
                                setores_unicos = df_rel['localizacao'].dropna().apply(lambda x: str(x).strip()).unique()
                                setores_unicos = [s for s in setores_unicos if s not in ["", "None", "nan"]]
                                total_setores = len(setores_unicos) if len(setores_unicos) > 0 else 1
                                
                                str_moveis = str(total_moveis).rjust(4)
                                str_setores = str(total_setores).zfill(4)
                                
                                pdf.set_font('Courier', '', 9)
                                pdf.ln(2)
                                pdf.cell(0, 4, f"    {str_moveis} Movel(is)  deste Setor mostrado(s).", 0, 1, 'L')
                                pdf.ln(1)
                                
                                pdf.line(10, pdf.get_y(), 287, pdf.get_y())
                                pdf.ln(1)
                                
                                pdf.cell(0, 4, f"    {str_setores} Setor(es) mostrado(s).", 0, 1, 'L')
                                pdf.ln(1)
                                
                                pdf.line(10, pdf.get_y(), 287, pdf.get_y())
                                
                                st.session_state.pdf_conam = get_pdf_bytes(pdf)
                    
                    if 'pdf_conam' in st.session_state:
                        st.success("Relatório processado!")
                        st.download_button(
                            label="📥 Baixar Relatório (PDF)", 
                            data=st.session_state.pdf_conam, 
                            file_name=f"Relatorio_CONAM_{datetime.now().strftime('%Y%m%d')}.pdf", 
                            mime="application/pdf", 
                            use_container_width=True
                        )

            # ==================================================================
            # --- RELATÓRIO FOTOGRÁFICO, ANALÍTICO E LEGAL ---
            # ==================================================================
            elif sub_aba == "📸 Relatório Fotográfico por Setor":
                st.markdown("#### Filtros do Catálogo Visual e Analítico")
                st.write("Gera um relatório completo contendo a análise do acervo, a legislação municipal vigente e as fotos detalhadas.")
                
                locais_foto = [loc for loc in df_patrimonio['localizacao'].unique() if pd.notnull(loc) and str(loc).strip() != ""]
                local_foto_selecionado = st.selectbox("Selecione o Ambiente para Análise:", ["Todos os Ambientes"] + sorted(locais_foto), key="sel_foto")
        
                # Filtra os dados conforme a seleção
                df_foto = df_patrimonio.copy()
                df_foto['codigo_num'] = pd.to_numeric(df_foto['codigo'], errors='coerce')
                df_foto = df_foto.sort_values(by=['localizacao', 'codigo_num', 'codigo'])
                
                if local_foto_selecionado != "Todos os Ambientes":
                    df_foto = df_foto[df_foto['localizacao'] == local_foto_selecionado]
        
                if df_foto.empty:
                    st.warning("Nenhum bem encontrado para este ambiente.")
                else:
                    # ==========================================================
                    # DASHBOARD INTERATIVO NA TELA (STREAMLIT)
                    # ==========================================================
                    st.markdown("---")
                    st.markdown(f"### 📊 Panorama do Acervo: {local_foto_selecionado}")
                    
                    col_met1, col_met2, col_met3, col_met4 = st.columns(4)
                    col_met1.metric("Total de Bens Listados", len(df_foto))
                    col_met2.metric("Bens em Ótimo/Bom/Novo Estado", len(df_foto[df_foto['estado'].isin(['Ótimo', 'Bom', 'Novo'])]))
                    col_met3.metric("Bens em Regular Estado", len(df_foto[df_foto['estado'].isin(['Regular'])]))
                    col_met4.metric("Bens Inservíveis/Ruins", len(df_foto[df_foto['estado'].isin(['Ruim', 'Inservível/Sucata'])]))
        
                    col_graf1, col_graf2 = st.columns(2)
                    with col_graf1:
                        st.markdown("**Condição dos Bens**")
                        estado_counts = df_foto['estado'].value_counts()
                        st.bar_chart(estado_counts, color="#3498db")
                        
                    with col_graf2:
                        st.markdown("**Distribuição por Setor**")
                        if local_foto_selecionado == "Todos os Ambientes":
                            setor_counts = df_foto['localizacao'].value_counts().head(10) # Top 10 para não poluir
                            st.bar_chart(setor_counts, color="#2ecc71")
                        else:
                            st.info("Gráfico de setores oculto pois a busca é específica de um único ambiente.")
        
                    st.markdown("---")
                    
                    # ==========================================================
                    # GERAÇÃO DO PDF
                    # ==========================================================
                    st.markdown("#### Geração do Documento Oficial")
                    if st.button("⚙️ Processar Relatório e Catálogo (PDF)", type="primary", use_container_width=True):
                        with st.spinner("Compilando estatísticas, legislação e imagens..."):
                            pdf = FPDF('P', 'mm', 'A4')
                            pdf.set_auto_page_break(auto=True, margin=15)
                            
                            # ---------------------------------------------------------
                            # 1. DESIGN DA CAPA
                            # ---------------------------------------------------------
                            pdf.add_page()
                            pdf.set_fill_color(33, 47, 61) 
                            pdf.rect(0, 0, 210, 297, 'F')
                            
                            pdf.set_fill_color(52, 152, 219)
                            pdf.rect(0, 140, 210, 3, 'F')
                            
                            pdf.set_text_color(255, 255, 255)
                            pdf.set_y(90)
                            pdf.set_font("Arial", "B", 26)
                            pdf.cell(0, 10, "RELATÓRIO PATRIMONIAL".encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
                            pdf.set_font("Arial", "", 16)
                            pdf.cell(0, 10, "ANÁLISE, LEGISLAÇÃO E CATÁLOGO".encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
                            
                            pdf.set_y(160)
                            pdf.set_font("Arial", "B", 14)
                            pdf.cell(0, 8, "CEIEF RAFAEL AFFONSO LEITE - LIMEIRA/SP", ln=True, align='C')
                            pdf.set_font("Arial", "", 12)
                            
                            titulo_local = local_foto_selecionado if local_foto_selecionado != "Todos os Ambientes" else "Geral (Todos os Ambientes)"
                            pdf.cell(0, 8, f"Ambiente: {titulo_local}".encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
                            pdf.cell(0, 8, f"Data de Emissão: {datetime.now().strftime('%d/%m/%Y')}", ln=True, align='C')
        
                            # ---------------------------------------------------------
                            # 2. PÁGINA DE ANÁLISE E ESTATÍSTICAS
                            # ---------------------------------------------------------
                            pdf.add_page()
                            pdf.set_text_color(0, 0, 0)
                            pdf.set_font("Arial", "B", 16)
                            pdf.cell(0, 10, "1. RESUMO ANALÍTICO DO ACERVO".encode('latin-1', 'replace').decode('latin-1'), ln=True, align='L')
                            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                            pdf.ln(5)
        
                            pdf.set_font("Arial", "", 12)
                            pdf.cell(0, 8, f"Total de bens listados neste relatório: {len(df_foto)}", ln=True)
                            pdf.ln(5)
        
                            # Tabela de Condições
                            pdf.set_font("Arial", "B", 12)
                            pdf.cell(0, 8, "Estado de Conservação dos Bens:".encode('latin-1', 'replace').decode('latin-1'), ln=True)
                            pdf.set_font("Arial", "", 11)
                            
                            for estado, qtd in estado_counts.items():
                                estado_clean = str(estado).encode('latin-1', 'replace').decode('latin-1')
                                pdf.cell(10, 6, "-", ln=0, align='C')
                                pdf.cell(80, 6, f"{estado_clean}:", ln=0)
                                pdf.cell(20, 6, f"{qtd} item(ns)", ln=True)
                            
                            pdf.ln(10)
        
                            # ---------------------------------------------------------
                            # 3. PÁGINA DE LEGISLAÇÃO (IN 04/2021)
                            # ---------------------------------------------------------
                            pdf.add_page()
                            pdf.set_font("Arial", "B", 16)
                            pdf.cell(0, 10, "2. BASE LEGAL E RESPONSABILIDADE".encode('latin-1', 'replace').decode('latin-1'), ln=True, align='L')
                            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                            pdf.ln(5)
        
                            pdf.set_font("Arial", "B", 10)
                            pdf.cell(0, 6, "PREFEITURA MUNICIPAL DE LIMEIRA - SECRETARIA MUNICIPAL DE ADMINISTRAÇÃO".encode('latin-1', 'replace').decode('latin-1'), ln=True)
                            pdf.cell(0, 6, "INSTRUÇÃO NORMATIVA N° 04/2021".encode('latin-1', 'replace').decode('latin-1'), ln=True)
                            pdf.ln(3)
        
                            texto_legislacao = """
        "Disciplina os procedimentos em casos de roubo, furto ou danos aos bens patrimoniais das Unidades da Administração Municipal e estabelece as responsabilidades dos Gestores."
        
        Art. 2°. A gestão dos bens patrimoniais alocados em cada Secretaria Municipal e seus Departamentos, Divisões e Setores é de responsabilidade exclusiva do detentor de carga patrimonial.
        Parágrafo único: Todo servidor público poderá ser responsabilizado administrativa e juridicamente pelo desaparecimento ou dano ao material que lhe for confiado para guarda e/ou uso, bem como pelo dano que, dolosa ou culposamente, causar a qualquer material.
        
        Art. 4°. Conceitos fundamentais:
        III. Extravio: É o desaparecimento de bens por furto, roubo ou por negligência do responsável pela guarda.
        VI. Avaria: Significa estrago ou danificação total ou parcial de bem, em virtude de mau uso ou sinistro.
        VIII. Detentor de Carga Patrimonial: servidor designado responsável pelo uso, guarda e conservação de bem patrimonial.
        
        Art. 5°. É dever dos utilizadores dos bens patrimoniais utilizar adequadamente os equipamentos e materiais, bem como adotar e propor providências que preservem a segurança e conservação dos bens móveis.
        
        Art. 6°. Em caso de furto, roubo ou danos aos bens públicos, a Secretaria Responsável deve providenciar o Boletim de Ocorrência (B.O) relatando à autoridade policial o ocorrido, contendo os números dos bens patrimoniais furtados/roubados.
                            """
                            pdf.set_font("Arial", "", 10)
                            texto_limpo = texto_legislacao.strip().encode('latin-1', 'replace').decode('latin-1')
                            pdf.multi_cell(0, 6, texto_limpo)
        
                            # ---------------------------------------------------------
                            # 4. CATÁLOGO FOTOGRÁFICO
                            # ---------------------------------------------------------
                            pdf.add_page()
                            pdf.set_font("Arial", "B", 16)
                            pdf.cell(0, 10, "3. CATÁLOGO FOTOGRÁFICO DOS BENS".encode('latin-1', 'replace').decode('latin-1'), ln=True, align='L')
                            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                            pdf.ln(5)
        
                            import tempfile
                            import os
                            
                            altura_card = 48
                            espacamento = 5
                            
                            for i, row in df_foto.iterrows():
                                codigo = str(row['codigo'])
                                nome = str(row.get('nome', '')).encode('latin-1', 'replace').decode('latin-1')
                                estado = str(row.get('estado', 'Não informado')).encode('latin-1', 'replace').decode('latin-1')
                                loc = str(row.get('localizacao', 'Não informado')).encode('latin-1', 'replace').decode('latin-1')
                                obs = str(row.get('observacao', '')).encode('latin-1', 'replace').decode('latin-1')
                                if obs in ['None', 'nan', '']: obs = "Nenhuma observação registrada."
                                
                                if pdf.get_y() + altura_card > 280:
                                    pdf.add_page()
                                
                                y_inicial = pdf.get_y()
                                
                                # Desenha o Card
                                pdf.set_fill_color(248, 249, 250)
                                pdf.set_draw_color(220, 220, 220)
                                pdf.rect(10, y_inicial, 190, altura_card, 'DF')
                                
                                pos_x_img = 12
                                pos_y_img = y_inicial + 2
                                img_size = 44
                                
                                if row.get('foto_base64') and pd.notnull(row['foto_base64']) and str(row['foto_base64']).strip() != "":
                                    try:
                                        img_data = base64.b64decode(row['foto_base64'])
                                        img = Image.open(io.BytesIO(img_data))
                                        
                                        if img.mode != 'RGB': 
                                            img = img.convert('RGB')
                                        
                                        img_cortada = ImageOps.fit(img, (400, 400), Image.Resampling.LANCZOS)
                                        
                                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                                            img_cortada.save(tmp_file, format="JPEG", quality=90)
                                            tmp_path = tmp_file.name
                                        
                                        pdf.image(tmp_path, x=pos_x_img, y=pos_y_img, w=img_size, h=img_size)
                                        os.unlink(tmp_path)
                                    except Exception:
                                        pdf.set_xy(pos_x_img, pos_y_img + 18)
                                        pdf.set_font("Arial", "I", 8)
                                        pdf.cell(img_size, 5, "[ERRO]", align='C')
                                else:
                                    pdf.set_fill_color(230, 230, 230)
                                    pdf.rect(pos_x_img, pos_y_img, img_size, img_size, 'F')
                                    pdf.set_xy(pos_x_img, pos_y_img + 18)
                                    pdf.set_font("Arial", "I", 8)
                                    pdf.set_text_color(150, 150, 150)
                                    pdf.cell(img_size, 5, "[SEM FOTO]", align='C')
                                    pdf.set_text_color(0, 0, 0)
                                
                                pos_x_texto = pos_x_img + img_size + 8
                                
                                pdf.set_xy(pos_x_texto, y_inicial + 4)
                                pdf.set_font("Arial", "B", 12)
                                nome_formatado = f"{codigo} - {nome}"
                                if len(nome_formatado) > 55: nome_formatado = nome_formatado[:52] + "..."
                                pdf.cell(130, 6, nome_formatado, ln=True)
                                
                                pdf.set_xy(pos_x_texto, pdf.get_y() + 2)
                                pdf.set_font("Arial", "B", 9)
                                pdf.set_text_color(80, 80, 80)
                                pdf.cell(15, 5, "Local:")
                                pdf.set_font("Arial", "", 9)
                                pdf.set_text_color(0, 0, 0)
                                pdf.cell(60, 5, loc)
                                
                                pdf.set_font("Arial", "B", 9)
                                pdf.set_text_color(80, 80, 80)
                                pdf.cell(15, 5, "Status:")
                                pdf.set_font("Arial", "", 9)
                                pdf.set_text_color(0, 0, 0)
                                pdf.cell(40, 5, estado, ln=True)
                                
                                pdf.set_xy(pos_x_texto, pdf.get_y() + 1)
                                pdf.set_font("Arial", "B", 9)
                                pdf.set_text_color(80, 80, 80)
                                pdf.cell(15, 5, "Obs:")
                                pdf.set_font("Arial", "", 8)
                                pdf.set_text_color(50, 50, 50)
                                
                                if len(obs) > 130: obs = obs[:127] + "..."
                                pdf.multi_cell(115, 4, obs)
                                
                                pdf.set_y(y_inicial + altura_card + espacamento)
                                
                            st.session_state.pdf_foto_final = get_pdf_bytes(pdf)
                            
                    if 'pdf_foto_final' in st.session_state:
                        st.success("Relatório gerado com sucesso!")
                        st.download_button(
                            label="📥 Baixar Relatório Oficial (PDF)", 
                            data=st.session_state.pdf_foto_final, 
                            file_name=f"Relatorio_Patrimonio_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf", 
                            mime="application/pdf", 
                            use_container_width=True
                        )
        
                    else:
                        st.info("Nenhum bem cadastrado no inventário ainda para gerar relatórios.")
